"""Generate the canonical /00_Copilot_Notebook_Demo_Guide.docx for the Zava hub.

Modeled on the PERSOL APAC Copilot Notebook Demo Guide pattern but generalized
for any of the 55 entries (43 industries + 12 departments). Each entry's
Notebook block on the hub references this file as the FIRST source so that the
Notebook itself contains the demo run-sheet alongside the entry's own data.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "files/00_Copilot_Notebook_Demo_Guide.docx"

doc = Document()

# Title page
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ZAVA CONGLOMERATE")
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F,0x3A,0x68)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Copilot Notebook — Demo Run-Sheet")
r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F,0x3A,0x68)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("One demo guide that works for every entry on the hub — Industry, Department, or General")
r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = RGBColor(0x4B,0x5C,0x77)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Internal demo asset · Reference only")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x6B,0x73,0x80)

doc.add_paragraph()

# How to use
h = doc.add_heading("📋 How to use this guide", level=2)
doc.add_paragraph(
    "This document is your complete presenter's run-sheet for any Copilot Notebook demo "
    "on the Zava hub. Pick any industry or department entry, open its detail page, click "
    "the 📓 Notebook tab — and follow the 5 sections below in order. Total demo run-time: "
    "20–30 minutes per entry."
).runs[0].font.size = Pt(11)

doc.add_paragraph(
    "✅  Sections 1–3 are pre-demo setup. Section 4 contains the typed prompts (run "
    "in sequence — each builds on the previous). Section 5 is the Quick Create finale "
    "(Page · Audio Overview · Word/PPT Quick Create) — one click each, no typing."
).runs[0].font.size = Pt(11)

# Section 1
doc.add_heading("SECTION 1 — What you need before you start", level=1)

doc.add_heading("Licence requirement", level=3)
p = doc.add_paragraph(
    "Copilot Notebook requires a Microsoft 365 Copilot licence. It is NOT available on "
    "standard Microsoft 365 without the Copilot add-on."
); p.runs[0].font.size = Pt(11)
for line in [
    "✅  Microsoft 365 Copilot (enterprise) — full access, recommended for demo",
    "✅  Microsoft Copilot Pro (personal/SMB) — full access",
    "❌  Microsoft 365 without Copilot add-on — Notebooks feature not available",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "⚠️  Check before the demo:  Confirm the presenting account has a Copilot licence at "
    "https://microsoft365.com. If 'Notebooks' does not appear in the left-hand "
    "navigation, the licence is missing — do not proceed without resolving this first."
).runs[0].font.size = Pt(10)

doc.add_heading("Files for this demo", level=3)
p = doc.add_paragraph(
    "Each Zava entry on the hub lists its own reference files in the Notebook setup "
    "block at the top of the 📓 Notebook tab. The pattern is always the same:"
); p.runs[0].font.size = Pt(11)
for line in [
    "1 Demo Guide (this file) — explains the flow, the Instructions field, and the prompts",
    "3–6 entry-specific data files (.xlsx) — the financial / operational tables to analyse",
    "1–3 entry-specific policy / framework files (.docx) — the governance context",
    "Total per Notebook: 5–10 files (limit is 20 files per notebook, 150 MB per file)",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "ℹ️  The hub renders the exact list of files for each entry — you do NOT have to "
    "guess. Open the entry → 📓 Notebook tab → the green 'Notebook setup' panel at the "
    "top lists every file the demo will reference. Upload all of them in one go before "
    "you start the prompt sequence."
).runs[0].font.size = Pt(10)

# Section 2
doc.add_heading("SECTION 2 — How to set up the notebook", level=1)

doc.add_heading("Step 1 — Create the notebook", level=3)
for i, line in enumerate([
    "Go to https://m365.cloud.microsoft and sign in with your Copilot-licensed account",
    "In the left navigation panel, click Notebooks → All Notebooks",
    "Click + New Notebook",
    "Name it after the entry you are demoing (e.g. 'Zava Banking — Group CFO War-Room')",
    "Click Create",
], start=1):
    p = doc.add_paragraph(f"{i}. {line}"); p.runs[0].font.size = Pt(11)

doc.add_heading("Step 2 — Add ALL files as references in one go", level=3)
p = doc.add_paragraph(
    "Method A — Upload directly (recommended for demo reliability):"
); p.runs[0].font.size = Pt(11); p.runs[0].font.bold = True
for line in [
    "In the open notebook, click Add references or the + icon in the Sources panel",
    "Select Upload file",
    "Upload this Demo Guide (00_Copilot_Notebook_Demo_Guide.docx) PLUS every entry-specific file listed in the green 'Notebook setup' panel on the hub — drag-drop them all at once",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
p = doc.add_paragraph("Method B — Link from OneDrive / SharePoint:"); p.runs[0].font.size = Pt(11); p.runs[0].font.bold = True
for line in [
    "Click Add references → Files → OneDrive (or SharePoint)",
    "Navigate to the folder where you stored the entry's files",
    "Select all files (including this guide) and click Add",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "✅  Verify sources loaded:  After adding, you should see all files listed in the "
    "Sources panel of the notebook. Click each file name to confirm it shows a preview "
    "before running any prompts."
).runs[0].font.size = Pt(10)

doc.add_heading("Step 3 — Paste the Copilot Instructions (critical — do not skip)", level=3)
p = doc.add_paragraph(
    "The Instructions field is a persistent system prompt — it tells Copilot its role "
    "and context for every query in this notebook. Without this, responses will be "
    "generic and uncalibrated."
); p.runs[0].font.size = Pt(11)
for line in [
    "Locate the Instructions field at the top of the notebook (may appear as 'Custom instructions' or behind a ⚙️ gear icon)",
    "Click to expand it",
    "Copy the Instructions text shown in the green 'Notebook setup' panel of the entry on the hub — the text is bespoke per entry (e.g. Banking CFO War-Room, ESG Disclosure Reviewer, BPK Audit Closure Lead, etc.)",
    "Paste into the Instructions field and click Save / Apply before proceeding",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "ℹ️  Why per-entry Instructions matter:  An IR Disclosure Reviewer Notebook needs "
    "different system framing than a Procurement Compliance Reviewer Notebook. The hub "
    "supplies entry-specific Instructions so Copilot answers in the right register, "
    "with the right regulators, the right job titles, and the right escalation paths."
).runs[0].font.size = Pt(10)

doc.add_heading("Step 4 — Pre-demo connectivity test", level=3)
p = doc.add_paragraph(
    "Before the live session, run this test prompt to verify the files loaded correctly:"
); p.runs[0].font.size = Pt(11)
p = doc.add_paragraph(
    "\"Connectivity test. List every file currently loaded in this notebook with its file "
    "type and the number of pages or sheets. Then in 1 sentence each, summarise what each "
    "file contains. Flag any file you cannot read.\""
); p.runs[0].font.size = Pt(10); p.runs[0].font.italic = True
doc.add_paragraph(
    "Expected: a numbered list of every file you uploaded with a short summary line each. "
    "If any file is missing or unreadable, re-upload it before the demo starts."
).runs[0].font.size = Pt(10)

# Section 3
doc.add_heading("SECTION 3 — Demo narrative & framing", level=1)
p = doc.add_paragraph(
    "Before typing any prompt, set the scene for the audience with this 60-second verbal "
    "introduction (adapt to the entry you are demoing):"
); p.runs[0].font.size = Pt(11)
p = doc.add_paragraph(
    "💬 Say this to open:  'I have just loaded a working pack into Copilot Notebook — "
    "it is the same pack a senior leader in this team would receive for a real war-room "
    "session. Several spreadsheets with the actual numbers, two or three policy / "
    "framework documents, and a system instruction that sets Copilot's role. Normally a "
    "team would spend a full day reading all of this before they could even start "
    "analysis. Watch what Copilot does with the same pack in 5 minutes.'"
); p.runs[0].font.size = Pt(11)
p = doc.add_paragraph(
    "Then open one of the entry-specific data spreadsheets on screen and scroll through "
    "the rows so the audience can see the raw data volume. Then say:"
); p.runs[0].font.size = Pt(11)
p = doc.add_paragraph(
    "💬 Then say:  'This is one source. There are several of them. Watch what happens "
    "when I give all of this to Copilot Notebook — it reads every document together "
    "and gives me the answer in under a minute. And because I set the Instructions field, "
    "Copilot answers in the voice of the senior role I am demoing — not as a generic AI.'"
); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "💡 Demo pacing tip:  Spend 30 seconds reading each Copilot response aloud and "
    "pointing out 2–3 specific data points before moving to the next prompt. The audience "
    "needs time to absorb what they see. Do not rush."
).runs[0].font.size = Pt(10)

# Section 4
doc.add_heading("SECTION 4 — Demo prompts (ready to copy-paste)", level=1)
p = doc.add_paragraph(
    "The hub renders the exact prompts for each entry on the 📓 Notebook tab. Each entry "
    "ships with 3–7 typed prompts that build on each other, plus 3 Quick Create buttons "
    "for the finale. The pattern is always the same:"
); p.runs[0].font.size = Pt(11)
for line in [
    "PROMPT 1 — Synthesis: read every source together and produce a structured brief",
    "PROMPT 2 — Comparison / scoring: rank, score, or stack-rank the entities in the data",
    "PROMPT 3 — Recommendation: take a position with evidence-based reasoning",
    "PROMPT 4 — Quick Create Page: turn the analysis into a shareable Copilot Page",
    "PROMPT 5 — Quick Create Word doc: generate the formal write-up",
    "PROMPT 6 — Quick Create PPT: generate the steering committee deck",
    "PROMPT 7 — Quick Create Audio Overview: 60–90 second podcast walkthrough",
]:
    p = doc.add_paragraph(line, style='List Bullet'); p.runs[0].font.size = Pt(11)
doc.add_paragraph(
    "⚠️  If Copilot returns an error or incomplete response:  Scope the prompt down "
    "further. Say: 'Look at file A and file B only. I will ask about the others "
    "separately.' Then combine the findings yourself. Large packs across many sources "
    "can hit Copilot's context limit in a single query."
).runs[0].font.size = Pt(10)
doc.add_paragraph(
    "🔍 Quick Create finale — no typing needed:  Use the Quick Create bar at the top "
    "of the notebook. Everything generated here is grounded in the uploaded files AND "
    "the full prompt conversation above."
).runs[0].font.size = Pt(10)

# Section 5
doc.add_heading("SECTION 5 — Common pitfalls & recovery", level=1)

for title, body in [
    ("Files not loading",
     "If a file shows as unsupported or unreadable, the most common cause is file size "
     "(over 150 MB) or a corrupted upload. Re-export the file with smaller pages / "
     "sheets and re-upload. Do not proceed without verifying every file is readable."),
    ("Generic, uncalibrated answers",
     "If Copilot gives generic responses, you skipped the Instructions field. Stop the "
     "demo, paste the entry-specific Instructions text from the hub, click Save, and "
     "re-run the prompt. Do not try to fix this with prompt engineering — set the "
     "Instructions field instead."),
    ("Context-limit errors on large packs",
     "Scope down: 'Look at FILE_A and FILE_B only. I will ask about FILE_C and FILE_D in "
     "the next prompt.' Then synthesise the findings yourself. Notebook can read all "
     "files but a single prompt cannot always reason across all of them at once."),
    ("Answers that feel hallucinated",
     "Add: 'Cite the source file name and tab/section after every claim. Do not "
     "speculate beyond the sources.' to the prompt. The Instructions field on the hub "
     "already includes this directive — re-paste it if responses drift."),
    ("Numbers don't match the spreadsheet",
     "Quote the file and tab name explicitly: 'In the Quarterly Earnings Tracker tab "
     "of /01_Zava_Group_Financial_Performance.xlsx, read the row labelled FY2025 Q4 "
     "EBITDA.' Specificity always improves accuracy."),
]:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph(body); p.runs[0].font.size = Pt(11)

# Closing
doc.add_heading("🏁 Closing line", level=2)
p = doc.add_paragraph(
    "'In under 30 minutes: an entire pack read, every document analysed, every gap "
    "surfaced, a steering committee deck produced, a Word write-up drafted, and an audio "
    "briefing for stakeholders who weren't in the room. The team spends their time on "
    "the decision — not the data processing.'"
); p.runs[0].font.size = Pt(11); p.runs[0].font.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— END OF DEMO GUIDE —")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x6B,0x73,0x80)

doc.save(OUT)
print(f"Wrote {OUT}")
