"""Generate the 316 missing industry-specific reference files referenced in data.js.

For each entry in data.js:
  - Read its id, name, tagline, scenario, files[]
  - For each file in files[]:
      .xlsx -> 3-sheet workbook (KPIs + 90-day plan + watchlist) tied to entry context
      .docx -> 4-section document (Executive Summary + Background + Detail + Actions) tied to entry context
  - Skip files that already exist in files/
  - Write to files/

Files are unique-per-entry (no copy-paste): the workbook headers, KPIs, and rows
all derive from the entry's name + tagline. The document paragraphs reference the
entry scenario.

Run: python gen_industry_files.py
"""

import re, os, json, random
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.abspath(__file__))
FILES_DIR = os.path.join(ROOT, 'files')
os.makedirs(FILES_DIR, exist_ok=True)

# ---- Parse data.js to extract per-entry context ------------------------------
data_js = open(os.path.join(ROOT, 'data.js'), 'r', encoding='utf-8').read()

def _parse_entries():
    """Return list of dicts {id, name, tagline, scenario, files[]}."""
    out = []
    # Match each top-level entry block
    entry_pat = re.compile(
        r"id:\s*'([\w-]+)',\s*\n\s*sectorId:\s*'[^']*',[^}]*?name:\s*'((?:[^'\\]|\\.)+)'",
        re.DOTALL)
    hits = list(entry_pat.finditer(data_js))
    for i, m in enumerate(hits):
        eid = m.group(1)
        name = m.group(2)
        s = m.start()
        e = hits[i+1].start() if i+1 < len(hits) else len(data_js)
        block = data_js[s:e]
        # tagline
        tag = re.search(r"tagline:\s*'((?:[^'\\]|\\.)*)'", block)
        tagline = tag.group(1).replace("\\'", "'") if tag else ''
        # scenario
        scn = re.search(r"scenario:\s*'((?:[^'\\]|\\.)*)'", block)
        scenario = scn.group(1).replace("\\'", "'") if scn else ''
        # files
        files = re.findall(r'(\b[A-Z]{2,5}_\d{2}_[A-Za-z0-9_&-]+\.(?:xlsx|docx|pdf))', block)
        out.append({
            'id': eid, 'name': name, 'tagline': tagline,
            'scenario': scenario, 'files': sorted(set(files))
        })
    return out

# ---- Workbook builder --------------------------------------------------------

HEADER_FILL = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
HEADER_FONT = Font(bold=True, color='FFFFFF', size=11)
TITLE_FONT  = Font(bold=True, color='1F4E78', size=14)

def _stem_topic(stem_lo):
    """Map filename stem (lowercased, underscored) to a workbook topic.
    Returns (sheet_name_a, sheet_name_b, sheet_name_c, kpi_name)."""
    # Heuristic mapping based on filename keywords
    if any(w in stem_lo for w in ['production','operations','metrics','oee','plant']):
        return ('Plant Metrics', 'Recovery Bridge', 'Watch List', 'Plant OEE')
    if any(w in stem_lo for w in ['financial','financials','financ','pl_','pnl','p_l','pl-']):
        return ('P&L Summary', 'Variance Detail', 'Recovery Plan', 'EBITDA Margin')
    if any(w in stem_lo for w in ['cashflow','cash_','liquidity']):
        return ('Cashflow', 'Working Capital', 'Watchlist', 'Free Cash Flow')
    if any(w in stem_lo for w in ['cost','opex','cost_']):
        return ('Cost Base', 'Variance Detail', 'Reset Plan', 'Cost-to-Income')
    if any(w in stem_lo for w in ['revenue','sales','top_line']):
        return ('Revenue Mix', 'Customer Cohort', 'Pipeline', 'Revenue Growth')
    if any(w in stem_lo for w in ['fraud','risk','npl','exposure']):
        return ('Exposure Snapshot', 'Loss Bridge', 'Watch List', 'Risk Index')
    if any(w in stem_lo for w in ['portfolio','asset','aum']):
        return ('Portfolio Snapshot', 'Performance Bridge', 'Watch List', 'AUM')
    if any(w in stem_lo for w in ['hr','attrition','headcount','workforce','people']):
        return ('Headcount', 'Attrition Bridge', 'Hot-Spots', 'Attrition %')
    if any(w in stem_lo for w in ['budget','absorption','spending']):
        return ('Programme Spend', 'Absorption Bridge', 'Programme Status', 'Absorption %')
    if any(w in stem_lo for w in ['compliance','audit','tracker','findings']):
        return ('Findings', 'Closure Plan', 'Owner Tracker', 'Open Findings')
    if any(w in stem_lo for w in ['route','flight','load','traffic']):
        return ('Route P&L', 'Load Factor', 'Watch List', 'Route Margin')
    if any(w in stem_lo for w in ['property','rent','occupancy','tenant']):
        return ('Property Returns', 'Occupancy Bridge', 'Tenant Watchlist', 'NOI Yield')
    if any(w in stem_lo for w in ['fleet','vessel','ship']):
        return ('Fleet Snapshot', 'Voyage P&L', 'Drydock Plan', 'Fleet Utilisation')
    if any(w in stem_lo for w in ['oee','quality','reject']):
        return ('Plant OEE', 'Quality Pareto', 'Watch List', 'OEE')
    if any(w in stem_lo for w in ['investor','ir_','disclosure']):
        return ('Holder Activity', 'Q&A Log', 'Disclosure Tracker', 'Active Holders')
    # default
    return ('Snapshot', 'Bridge', 'Watch List', 'Index')

def _industry_columns(name, prefix_lo):
    """Pick a few industry-flavoured KPI columns based on entry name."""
    n = name.lower()
    if 'bank' in n or prefix_lo == 'bnk':
        return ['NIM %', 'CIR %', 'NPL %', 'CET1 %', 'ROE %']
    if 'aviation' in n or 'airline' in n or 'air' == prefix_lo or prefix_lo == 'avn':
        return ['Load Factor %', 'CASK', 'RASK', 'On-time %', 'Fuel Hedge %']
    if 'hospital' in n or 'health' in n or prefix_lo in ('hc',):
        return ['Bed Occupancy %', 'ALOS days', 'Margin %', 'Patient NPS', 'Pharmacy %']
    if 'property' in n or 'real estate' in n or prefix_lo in ('prop','re'):
        return ['Occupancy %', 'NOI Yield %', 'Rent Growth %', 'Capex Ratio %', 'WALE yrs']
    if 'retail' in n or 'fmcg' in n or prefix_lo in ('rt','fmcg'):
        return ['SSSG %', 'Gross Margin %', 'Stock Turn x', 'Footfall', 'Basket Size']
    if 'plantation' in n or prefix_lo == 'plt':
        return ['CPO MT', 'Yield t/ha', 'Cost/MT', 'OER %', 'KER %']
    if 'media' in n or 'telco' in n or prefix_lo in ('me','tc'):
        return ['ARPU', 'Subs k', 'Churn %', 'EBITDA Margin', 'Capex/Rev %']
    if 'insur' in n or prefix_lo in ('li','gi','ib'):
        return ['Combined Ratio %', 'Loss Ratio %', 'Expense Ratio %', 'GWP', 'Embedded Value']
    if 'fintech' in n or 'remit' in n or prefix_lo in ('ft','remit'):
        return ['TPV', 'Fraud bps', 'Take Rate %', 'MAU k', 'GMV']
    if 'logistic' in n or 'shipping' in n or prefix_lo in ('log',):
        return ['On-time %', 'Cost/TEU', 'Fleet Util %', 'Empty Mile %', 'OTIF %']
    if 'manufactur' in n or 'tyre' in n or 'auto' in n or 'glove' in n or 'semi' in n or prefix_lo in ('mfg','semi','autom','auto','glove'):
        return ['OEE %', 'Yield %', 'Reject ppm', 'Capacity Util %', 'Cost/unit']
    if 'oil' in n or 'gas' in n or prefix_lo in ('ogd','ogu'):
        return ['Production kbpd', 'Refining Margin', 'OEE %', 'Maintenance Days', 'Reserves yrs']
    if 'rare' in n or 'mining' in n or prefix_lo in ('coal','re'):
        return ['Production kt', 'Cash Cost/t', 'Strip Ratio', 'Recovery %', 'Reserves yrs']
    if 'bpo' in n or prefix_lo == 'bpo':
        return ['Utilisation %', 'CSAT', 'Attrition %', 'Revenue/FTE', 'Margin %']
    if 'government' in n or 'glc' in n or prefix_lo in ('gov','glc','reg'):
        return ['Absorption %', 'Findings Open', 'KPI Met %', 'Citizen NPS', 'Programme Status']
    # default
    return ['KPI 1', 'KPI 2', 'KPI 3', 'KPI 4', 'KPI 5']

def _make_xlsx(path, entry, fname):
    wb = Workbook()
    stem = fname.replace('.xlsx', '')
    parts = stem.split('_', 2)
    prefix = parts[0]
    title_part = parts[2].replace('_', ' ') if len(parts) >= 3 else stem
    sa, sb, sc, kpi = _stem_topic(stem.lower())
    cols = _industry_columns(entry['name'], prefix.lower())

    # --- Sheet A: KPI snapshot ---
    ws = wb.active
    ws.title = sa
    ws['A1'] = f"{entry['name']} — {title_part}"
    ws['A1'].font = TITLE_FONT
    ws.merge_cells('A1:G1')
    ws['A2'] = entry['tagline'][:120] if entry['tagline'] else ''
    ws['A2'].font = Font(italic=True, color='555555', size=10)
    ws.merge_cells('A2:G2')
    headers = ['Unit / Segment', 'FY24 Actual'] + cols
    for j, h in enumerate(headers, 1):
        c = ws.cell(row=4, column=j, value=h)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal='center')
    rng = list(range(5, 13))
    units = [
        f"{entry['name']} — Division 1", f"{entry['name']} — Division 2",
        f"{entry['name']} — Division 3", f"{entry['name']} — Division 4",
        f"{entry['name']} — Division 5", f"{entry['name']} — Division 6",
        'Group Total', 'Variance vs Budget'
    ]
    rng_seed = sum(ord(c) for c in entry['id']) + sum(ord(c) for c in fname)
    rnd = random.Random(rng_seed)
    for r, u in zip(rng, units):
        ws.cell(row=r, column=1, value=u).font = Font(bold=(u in ('Group Total','Variance vs Budget')))
        ws.cell(row=r, column=2, value=round(rnd.uniform(120, 980), 1))
        for k in range(len(cols)):
            ws.cell(row=r, column=3+k, value=round(rnd.uniform(0.4, 99.0), 1))
    ws.column_dimensions['A'].width = 36
    for j in range(2, len(headers)+1):
        ws.column_dimensions[get_column_letter(j)].width = 14

    # --- Sheet B: Recovery / variance bridge ---
    wb.create_sheet(sb)
    ws2 = wb[sb]
    ws2['A1'] = f"{title_part} — {sb}"
    ws2['A1'].font = TITLE_FONT
    ws2.merge_cells('A1:E1')
    headers2 = ['Driver', 'FY25 Impact', 'FY26 Plan', 'Owner', 'Status']
    for j, h in enumerate(headers2, 1):
        c = ws2.cell(row=3, column=j, value=h)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
    drivers = [
        ('Volume / Mix', 'Pricing & Discount', 'Cost Base Reset', 'Capex Re-baseline',
         'Working Capital', 'Regulatory Provisions', 'FX & Hedging',
         'Productivity / Automation')
    ][0]
    owners = ['CFO', 'COO', 'CRO', 'CSO', 'CTO']
    statuses = ['On track', 'At risk', 'Behind', 'Recovered', 'Mitigated']
    for i, d in enumerate(drivers, start=4):
        ws2.cell(row=i, column=1, value=d)
        ws2.cell(row=i, column=2, value=round(rnd.uniform(-180, 240), 1))
        ws2.cell(row=i, column=3, value=round(rnd.uniform(50, 300), 1))
        ws2.cell(row=i, column=4, value=rnd.choice(owners))
        ws2.cell(row=i, column=5, value=rnd.choice(statuses))
    for j, w in enumerate([28, 14, 14, 12, 14], start=1):
        ws2.column_dimensions[get_column_letter(j)].width = w

    # --- Sheet C: Watch list ---
    wb.create_sheet(sc)
    ws3 = wb[sc]
    ws3['A1'] = f"{title_part} — {sc}"
    ws3['A1'].font = TITLE_FONT
    ws3.merge_cells('A1:F1')
    headers3 = ['Item', 'Indicator', 'FY25', 'Threshold', 'Trigger?', 'Owner']
    for j, h in enumerate(headers3, 1):
        c = ws3.cell(row=3, column=j, value=h)
        c.font = HEADER_FONT; c.fill = HEADER_FILL
    items = [
        f"{entry['name']} top customer", f"{entry['name']} top vendor",
        f"{entry['name']} flagship product", f"{entry['name']} regulator pack",
        f"{entry['name']} largest contract", f"{entry['name']} board KPI #1",
        f"{entry['name']} board KPI #2", f"{entry['name']} ESG signal"
    ]
    for i, it in enumerate(items, start=4):
        ws3.cell(row=i, column=1, value=it)
        ws3.cell(row=i, column=2, value=rnd.choice(cols))
        ws3.cell(row=i, column=3, value=round(rnd.uniform(40, 95), 1))
        ws3.cell(row=i, column=4, value=round(rnd.uniform(60, 90), 1))
        ws3.cell(row=i, column=5, value=rnd.choice(['No','Yes','Watch']))
        ws3.cell(row=i, column=6, value=rnd.choice(owners))
    for j, w in enumerate([34, 14, 10, 12, 10, 12], start=1):
        ws3.column_dimensions[get_column_letter(j)].width = w

    wb.save(path)


# ---- Document builder --------------------------------------------------------

def _make_docx(path, entry, fname):
    doc = Document()
    stem = fname.replace('.docx', '')
    parts = stem.split('_', 2)
    prefix = parts[0]
    title_part = parts[2].replace('_', ' ') if len(parts) >= 3 else stem
    title = f"{entry['name']} — {title_part}"
    h = doc.add_heading(title, level=0)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p = doc.add_paragraph()
    p.add_run(f"Document owner: {entry['name']} Group Office").italic = True
    p.add_run(f" · Version: FY2025-A").italic = True

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        entry['tagline'] or
        f"This document summarises the {title_part.lower()} for {entry['name']}.")
    doc.add_paragraph(
        f"This {title_part.lower()} is grounded on the FY2025 working dataset for "
        f"{entry['name']} and informs the Board, the regulator, and the operating "
        f"committee of the corrective actions agreed under the recovery plan.")

    doc.add_heading('2. Scenario & Background', level=1)
    scn = entry['scenario'] or (entry['tagline'] + '.')
    doc.add_paragraph(scn)
    doc.add_paragraph(
        f"The team has triangulated {entry['name']}'s most recent reported "
        f"performance against three peer benchmarks and against the regulator's "
        f"published thresholds for the sector. Findings carried over from the "
        f"prior cycle have been re-examined to test whether the original "
        f"assumptions still hold.")

    doc.add_heading('3. Detail & Working', level=1)
    sub_heads = [
        '3.1 Position at end of period',
        '3.2 Key drivers',
        '3.3 Sensitivities and assumptions',
        '3.4 Open issues and unknowns',
    ]
    for sh in sub_heads:
        doc.add_heading(sh, level=2)
        doc.add_paragraph(
            f"{sh.split(' ',1)[1]} for {entry['name']} as it stands on the "
            f"working dataset for this {title_part.lower()}. Reference figures "
            f"and supporting tabs are stored alongside this document in the "
            f"Group repository under the {entry['id'].upper()} folder.")

    doc.add_heading('4. Recommended Actions', level=1)
    bullets = [
        f"Confirm the {title_part.lower()} narrative with the {entry['name']} CFO before Board cut-off.",
        f"Lock the assumption set for the recovery plan and circulate to division MDs.",
        f"Pre-brief the regulator-facing officer on the 30/60/90-day commitments.",
        f"Schedule the cross-functional review with risk, legal, and IR.",
        f"Capture decisions in the Group decision log with named owners and due dates.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('5. Appendix — Source files', level=1)
    for f in entry['files']:
        doc.add_paragraph(f, style='List Bullet')

    doc.save(path)


# ---- Main ---------------------------------------------------------------------

def main():
    entries = _parse_entries()
    print(f'Parsed {len(entries)} entries from data.js')
    existing = set(os.listdir(FILES_DIR))
    generated = 0
    skipped = 0
    errors = []
    for e in entries:
        for fname in e['files']:
            if fname in existing:
                skipped += 1
                continue
            path = os.path.join(FILES_DIR, fname)
            try:
                if fname.endswith('.xlsx'):
                    _make_xlsx(path, e, fname)
                elif fname.endswith('.docx'):
                    _make_docx(path, e, fname)
                elif fname.endswith('.pdf'):
                    # Generate as docx in pdf clothing — write a placeholder docx
                    _make_docx(path.replace('.pdf', '.docx'), e, fname)
                    # then ignore; pdf alt skipped
                    continue
                else:
                    continue
                generated += 1
                if generated % 50 == 0:
                    print(f'  {generated} generated...')
            except Exception as ex:
                errors.append((fname, str(ex)))
    print(f'\nGenerated: {generated}')
    print(f'Skipped (already exist): {skipped}')
    if errors:
        print(f'Errors: {len(errors)}')
        for f, msg in errors[:5]:
            print(f'  {f}: {msg}')

if __name__ == '__main__':
    main()
