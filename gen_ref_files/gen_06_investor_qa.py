"""Generates 06_Zava_Investor_QA_FY2025.docx — 12 sections, ~50-70 anticipated Q&A pairs.

Sections:
  1. Group Performance Overview
  2. EBITDA Miss Drivers
  3. Properties Division Deep-Dive
  4. Healthcare Division Deep-Dive
  5. Retail Division Deep-Dive
  6. Other Divisions (Chem, Agri, Mfg, FS, BPO, Trading, Pharma, Treasury) — combined
  7. Capital Allocation & Capex
  8. Dividend Policy & Distributions
  9. M&A Pipeline
 10. ESG & Climate Disclosures
 11. Regulatory & Leadership
 12. FY2026 Guidance & Outlook
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gen_ref_files.common import DIVISIONS, FY25, fy25_eb

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)


def H(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20 if level == 1 else 14 if level == 2 else 11)
    run.font.name = 'Aptos'
    run.font.color.rgb = NAVY


def P(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    if color:
        run.font.color.rgb = color


def Q(doc, num, question, answer, source=None):
    """Render a Q&A pair with question in bold, answer below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'Q{num}.  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = 'Aptos'
    run = p.add_run(question)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run('A.  ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    run.font.name = 'Aptos'
    run = p.add_run(answer)
    run.font.size = Pt(10)
    run.font.name = 'Aptos'

    if source:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f'Source / Reference: {source}')
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = GREY
        run.font.name = 'Aptos'


# ── Cover ──────────────────────────────────────────────────────────────────
def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    run = p.add_run('ZAVA GROUP\n')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'
    run = p.add_run('FY2025 Investor Q&A Anticipation Pack\n')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nFor use by IR Team only — Pre-results briefing pack\n'
                    'Anticipated questions from sell-side analysts and top-30 institutional holders\n'
                    'Source: Group IR with input from Group FP&A and Division CFOs')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = 'Aptos'

    doc.add_page_break()


# ── Section 1 ──────────────────────────────────────────────────────────────
def s1(doc):
    H(doc, '1. Group Performance Overview')
    P(doc, 'This section addresses the headline FY2025 outcome, the bridge from the FY2024 base, '
           'and the sequencing of the recovery trajectory.', italic=True)

    Q(doc, 1, 'Group revenue MYR 42.4B vs MYR 44.2B budget — the 4.1% top-line miss looks contained, but the EBITDA miss is 18%. Why the operating-leverage drag?',
      'Three reasons. First, margin compression in Properties, Healthcare and Retail was disproportionate to revenue softness — the negative-EBITDA outcome in those three divisions is the single largest driver. Second, fixed-cost absorption deteriorated where utilisation slipped (Subang Healthcare at 58%, Penang chemical plant downtime). Third, IDR/USD adverse FX hit reported EBITDA by ~MYR 195M of pure translation. Net of these three buckets we account for ~MYR 1.1B of the MYR 1.3B EBITDA gap; the remaining ~MYR 200M reflects the Hindustan Bulk Trading-division event (~MYR 65M) plus smaller one-offs.',
      'Group P&L Summary tab; Variance Drill-Down tab')

    Q(doc, 2, 'Is the FY2025 outcome a structural break or a cyclical air-pocket?',
      'Eight of eleven divisions are within ±5% of their plan EBITDA. The shortfall is concentrated in three specific divisions, each with identifiable, mostly cyclical drivers — Mid-Valley delay, Penang Tower writedown, post-pandemic patient-volume return, retail footfall correction. We see this as an acute cyclical event with structural overlay (online/format pressure in Retail, specialist mix in Healthcare). The recovery path is anchored on division-specific actions, not on hoping the cycle turns.',
      None)

    Q(doc, 3, 'How does FY2025 compare to your last comparable down-cycle?',
      'In FY2020 (COVID year) Group EBITDA fell 22% YoY, recovering 28% in FY2021. The FY2025 EBITDA decline of 22% (vs FY2024) is the same order of magnitude. The FY2026 recovery is not expected to be as sharp — we are guiding 12-15% EBITDA recovery — because the property and retail formats need 18-24 months of execution to reset.',
      None)

    Q(doc, 4, 'What is the impact on dividends?',
      'The Board has resolved to maintain the FY2025 final DPS at MYR 0.18 per share, in line with the dividend-policy target payout ratio of 40-60% of normalised PATAMI. The full-year DPS is therefore MYR 0.30. We did not pull from the dividend reserve; FY2025 PATAMI of MYR 2.1B comfortably covers the MYR 1.05B distribution.',
      'Dividend policy: §8 of this document')

    Q(doc, 5, 'Has the FY2025 outcome triggered any covenant tightness?',
      'Net Debt / EBITDA stands at 3.1× at year-end, against the most-restrictive covenant of 3.5×. We have approximately 11% headroom, in the Amber zone. We have proactively engaged all 18 lenders; no waiver has been required. The L11 (UOB MYR Term Loan) was the only facility where we sought a one-time waiver in February 2024 for the Penang Tower writedown specifically; that has been closed. We have a quarterly call scheduled with the lender group in Q1 FY2026.',
      'Lender Covenant Tracker; Waiver History tab')

    Q(doc, 6, 'How do we read free cash flow this year?',
      'FCF of MYR 2.18B is below the MYR 3.25B plan. The shortfall is ~75% EBITDA-driven and ~25% working-capital-driven (longer customer terms in Trading, inventory build-up in Retail and Properties). We expect working capital to release ~MYR 400M in H1 FY2026 as inventory normalises.',
      None)
    doc.add_page_break()


# ── Section 2 ──────────────────────────────────────────────────────────────
def s2(doc):
    H(doc, '2. EBITDA Miss Drivers')
    P(doc, 'Detailed view on the MYR 1.3B EBITDA gap.', italic=True)

    Q(doc, 7, 'Can you walk us through the EBITDA bridge from budget to actual?',
      'The bridge is approximately: Properties contribution shortfall MYR 305M; Healthcare shortfall MYR 195M; Retail shortfall MYR 245M; Chemicals shortfall MYR 95M; Manufacturing shortfall MYR 85M; Agribusiness shortfall MYR 130M; Trading shortfall MYR 65M (Hindustan Bulk one-off); Treasury shortfall MYR 90M (mark-down on investment portfolio); FX translation drag MYR 195M; partially offset by Financial Services beat MYR 65M and BPO beat MYR 35M. Net: MYR 1.317B miss.',
      'Variance Summary tab; Driver Tree tab')

    Q(doc, 8, 'Were there any unbudgeted one-offs that should be normalised out?',
      'Three material one-offs: Penang Tower writedown MYR 320M (in Properties); Mid-Valley delay-related cost overrun MYR 110M; and the Hindustan Bulk counterparty event in Trading MYR 65M. Together MYR 495M. Stripping these, the underlying EBITDA miss would be ~MYR 822M, or ~11% of budget — still a miss, but more reflective of cyclical conditions than structural deterioration.',
      None)

    Q(doc, 9, 'How material is the FX drag and how should we model it for FY2026?',
      'IDR/USD weakened ~9% YoY; this hit our consolidated MYR-reported EBITDA by ~MYR 195M of pure translation in FY2025 (no operational change in IDR-denominated subsidiaries). For FY2026 we are modelling 8.0-8.5% additional IDR weakness in our base case, implying additional translation drag of ~MYR 80-110M. We are progressively expanding our hedging programme but local-currency funding remains the most economical hedge for translation exposure.',
      None)

    Q(doc, 10, 'How concentrated is the EBITDA exposure across divisions?',
      'No single division accounts for more than 19% of Group EBITDA. The top-3 (FS 15%, BPO 14%, Pharma 11%) account for 40% — diversified enough to absorb single-division shocks. The bottom-3 (Properties, Healthcare, Retail) currently subtract from Group EBITDA — this is the inverse problem: not concentration risk, but a need to recover three under-performers simultaneously.',
      None)

    Q(doc, 11, 'How are you confident that Properties, Healthcare and Retail recover by FY2027 rather than continuing to deteriorate?',
      'Each division has a 90-day stabilisation plan now in execution; a 12-month recovery plan tracking against quantified milestones (sell-down rates for Properties, occupancy and case-mix for Healthcare, comp-store sales and online conversion for Retail); and a 3-year transformation plan with capital reallocation. The Group ExCo reviews progress monthly; the Audit & Risk Committee reviews quarterly. The combined FY2026 recovery target is MYR 330M positive EBITDA contribution from these three divisions.',
      'Strategy §3 — Pillar 1; Recovery Levers tab')

    Q(doc, 12, 'Are there scenarios in which you take a strategic-action call (write-down, exit, demerger) on any of the three Red divisions?',
      'We do not see Properties, Healthcare or Retail as candidates for full exit at this point. Each retains meaningful franchise value (Properties land-bank quality; Healthcare specialist network; Retail brand recognition). We have approved selective rationalisation: 18 underperforming retail stores to close; non-core property land-bank disposal of MYR 600M; conversion of underutilised Healthcare floors. A more comprehensive strategic review is built into the standard 24-month portfolio-review cadence — the next is scheduled Q3 FY2026.',
      None)
    doc.add_page_break()


# ── Section 3 ──────────────────────────────────────────────────────────────
def s3(doc):
    H(doc, '3. Properties Division Deep-Dive')
    Q(doc, 13, 'The Mid-Valley Phase 2 completion delay — what is the realistic completion window now?',
      'Phase 2 is now expected to complete by Q3 FY2026, an 8-month delay vs the FY2025 plan. The delay is largely contractor-driven; we have replaced two trade contractors and added a project-recovery team. We have provided for the delay impact in our FY2025 EBITDA. Phase 3 timing remains FY2027 H2.',
      None)

    Q(doc, 14, 'Penang Tower writedown of MYR 320M — what is the residual book value, and what is the recovery path?',
      'Residual book value post-writedown is MYR 480M. We are repositioning the asset from pure-office to mixed-use (office + serviced apartments + retail podium). Concept design is complete; planning approval expected Q2 FY2026. Recovery EBITDA contribution targeted FY2027 onwards. We have stress-tested for a further MYR 80M writedown (1-in-20 scenario) and the capital position absorbs it.',
      None)

    Q(doc, 15, 'How should we read the unsold inventory of 320 units in Klang Valley?',
      'At current sales velocity of ~22 units/month group-wide, this is ~14 months of supply. We have reduced new launches and are accelerating sell-down via bundled financing (in partnership with Group Financial Services arm), targeted price discipline (5-7% reductions on selective stock), and incentive packages. The target is to reduce unsold inventory below 200 units by end-FY2026.',
      None)

    Q(doc, 16, 'Are you considering REIT-isation of any of the property portfolio?',
      'We have evaluated REIT-isation periodically; at current cap-rate and yield environment, a REIT vehicle would be sub-optimal versus continuing to hold and recycle assets directly. We have shortlisted approximately MYR 1.4B of stabilised income-producing assets that could form the basis of a future REIT, but this is not on the FY2026 capital-markets agenda.',
      None)

    Q(doc, 17, 'What is the breakdown of the MYR 600M asset-disposal target?',
      'Three components: (a) MYR 350M of land-bank disposal (3 parcels in Klang Valley, IM out to interested parties); (b) MYR 180M Penang Tower Phase 2 strategic exit if pricing is achieved; (c) MYR 70M legacy retail real-estate. Proceeds will be reinvested into core development pipeline and partly used to delever.',
      'Strategy §4 — Divestment Pipeline')
    doc.add_page_break()


# ── Section 4 ──────────────────────────────────────────────────────────────
def s4(doc):
    H(doc, '4. Healthcare Division Deep-Dive')
    Q(doc, 18, 'Patient volumes are 11% below pre-pandemic — is this the new normal?',
      'Our reading is that ~6-7% of the gap is structural (telemedicine substitution for low-acuity consultations; reduced elective procedure backlog from peak post-pandemic restoration) and ~4-5% is cyclical (cost-of-living-driven postponement of elective procedures and aesthetic services). Specialist procedural volume is more resilient than primary care; this informs our format conversion.',
      None)

    Q(doc, 19, 'The Subang facility at 58% utilisation — what is the format-conversion plan?',
      'We are converting two underutilised inpatient floors to outpatient day-surgery and chronic-care formats. Capex required is MYR 35M; planning approval from MOH targeted Q2 FY2026. Once converted, we expect blended utilisation to lift to 78% by FY2027 and EBITDA contribution to recover MYR 65M.',
      'Recovery Levers tab; Strategy §3.x Healthcare')

    Q(doc, 20, 'How do you view the specialist consultant cost trajectory? It seems to be the single biggest pressure.',
      'Specialist cost grew 7% YoY against revenue growth of 1%. We are renegotiating bundled-procedure contracts vs fee-per-service for high-volume specialties (cardiology, orthopaedics, ophthalmology) — early data shows 3-5% cost reduction with no impact on quality. We are also building junior consultant career paths to reduce reliance on senior locum specialists. Steady-state target: specialist cost/revenue ratio at 22% (FY25 actual: 26%).',
      None)

    Q(doc, 21, 'Is bundled-product partnership with the FS arm just a slogan or has it real economic content?',
      'It has economic content. We are designing a wellness-bundled health-insurance product with the Group FS arm, targeting middle-class urban customers in Malaysia. The bundle includes annual screenings at our hospitals, telemedicine access, and a managed condition path. Pricing modelled at MYR 180-220 per month per family. Soft-launch Q2 FY2026, target 25,000 subscribers by end-FY2027.',
      None)
    doc.add_page_break()


# ── Section 5 ──────────────────────────────────────────────────────────────
def s5(doc):
    H(doc, '5. Retail Division Deep-Dive')
    Q(doc, 22, 'The footfall decline of 14% — is this digital cannibalisation, or is footfall structurally declining for the format you operate?',
      'Diagnostic data suggests both. Approximately 60% of the footfall decline correlates with online conversion (the same customer is buying online instead of physical); the remaining 40% is footfall lost outright (post-pandemic mall avoidance, fuel-cost-constrained discretionary travel, local catchment dilution). We are responding with omni-channel investment for the first 60%, and store-portfolio rationalisation for the last 40%.',
      None)

    Q(doc, 23, 'You are closing 18 stores. What is the criteria, and what is the saved-cost reinvestment thesis?',
      'Criteria: 4-wall EBITDA margin < 5%, comp-sales decline > 10% for 3 consecutive years, and lease-renewal renegotiation likelihood low. Saved cost target MYR 95M annualised. We will reinvest 35% of saved cost into omni-channel platform (digital fulfilment, in-store pickup, dynamic pricing) and 25% into category-killer pilot (beauty + electronics flagship). The remaining 40% goes to deleveraging.',
      'Recovery Levers tab')

    Q(doc, 24, 'The MYR 110M inventory writedown — is this a one-off or do you expect more in FY2026?',
      'The MYR 110M was specifically against FW2024 carryover stock and Cat A surplus. We have tightened buying-margin discipline, deployed dynamic-markdown algorithms, and reduced open-to-buy budgets by 8%. We do not expect a repeat of this scale. FY2026 inventory provisioning is back to normalised level of ~1.5% of revenue.',
      None)

    Q(doc, 25, 'How much of the omni-channel investment converts to actual gross-margin uplift?',
      'Our internal modelling assumes a 130-180 bps gross-margin uplift over 24-36 months, driven by mix shift to higher-margin own-label, dynamic-pricing capture, and reduced markdown depth. We will report progress as a separately broken-out line in the FY2026 results.',
      None)
    doc.add_page_break()


# ── Section 6 ──────────────────────────────────────────────────────────────
def s6(doc):
    H(doc, '6. Other Divisions — Brief Updates')
    Q(doc, 26, 'Chemicals: how exposed are you to a prolonged China demand softness?',
      'Approximately 28% of Chemicals revenue is China-exposed (direct exports + indirect via ASEAN customers servicing China). We are accelerating the mix shift to specialty chemicals (target 35% specialty by FY2027 vs 22% FY25); the Pengerang capacity expansion is targeted at this shift. Domestic and intra-ASEAN demand remains resilient.',
      None)

    Q(doc, 27, 'Agribusiness: with CPO down 18%, what is your base-case for FY2026?',
      'We are modelling MYR 3,400-3,500 / MT for FY2026 (vs MYR 3,750 average FY2025), a further 5-7% decline in our base case. Our cost-of-production at refined-level is MYR 2,750 / MT, so we remain profitable at our base case. The yield-recovery programme in Sumatra contributes ~MYR 32M of FY2026 EBITDA.',
      None)

    Q(doc, 28, 'Manufacturing: is the auto-OEM softness cyclical?',
      'Yes, our reading is cyclical, primarily driven by ASEAN auto-OEM inventory destocking. We expect normalisation H2 FY2026. The Penang plant downtime is being addressed via mechanical-reliability programme; aerospace contract pipeline is intact.',
      None)

    Q(doc, 29, 'Financial Services: what is driving the beat? Is it sustainable?',
      'Three drivers: margin lending growth (book +18% YoY), improved investment income (rates environment), and underwriting-fee beat (3 mandates closed in H2). The first two are sustainable into FY2026; the third is mandate-dependent. We are building wealth-management capability to add a structural growth pillar.',
      None)

    Q(doc, 30, 'BPO: how much of the AI productivity will the BPO industry ultimately have to give back to clients?',
      'Approximately 35-50% in our experience. The other side of that is: our AI-leveraged delivery cost falls faster, opening pricing room for outcome-based models that capture more value per customer. Our pricing strategy is shifting to vertical-specialised, outcome-based contracts where we share in customer business outcomes. New US client wins in FY2025 reflect this transition.',
      None)

    Q(doc, 31, 'Trading: post-Hindustan Bulk, is there any further counterparty risk in the energy book?',
      'We have completed full review of the top-30 counterparties; no other counterparty shows similar early-warning signs. Trade-credit insurance has been mandated for all single-counterparty exposures > MYR 30M (with effect from Sep 2024). We are also tightening credit-limit reviews to monthly cadence.',
      None)

    Q(doc, 32, 'Pharmaceuticals: how do you balance generics speed-to-market with the new branded portfolio investment?',
      'They use largely separate teams. Generics is a tender-and-volume game — our 12 ANDA-per-year filing target is independent of the branded build. Branded portfolio is being built selectively in two niches (cardiometabolic and women\'s health) where we have established regulatory and channel strengths.',
      None)
    doc.add_page_break()


# ── Section 7 ──────────────────────────────────────────────────────────────
def s7(doc):
    H(doc, '7. Capital Allocation & Capex')
    Q(doc, 33, 'The MYR 12.4B 3-year capex programme — is this affordable given the FY2025 cash flow?',
      'Affordability is built around three sources: (a) operating cash flow of MYR 8.5-9.5B over the 3-year period (after working-capital, tax, finance costs); (b) net divestment proceeds of MYR 1.5-2.0B; (c) incremental borrowings of MYR 1.5-2.0B (preferentially sustainability-linked). Net Debt/EBITDA will move temporarily towards 3.3-3.4× before returning under 3.0× by end-FY2027.',
      'Strategy §5')

    Q(doc, 34, 'Why MYR 2.4B of energy-transition capex specifically — what is the IRR?',
      'The MYR 2.4B includes (i) solar + storage at all manufacturing/refinery sites (MYR 1.4B, IRR 9-12% before sustainability-linked-loan margin step-down benefits); (ii) electrification of agri operations (MYR 0.5B, IRR 7-10%); (iii) methane abatement and minor process upgrades (MYR 0.5B, IRR 11-15%). Blended IRR ~10%, before the regulatory cost-avoidance benefit (carbon pricing in MY/ID by 2026-27) which adds 200-300 bps.',
      None)

    Q(doc, 35, 'How do you think about the relative claim of the three Red divisions on capex?',
      'They get minimum-required maintenance capex (~MYR 1.6B over 3 years, 12.9% of total) but no growth capex until ROIC > 8%. This discipline is a deliberate design feature of the recovery programme. Releasing growth capex too early would not change the recovery trajectory and would dilute capital available for the compound engines.',
      None)

    Q(doc, 36, 'Could you defer significant capex if cash-flow pressures intensify?',
      'Yes. Approximately MYR 2.8B of the 3-year programme is sequenced FY2026-2027 and is deferrable up to 9-12 months without breaking strategic milestones. Maintenance capex of MYR 3.2B is essentially non-deferrable. Energy-transition capex is partly deferrable but with regulatory-cost trade-off. We have built explicit deferral playbooks for two stress scenarios.',
      None)
    doc.add_page_break()


# ── Section 8-12 ──────────────────────────────────────────────────────────
def s8_12(doc):
    H(doc, '8. Dividend Policy & Distributions')
    Q(doc, 37, 'Confirm the FY2025 dividend.',
      'Interim DPS MYR 0.12 (paid Sep 2024); Final DPS MYR 0.18 (proposed; subject to AGM approval). Full-year DPS MYR 0.30, unchanged YoY. Total distribution MYR 1.05B against PATAMI of MYR 2.1B — payout ratio 50%.',
      None)

    Q(doc, 38, 'Will you maintain MYR 0.30 DPS through FY2027 if EBITDA recovery is delayed?',
      'Our policy is 40-60% payout of normalised PATAMI. If FY2026 PATAMI recovers per plan to MYR 2.6B, the policy supports MYR 0.30-0.40 DPS. In a downside scenario (FY2026 PATAMI flat at MYR 2.1B), we would maintain MYR 0.30. We would only consider DPS reduction in a more severe scenario, which is not in our base-case planning.',
      None)

    Q(doc, 39, 'Any plans for special dividends or buybacks?',
      'Not in the FY2025-2027 horizon. Capital deployment priorities are (i) the MYR 12.4B capex programme, (ii) selective M&A, (iii) maintaining the dividend policy. Buyback consideration would only arise post-FY2027 if capital structure has surplus.',
      None)

    H(doc, '9. M&A Pipeline')
    Q(doc, 40, 'Walk us through the BPO M&A pipeline (M-1 and M-2).',
      'M-1 (US Project Anvil) is in NDA + early DD; we are interested but the EV expectation needs to compress 8-12% to meet our IRR hurdle. We will not chase. Realistic close H1 FY2026 if pricing aligns. M-2 (EU Project Mercury) is on long-list; market has more potential targets — we expect to identify a preferred candidate by Q2 FY2026.',
      'Strategy §4 — Acquisition Pipeline')

    Q(doc, 41, 'Wealth management arm — build vs buy?',
      'Hybrid. We are building distribution and advisor team organically (with Group customer-base advantage), but we are also evaluating a tactical acquisition of an established wealth manager (M-3 Project Pearl, Indonesia) to accelerate AUM scale. Decision Q2 FY2026.',
      None)

    Q(doc, 42, 'What is your EV / IRR discipline for M&A?',
      '5-year IRR ≥ 12%, EPS-accretion within 24 months, post-deal Net Debt/EBITDA < 3.0×, and demonstrable strategic fit with the five pillars. We have walked away from 3 deals in FY2025 where pricing did not meet hurdles. We will continue to walk away.',
      None)

    H(doc, '10. ESG & Climate Disclosures')
    Q(doc, 43, 'Are you on track for your SBTi Scope 1+2 −45% by 2030 commitment?',
      'Currently at −18% vs FY2019 baseline. Trajectory requires step-up; the renewables portion of the MYR 2.4B energy-transition capex is the primary lever. Interim target FY2027 is −32%. We will report progress in the FY2025 Sustainability Report (publication Apr 2026).',
      None)

    Q(doc, 44, 'Sustainability-linked loan KPI — did you meet FY2024 targets?',
      'Yes, both KPIs (Scope 1+2 reduction milestone, renewable-electricity %) were met. The 12.5 bps step-down clauses applied. FY2025 KPIs are challenging but on-track based on YTD performance.',
      None)

    Q(doc, 45, 'How exposed are you to ASEAN carbon pricing?',
      'Modelled exposure at MYR 50/tCO2e is approximately MYR 280M annual cost (Scope 1+2). Our abatement pathway through FY2027 reduces this by ~MYR 115M. We are not lobbying against carbon pricing; we believe early implementation gives the Group competitive advantage relative to peers.',
      None)

    Q(doc, 46, 'How does the climate-linked compensation work?',
      '15% of senior-executive PSU vesting is tied to Scope 1+2 trajectory. If the Group misses the FY2027 milestone (−32%), 15% of the granted PSUs lapse. This applies to ~120 senior leaders across the Group.',
      'Strategy §6.2; Policy Handbook §3')

    H(doc, '11. Regulatory & Leadership')
    Q(doc, 47, 'Indonesia UU PDP and Malaysia PDPA 2024 amendments — material readiness?',
      'Group DPO appointed FY2024 (covering all jurisdictions). Mandatory breach-notification process operational. Cross-border data-transfer impact assessments completed for all material flows. Staff training rolled out FY2024 (97% completion). We do not see material residual readiness gap.',
      'Policy Handbook §6')

    Q(doc, 48, 'CFO succession — any update?',
      'No update at this time. Group CFO is in role and is leading the FY2025-2027 strategy execution. Board has identified three internal successors as part of the standard top-100 succession depth programme.',
      None)

    Q(doc, 49, 'Any AGM or EGM matters investors should be aware of?',
      'Standard AGM Apr 2026 with usual resolutions. No EGM scheduled. Re-election of two Independent Directors will be subject to shareholder approval per the standard 9-year rule.',
      None)

    H(doc, '12. FY2026 Guidance & Outlook')
    Q(doc, 50, 'What is the FY2026 guidance?',
      'We are guiding (i) Group revenue growth 6-8% YoY (MYR 44.9-45.8B); (ii) EBITDA recovery of 12-15% YoY (MYR 6.8-7.0B), driven by the three Red divisions returning to positive territory and the compound engines continuing to grow; (iii) FCF MYR 2.7-3.0B; (iv) Net Debt/EBITDA returning under 3.0× by end-FY2026. Guidance reflects a base-case view; downside and upside scenarios are tracked internally.',
      None)

    Q(doc, 51, 'What are the key assumptions in the FY2026 guidance?',
      'CPO MYR 3,450 / MT; IDR/USD 16,800; Brent USD 78/bbl; ASEAN GDP growth 4.6%; benchmark rates flat from current; no major M&A in the year (M-1 / M-3 closes assumed mid-year, modest H2 contribution).',
      None)

    Q(doc, 52, 'How do you bridge the FY2026 EBITDA recovery to MYR 6.8-7.0B?',
      'Bridge components: Properties recovery +MYR 180M; Healthcare recovery +MYR 95M; Retail recovery +MYR 70M; FS growth +MYR 230M; BPO growth +MYR 300M; Pharma growth +MYR 180M; Other 4 divisions stable to slightly up; FX neutral assuming our forecast holds. Total bridge ~MYR 800M-1,000M, against the FY2025 base of MYR 6.1B → MYR 6.9-7.1B (within guided range).',
      None)

    Q(doc, 53, 'When do you next provide an interim update?',
      'Q1 FY2026 trading update Apr 2026; H1 FY2026 results Aug 2026.',
      None)


def s13_regulatory(doc):
    H(doc, '13. Regulatory, Listing & Disclosure Q&A')
    Q(doc, 54, 'Has Bursa Malaysia or the Securities Commission Malaysia raised any query on the FY2025 result or its disclosure?',
      'No formal query letter has been received as at result-release date. We pre-engaged Bursa Listing Division and the SC Corporate Surveillance Department in November 2024 with a confidential briefing on the EBITDA shortfall trajectory and the disclosure plan, including the timing of the Q4 trading update and the announcement of the three Red divisions. The trading update issued 31 Dec 2024 satisfied Chapter 9 immediate-disclosure requirements (Bursa Listing Requirements). We will continue normal engagement and remain compliant with MCCG 2021 — including the Audit Committee Report and Statement on Risk Management & Internal Control sections of the Annual Report.',
      'Bursa LR Chapter 9; SC CMSA 2007; MCCG 2021')

    Q(doc, 55, 'For PT Zava Pharmaceuticals Tbk and PT Zava Agribusiness Tbk on IDX — any OJK or BEI matters?',
      'Both Indonesia-listed subsidiaries filed their Q3 2024 LK on time with OJK (POJK 14/2022) and IDX. No OJK enforcement notice or BEI suspension has been received. Related-party transactions with the Group parent are disclosed under POJK 19/2017 with independent appraiser fairness opinions where required. Both entities remain on the IDX Main Board with no special-monitoring designation.',
      'POJK 14/2022; POJK 19/2017')

    Q(doc, 56, 'Bank Negara Malaysia and Bank Indonesia — any regulatory issues at the Financial Services division?',
      'Zava Financial Services (M) Bhd received the FY2024 BNM RMiT thematic review outcome in Q1 FY2025 with no material findings; one observation on third-party-cyber-controls is being closed in FY2025 (no impact on capital adequacy or licence). PT Zava Multi Finance is current on POJK 35/2018 health-of-finance ratios and on UU 4/2023 (P2SK) implementation milestones. Capital adequacy ratios at the FS Division are well above Basel III minima (CET1 14.2%, Total CAR 17.8%).',
      'BNM RMiT; BNM RMCG; POJK 35/2018; UU 4/2023')

    Q(doc, 57, 'Are you considering early adoption of IFRS S1 / S2 climate disclosure ahead of the NACRA / Bursa-mandated FY2027 timeline?',
      'We are running a phased adoption: gap-assessment in FY2025, parallel-run in FY2026, full disclosure in FY2027 in line with the National Sustainability Reporting Framework (NSRF) and the Bursa Sustainability Reporting Guide (3rd edition). We already produce TCFD-aligned climate disclosure annually since FY2024 (4 pillars × 11 recommendations), so the incremental content for IFRS S2 is mostly in scenario analysis depth and Scope 3 financed-emissions for the FS division. We see no benefit in a year-early adoption that would not be peer-comparable.',
      'TCFD; IFRS S1+S2; NACRA; Bursa Sustainability Guide')

    Q(doc, 58, 'Are FY2025 results prepared on the same MFRS basis? Any new standards adopted?',
      'Yes — FY2025 financials are MFRS-compliant (MFRS being equivalent to IFRS as adopted in Malaysia). New standards effective FY2025 are MFRS 17 (Insurance Contracts — limited Group impact via FS), and amendments to MFRS 101 on classification of liabilities. The Indonesia subsidiaries report under PSAK and reconcile to MFRS at consolidation; the reconciliation note is in the consolidated Annual Report. The Penang Tower MYR 320M writedown was processed under MFRS 136 impairment testing and fully audited.',
      'MFRS 17, MFRS 101, MFRS 136; PSAK reconciliation')

    Q(doc, 59, 'On data privacy — what is your readiness for PDPA Amendment Act 2024 (effective 1 June 2025) and UU PDP (full enforcement Oct 2024)?',
      'Group Data Protection Officer was appointed in FY2024, with division DPOs in MY and ID entities. The MY DPO registration with JPDP / Jabatan Perlindungan Data Peribadi has been completed; the breach-notification protocol meets the 72-hour requirement under the amended PDPA. In Indonesia, the 3×24-hour breach notification protocol to Lembaga PDP is operational; standard-contractual-clauses for cross-border transfers from PT Zava entities to the Group parent are signed. Staff training is at 97% completion across MY/ID/SG/TH/VN. We have not had a reportable breach in FY2024 or FY2025 to date.',
      'PDPA 2010 + Act A1709/2024; UU 27/2022 (UU PDP)')

    Q(doc, 60, 'MCCG 2021 corporate-governance compliance — any departures?',
      'We comply with all 32 MCCG 2021 Practices on an "apply" basis except for one Step-Up Practice (independent Chairman) which we currently meet through a Senior Independent Director arrangement. The Step-Up Practice on at-least-30%-women on Board is met (current Board composition 4 of 11 = 36%). Departure rationales are disclosed in the Corporate Governance Report.',
      'MCCG 2021')

    Q(doc, 61, 'Halal certification status across MY/ID Pharmaceuticals and Agribusiness?',
      'Zava Pharmaceuticals (MY) — Halal certification under JAKIM is current; renewal underway for two Penang production lines. Zava Pharmaceuticals (ID) — BPJPH certification under UU JPH 33/2014 is current. Agribusiness palm oil supply chain — RSPO P&C 2018 certified since 2020; FY2024 ASA-1 surveillance audit closed without major NCR; smallholder traceability programme covers 78% of FFB throughput, target 95% by FY2027.',
      'JAKIM; BPJPH (UU JPH 33/2014); RSPO P&C 2018')


def main(out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    cover(doc)
    s1(doc)
    s2(doc)
    s3(doc)
    s4(doc)
    s5(doc)
    s6(doc)
    s7(doc)
    s8_12(doc)
    s13_regulatory(doc)
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else 'files/06_Zava_Investor_QA_FY2025.docx'
    main(out)
