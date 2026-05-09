"""Generates 03_Zava_Group_Strategy_Framework.docx — "ZAVA FORWARD 2030".

Sections:
  - Cover + TOC + Foreword
  - 1. Strategic Context (FY25 reality + macro)
  - 2. Five Pillars of Zava Forward 2030
  - 3. Division-by-Division Strategy (all 11)
  - 4. M&A and Portfolio Pipeline
  - 5. Capital Allocation FY2025-2027 (MYR 12.4B)
  - 6. People & Culture Agenda
  - 7. Strategic Risks & Mitigations
  - 8. Strategy Cascade & Governance
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from gen_ref_files.common import DIVISIONS, FY25, fy25_eb, hist_rev

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xBF, 0x8F, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x00, 0x70, 0x32)
RED = RGBColor(0xC0, 0x00, 0x00)


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def H(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20 if level == 1 else 14 if level == 2 else 11)
    run.font.name = 'Aptos'
    run.font.color.rgb = color
    return p


def P(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    if color:
        run.font.color.rgb = color
    return p


def B(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if not p.runs:
        run = p.add_run(text)
    else:
        p.runs[0].text = text
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Aptos'


def N(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if not p.runs:
        run = p.add_run(text)
    else:
        p.runs[0].text = text
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Aptos'


def T(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Aptos'
        shade(cell, header_color)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Aptos'
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    return t


# ── Cover ──────────────────────────────────────────────────────────────────
def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run('ZAVA FORWARD\n2030\n')
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'
    run = p.add_run('\nGroup Strategy Framework')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\nVersion 4.0 | Board-approved 12 December FY2024\n')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = 'Aptos'
    run = p.add_run('Confidential — Group ExCo, Division CEOs, Board Members only')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    run.font.name = 'Aptos'

    doc.add_page_break()
    H(doc, 'Foreword from the Group Chief Executive', level=1)
    P(doc, 'FY2025 closes a difficult year for Zava Group. Group revenue of MYR 42.4 billion '
           'against a budget of MYR 44.2 billion represents a 4.1% top-line shortfall, but the '
           'EBITDA picture is the harder story: MYR 6.1 billion against MYR 7.4 billion budget, an '
           '18% miss that reduces our net debt headroom and pressures three divisions — '
           'Properties, Healthcare and Retail — into negative-EBITDA territory for the first time '
           'in the Group\'s 38-year history.')
    P(doc, 'I do not see this as a story of strategic failure. The shortfall has six identifiable '
           'drivers — Property completion delays, Healthcare patient-volume softness, Retail '
           'footfall decline, China chemicals demand weakness, IDR/USD adverse FX, and the '
           'Hindustan Bulk counterparty event in Trading. Each has been studied; each has a '
           'remediation programme; and the underlying portfolio strength remains intact. '
           'Financial Services beat budget by 7%, BPO beat by 4%, and the four other divisions '
           'are operating within plan tolerance.')
    P(doc, 'But it does sharpen the case for Zava Forward 2030 — the strategic framework set '
           'out in this document. The next five years require us to: (i) recover the cyclical '
           'divisions decisively; (ii) compound advantage in our growth engines; (iii) decarbonise '
           'and digitalise across the enterprise; (iv) sharpen our portfolio with disciplined '
           'M&A and selective exits; and (v) build the talent and capability to execute. The '
           'framework is built around five strategic pillars and is grounded in clear capital-'
           'allocation discipline: MYR 12.4 billion of capex through FY2027, a fix-first lens on '
           'the three Red divisions, and an unwavering commitment to capital-structure '
           'protection (Net Debt / EBITDA < 3.0× through cycle).')
    P(doc, 'I commend this framework to the Board, the ExCo, and to every division leader. '
           'Execution is now the priority.', italic=True)
    P(doc, '— Group CEO, December FY2024', italic=True)
    doc.add_page_break()


# ── Section 1: Strategic Context ───────────────────────────────────────────
def s1_context(doc):
    H(doc, '1. Strategic Context — Where We Stand', level=1)

    H(doc, '1.1 FY2025 Performance Reality', level=2)
    P(doc, 'The FY2025 outcome reshapes the starting position for our 2030 strategy:')
    T(doc, ['Metric', 'FY2024A', 'FY2025 Bud', 'FY2025 Act', 'Variance', 'Status'],
      [
          ['Revenue (MYR M)', f'{sum(hist_rev(d, 2) for d in [r[0] for r in DIVISIONS]):,.0f}',
           '44,200', '42,380', '−1,820 (−4.1%)', 'Below'],
          ['EBITDA (MYR M)', '7,890', '7,420', '6,103', '−1,317 (−17.8%)', 'RED'],
          ['EBITDA Margin %', '17.6%', '16.8%', '14.4%', '−240 bps', 'RED'],
          ['Net Debt / EBITDA', '2.4×', '2.5×', '3.1×', '+0.6×', 'AMBER'],
          ['Free Cash Flow (MYR M)', '3,420', '3,250', '2,180', '−1,070 (−33%)', 'RED'],
          ['Divisions in negative EBITDA', '0', '0', '3', '+3', 'RED'],
      ], col_widths=[4.5, 2.5, 2.5, 2.5, 3.0, 1.5])

    H(doc, '1.2 The Three Red Divisions', level=2)
    P(doc, 'Three divisions reported negative EBITDA in FY2025 — a first for the Group. Each '
           'has a distinct root cause:')
    B(doc, 'Properties (act margin −0.5% on rev MYR 1,800M): Mid-Valley project completion delay, '
           'Penang Tower writedown of MYR 320M, slowing pre-sales velocity in Klang Valley, and '
           'unsold inventory carrying cost. Cyclical, not structural.')
    B(doc, 'Healthcare (act margin −0.5% on rev MYR 1,200M): Patient volume below pre-pandemic '
           'baseline by 11%, specialist consultant cost growth outpacing revenue, Subang facility '
           'utilisation at 58%. Mix of cyclical and emerging structural issues.')
    B(doc, 'Retail (act margin −0.5% on rev MYR 2,200M): Online cannibalisation of physical '
           'stores, footfall down 14% at flagship locations, lease cost inflation, and a one-time '
           'inventory writedown of MYR 110M. Structural pressure on the format mix.')

    H(doc, '1.3 The Six Outperformers', level=2)
    P(doc, 'Eight of eleven divisions are at-or-better than plan. Two are notable beats:')
    B(doc, 'Financial Services beat budget by 7% on EBITDA, driven by margin lending growth and improved investment income.')
    B(doc, 'BPO beat budget by 4%, with three new US captive deals closed in H2 and pricing escalators kicking in.')
    P(doc, 'Six others — Chemicals, Manufacturing, Treasury Holdings, Pharmaceuticals, '
           'Agribusiness and Trading — are within ±5% of plan, with discrete recovery items '
           'identified for FY2026.')

    H(doc, '1.4 Macro & Industry Backdrop', level=2)
    P(doc, 'The next five years sit at the intersection of structural shifts the Group must '
           'navigate:')
    B(doc, 'China growth slowdown and demand re-routing: olefins demand softness affecting our chemicals book; partial offset from RCEP intra-ASEAN trade flows.')
    B(doc, 'Energy transition acceleration: ASEAN-wide carbon pricing on the horizon (CBAM extension to ASEAN exporters by 2027, MY domestic carbon pricing by 2026, ID early signals).')
    B(doc, 'Persistent IDR/USD pressure: IDR has weakened ~9% YoY against USD; FX translation drag is structural, not cyclical, for the next 12-24 months.')
    B(doc, 'Generative AI productivity wave: BPO division must convert AI productivity into pricing leverage rather than allowing client share-of-savings claims.')
    B(doc, 'ASEAN healthcare demand recovery is uneven — cosmetic and elective procedures returning faster than chronic disease management.')
    B(doc, 'Property: Klang Valley overhang likely persists 18-24 months; Penang and Iskandar showing signs of recovery.')
    doc.add_page_break()


# ── Section 2: Five Pillars ────────────────────────────────────────────────
def s2_pillars(doc):
    H(doc, '2. Five Pillars of Zava Forward 2030', level=1)
    P(doc, 'Zava Forward 2030 organises the Group strategy around five reinforcing pillars. '
           'Each pillar has measurable outcomes, capital implications and division-level cascades.')

    pillars = [
        {
            'num': '2.1',
            'name': 'Pillar 1 — Recover & Reinvent',
            'colour': 'BF8F00',
            'tagline': 'Restore the three Red divisions to positive EBITDA by FY2027.',
            'desc': 'The three negative-EBITDA divisions (Properties, Healthcare, Retail) are not '
                   'candidates for managed decline. Properties and Healthcare have core competitive '
                   'positions worth defending; Retail requires format reinvention. Each division '
                   'has a 90-day stabilisation plan, a 12-month recovery plan, and a 3-year '
                   'transformation plan, with monthly tracking at Group ExCo level and quarterly '
                   'tracking at Board level.',
            'kpis': [
                ('Properties EBITDA', 'MYR -9M', 'MYR +180M', 'MYR +320M', 'FY26', 'FY27'),
                ('Healthcare EBITDA', 'MYR -6M', 'MYR +90M', 'MYR +180M', 'FY26', 'FY27'),
                ('Retail EBITDA', 'MYR -11M', 'MYR +60M', 'MYR +140M', 'FY26', 'FY27'),
                ('Combined EBITDA contribution', 'MYR -26M', 'MYR +330M', 'MYR +640M', '—', '—'),
            ],
            'commitments': [
                'Asset-rationalisation programme: dispose of MYR 800M of non-core property and retail real-estate by FY2027',
                'Healthcare format refresh: convert two underutilised inpatient floors at Subang into outpatient day-surgery and chronic-care formats',
                'Retail digital-first reset: close 18 underperforming stores; redeploy 35% of saved cost into omni-channel and category-killer formats',
                'Capital reallocation: hold growth capex at minimum-required levels in these three divisions until ROIC > 8%',
            ],
        },
        {
            'num': '2.2',
            'name': 'Pillar 2 — Compound the Engines',
            'colour': '00703C',
            'tagline': 'Double down on Financial Services, BPO and Pharmaceuticals.',
            'desc': 'Three divisions have demonstrated structural advantage — Financial Services '
                   '(strong asset-quality and rising fee streams), BPO (AI-leveraged offshore '
                   'delivery with US/EU customer base), and Pharmaceuticals (generics tender wins '
                   'across Indonesia and the Philippines, plus emerging branded portfolio). These '
                   'engines justify accelerated capital deployment.',
            'kpis': [
                ('FS EBITDA', 'MYR 920M', 'MYR 1,150M', 'MYR 1,400M', 'FY27', 'FY30'),
                ('BPO EBITDA', 'MYR 880M', 'MYR 1,180M', 'MYR 1,500M', 'FY27', 'FY30'),
                ('Pharma EBITDA', 'MYR 700M', 'MYR 880M', 'MYR 1,100M', 'FY27', 'FY30'),
                ('Combined EBITDA contribution', 'MYR 2,500M', 'MYR 3,210M', 'MYR 4,000M', '—', '—'),
            ],
            'commitments': [
                'Financial Services: launch wealth-management arm targeting MYR 25B AUM by FY2028',
                'BPO: bolt-on acquisitions of two US-based mid-tier BPOs by FY2027 (target combined revenue MYR 800M, EV MYR 1.5-2B)',
                'Pharmaceuticals: file 12 new ANDA dossiers per year through FY2027; build branded portfolio in cardiometabolic and women\'s-health',
                'Capital tilt: 35% of FY2025-2027 growth capex earmarked for these three divisions',
            ],
        },
        {
            'num': '2.3',
            'name': 'Pillar 3 — Decarbonise & Digitise',
            'colour': '2E75B6',
            'tagline': 'Cut Scope 1+2 by 45% and embed AI in every core workflow.',
            'desc': 'Zava\'s 2030 climate commitment (Scope 1+2 −45% vs FY2019 baseline) is '
                   'science-based-target validated and capital-backed. In parallel, the Group is '
                   'investing in enterprise digital and AI platform foundations to convert the '
                   'productivity wave into competitive advantage.',
            'kpis': [
                ('Scope 1+2 reduction', '−18%', '−32%', '−45%', 'FY27', 'FY30'),
                ('Renewable electricity %', '34%', '55%', '70%', 'FY27', 'FY30'),
                ('Digital revenue %', '8%', '15%', '25%', 'FY27', 'FY30'),
                ('AI use-cases in production', '14', '60', '150+', 'FY27', 'FY30'),
                ('Annual digital productivity savings', 'MYR 80M', 'MYR 280M', 'MYR 500M', 'FY27', 'FY30'),
            ],
            'commitments': [
                'Group-wide energy-transition capex of MYR 2.4 billion through FY2027',
                'Solar and storage at all manufacturing and refinery sites by FY2027 (15-year PPAs preferred)',
                'M365 Copilot Pro+ deployment to 12,000 knowledge workers by mid-FY2026',
                'Sector-specific AI platforms: BPO Cognitive Workbench, Pharma Regulatory Co-pilot, FS Risk Brain, Healthcare Diagnostic Assistant',
                'Climate-linked compensation: 15% of senior-executive PSU vest tied to Scope 1+2 trajectory',
            ],
        },
        {
            'num': '2.4',
            'name': 'Pillar 4 — Sharpen the Portfolio',
            'colour': '7030A0',
            'tagline': 'Disciplined M&A, selective exits, sharper governance.',
            'desc': 'The Group portfolio of 11 divisions is right-sized but under-curated. '
                   'Pillar 4 commits to disciplined portfolio management: clear identification of '
                   'core, growth, and divestment candidates; M&A discipline anchored in 5-year '
                   'IRR ≥ 12% and accretion within 24 months; selective exits where ROIC trails '
                   'cost of capital persistently.',
            'kpis': [
                ('M&A deployed (MYR B)', '0.0', '2.5', '4.5', 'FY27', 'FY30'),
                ('Divestments completed (MYR M)', '0', '1,200', '2,500', 'FY27', 'FY30'),
                ('Group ROIC', '8.4%', '10.5%', '12.5%', 'FY27', 'FY30'),
                ('No. of bolt-ons closed', '0', '5', '12', 'FY27', 'FY30'),
            ],
            'commitments': [
                'M&A pipeline reviewed quarterly by Investment Committee; minimum hurdle rate 12% IRR over 5 years',
                'Divestment candidates: non-core trading subsidiaries, underperforming retail real-estate, and the Surabaya manufacturing JV (50% stake, MYR 280M)',
                'Portfolio-review cadence: deep-dive on every division every 24 months at Board level',
                'Capital-recycling target: net divestments to fund 30%+ of M&A spend',
            ],
        },
        {
            'num': '2.5',
            'name': 'Pillar 5 — Talent & Trust',
            'colour': 'C00000',
            'tagline': 'Build the people, the culture and the licence to operate for 2030.',
            'desc': 'The Group cannot deliver pillars 1-4 without parallel investment in talent, '
                   'leadership and stakeholder trust. Pillar 5 is the enabling pillar: '
                   'compensation modernisation, leadership pipeline depth, ESG-grounded community '
                   'investment, and governance maturity.',
            'kpis': [
                ('Voluntary attrition rate', '12.4%', '< 10%', '< 8%', 'FY27', 'FY30'),
                ('Engagement score (Pulse)', '74', '78', '82', 'FY27', 'FY30'),
                ('Internal-fill rate (senior roles)', '52%', '65%', '75%', 'FY27', 'FY30'),
                ('% Women in management JG6+', '28.4%', '32%', '40%', 'FY27', 'FY30'),
                ('Community investment (MYR M)', '50', '60', '85', 'FY27', 'FY30'),
                ('Independent ESG rating (Sustainalytics)', 'Med-Risk 24', 'Low-Risk 18', 'Low-Risk < 15', 'FY27', 'FY30'),
            ],
            'commitments': [
                'Reskilling fund of MYR 60M committed FY2025-2027 (priority: AI, data, sustainability, digital)',
                'Living wage certification for direct workforce by FY2026',
                'Top-100 succession depth: every Group ExCo + Division CEO role to have ≥ 2 ready-now successors by FY2026',
                'Community SROI ≥ 3.0× across the Group',
            ],
        },
    ]

    for p_data in pillars:
        H(doc, f'{p_data["num"]} {p_data["name"]}', level=2, color=NAVY)
        P(doc, p_data['tagline'], italic=True, size=11, color=GREY)
        P(doc, p_data['desc'])
        P(doc, 'Outcome metrics:', bold=True, size=10)
        T(doc, ['Metric', 'FY2025', f'{p_data["kpis"][0][3] if "FY27" in str(p_data["kpis"][0]) else "Mid"}',
                'Long-term', 'Mid by', 'Long by'],
          [list(k) for k in p_data['kpis']],
          col_widths=[5.5, 2.5, 2.5, 2.5, 1.5, 1.5])
        P(doc, 'Strategic commitments:', bold=True, size=10)
        for c in p_data['commitments']:
            B(doc, c)
        doc.add_page_break()


# ── Section 3: Division-by-division strategy ──────────────────────────────
def s3_divisions(doc):
    H(doc, '3. Division-by-Division Strategy', level=1)
    P(doc, 'Each of the 11 divisions has a defined strategic role within Zava Forward 2030. '
           'The role drives capital allocation, M&A appetite, and the ambition level on cost, '
           'growth and sustainability.')
    T(doc, ['Division', 'Strategic Role', 'FY25 EBITDA', 'FY27 Target', 'M&A Posture'],
      [
          ['Chemicals', 'Defend & extend', 'MYR 480M', 'MYR 600M', 'Selective bolt-on (specialty)'],
          ['Agribusiness & Plantations', 'Operational excellence', 'MYR 320M', 'MYR 410M', 'Hold; exit non-strategic estates'],
          ['Manufacturing', 'Digital transformation', 'MYR 360M', 'MYR 480M', 'Selective (advanced manuf.)'],
          ['Financial Services', 'Compound', 'MYR 920M', 'MYR 1,150M', 'Active (wealth, fintech)'],
          ['Properties', 'Recover', 'MYR -9M', 'MYR 320M', 'Hold; recycle MYR 600M assets'],
          ['BPO', 'Compound', 'MYR 880M', 'MYR 1,180M', 'Active (US/EU mid-tier)'],
          ['Trading', 'Risk-managed run-off', 'MYR 110M', 'MYR 140M', 'Divest non-core'],
          ['Pharmaceuticals', 'Compound', 'MYR 700M', 'MYR 880M', 'Active (branded niches)'],
          ['Healthcare', 'Reinvent format', 'MYR -6M', 'MYR 180M', 'Hold; recycle low-utilisation assets'],
          ['Retail', 'Format reset', 'MYR -11M', 'MYR 140M', 'Hold; close underperformers'],
          ['Treasury Holdings', 'Capital allocator', 'MYR 39M', 'MYR 90M', 'Group capital allocator'],
      ], col_widths=[3.5, 4.0, 2.0, 2.5, 5.0])

    div_strategies = [
        ('Chemicals', 'Anchor commodity-grade book on cost-leadership; tilt mix to specialty chemicals (electronics, EV battery materials, semiconductor packaging) targeting 35% of revenue by FY2027 (FY25 actual: 22%). Capacity expansion at Pengerang via brownfield to capture incremental ASEAN demand. Carbon-cost preparation: methane abatement programme and 100MW captive solar.'),
        ('Agribusiness & Plantations', 'Yield-recovery programme across Sumatran estates (target 23 t/ha by FY2027, vs 19 t/ha FY2025). Replanting acceleration funded by sustainability-linked sukuk. Exit two non-core upstream estates (Kalimantan Barat) and reinvest into refinery capacity. NDPE compliance to 100% by FY2027.'),
        ('Manufacturing', 'Industrial digital transformation across 14 plants; target 30% labour-productivity uplift by FY2027 via AI-assisted predictive maintenance and adaptive scheduling. Aerospace contract pipeline (MRO + parts) to 22% of revenue. Two plant closures (Kuala Lumpur + Subang Jaya) and concentration into the Senai mega-site.'),
        ('Financial Services', 'Compound advantage. Wealth-management arm launch FY2026 (target MYR 25B AUM by FY2028). Margin-lending and securities-financing scale-up. Deeper IB underwriting capability across ASEAN ECM/DCM. Asset quality discipline: NPL < 1.5% through cycle. Selective wealth + fintech bolt-ons.'),
        ('Properties', 'Recovery anchored on three pillars: (1) sell-down of 320 unsold units at Mid-Valley with sharper pricing and bundled financing; (2) Penang Tower recovery via mixed-use repositioning to office + serviced apartments; (3) Disposal of MYR 600M non-core land bank. New launches paused until inventory months-of-supply < 8.'),
        ('BPO', 'Compounding engine. Two US bolt-ons targeted FY2026/27 (combined revenue ~MYR 800M, EV MYR 1.5-2B). AI-led productivity converted into outcome-based pricing models (vs FTE-based). Centre expansion: Cebu (PH), Ho Chi Minh City (VN). New Cognitive Workbench platform deployed across all delivery centres.'),
        ('Trading', 'Risk-managed run-off of energy-trading book post-Hindustan Bulk event; tighten counterparty limits, mandate trade-credit insurance for all single-counterparty exposure > MYR 30M. Pivot to higher-margin specialty trading (rare earths, agri-soft). Divest non-core agri-trading subsidiary in Singapore (estimated proceeds MYR 280M).'),
        ('Pharmaceuticals', 'Compound. Indonesia generics: file 12 ANDA dossiers per year; target 5 launches per year. Branded portfolio build in cardiometabolic and women\'s-health (Malaysia, Indonesia, Philippines). API backward-integration evaluation FY2026. Regulatory Co-pilot platform to compress dossier preparation time by 35%.'),
        ('Healthcare', 'Reinvent. Convert Subang underutilised inpatient floors into day-surgery and outpatient chronic-care. Specialist-network pivot: more procedural volume, less low-acuity admission. Insurance partnership with Group Financial Services for bundled wellness products. Single-network branding refresh.'),
        ('Retail', 'Format reset. Close 18 underperforming stores by mid-FY2026; redeploy 35% of saved cost into omni-channel platform investment. Category-killer pilot in beauty + electronics (selectively bigger basket sizes). Dynamic pricing platform deployment. Loyalty programme refresh with fintech-enabled instalments.'),
        ('Treasury Holdings', 'Group capital allocator. Continue acting as the in-house bank for all 11 divisions. Optimise the investment portfolio for through-cycle yield. Manage the FY2025-2027 MYR 12.4B capex programme allocation. Maintain Net Debt / EBITDA < 3.0× through cycle.'),
    ]

    for div_name, div_text in div_strategies:
        H(doc, f'3.x  {div_name}', level=2)
        P(doc, div_text)
    doc.add_page_break()


# ── Section 4: M&A pipeline ────────────────────────────────────────────────
def s4_ma(doc):
    H(doc, '4. M&A and Portfolio Pipeline', level=1)
    P(doc, 'The Group has assembled a disciplined, division-led M&A pipeline. Each opportunity '
           'is assessed against minimum hurdles: 5-year IRR ≥ 12%, EPS-accretion within 24 '
           'months, post-deal Net Debt/EBITDA < 3.0×, and demonstrable strategic fit with the '
           'five pillars of Zava Forward 2030.')

    H(doc, '4.1 Acquisition Pipeline (FY2025-2027)', level=2)
    T(doc, ['Code', 'Target', 'Division', 'Rationale', 'EV (MYR M)', 'Stage', 'Close Target'],
      [
          ['M-1', 'US BPO Project Anvil', 'BPO', 'Mid-tier US BPO; ~MYR 500M revenue; banking/insurance vertical', '900-1,100', 'NDA + early due diligence', 'H1 FY2026'],
          ['M-2', 'EU BPO Project Mercury', 'BPO', 'European mid-tier; ~MYR 300M revenue; healthcare/pharma vertical', '600-800', 'Long-list', 'H2 FY2026'],
          ['M-3', 'Indonesia Wealth Project Pearl', 'Financial Services', 'Independent wealth manager; AUM ~MYR 6B; OJK licence', '450-550', 'NDA signed', 'H1 FY2026'],
          ['M-4', 'Vietnam SME Lender Project Tiger', 'Financial Services', 'Small SME lender, MYR 1.2B portfolio; gateway licence', '320-380', 'Term sheet', 'Q4 FY2025'],
          ['M-5', 'Branded Cardio API Maker', 'Pharmaceuticals', 'Indonesia branded API; 4 products; cGMP-certified', '180-240', 'Indicative offer', 'H2 FY2026'],
          ['M-6', 'Women\'s Health Brand SEA', 'Pharmaceuticals', 'Women\'s-health branded portfolio; 6 SKUs', '120-180', 'Confidential discussion', 'FY2027'],
          ['M-7', 'Specialty Chemicals Bolt-on', 'Chemicals', 'EV battery additives; ~MYR 200M revenue', '350-450', 'Long-list', 'FY2027'],
          ['M-8', 'Day-surgery Chain', 'Healthcare', '4-clinic ambulatory chain in MY; complement to inpatient network', '180-220', 'Long-list', 'FY2027'],
      ], col_widths=[1.2, 4.5, 2.8, 5.5, 2.0, 2.5, 2.0])

    H(doc, '4.2 Divestment Pipeline (FY2025-2027)', level=2)
    T(doc, ['Code', 'Asset', 'Division', 'Rationale', 'Expected Proceeds (MYR M)', 'Stage', 'Close Target'],
      [
          ['D-1', 'Klang Valley Land Bank (3 parcels)', 'Properties', 'Non-core; long-dated holding cost', '450-550', 'IM out', 'H1 FY2026'],
          ['D-2', 'Penang Tower Phase 2', 'Properties', 'Strategic exit post-writedown', '180-220', 'Discussions', 'H2 FY2026'],
          ['D-3', 'Surabaya Mfg JV (50%)', 'Manufacturing', 'Non-core; partner ROFR exercised', '260-320', 'Negotiating', 'Q3 FY2025'],
          ['D-4', 'Singapore Agri-Trading Sub', 'Trading', 'Non-core post-Hindustan Bulk event', '240-300', 'Auction prep', 'H2 FY2026'],
          ['D-5', '2 Kalimantan Barat Estates', 'Agribusiness', 'Non-strategic; recycle into refinery capacity', '350-400', 'Beauty parade', 'H1 FY2026'],
          ['D-6', '18 Underperforming Retail Stores', 'Retail', 'Format reset', 'Lease exit; net-positive cash', 'In execution', 'Through FY2026'],
      ], col_widths=[1.2, 5.5, 2.5, 5.5, 3.0, 2.0, 2.0])

    P(doc, 'Total expected M&A deployment FY2025-2027: MYR 2.5-3.5 billion. Total expected '
           'divestment proceeds FY2025-2027: MYR 1.5-2.0 billion. Net deployment: MYR 1.0-1.5 '
           'billion to be funded from operating cash flow and existing facility headroom.')
    doc.add_page_break()


# ── Section 5: Capital Allocation ──────────────────────────────────────────
def s5_capex(doc):
    H(doc, '5. Capital Allocation FY2025-2027', level=1)
    P(doc, 'The Group has approved a MYR 12.4 billion capital expenditure programme for the '
           'three-year period FY2025-2027, reflecting the strategic priorities set out in '
           'pillars 1-5.')

    H(doc, '5.1 Capex by Division (MYR M)', level=2)
    T(doc, ['Division', 'FY2025 Plan', 'FY2026 Plan', 'FY2027 Plan', '3-yr Total', '% of Group'],
      [
          ['Chemicals', '480', '520', '450', '1,450', '11.7%'],
          ['Agribusiness & Plantations', '320', '380', '350', '1,050', '8.5%'],
          ['Manufacturing', '410', '460', '380', '1,250', '10.1%'],
          ['Financial Services', '180', '220', '250', '650', '5.2%'],
          ['Properties', '280', '180', '160', '620', '5.0%'],
          ['BPO', '320', '380', '420', '1,120', '9.0%'],
          ['Trading', '40', '40', '40', '120', '1.0%'],
          ['Pharmaceuticals', '240', '280', '320', '840', '6.8%'],
          ['Healthcare', '180', '160', '180', '520', '4.2%'],
          ['Retail', '120', '160', '180', '460', '3.7%'],
          ['Treasury Holdings', '20', '25', '25', '70', '0.6%'],
          ['Group Energy Transition (cross-cutting)', '720', '820', '880', '2,420', '19.5%'],
          ['Group Digital & AI (cross-cutting)', '420', '480', '520', '1,420', '11.5%'],
          ['Group Other (HQ, real estate)', '120', '180', '120', '420', '3.4%'],
          ['TOTAL', '4,070', '4,205', '4,125', '12,400', '100.0%'],
      ], col_widths=[5.0, 2.4, 2.4, 2.4, 2.4, 2.4])

    H(doc, '5.2 Capex by Strategic Theme', level=2)
    T(doc, ['Theme', '3-yr Total (MYR M)', '% of Total', 'Anchor Pillar'],
      [
          ['Energy Transition (Scope 1+2 abatement)', '2,420', '19.5%', 'Pillar 3'],
          ['Digital & AI Platform', '1,420', '11.5%', 'Pillar 3'],
          ['Growth — Compound Engines (FS, BPO, Pharma)', '2,610', '21.0%', 'Pillar 2'],
          ['Maintenance Capex (all divisions)', '3,200', '25.8%', 'Operational'],
          ['Recovery & Reinvent (Properties, HC, Retail)', '1,600', '12.9%', 'Pillar 1'],
          ['Operational Excellence (Chem, Agri, Mfg, Trading)', '1,150', '9.3%', 'Pillar 1/2'],
          ['Group HQ, Real Estate, Other', '0', '0.0%', '—'],
      ], col_widths=[6.5, 2.5, 2.0, 2.5])

    H(doc, '5.3 Funding & Capital Structure', level=2)
    P(doc, 'The MYR 12.4B programme is funded via:')
    B(doc, 'Operating Free Cash Flow: MYR 8.5-9.5B over the period (after working-capital, tax, finance costs)')
    B(doc, 'Net Divestment Proceeds: MYR 1.5-2.0B')
    B(doc, 'Additional Borrowing: MYR 1.5-2.0B (incremental sustainability-linked instruments preferred)')
    B(doc, 'Dividend Reinvestment / Scrip: MYR 0.5-0.8B contingency')
    P(doc, 'Capital-structure target: Net Debt / EBITDA < 3.0× through cycle (hard covenant '
           'limit 3.5×). Interest cover > 4.0×. Cash + undrawn committed facilities ≥ 12 months '
           'of debt maturity at all times.')
    doc.add_page_break()


# ── Section 6-8: People, Risks, Governance ─────────────────────────────────
def s6_people(doc):
    H(doc, '6. People & Culture Agenda', level=1)
    P(doc, 'Zava Forward 2030 cannot be delivered without parallel investment in talent, '
           'leadership pipeline depth, and culture.')
    H(doc, '6.1 Talent Initiatives', level=2)
    B(doc, 'Re-skilling fund of MYR 60M committed FY2025-2027 — priority capabilities: AI, data, sustainability, digital product, advanced manufacturing')
    B(doc, 'Top-100 succession depth: every Group ExCo and Division CEO role to have ≥ 2 ready-now successors by FY2026 (current: 71% have at least one identified successor)')
    B(doc, 'Internal-fill rate of senior roles: target > 65% by FY2027 (FY25 baseline 52%)')
    B(doc, 'Leadership pipeline: Zava Academy Senior Leader Programme launched for Top-300 (cohort size 60 per year)')

    H(doc, '6.2 Culture & Engagement', level=2)
    B(doc, 'Engagement (Pulse) target 78 by FY2027 (FY25 baseline 74). Top driver: career growth perception.')
    B(doc, 'Voluntary attrition target < 10% by FY2027 (FY25 baseline 12.4%). Focus on digital and sustainability roles where attrition is highest.')
    B(doc, 'Reward modernisation: ESG-linked PSU 15% weight (already implemented FY2024); digital-skills premium for critical-shortage roles')
    B(doc, 'Hybrid working: 3-2 model (3 days office, 2 days flexible) maintained; revisit FY2026 based on productivity and engagement data')

    H(doc, '6.3 Diversity, Equity & Inclusion', level=2)
    B(doc, 'Women in management JG6+: 32% by FY2027, 40% by FY2030 (FY25 baseline 28.4%)')
    B(doc, 'Persons with disability: 1.2% by FY2027, 2.0% by FY2030 (FY25 baseline 0.8%)')
    B(doc, 'Local-national senior managers in each ASEAN country > 78% by FY2027')
    B(doc, 'Pay equity (adjusted): < 1% gender gap maintained year-on-year')

    doc.add_page_break()


def s7_risks(doc):
    H(doc, '7. Strategic Risks & Mitigations', level=1)
    T(doc, ['Risk', 'Impact', 'Likelihood', 'Mitigation Focus'],
      [
          ['Recovery in Properties slower than plan', 'High', 'High', 'Aggressive sell-down + asset disposal; format conversion'],
          ['ASEAN carbon pricing earlier or steeper than expected', 'High', 'Medium', 'Accelerated abatement capex; nature-based offsets reserve'],
          ['IDR/USD continues to weaken', 'Medium', 'High', 'Layered hedging programme; pricing pass-through; ID localisation'],
          ['BPO AI productivity → client share-of-savings claims', 'Medium', 'High', 'Outcome-based pricing pivot; vertical specialisation'],
          ['Healthcare patient volumes structurally lower', 'High', 'Medium', 'Format conversion to outpatient/day-surgery; ambulatory M&A'],
          ['M&A pipeline mis-execution (overpay, integration)', 'High', 'Medium', 'Disciplined hurdle rates; pre-deal integration plans; governance'],
          ['Cyber incident / ransomware', 'High', 'Medium', 'Zero-trust roll-out FY2025; quarterly red team; cyber insurance'],
          ['Geopolitical (US-China supply-chain)', 'Medium', 'Medium', 'Multi-source critical components; ASEAN+1 supplier diversification'],
          ['Property writedowns (Penang Tower, Mid-Valley)', 'Medium', 'High', 'Conservative carrying values; quarterly impairment testing'],
          ['Regulatory change (PDPA MY 2024 amend, UU PDP ID)', 'Medium', 'High', 'DPO appointed; mandatory breach process; staff training'],
      ], col_widths=[5.0, 1.8, 1.8, 8.5])
    doc.add_page_break()


def s8_governance(doc):
    H(doc, '8. Strategy Cascade & Governance', level=1)
    P(doc, 'Zava Forward 2030 is operationalised through a structured cascade and governance '
           'rhythm.')
    H(doc, '8.1 KPI Cascade', level=2)
    B(doc, 'Group strategic KPIs (15) — owned by Group ExCo, tracked monthly')
    B(doc, 'Division strategic KPIs (8-10 per division) — owned by Division CEO, cascaded to function heads')
    B(doc, 'Function strategic KPIs (5-8 per function) — owned by Function head, cascaded to teams')
    B(doc, 'Individual KPIs — every JG5+ employee has minimum two strategy-aligned KPIs in their annual plan')

    H(doc, '8.2 Governance Cadence', level=2)
    T(doc, ['Forum', 'Frequency', 'Lead', 'Focus'],
      [
          ['Group ExCo', 'Weekly', 'Group CEO', 'Operational rhythm + emerging issues'],
          ['Group Strategy Council', 'Monthly', 'Group CEO + Strategy Director', 'KPI tracking, recovery plans, M&A pipeline'],
          ['Investment Committee', 'Monthly', 'Group CFO', 'Capex > MYR 100M, M&A, divestments'],
          ['Audit & Risk Committee', 'Quarterly', 'INED Chair', 'Risk register, internal audit, financial integrity'],
          ['Sustainability Committee', 'Quarterly', 'INED Chair', 'Climate KPIs, ESG disclosures, social agenda'],
          ['Board', 'Bi-monthly', 'Chairman', 'Strategy review, major decisions, regulatory'],
          ['Strategy Refresh', 'Annual (Sept)', 'Group Strategy Director', 'Full Zava Forward 2030 refresh + cascade'],
          ['Strategy Deep-Refresh', '3-yearly', 'Group CEO + Board', 'Major recalibration of pillars and capital allocation'],
      ], col_widths=[4.5, 2.0, 4.0, 6.5])

    P(doc, 'Strategy delivery is the single most important deliverable of every senior leader. '
           'It is reflected in compensation, career outcomes, and accountability conversations. '
           'Zava Forward 2030 is not a document; it is the operating system of the Group.', italic=True)


def main(out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    cover(doc)
    s1_context(doc)
    s2_pillars(doc)
    s3_divisions(doc)
    s4_ma(doc)
    s5_capex(doc)
    s6_people(doc)
    s7_risks(doc)
    s8_governance(doc)
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else 'files/03_Zava_Group_Strategy_Framework.docx'
    main(out)
