"""Generates 02_Zava_Group_Policy_Handbook.docx — 10 sections × ~3-4 pages each.

Sections:
  1. Corporate Governance
  2. Code of Ethics & Business Conduct
  3. Human Resources Policy
  4. Financial Governance & Delegation of Authority
  5. Procurement & Vendor Management
  6. Data Privacy & Protection (PDPA MY + UU PDP ID)
  7. Environmental, Social & Governance (ESG)
  8. Health, Safety & Environment (HSE)
  9. Legal & Compliance
 10. Enterprise Risk Management
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from gen_ref_files.common import DIVISIONS

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20 if level == 1 else 14 if level == 2 else 11)
    run.font.name = 'Aptos'
    run.font.color.rgb = color
    return p


def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Aptos'
    if not p.runs:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Aptos'
    else:
        p.runs[0].text = text


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if not p.runs:
        run = p.add_run(text)
    else:
        p.runs[0].text = text
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Aptos'


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Aptos'
        shade(cell, '1F4E79')
    # Data
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Aptos'
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    return t


# ── Cover & TOC ─────────────────────────────────────────────────────────────
def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('ZAVA GROUP\n')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'
    run = p.add_run('Group Policy Handbook')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    run.font.name = 'Aptos'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Version 8.2 | Effective 1 January FY2025\nBoard-Approved | Group-wide application')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = 'Aptos'

    doc.add_page_break()
    add_heading(doc, 'Table of Contents', level=1)
    sections = [
        '1. Corporate Governance',
        '2. Code of Ethics & Business Conduct',
        '3. Human Resources Policy',
        '4. Financial Governance & Delegation of Authority',
        '5. Procurement & Vendor Management',
        '6. Data Privacy & Protection (PDPA MY + UU PDP ID)',
        '7. Environmental, Social & Governance (ESG)',
        '8. Health, Safety & Environment (HSE)',
        '9. Legal & Compliance',
        '10. Enterprise Risk Management',
        'Appendix A — Approval Matrix',
        'Appendix B — Definitions & Acronyms',
        'Appendix C — Policy Owner & Review Cycle',
    ]
    for s in sections:
        add_para(doc, s, size=11)

    add_para(doc, '\nThis Handbook supersedes Version 8.1 (FY2024). All Zava Group employees, '
                  'contractors and authorised representatives are required to comply with the '
                  'policies set out herein, as adapted to local regulatory requirements in '
                  'Malaysia, Indonesia, Singapore, Vietnam, Thailand and the Philippines.', italic=True, size=10)
    doc.add_page_break()


# ── Section 1: Corporate Governance ───────────────────────────────────────
def section_1(doc):
    add_heading(doc, '1. Corporate Governance', level=1)
    add_para(doc, '1.1 Purpose & Scope', bold=True, size=12)
    add_para(doc, 'This Section establishes the corporate-governance framework for Zava Group '
                  'and all subsidiaries (collectively, the "Group"). It is informed by the Malaysian '
                  'Code on Corporate Governance (MCCG 2021), the Indonesia Financial Services '
                  'Authority Regulation No. 21/POJK.04/2015 on Public Company Governance, and '
                  'aligns with the OECD Principles of Corporate Governance.')

    add_para(doc, '1.2 Board of Directors', bold=True, size=12)
    add_para(doc, 'The Group is led by a unitary Board of nine (9) Directors, comprising:')
    add_bullet(doc, 'One (1) Non-Executive Group Chairman (independent)')
    add_bullet(doc, 'One (1) Group Chief Executive (executive)')
    add_bullet(doc, 'Four (4) Independent Non-Executive Directors')
    add_bullet(doc, 'Three (3) Non-Independent Non-Executive Directors representing major shareholders')
    add_para(doc, 'Independence is reviewed annually against MCCG and Bursa Malaysia Listing '
                  'Requirements; any Director serving more than nine (9) cumulative years is '
                  'subject to shareholder re-election with two-tier voting.')

    add_para(doc, '1.3 Board Committees', bold=True, size=12)
    add_table(doc,
        ['Committee', 'Mandate', 'Composition', 'Quorum', 'Frequency'],
        [
            ['Audit & Risk', 'External audit, internal audit, ERM, financial integrity', '3 INED + Group CFO (in attendance)', '2 INED', 'Quarterly + ad-hoc'],
            ['Nomination & Remuneration', 'Director nomination, succession, executive remuneration', '3 INED + Chairman', '2 INED', 'Half-yearly'],
            ['Sustainability & ESG', 'Climate, social, governance disclosures and TCFD', '2 INED + 1 NED + Group CSO', '2', 'Quarterly'],
            ['Investment & Capital', 'Capex >MYR 100M, M&A, divestments, capital structure', '3 NED + Group CFO', '3', 'Monthly'],
            ['Related Party Transactions', 'RPT review and disclosure per MFRS 124 / PSAK 7', '3 INED', '2', 'Quarterly + ad-hoc'],
        ], col_widths=[3.5, 5.0, 4.0, 1.8, 2.5])

    add_para(doc, '1.4 Director Duties', bold=True, size=12)
    add_para(doc, 'Each Director shall:')
    add_numbered(doc, 'Act in good faith and in the best interests of the Group, exercising independent judgement.')
    add_numbered(doc, 'Disclose any direct or indirect interest in any transaction, contract or arrangement involving the Group, in writing, prior to deliberation.')
    add_numbered(doc, 'Devote sufficient time to discharge the role; cumulative directorships in listed entities shall not exceed five (5) at any time.')
    add_numbered(doc, 'Attend not fewer than 75% of scheduled Board meetings annually.')
    add_numbered(doc, 'Complete a minimum of 25 CPD/CPE hours per annum on relevant topics including ESG, cyber risk, AI governance and sector developments.')

    add_para(doc, '1.5 Reserved Matters for the Board', bold=True, size=12)
    add_para(doc, 'The following matters are reserved exclusively for the Board and may not be '
                  'delegated to management or any sub-committee:')
    matters = [
        ('Capital expenditure or commitment', '> MYR 250 million single transaction'),
        ('Mergers, acquisitions, joint ventures, divestments', 'Any transaction > MYR 200 million'),
        ('Borrowings (incremental)', '> MYR 500 million single facility'),
        ('Dividend declaration', 'All interim and final dividends'),
        ('Issuance of equity / hybrid instruments', 'All issuances'),
        ('Annual budget and 5-year strategic plan approval', 'Both annual'),
        ('Material litigation', 'Claim or counterclaim > MYR 50 million'),
        ('Senior executive appointments', 'Group CEO, Group CFO, Division CEOs, Heads of Internal Audit & Compliance'),
        ('Material related-party transactions', 'Aggregate exceeding 5% of NTA'),
        ('Whistleblowing investigations involving senior executives', 'All cases'),
    ]
    add_table(doc, ['Matter', 'Threshold / Description'], matters, col_widths=[7.0, 10.0])

    add_para(doc, '1.6 Annual Board Effectiveness Review', bold=True, size=12)
    add_para(doc, 'A formal annual board-effectiveness review is undertaken, with an external '
                  'reviewer engaged at least once every three (3) years. The 2024 review (conducted '
                  'by KPMG Governance Advisory) identified four priority focus areas for FY2025: '
                  '(i) succession depth for the Group CEO and CFO roles; (ii) climate-risk literacy '
                  'across the Board; (iii) cybersecurity oversight maturity; and (iv) ESG-linked '
                  'executive compensation design. Action plans against each are tracked by the '
                  'Nomination & Remuneration Committee on a quarterly basis.')
    doc.add_page_break()


# ── Section 2: Code of Ethics ──────────────────────────────────────────────
def section_2(doc):
    add_heading(doc, '2. Code of Ethics & Business Conduct', level=1)
    add_para(doc, '2.1 Standards', bold=True, size=12)
    add_para(doc, 'All Zava Group employees, officers, Directors, contractors and authorised '
                  'representatives ("Personnel") shall conduct themselves with integrity, fairness, '
                  'respect and accountability. The standards in this Code are non-negotiable and '
                  'apply globally without exception.')

    add_para(doc, '2.2 Anti-Bribery & Corruption', bold=True, size=12)
    add_para(doc, 'Zava Group operates a strict zero-tolerance policy against bribery and '
                  'corruption, in compliance with the Malaysian Anti-Corruption Commission Act 2009 '
                  '(MACC Act, including Section 17A on corporate liability), the Indonesian Law No. '
                  '20/2001 on Corruption Eradication, the UK Bribery Act 2010 (extra-territorial), '
                  'and the US Foreign Corrupt Practices Act (FCPA).')
    add_para(doc, 'Personnel shall not, directly or indirectly:')
    add_bullet(doc, 'Offer, promise, give, demand or accept any financial or other advantage to or from any person, public official or commercial counterparty in connection with Zava Group business;')
    add_bullet(doc, 'Engage in facilitation payments, regardless of local custom or practice;')
    add_bullet(doc, 'Use cash or off-the-books accounts for any business transaction;')
    add_bullet(doc, 'Make political donations on behalf of the Group without prior approval of the Board;')
    add_bullet(doc, 'Provide gifts, hospitality or entertainment exceeding the thresholds in §2.3.')

    add_para(doc, '2.3 Gifts, Hospitality & Entertainment Threshold Matrix', bold=True, size=12)
    add_table(doc,
        ['Counterparty Type', 'Single-event Gift', 'Single-event Hospitality', 'Annual Cumulative', 'Pre-approval required'],
        [
            ['Public official (any country)', 'MYR 0', 'MYR 0', 'MYR 0', 'PROHIBITED — escalate to GC'],
            ['Regulator under active engagement', 'MYR 0', 'MYR 200 (light meal only)', 'MYR 200', 'Yes — Compliance'],
            ['Commercial supplier (active tender)', 'MYR 0', 'MYR 0', 'MYR 0', 'PROHIBITED'],
            ['Commercial supplier (BAU)', 'MYR 200', 'MYR 500', 'MYR 1,500', 'Above threshold — Line Manager'],
            ['Customer (private sector)', 'MYR 500', 'MYR 1,500', 'MYR 5,000', 'Above threshold — Line Manager'],
            ['Industry association / community', 'MYR 250', 'MYR 750', 'MYR 2,000', 'Above threshold — Line Manager'],
        ], col_widths=[4.5, 2.5, 3.0, 3.0, 3.5])
    add_para(doc, 'All gifts and hospitality, regardless of value, shall be logged in the Group '
                  'Gifts & Hospitality Register within five (5) business days of the event, via the '
                  'Compliance Portal.')

    add_para(doc, '2.4 Conflicts of Interest', bold=True, size=12)
    add_para(doc, 'Personnel must declare any actual, potential or perceived conflict of interest '
                  'on engagement and annually thereafter (or sooner if circumstances change). '
                  'Examples include: family members in supplier or customer organisations; '
                  'directorships in unrelated businesses; private investments in companies with '
                  'which Zava Group transacts; secondary employment or consulting arrangements.')

    add_para(doc, '2.5 Whistleblowing & Speak-up', bold=True, size=12)
    add_para(doc, 'Zava Group operates an independent Speak-Up channel administered by Deloitte '
                  'Forensic, accessible 24/7 in English, Bahasa Malaysia, Bahasa Indonesia, '
                  'Mandarin, Tamil, Thai, Vietnamese and Tagalog. Reports may be submitted '
                  'anonymously. Retaliation against any person making a good-faith report is itself '
                  'a serious disciplinary offence and a criminal offence under the Malaysian '
                  'Whistleblower Protection Act 2010 and equivalent ID/SG/VN/TH/PH legislation.')
    add_para(doc, 'Speak-Up channel: speakup.zava.com | +603-2200-0911 (MY) | +6221-2922-0911 (ID) | speakup@zava.com')

    add_para(doc, '2.6 Sanctions, Trade Compliance & Export Controls', bold=True, size=12)
    add_para(doc, 'Zava Group complies with applicable sanctions regimes including UN Security '
                  'Council, EU, OFAC, HMT and ASEAN-specific designations. The Sanctions Compliance '
                  'Programme covers customer screening, supplier screening, dual-use goods '
                  'classification, end-use certification, and a quarterly sanctions-list refresh '
                  'across all transactional systems via the Group Refinitiv World-Check feed.')

    add_para(doc, '2.7 Disciplinary Consequences', bold=True, size=12)
    add_para(doc, 'Breaches of this Code attract disciplinary action up to and including '
                  'termination of employment, civil recovery of losses, and criminal referral. '
                  'Senior management breaches are reported to the Audit & Risk Committee.')
    doc.add_page_break()


# ── Section 3: HR Policy ───────────────────────────────────────────────────
def section_3(doc):
    add_heading(doc, '3. Human Resources Policy', level=1)
    add_para(doc, '3.1 Recruitment & Onboarding', bold=True, size=12)
    add_para(doc, 'Recruitment shall be conducted on the basis of merit, qualification and fit '
                  'for role, free from discrimination on the basis of race, religion, gender, age, '
                  'disability, marital status, sexual orientation or political affiliation. All '
                  'Group recruitment uses the Workday Recruiting platform and is subject to the '
                  'four-eyes principle (recruiter + hiring manager) at offer stage.')

    add_para(doc, '3.2 Compensation & Benefits Framework', bold=True, size=12)
    add_para(doc, 'Compensation comprises:')
    add_bullet(doc, 'Fixed: base salary, contractual allowances, statutory contributions (EPF/SOCSO/EIS for MY; BPJS for ID; CPF for SG)')
    add_bullet(doc, 'Variable short-term: annual bonus tied to Group, Division and Individual KPIs (typically 8-25% of base for non-executives, up to 100% for senior executives)')
    add_bullet(doc, 'Variable long-term: Restricted Share Units (RSUs) and Performance Share Units (PSUs) for Job Grade JG7+ with 3-year vesting and ESG-linked performance modifier (15% weight)')
    add_bullet(doc, 'Benefits: medical (incl. spouse + 4 children), dental, optical, GTL/GHS insurance, parental leave (16 weeks maternity, 4 weeks paternity), bereavement, study assistance, retirement preparation programme')

    add_para(doc, '3.3 Pay Equity & Gender Pay Gap', bold=True, size=12)
    add_para(doc, 'Group conducts an annual pay-equity audit using Mercer methodology. The FY2024 '
                  'audit closed the unadjusted gender pay gap to 4.1% (from 6.8% in FY2022). '
                  'Adjusted gap (controlling for grade, function, geography, tenure) was 0.7%. '
                  'Target FY2027: unadjusted gap < 2%.')

    add_para(doc, '3.4 Performance Management', bold=True, size=12)
    add_para(doc, 'Annual performance is assessed against:')
    add_bullet(doc, 'Objectives & Key Results (50%) — set in October, reviewed at March mid-year, finalised in September')
    add_bullet(doc, 'Behaviours against Zava Values (30%) — Care, Courage, Collaboration, Curiosity')
    add_bullet(doc, 'Stakeholder feedback (20%) — 360-degree input for JG6+; manager and peer feedback for JG5 and below')
    add_para(doc, 'Performance is calibrated by Function-level Talent Review, then Group People Committee. '
                  'Forced ranking is not applied; distributions are guided but not enforced.')

    add_para(doc, '3.5 Talent Development', bold=True, size=12)
    add_para(doc, 'All Personnel are entitled to a minimum 40 hours of formal learning per annum. '
                  'The Zava Academy operates four schools: Leadership, Functional Excellence, '
                  'Digital & Data, and ESG. Senior leaders (JG7+) are required to undertake an '
                  'executive development programme (e.g., INSEAD AMP, Wharton Advanced Management) '
                  'within 18 months of appointment.')

    add_para(doc, '3.6 Diversity, Equity & Inclusion', bold=True, size=12)
    add_table(doc,
        ['DEI Metric', 'FY2022', 'FY2023', 'FY2024', 'FY2025 Target', 'FY2027 Aspiration'],
        [
            ['% Women in workforce (Group)', '38.4%', '40.1%', '41.7%', '43.0%', '45.0%'],
            ['% Women in management (JG6+)', '24.1%', '26.7%', '28.4%', '32.0%', '40.0%'],
            ['% Women on Board', '22.2%', '22.2%', '33.3%', '33.3%', '40.0%'],
            ['Persons with disability', '0.4%', '0.6%', '0.8%', '1.2%', '2.0%'],
            ['Indigenous / Bumiputera (MY)', '52.1%', '52.8%', '53.4%', '54.0%', 'Maintain'],
            ['Local nationals in Sr Mgmt (each ASEAN country)', '67%', '71%', '74%', '78%', '> 80%'],
        ], col_widths=[5.0, 1.8, 1.8, 1.8, 2.4, 2.7])

    add_para(doc, '3.7 Grievance & Disciplinary', bold=True, size=12)
    add_para(doc, 'A formal grievance may be raised with the Line Manager, HR Business Partner, '
                  'Division Head of HR, or Group Chief Human Resources Officer. Grievances must be '
                  'acknowledged within five (5) business days and concluded within thirty (30) '
                  'business days. Disciplinary action follows a four-stage process: verbal warning, '
                  'written warning, final written warning, and termination — with a domestic '
                  'inquiry preceding any termination for misconduct.')

    add_para(doc, '3.8 Workplace Wellbeing', bold=True, size=12)
    add_para(doc, 'The Group provides 24/7 access to the EAP (Employee Assistance Programme) via '
                  'Lyra Health, covering up to twelve (12) sessions of confidential counselling per '
                  'employee per annum, plus extended support for severe mental-health conditions '
                  'as clinically indicated. Workload sustainability is monitored via the quarterly '
                  'Pulse Survey, with red-flag teams escalated to Division CEOs.')
    doc.add_page_break()


# ── Section 4: Financial Governance ────────────────────────────────────────
def section_4(doc):
    add_heading(doc, '4. Financial Governance & Delegation of Authority', level=1)
    add_para(doc, '4.1 Approval Authority Matrix (excerpt — full version in Appendix A)', bold=True, size=12)
    add_table(doc,
        ['Decision Type', 'JG5/Mgr', 'JG6/Snr Mgr', 'JG7/AGM', 'JG8/GM', 'Div CFO', 'Div CEO', 'Group CFO', 'Group CEO', 'Inv. Cmte', 'Board'],
        [
            ['Operating expenditure', '50K', '250K', '1M', '5M', '15M', '30M', '75M', '150M', '—', '> 250M'],
            ['Capital expenditure', '—', '100K', '500K', '2M', '8M', '20M', '50M', '100M', '250M', '> 250M'],
            ['Hiring (annual fixed cost)', '—', '300K', '1M', '3M', '8M', '15M', '40M', '80M', '—', '> 100M'],
            ['Marketing campaigns', '20K', '100K', '500K', '2M', '5M', '10M', '25M', '50M', '—', '> 75M'],
            ['Customer credit limit', '—', '500K', '2M', '5M', '15M', '30M', '60M', '100M', '—', '> 150M'],
            ['Bad debt write-off', '—', '50K', '250K', '1M', '5M', '10M', '25M', '50M', '—', '> 75M'],
            ['Litigation settlement', '—', '100K', '500K', '2M', '5M', '15M', '40M', '80M', '—', '> 100M'],
            ['M&A / divestment', '—', '—', '—', '—', '—', '—', '—', '—', '< 200M', '> 200M'],
            ['New borrowing facility', '—', '—', '—', '—', '—', '—', '—', '< 200M', '< 500M', '> 500M'],
            ['Currency hedging (notional)', '—', '—', '5M', '20M', '50M', '100M', '250M', '500M', '—', '> 500M'],
        ], col_widths=[3.0, 1.0, 1.2, 1.0, 1.0, 1.2, 1.2, 1.4, 1.4, 1.4, 1.4])
    add_para(doc, 'All amounts in MYR. Equivalents in IDR/SGD/THB/VND/PHP apply per the Group '
                  'monthly FX policy rate.')

    add_para(doc, '4.2 Budgeting & Forecasting Cycle', bold=True, size=12)
    add_para(doc, 'The Group operates an integrated annual planning cycle:')
    add_bullet(doc, 'July-August: 5-year strategic-plan refresh (Group Strategy + Division Heads)')
    add_bullet(doc, 'September: Group budget guidelines issued (volume, FX, inflation, tax assumptions)')
    add_bullet(doc, 'October: Division budget submissions (bottom-up)')
    add_bullet(doc, 'November: Group consolidation, challenge sessions, draft Board paper')
    add_bullet(doc, 'December: Board approval; cascade to KPI scorecards')
    add_bullet(doc, 'Quarterly: Re-forecast with rolling 4-quarter view (R4Q); variance > ±5% triggers root-cause review')

    add_para(doc, '4.3 Financial Reporting Standards', bold=True, size=12)
    add_para(doc, 'The Group prepares consolidated financial statements under MFRS (Malaysian '
                  'Financial Reporting Standards), which are equivalent to IFRS. Subsidiary '
                  'reporting under PSAK (Indonesia), TFRS (Thailand), VAS (Vietnam), PFRS '
                  '(Philippines) is reconciled monthly through the Oracle EPM consolidation engine. '
                  'The Group has elected MFRS 16 (Leases) full retrospective method since FY2019, '
                  'and applies MFRS 9 expected-credit-loss model for all financial instruments.')

    add_para(doc, '4.4 Treasury & Cash Management', bold=True, size=12)
    add_para(doc, 'Treasury operates a centralised in-house bank (IHB) covering 11 divisions and '
                  '143 entities. Counterparty bank exposure is limited per the Group Counterparty '
                  'Risk Policy; aggregate exposure to any single bank shall not exceed 20% of total '
                  'cash and short-term investments unless approved by the Investment Committee. '
                  'Cash forecasting is on a rolling 13-week basis, with weekly variance review at '
                  'Group Treasurer level.')
    doc.add_page_break()


# ── Section 5: Procurement ─────────────────────────────────────────────────
def section_5(doc):
    add_heading(doc, '5. Procurement & Vendor Management', level=1)
    add_para(doc, '5.1 Sourcing Principles', bold=True, size=12)
    add_para(doc, 'All procurement shall be conducted under the principles of fair competition, '
                  'value for money, ethical sourcing, supply continuity, and total cost of '
                  'ownership over headline price. The Source-to-Pay (S2P) platform is SAP Ariba, '
                  'integrated with Workday and Oracle EBS.')

    add_para(doc, '5.2 Sourcing Method by Spend Threshold', bold=True, size=12)
    add_table(doc,
        ['Spend (per contract)', 'Sourcing Method', 'Min. Quotes', 'Approval', 'Cycle Time SLA'],
        [
            ['< MYR 50K', 'Direct purchase / catalogue', '1', 'Requestor + Line Mgr', '3 days'],
            ['MYR 50K — 250K', 'Three-quote sourcing (RFQ)', '3', 'Category Manager', '10 days'],
            ['MYR 250K — 2M', 'Limited tender', '5', 'Procurement Director', '20 days'],
            ['MYR 2M — 25M', 'Open tender + technical evaluation', '5+', 'Procurement Cmte', '40 days'],
            ['> MYR 25M', 'Open tender + Investment Cmte review', '7+', 'Investment Cmte / Board', '60 days'],
            ['Sole source (any value)', 'Justification memo + 2nd-line review', 'N/A', 'Group CPO + GC', 'Per case'],
        ], col_widths=[3.0, 4.5, 1.5, 4.0, 2.0])

    add_para(doc, '5.3 Vendor Onboarding & Due Diligence', bold=True, size=12)
    add_para(doc, 'All new vendors are subject to:')
    add_bullet(doc, 'Identity & beneficial ownership verification (KYV) per AML / CFT standards')
    add_bullet(doc, 'Sanctions screening (Refinitiv World-Check) — refreshed quarterly throughout the relationship')
    add_bullet(doc, 'Anti-Bribery & Corruption certification (signed)')
    add_bullet(doc, 'Modern Slavery & Forced Labour declaration (UK MSA aligned)')
    add_bullet(doc, 'ESG self-assessment via EcoVadis (mandatory for vendors > MYR 500K p.a.)')
    add_bullet(doc, 'Cybersecurity assessment (mandatory for vendors handling Group data — SIG Lite questionnaire)')
    add_bullet(doc, 'Financial-health screening (Bureau Van Dijk / Dun & Bradstreet)')

    add_para(doc, '5.4 Conflict of Interest in Procurement', bold=True, size=12)
    add_para(doc, 'Personnel involved in any procurement decision shall declare any relationship '
                  '— direct or through immediate family — with any participating vendor. '
                  'Undeclared conflicts are a serious disciplinary offence (refer §2.7).')

    add_para(doc, '5.5 Supplier Performance Management', bold=True, size=12)
    add_para(doc, 'Strategic and Critical vendors (Tier 1, accounting for ~80% of spend) are '
                  'subject to quarterly business reviews against contractual KPIs covering quality, '
                  'delivery, cost, innovation, ESG and security. Underperformance triggers a 90-day '
                  'remediation plan; persistent underperformance triggers debarment review.')
    doc.add_page_break()


# ── Section 6: Data Privacy ────────────────────────────────────────────────
def section_6(doc):
    add_heading(doc, '6. Data Privacy & Protection', level=1)
    add_para(doc, '6.1 Regulatory Framework', bold=True, size=12)
    add_para(doc, 'The Group complies with applicable personal-data legislation in each operating '
                  'country, including:')
    add_bullet(doc, 'Malaysia: Personal Data Protection Act 2010 (PDPA), as amended by the PDP Amendment Act 2024 (effective 1 June 2025), including the new Data Protection Officer and Mandatory Breach Notification provisions')
    add_bullet(doc, 'Indonesia: Undang-Undang Pelindungan Data Pribadi (UU PDP) Law No. 27/2022 (full enforcement October 2024)')
    add_bullet(doc, 'Singapore: Personal Data Protection Act 2012 (revised 2020), including the Do Not Call provisions')
    add_bullet(doc, 'Thailand: Personal Data Protection Act 2019 (PDPA Thailand)')
    add_bullet(doc, 'Vietnam: Decree 13/2023/ND-CP on Personal Data Protection')
    add_bullet(doc, 'Philippines: Data Privacy Act 2012 (RA 10173)')
    add_bullet(doc, 'European Economic Area: GDPR (where applicable to EU data subjects)')

    add_para(doc, '6.2 Data Protection Principles', bold=True, size=12)
    add_para(doc, 'All processing of personal data shall be:')
    add_numbered(doc, 'Lawful — supported by a legitimate basis (consent, contract necessity, legal obligation, vital interest, public task or legitimate interest)')
    add_numbered(doc, 'Fair & transparent — supported by clear privacy notices in the relevant local language')
    add_numbered(doc, 'Purpose-limited — used only for stated and compatible purposes')
    add_numbered(doc, 'Data-minimised — limited to what is necessary for the stated purpose')
    add_numbered(doc, 'Accurate — kept up to date with reasonable steps to correct errors')
    add_numbered(doc, 'Time-bound — retained no longer than necessary per the Records Retention Schedule')
    add_numbered(doc, 'Secure — protected by appropriate technical and organisational measures')
    add_numbered(doc, 'Accountable — documented and demonstrable to regulators on request')

    add_para(doc, '6.3 Data Subject Rights', bold=True, size=12)
    add_table(doc,
        ['Right', 'PDPA MY', 'UU PDP ID', 'PDPA SG', 'GDPR'],
        [
            ['Access', 'Yes (21 days)', 'Yes (3 days)', 'Yes (30 days)', 'Yes (30 days)'],
            ['Correction / Rectification', 'Yes (21 days)', 'Yes (3 days)', 'Yes', 'Yes'],
            ['Erasure / Right to be Forgotten', 'Limited', 'Yes', 'Limited', 'Yes'],
            ['Restriction of Processing', 'Limited', 'Yes', 'No express right', 'Yes'],
            ['Data Portability', 'No express right', 'Yes', 'Yes (in transition)', 'Yes'],
            ['Object to Processing', 'Yes', 'Yes', 'Limited', 'Yes'],
            ['Object to Automated Decision-Making', 'Limited', 'Yes', 'Limited', 'Yes'],
            ['Withdraw Consent', 'Yes', 'Yes', 'Yes', 'Yes'],
        ], col_widths=[5.0, 2.8, 2.8, 2.8, 2.8])

    add_para(doc, '6.4 Data Breach Notification', bold=True, size=12)
    add_para(doc, 'Any actual or suspected personal-data breach must be reported to the Group '
                  'Data Protection Office within twenty-four (24) hours of discovery. Onward '
                  'regulatory notification timelines are:')
    add_bullet(doc, 'Malaysia (PDPA 2024 amendment): 72 hours to JPDP if there is significant harm')
    add_bullet(doc, 'Indonesia (UU PDP): 3 × 24 hours to the PDP Authority')
    add_bullet(doc, 'Singapore: 72 hours to PDPC if significant scale or significant harm')
    add_bullet(doc, 'EEA: 72 hours to lead supervisory authority (GDPR Art. 33)')

    add_para(doc, '6.5 Cross-Border Data Transfer', bold=True, size=12)
    add_para(doc, 'Transfers outside the country of origin require: a legitimate basis; an '
                  'adequate level of protection in the destination country (white-list reliance, '
                  'where available); or an appropriate safeguard such as Group Binding Corporate '
                  'Rules (BCR), Standard Contractual Clauses (SCC), or explicit informed consent. '
                  'A Transfer Impact Assessment (TIA) is required for high-risk data and high-risk '
                  'destinations.')

    add_para(doc, '6.6 AI & Algorithmic Accountability', bold=True, size=12)
    add_para(doc, 'AI-assisted processing of personal data, including generative AI tools and '
                  'automated decision-making, is governed by the Group AI Use Policy (Appendix to '
                  'Section 6). Key requirements: AI Use Case Register entry prior to deployment, '
                  'human-in-the-loop for material decisions, fairness testing for protected '
                  'characteristics, and explainability standards aligned with the EU AI Act risk '
                  'tiering.')
    doc.add_page_break()


# ── Section 7: ESG ─────────────────────────────────────────────────────────
def section_7(doc):
    add_heading(doc, '7. Environmental, Social & Governance (ESG)', level=1)
    add_para(doc, '7.1 Sustainability Strategy — "Zava Forward 2030"', bold=True, size=12)
    add_para(doc, 'The Group has adopted "Zava Forward 2030" as its sustainability strategy, with '
                  'three pillars: Climate Action, Inclusive Prosperity, and Trusted Stewardship. '
                  'The strategy is aligned with TCFD, GRI Standards 2021, IFRS S1 / S2 (effective '
                  'FY2025 reporting), Bursa Sustainability Reporting Guide v3, and the OJK '
                  'Indonesia Sustainability Report regulations.')

    add_para(doc, '7.2 Climate Targets & Pathways', bold=True, size=12)
    add_table(doc,
        ['Target', 'Baseline', 'Interim (2030)', 'Long-term (2050)', 'Status FY2024', 'Validation'],
        [
            ['Scope 1+2 GHG (tCO2e)', 'FY2019: 8.4m', '−45% vs base', 'Net Zero', '−18%', 'SBTi 1.5°C aligned'],
            ['Scope 3 GHG (tCO2e)', 'FY2021: 24.1m', '−25% vs base', 'Net Zero', '−7%', 'SBTi-FLAG'],
            ['Renewable electricity', '11% (FY2019)', '70%', '100%', '34%', 'RE100 commitment'],
            ['Water withdrawal (high-stress areas)', 'FY2019: 12.4 ML', '−35%', '−50%', '−14%', 'CDP Water'],
            ['Waste diverted from landfill', 'FY2019: 41%', '85%', '95%', '67%', 'Internal'],
            ['Plantation NDPE compliance (Agri div)', '78% (FY2020)', '100%', 'Maintain', '94%', 'RSPO + ZSL SPOTT'],
        ], col_widths=[5.0, 2.5, 2.0, 2.5, 2.0, 3.0])

    add_para(doc, '7.3 Just Transition & Social Pillar', bold=True, size=12)
    add_para(doc, 'The Group is committed to a just transition that prioritises:')
    add_bullet(doc, 'Reskilling at-risk roles (e.g., legacy energy operations) through the Zava Reskilling Fund (MYR 60M committed FY2025-2027)')
    add_bullet(doc, 'Living wage certification across direct workforce by FY2026 (Fair Wage Network methodology)')
    add_bullet(doc, 'Community development MYR 50M annual minimum (target ROI > 3× via SROI methodology)')
    add_bullet(doc, 'Indigenous community consent (FPIC) for any plantation, infrastructure or extractive project')

    add_para(doc, '7.4 Climate-related Financial Disclosures', bold=True, size=12)
    add_para(doc, 'The Group discloses climate-related financial information per TCFD across '
                  'Governance, Strategy, Risk Management and Metrics & Targets. Scenario analysis '
                  'has been conducted under three pathways: orderly transition (NGFS Net Zero '
                  '2050), disorderly transition (NGFS Delayed), and hot-house (NGFS Current '
                  'Policies), at 2030 and 2050 horizons. Identified financial impacts are reflected '
                  'in long-range capital allocation and division-level strategy reviews.')

    add_para(doc, '7.5 Sustainability-Linked Financing', bold=True, size=12)
    add_para(doc, 'The Group has issued / drawn the following sustainability-linked instruments:')
    add_bullet(doc, 'MYR 1.5 billion Sustainability-Linked Loan (FY2023, 5-year), pricing tied to Scope 1+2 reduction milestones (12.5 bps step-up if missed)')
    add_bullet(doc, 'MYR 600 million Green Sukuk (FY2024), proceeds dedicated to renewable-energy capex (audited by Sustainalytics)')
    add_bullet(doc, 'USD 200 million Transition Bond (FY2024), proceeds dedicated to low-carbon refinery upgrades and electrification of agri operations')
    doc.add_page_break()


# ── Section 8: HSE ─────────────────────────────────────────────────────────
def section_8(doc):
    add_heading(doc, '8. Health, Safety & Environment (HSE)', level=1)
    add_para(doc, '8.1 Vision & Targets', bold=True, size=12)
    add_para(doc, 'Zero harm to people, environment and assets is the foundational HSE commitment '
                  'of the Group. The HSE Management System is certified to ISO 45001 (Health & '
                  'Safety) and ISO 14001 (Environment) at all manufacturing, agri, oil & gas, and '
                  'healthcare sites.')

    add_para(doc, '8.2 Key HSE Performance Indicators', bold=True, size=12)
    add_table(doc,
        ['Indicator', 'FY2022', 'FY2023', 'FY2024', 'FY2025 Target', 'Long-term Aspiration'],
        [
            ['Total Recordable Injury Rate (per 200K hrs)', '0.42', '0.36', '0.31', '< 0.25', '< 0.10'],
            ['Lost Time Injury Frequency Rate', '0.18', '0.14', '0.11', '< 0.08', '< 0.05'],
            ['Fatalities (employees + contractors)', '2', '1', '0', '0', '0'],
            ['High-potential incidents (HiPo)', '47', '39', '32', '< 25', '< 10'],
            ['Process Safety Tier 1', '4', '3', '2', '< 2', '0'],
            ['Process Safety Tier 2', '14', '11', '8', '< 6', '< 3'],
            ['Environmental incidents (reportable)', '6', '4', '3', '< 2', '0'],
            ['HSE training completion (%)', '93.4%', '95.6%', '97.2%', '> 98%', '100%'],
        ], col_widths=[5.5, 1.6, 1.6, 1.6, 2.4, 3.0])

    add_para(doc, '8.3 Permit-to-Work & Lifesaving Rules', bold=True, size=12)
    add_para(doc, 'The Group operates a tiered permit-to-work system covering hot work, confined '
                  'space, work-at-height, energy isolation, hot oil / hot gas, ground breaking, and '
                  'lifting operations. Eight (8) Lifesaving Rules apply across the Group; breach is '
                  'a summary disciplinary offence and may result in immediate termination.')
    add_bullet(doc, 'Always work with a valid Permit-to-Work where required')
    add_bullet(doc, 'Always follow Energy Isolation procedures')
    add_bullet(doc, 'Always wear the prescribed Personal Protective Equipment')
    add_bullet(doc, 'Never enter a Confined Space without a valid permit and a Standby Person')
    add_bullet(doc, 'Never work under a suspended load')
    add_bullet(doc, 'Never bypass critical safety devices or interlocks')
    add_bullet(doc, 'Always test for atmosphere before entering any potentially hazardous space')
    add_bullet(doc, 'Always travel safely — seatbelts, no-mobile, defensive driving, work-rest cycles')

    add_para(doc, '8.4 Process Safety Management', bold=True, size=12)
    add_para(doc, 'For high-hazard installations (refineries, petrochemical plants, '
                  'fertiliser plants, gas-processing plants), the Group applies a 14-element '
                  'Process Safety Management framework based on CCPS Risk-Based Process Safety, '
                  'including: PHA/HAZOP review every 5 years; Mechanical Integrity programme with '
                  'risk-based inspection; Management of Change with multi-disciplinary review; '
                  'Operating Procedures with annual recertification; Pre-Start-up Safety Review for '
                  'all new and modified installations; Emergency Preparedness with annual full-'
                  'scale drills.')

    add_para(doc, '8.5 Contractor Safety', bold=True, size=12)
    add_para(doc, 'Contractor safety performance is integrated into the Group HSE scorecard. '
                  'Contractor incidents are weighted equivalent to employee incidents in monthly '
                  'reporting. Contractors must achieve EcoVadis Silver or equivalent for HSE-'
                  'critical work; pre-qualification includes Avetta or ISNetworld registration.')
    doc.add_page_break()


# ── Section 9: Legal & Compliance ──────────────────────────────────────────
def section_9(doc):
    add_heading(doc, '9. Legal & Compliance', level=1)
    add_para(doc, '9.1 Group Compliance Framework', bold=True, size=12)
    add_para(doc, 'The Group operates a three-lines-of-defence model:')
    add_numbered(doc, 'First Line: Operations & Business — primary responsibility for risk and compliance ownership in day-to-day activities')
    add_numbered(doc, 'Second Line: Group Compliance, Risk, Legal, Data Privacy, HSE — provides framework, training, monitoring and challenge')
    add_numbered(doc, 'Third Line: Group Internal Audit — provides independent assurance to the Board via the Audit & Risk Committee')

    add_para(doc, '9.2 Compliance Risk Universe', bold=True, size=12)
    add_table(doc,
        ['Compliance Risk Domain', 'Lead 2nd-Line Function', 'Key External Reference'],
        [
            ['Anti-Bribery & Corruption', 'Group Compliance', 'MACC Act 2009 §17A; FCPA; UK Bribery Act'],
            ['Sanctions & Trade', 'Group Compliance', 'OFAC, UN, EU, HMT, BNM ECF'],
            ['Anti-Money Laundering / CFT', 'Group Compliance + FS Compliance', 'AMLA 2001 (MY); UU TPPU (ID); FATF 40 Recs'],
            ['Competition / Antitrust', 'Group Legal', 'Competition Act 2010 (MY); UU 5/1999 (ID)'],
            ['Tax', 'Group Tax', 'IRBM rulings; DGT regulations; OECD BEPS 2.0 Pillar 2'],
            ['Listing & Disclosure', 'Company Secretary', 'Bursa LR; OJK POJK; SGX Listing Manual'],
            ['Securities & Insider Trading', 'Group Legal + CoSec', 'CMSA 2007 (MY); UU PM (ID)'],
            ['Data Privacy', 'Group DPO', 'PDPA MY (2024 amend); UU PDP ID; GDPR'],
            ['Environmental', 'Group HSE', 'Env Quality Act 1974 (MY); UU 32/2009 (ID)'],
            ['Sectoral / Licensing', 'Division Legal', 'BNM, OJK, MOH, MITI, MIDA, BPOM, etc.'],
        ], col_widths=[5.0, 4.5, 7.5])

    add_para(doc, '9.3 Legal & Litigation Management', bold=True, size=12)
    add_para(doc, 'All legal matters with potential exposure greater than MYR 5 million must be '
                  'reported to the Group General Counsel within five (5) business days of '
                  'becoming aware. Matters above MYR 50 million are reported to the Audit & Risk '
                  'Committee. The Group operates an enterprise legal management system (Legalweek '
                  'Matter Management) which tracks all active matters, panel firm engagement, '
                  'budget vs actual fees, and key milestones.')

    add_para(doc, '9.4 Contract Management', bold=True, size=12)
    add_para(doc, 'All contracts are drafted from Group Standard Templates wherever applicable. '
                  'Material variations require Group Legal review. The contract management system '
                  '(Icertis) holds the executed copy of every contract above MYR 250K, with '
                  'auto-extracted obligations, expiry alerts and renewal management.')

    add_para(doc, '9.5 Regulatory Engagement', bold=True, size=12)
    add_para(doc, 'The Group maintains constructive, transparent relationships with all relevant '
                  'regulators. Engagement principles: single point of accountability per regulator, '
                  'proactive disclosure of material matters, prompt response to information '
                  'requests (within stated deadlines), and consistent senior representation in '
                  'formal meetings.')
    doc.add_page_break()


# ── Section 10: ERM ────────────────────────────────────────────────────────
def section_10(doc):
    add_heading(doc, '10. Enterprise Risk Management', level=1)
    add_para(doc, '10.1 Risk Appetite Statement', bold=True, size=12)
    add_para(doc, 'The Group accepts measured risks in pursuit of strategic objectives, balanced '
                  'against capital strength and stakeholder commitments. The Board-approved Risk '
                  'Appetite Statement (FY2025) sets quantitative tolerances for:')
    add_bullet(doc, 'Capital adequacy: Net Debt / EBITDA < 3.0× through-cycle (hard limit 3.5×)')
    add_bullet(doc, 'Liquidity: minimum 3 months of operating cash, plus undrawn committed facility ≥ 12 months of debt maturity')
    add_bullet(doc, 'Earnings volatility: 1-in-20 year EBITDA shock < 25% of base year EBITDA')
    add_bullet(doc, 'Single counterparty exposure: < 15% of group equity to any single non-bank counterparty')
    add_bullet(doc, 'Concentration: no single industry > 30% of group EBITDA at SP&L level (3-year average)')
    add_bullet(doc, 'Reputation: zero tolerance for financial crime, fatal HSE events, or material data breach')

    add_para(doc, '10.2 Top Group Risks (FY2025 Risk Register summary)', bold=True, size=12)
    add_table(doc,
        ['Risk', 'Category', 'Impact', 'Likelihood', 'Trend', 'Owner', 'Top Mitigation'],
        [
            ['Property division margin recovery', 'Strategic', 'High', 'High', '↑', 'Div CEO Properties', 'Asset rationalisation programme'],
            ['Climate transition (carbon pricing)', 'ESG', 'High', 'Medium', '↑', 'Group CSO', 'Scope 1+2 −45% by 2030'],
            ['Cybersecurity incident', 'Operational', 'High', 'Medium', '↑', 'Group CIO', 'Zero-trust roll-out FY25 + IRP'],
            ['Geopolitical / supply chain', 'Strategic', 'High', 'Medium', '↑', 'Group COO', 'Multi-source critical components'],
            ['FX volatility (IDR, VND, THB)', 'Financial', 'Medium', 'High', '→', 'Group Treasurer', 'Layered hedging programme'],
            ['Interest-rate risk', 'Financial', 'Medium', 'High', '↑', 'Group Treasurer', 'Fix:float ratio target 60:40'],
            ['Talent attrition (digital roles)', 'Operational', 'Medium', 'High', '↑', 'Group CHRO', 'Reskilling + retention programme'],
            ['AI governance & misuse', 'Governance', 'Medium', 'Medium', '↑', 'Group CIO + DPO', 'AI Use Policy + Use Case Register'],
            ['Healthcare division performance', 'Strategic', 'High', 'High', '↑', 'Div CEO Healthcare', 'Specialist mix + outpatient growth'],
            ['Regulatory: Indonesia UU PDP', 'Compliance', 'Medium', 'Medium', '→', 'Group DPO', 'UU PDP readiness programme'],
            ['Regulatory: PDPA MY 2024 amend.', 'Compliance', 'Medium', 'High', '↑', 'Group DPO', 'DPO appointment + breach process'],
            ['Counterparty / credit', 'Financial', 'Medium', 'Medium', '→', 'Div CFO Trading', 'Credit insurance + LC discipline'],
        ], col_widths=[4.5, 1.8, 1.5, 1.8, 1.0, 3.0, 5.0])

    add_para(doc, '10.3 Risk Reporting & Escalation', bold=True, size=12)
    add_para(doc, 'Risks are reported on a quarterly basis to the Group Risk Committee, and to '
                  'the Audit & Risk Committee of the Board. Material new or escalating risks are '
                  'escalated within ten (10) business days of identification. Crisis management is '
                  'governed by the Group Crisis Management Plan, with annual full-scale exercises.')

    add_para(doc, '10.4 Insurance Programme', bold=True, size=12)
    add_para(doc, 'The Group maintains a coordinated insurance programme covering property '
                  'damage / business interruption, public & products liability, directors & '
                  'officers, cyber, marine, environmental, employee benefit, and professional '
                  'indemnity. The programme is renewed annually, with brokerage handled via two '
                  'panel brokers (Marsh and Aon) with a competitive split per renewal cycle.')

    # Appendices
    doc.add_page_break()
    add_heading(doc, 'Appendix A — Approval Matrix (extended)', level=1)
    add_para(doc, 'Refer to the full Approval Matrix maintained on the Group Compliance Portal at '
                  'compliance.zava.com/approval-matrix. The matrix is reviewed annually by the '
                  'Group Compliance Council and updated with Board approval.', italic=True)

    add_heading(doc, 'Appendix B — Definitions & Acronyms', level=1)
    glossary = [
        ('AGM', 'Assistant General Manager (JG7)'),
        ('AMLA', 'Anti-Money Laundering, Anti-Terrorism Financing & Proceeds of Unlawful Activities Act 2001 (Malaysia)'),
        ('BNM', 'Bank Negara Malaysia'),
        ('BPJS', 'Badan Penyelenggara Jaminan Sosial (Indonesia social security)'),
        ('BCR', 'Binding Corporate Rules'),
        ('CCPS', 'Center for Chemical Process Safety'),
        ('CFT', 'Combating the Financing of Terrorism'),
        ('CPD/CPE', 'Continuing Professional Development / Education'),
        ('DPO', 'Data Protection Officer'),
        ('EBITDA', 'Earnings before Interest, Tax, Depreciation, Amortisation'),
        ('EPM', 'Enterprise Performance Management'),
        ('ERM', 'Enterprise Risk Management'),
        ('FATF', 'Financial Action Task Force'),
        ('FPIC', 'Free, Prior and Informed Consent'),
        ('FCPA', 'US Foreign Corrupt Practices Act'),
        ('GDPR', 'General Data Protection Regulation (EU)'),
        ('GHS', 'Group Hospital & Surgical insurance'),
        ('GTL', 'Group Term Life insurance'),
        ('HiPo', 'High-potential incident'),
        ('IHB', 'In-House Bank'),
        ('INED', 'Independent Non-Executive Director'),
        ('IRBM', 'Inland Revenue Board Malaysia'),
        ('JG', 'Job Grade'),
        ('JPDP', 'Jabatan Perlindungan Data Peribadi (Malaysia DP Authority)'),
        ('KYV', 'Know Your Vendor'),
        ('LC', 'Letter of Credit'),
        ('MACC', 'Malaysian Anti-Corruption Commission'),
        ('MFRS', 'Malaysian Financial Reporting Standards (IFRS-equivalent)'),
        ('MITI', 'Ministry of International Trade & Industry (MY)'),
        ('NDPE', 'No Deforestation, No Peat, No Exploitation'),
        ('OECD', 'Organisation for Economic Co-operation and Development'),
        ('OJK', 'Otoritas Jasa Keuangan (Indonesia FSA)'),
        ('PDPA', 'Personal Data Protection Act'),
        ('PHA', 'Process Hazard Analysis'),
        ('PSAK', 'Pernyataan Standar Akuntansi Keuangan (Indonesia GAAP)'),
        ('PSU', 'Performance Share Unit'),
        ('RPT', 'Related Party Transaction'),
        ('RSPO', 'Roundtable on Sustainable Palm Oil'),
        ('RSU', 'Restricted Share Unit'),
        ('SBTi', 'Science Based Targets initiative'),
        ('SCC', 'Standard Contractual Clauses'),
        ('SROI', 'Social Return on Investment'),
        ('TIA', 'Transfer Impact Assessment'),
        ('UU PDP', 'Undang-Undang Pelindungan Data Pribadi (Indonesia DP Law 27/2022)'),
    ]
    add_table(doc, ['Term', 'Definition'], glossary, col_widths=[3.5, 13.5])

    add_heading(doc, 'Appendix C — Policy Owner & Review Cycle', level=1)
    review_cycle = [
        ['1. Corporate Governance', 'Company Secretary', 'Annually', 'NRC + Board'],
        ['2. Code of Ethics', 'Group Compliance', 'Bi-annually', 'ARC + Board'],
        ['3. HR Policy', 'Group CHRO', 'Annually', 'NRC + Board'],
        ['4. Financial Governance', 'Group CFO', 'Annually', 'ARC + Board'],
        ['5. Procurement', 'Group CPO', 'Annually', 'Group ExCo'],
        ['6. Data Privacy', 'Group DPO', 'Annually', 'ARC + Board'],
        ['7. ESG', 'Group CSO', 'Annually', 'Sustainability Cmte + Board'],
        ['8. HSE', 'Group HSE Director', 'Annually', 'Sustainability Cmte + Board'],
        ['9. Legal & Compliance', 'Group GC', 'Annually', 'ARC + Board'],
        ['10. ERM', 'Group CRO', 'Annually', 'ARC + Board'],
    ]
    add_table(doc, ['Section', 'Policy Owner', 'Review Frequency', 'Approving Authority'],
              review_cycle, col_widths=[5.5, 4.0, 3.5, 4.0])


def main(out_path):
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(10)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    cover(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else 'files/02_Zava_Group_Policy_Handbook.docx'
    main(out)
