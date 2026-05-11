#!/usr/bin/env python
"""gen_cowork_files.py v2 - 75 per-archetype builders for Zava Conglomerate sample files.

Each filename gets its own builder that produces realistic, name-matched content.
Run from the repo root with: python gen_cowork_files.py
"""
import random
import datetime as dt
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

from PIL import Image, ImageDraw, ImageFont

random.seed(42)

ROOT = Path(__file__).parent
OUT = ROOT / 'files'
OUT.mkdir(exist_ok=True)

HEADER_FILL = PatternFill('solid', fgColor='1F4E78')
ALT_FILL = PatternFill('solid', fgColor='F2F2F2')
HEADER_FONT = Font(bold=True, color='FFFFFF', size=11, name='Aptos')
BODY_FONT = Font(size=10, name='Aptos')
TITLE_FONT = Font(bold=True, color='1F4E78', size=14, name='Aptos')
BORDER = Border(
    left=Side(style='thin', color='D0D7DE'),
    right=Side(style='thin', color='D0D7DE'),
    top=Side(style='thin', color='D0D7DE'),
    bottom=Side(style='thin', color='D0D7DE'),
)

CUSTOMERS = [
    'Bumi Konsortium Sdn Bhd', 'Sapphire Holdings Bhd', 'Tropika Sdn Bhd',
    'Mawar Group', 'Cendana Properties', 'Hijau Industries',
    'Tanjung Resources', 'Bayu Logistics', 'Sinar Capital', 'Permata Trading',
    'Anggerik Manufacturing', 'Pertiwi Energy', 'Gemilang Plantations',
]
PEOPLE = [
    'Ahmad bin Hassan', 'Tan Wei Ming', 'Lim Mei Ling', 'Kumar Subramaniam',
    'Siti Aishah', 'Daichi Maruyama', 'Hadar Caspit', 'Sasha Ouellet',
    'Mod Admin', 'Nurul Huda', 'Vijay Raja', 'Chong Kah Wai', 'Farah Idris',
]
LOCATIONS = [
    'KL-Sentral', 'Penang', 'Johor Bahru', 'Kota Kinabalu', 'Kuching',
    'Cyberjaya', 'Iskandar', 'Ipoh', 'Shah Alam', 'Klang', 'Putrajaya',
]


def _xlsx(path, sheets):
    """sheets = [(sheet_name, title, headers, rows)]"""
    wb = Workbook()
    wb.remove(wb.active)
    for sname, title, headers, rows in sheets:
        ws = wb.create_sheet(sname[:31])
        ws['A1'] = title
        ws['A1'].font = TITLE_FONT
        if headers:
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=3, column=c, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = BORDER
        for r_idx, row in enumerate(rows, 4):
            for c_idx, val in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                cell.font = BODY_FONT
                cell.border = BORDER
                if r_idx % 2 == 0:
                    cell.fill = ALT_FILL
        for c in range(1, max(len(headers), 1) + 1):
            ws.column_dimensions[get_column_letter(c)].width = 19
        if headers:
            ws.freeze_panes = 'A4'
    wb.save(path)


def _docx(path, title, sections):
    doc = Document()
    for style in doc.styles:
        try:
            if style.name == 'Normal':
                style.font.name = 'Aptos'
                style.font.size = Pt(11)
        except Exception:
            pass
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    for kind, content in sections:
        if kind == 'h1':
            doc.add_heading(content, level=1)
        elif kind == 'h2':
            doc.add_heading(content, level=2)
        elif kind == 'p':
            doc.add_paragraph(content)
        elif kind == 'bullets':
            for b in content:
                doc.add_paragraph(b, style='List Bullet')
        elif kind == 'numbered':
            for n in content:
                doc.add_paragraph(n, style='List Number')
        elif kind == 'table':
            headers = content[0]
            rows = content[1:]
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = 'Light Grid Accent 1'
            for c, hdr in enumerate(headers):
                cell = t.rows[0].cells[c]
                cell.text = str(hdr)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for r_idx, row in enumerate(rows, 1):
                for c_idx, val in enumerate(row):
                    t.rows[r_idx].cells[c_idx].text = str(val)
    doc.save(path)


def _pdf(path, title, sections):
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    s = getSampleStyleSheet()
    h0 = ParagraphStyle('h0', parent=s['Title'], textColor=rl_colors.HexColor('#1F4E78'),
                        fontSize=18, spaceAfter=12)
    h1 = ParagraphStyle('h1', parent=s['Heading1'], textColor=rl_colors.HexColor('#1F4E78'),
                        fontSize=14, spaceAfter=8)
    h2 = ParagraphStyle('h2', parent=s['Heading2'], textColor=rl_colors.HexColor('#1F4E78'),
                        fontSize=12, spaceAfter=6)
    bod = ParagraphStyle('bod', parent=s['BodyText'], fontSize=10, leading=14, spaceAfter=6)
    story = [Paragraph(title, h0), Spacer(1, 6)]
    for kind, content in sections:
        if kind == 'h1':
            story.append(Paragraph(content, h1))
        elif kind == 'h2':
            story.append(Paragraph(content, h2))
        elif kind == 'p':
            story.append(Paragraph(content, bod))
        elif kind == 'bullets':
            for b in content:
                story.append(Paragraph(f'&bull; {b}', bod))
        elif kind == 'table':
            tbl = Table(content, hAlign='LEFT')
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#1F4E78')),
                ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.25, rl_colors.HexColor('#D0D7DE')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [rl_colors.white, rl_colors.HexColor('#F2F2F2')]),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 6))
        elif kind == 'pagebreak':
            story.append(PageBreak())
    doc.build(story)


FILE_HANDLERS = {}


def reg(name):
    def deco(fn):
        FILE_HANDLERS[name] = fn
        return fn
    return deco


# ============================================================================
# XLSX archetype builders (41)
# ============================================================================

TODAY = dt.date(2026, 5, 15)


def _d(off):
    return (TODAY + dt.timedelta(days=off)).isoformat()


@reg('AIR_Cancelled_Flights.xlsx')
def _b_air_cancelled(path):
    reasons = ['Tech-AOG', 'Crew-shortage', 'Weather-TS', 'ATC-flow', 'Slot-conflict', 'Bird-strike']
    rows = []
    for i in range(40):
        flt = f'ZA{random.randint(100, 999)}'
        org = random.choice(['KUL', 'PEN', 'BKI', 'KCH', 'JHB'])
        dst = random.choice(['SIN', 'BKK', 'HKG', 'CGK', 'NRT', 'SYD'])
        rows.append([_d(-i // 2), flt, org, dst, random.choice(reasons),
                     random.randint(80, 280), random.choice(['Refund-OK', 'Rebooked', 'In-progress']),
                     f'MYR {random.randint(40, 380) * 1000:,}'])
    _xlsx(path, [('Cancellations', 'Zava Aviation - Cancelled Flights (last 30 days)',
                  ['Date', 'Flight', 'Origin', 'Destination', 'Root cause',
                   'Pax affected', 'Disposition', 'Cost exposure'], rows)])


@reg('AUTO_Affected_VIN_List.xlsx')
def _b_auto_vin(path):
    rows = []
    for i in range(60):
        vin = 'JT' + ''.join(random.choices('ABCDEFGHJKLMNPRSTUVWXYZ0123456789', k=15))
        rows.append([vin, random.choice(['Civic', 'CR-V', 'Accord', 'HR-V']),
                     random.choice([2022, 2023, 2024]),
                     random.choice(['Brake-vacuum', 'Fuel-pump', 'ECU-firmware']),
                     random.choice(LOCATIONS) + ' Dealer',
                     random.choice(['Notified', 'Booked', 'Repaired', 'Pending'])])
    _xlsx(path, [('Recall VINs', 'Recall Campaign R-2026-04 - Affected VIN list',
                  ['VIN', 'Model', 'Year', 'Defect', 'Assigned dealer', 'Status'], rows)])


@reg('AVI_Slot_Request_Master.xlsx')
def _b_avi_slot(path):
    rows = []
    for i in range(50):
        carrier = random.choice(['ZA', 'MH', 'AK', 'OD', 'D7', 'SQ', 'CX'])
        rows.append([f'SLOT-{2026000 + i}', carrier,
                     f'{random.randint(5, 22):02d}:{random.choice([0, 15, 30, 45]):02d}',
                     random.choice(['KUL', 'PEN', 'BKI']),
                     random.choice(['IATA-S26', 'IATA-W26']),
                     random.choice(['Requested', 'Approved', 'Conditional', 'Declined']),
                     random.choice(PEOPLE)])
    _xlsx(path, [('Slots', 'IATA Slot Coordination - Request Master',
                  ['Slot Ref', 'Carrier', 'Time', 'Airport', 'Season', 'Status', 'Coordinator'], rows)])


@reg('BNK_Borrower_Financial_Statements_3yr.xlsx')
def _b_bnk_borrower(path):
    name = 'Bumi Konsortium Sdn Bhd'
    yrs = [2023, 2024, 2025]
    pnl_rows = [
        ['Revenue', '418,500,000', '462,300,000', '441,800,000'],
        ['Cost of sales', '(312,400,000)', '(345,700,000)', '(338,200,000)'],
        ['Gross profit', '106,100,000', '116,600,000', '103,600,000'],
        ['Opex', '(58,200,000)', '(63,100,000)', '(67,800,000)'],
        ['EBITDA', '47,900,000', '53,500,000', '35,800,000'],
        ['D&A', '(12,400,000)', '(13,200,000)', '(14,100,000)'],
        ['Finance cost', '(9,800,000)', '(11,300,000)', '(13,900,000)'],
        ['PBT', '25,700,000', '29,000,000', '7,800,000'],
        ['Tax', '(6,200,000)', '(7,000,000)', '(1,900,000)'],
        ['PAT', '19,500,000', '22,000,000', '5,900,000'],
    ]
    bs_rows = [
        ['Cash and equivalents', '32,400,000', '28,900,000', '18,200,000'],
        ['Trade receivables', '74,200,000', '82,500,000', '93,400,000'],
        ['Inventory', '58,900,000', '67,200,000', '79,300,000'],
        ['PPE', '187,200,000', '198,400,000', '212,500,000'],
        ['Total assets', '352,700,000', '377,000,000', '403,400,000'],
        ['Trade payables', '52,300,000', '58,700,000', '68,200,000'],
        ['Short-term debt', '38,500,000', '47,200,000', '62,400,000'],
        ['Long-term debt', '95,400,000', '102,300,000', '128,700,000'],
        ['Total liabilities', '186,200,000', '208,200,000', '259,300,000'],
        ['Equity', '166,500,000', '168,800,000', '144,100,000'],
    ]
    ratios = [
        ['Gross margin', '25.4%', '25.2%', '23.4%'],
        ['EBITDA margin', '11.4%', '11.6%', '8.1%'],
        ['Net debt / EBITDA', '2.10x', '2.25x', '4.84x'],
        ['Interest cover', '4.89x', '4.73x', '2.58x'],
        ['Current ratio', '1.61x', '1.52x', '1.31x'],
        ['DSO (days)', '65', '65', '77'],
        ['DIO (days)', '69', '71', '86'],
    ]
    _xlsx(path, [
        ('Cover', f'{name} - 3 Year Financial Statements',
         ['Field', 'Value'],
         [['Borrower', name], ['CCRIS ID', 'CCRIS-2026-44218'],
          ['Period covered', 'FY2023 - FY2025 (audited)'],
          ['Auditor', 'Permata & Co. Chartered Accountants'],
          ['Submitted', _d(-3)], ['Reviewing analyst', 'Hadar Caspit, Credit-Mid-Market']]),
        ('P&L', 'Income Statement (MYR)', ['Line item'] + [str(y) for y in yrs], pnl_rows),
        ('Balance Sheet', 'Balance Sheet (MYR)', ['Line item'] + [str(y) for y in yrs], bs_rows),
        ('Ratios', 'Credit Ratios', ['Ratio'] + [str(y) for y in yrs], ratios),
    ])


@reg('BNM_Loan_Book_Quarter_End.xlsx')
def _b_bnm_loan(path):
    segments = ['Retail-Mortgage', 'Retail-Auto', 'Retail-Personal', 'SME-Working-Cap',
                'SME-Term', 'Corporate-Term', 'Corporate-RCF', 'Trade-Finance']
    rows = []
    for s in segments:
        bal = random.randint(800, 9800) * 1_000_000
        st1 = round(bal * random.uniform(0.78, 0.93))
        st2 = round(bal * random.uniform(0.04, 0.12))
        st3 = bal - st1 - st2
        rows.append([s, f'{bal:,}', f'{st1:,}', f'{st2:,}', f'{st3:,}',
                     f'{st3 / bal * 100:.2f}%',
                     f'{random.randint(280, 920) * 1000:,}',
                     random.choice(['Stable', 'Watch', 'Improving', 'Deteriorating'])])
    _xlsx(path, [('Loan Book QE', 'BNM Quarterly Loan Book - Q1 FY2026',
                  ['Segment', 'Gross balance (MYR)', 'Stage 1', 'Stage 2', 'Stage 3',
                   'NPL %', 'ECL provision', 'Outlook'], rows)])


@reg('BPO_Client_SLA_Scorecard.xlsx')
def _b_bpo_sla(path):
    clients = ['Apex Insurance', 'Pacific Bank', 'Asianova Telco', 'Tropika Retail',
               'Cendana Power', 'Sapphire Holdings']
    metrics = ['AHT (sec)', 'FCR %', 'CSAT', 'NPS', 'Schedule adherence %',
               'Abandon rate %', 'Quality score', 'SLA met %']
    rows = []
    for c in clients:
        rows.append([c, random.randint(220, 380), f'{random.randint(72, 91)}%',
                     f'{random.uniform(3.6, 4.7):.2f}',
                     random.randint(28, 62), f'{random.randint(88, 97)}%',
                     f'{random.uniform(2.1, 4.8):.1f}%',
                     f'{random.uniform(88, 97):.1f}%',
                     f'{random.uniform(91, 99):.1f}%',
                     random.choice(['Green', 'Amber', 'Red'])])
    _xlsx(path, [('SLA Scorecard', 'BPO Client SLA Scorecard - Q1 FY2026',
                  ['Client'] + metrics + ['RAG'], rows)])


@reg('BRD_Q4_Financial_Pack.xlsx')
def _b_brd_q4(path):
    divs = ['Banking', 'Insurance', 'Property', 'Plantation', 'Power', 'Retail',
            'Manufacturing', 'Telco', 'Logistics']
    rev = [['Division', 'Q4 Actual (MYR M)', 'Q4 Budget (MYR M)', 'Variance %', 'YoY %']]
    ebd = [['Division', 'Q4 EBITDA (MYR M)', 'Q4 Budget (MYR M)', 'Variance %', 'Margin %']]
    for d in divs:
        a = random.randint(280, 2400)
        b = round(a * random.uniform(0.88, 1.08))
        v = (a - b) / b * 100
        e = round(a * random.uniform(0.08, 0.22))
        eb = round(e * random.uniform(0.88, 1.12))
        ev = (e - eb) / eb * 100
        rev.append([d, a, b, f'{v:+.1f}%', f'{random.uniform(-6, 14):+.1f}%'])
        ebd.append([d, e, eb, f'{ev:+.1f}%', f'{e / a * 100:.1f}%'])
    _xlsx(path, [
        ('Revenue', 'Group Q4 FY2025 Revenue Pack', rev[0], rev[1:]),
        ('EBITDA', 'Group Q4 FY2025 EBITDA Pack', ebd[0], ebd[1:]),
    ])


@reg('Capex_Model_Base.xlsx')
def _b_capex_model(path):
    projects = ['Bandar Zava Phase 2', 'Iskandar Logistics Hub', 'Penang Data Centre',
                'KK Hospital Wing C', 'Kuching Solar Farm 90MW', 'Klang Port Terminal 3',
                'Cyberjaya Tower D', 'Shah Alam Warehouse', 'JB Petrochem Upgrade']
    rows = []
    for p in projects:
        cost = random.randint(40, 980) * 1_000_000
        npv = round(cost * random.uniform(-0.08, 0.42))
        irr = random.uniform(6.4, 22.8)
        pb = random.uniform(3.2, 11.5)
        rows.append([p, f'MYR {cost / 1e6:,.0f}M', f'MYR {npv / 1e6:+,.0f}M',
                     f'{irr:.1f}%', f'{pb:.1f} yrs',
                     random.choice(['Approved', 'In FID', 'Awaiting board', 'Deferred']),
                     random.choice(PEOPLE)])
    _xlsx(path, [('Capex Pipeline', 'Group Capex Model - Base Case FY2026-2028',
                  ['Project', 'Total cost', 'NPV', 'IRR', 'Payback',
                   'Gate status', 'Sponsor'], rows)])


@reg('COAL_Mine_Production_Plan.xlsx')
def _b_coal_plan(path):
    months = ['Jun-26', 'Jul-26', 'Aug-26', 'Sep-26', 'Oct-26', 'Nov-26', 'Dec-26']
    pits = ['North Pit', 'South Pit', 'East Cut', 'West Cut']
    rows = []
    for p in pits:
        for m in months:
            tons = random.randint(180_000, 420_000)
            sr = random.uniform(3.2, 5.8)
            cal = random.randint(5400, 6400)
            rows.append([m, p, f'{tons:,}', f'{sr:.2f}', cal,
                         f'{random.randint(7800, 9800):,}', f'USD {random.randint(72, 118)}/t'])
    _xlsx(path, [('Production Plan', 'Coal Mine - 7-Month Production Plan',
                  ['Month', 'Pit', 'Tonnes ROM', 'Strip ratio', 'CV (kcal/kg)',
                   'OB removal (bcm)', 'Realised price'], rows)])


@reg('EDU_Cohort_Performance.xlsx')
def _b_edu_cohort(path):
    cohorts = ['BSc CS 2023', 'BSc CS 2024', 'BSc CS 2025', 'BBA 2023', 'BBA 2024',
               'BBA 2025', 'BEng EE 2023', 'BEng EE 2024', 'BEng ME 2024']
    rows = []
    for c in cohorts:
        rows.append([c, random.randint(82, 240),
                     f'{random.uniform(2.8, 3.7):.2f}',
                     f'{random.randint(72, 95)}%',
                     f'{random.randint(58, 92)}%',
                     f'{random.randint(60, 88)}%',
                     random.randint(4, 18)])
    _xlsx(path, [('Cohort KPIs', 'Cohort Performance Dashboard - FY2025',
                  ['Cohort', 'Headcount', 'Avg CGPA', 'Retention %',
                   'Placement %', 'On-time graduation %', 'Disciplinary cases'], rows)])


@reg('Emissions_Inventory.xlsx')
def _b_emissions(path):
    sites = ['KL HQ', 'Penang Plant', 'Iskandar Refinery', 'Kuching Mill',
             'JB Logistics', 'KK Hotel', 'Shah Alam Warehouse']
    sc1 = [['Site', 'Stationary (tCO2e)', 'Mobile (tCO2e)', 'Process (tCO2e)', 'Fugitive (tCO2e)', 'Total Scope 1']]
    sc2 = [['Site', 'Grid electricity (MWh)', 'Emission factor', 'Scope 2 location-based', 'Scope 2 market-based']]
    sc3 = [['Category', 'tCO2e', 'Methodology', 'Data quality']]
    for s in sites:
        a, b, c, d = (random.randint(120, 4800) for _ in range(4))
        sc1.append([s, a, b, c, d, a + b + c + d])
        mwh = random.randint(800, 18_000)
        sc2.append([s, mwh, '0.585 tCO2/MWh', round(mwh * 0.585), round(mwh * 0.412)])
    for cat in ['1. Purchased goods', '3. Fuel & energy', '4. Upstream transport',
                '6. Business travel', '7. Employee commute', '11. Use of sold products',
                '12. End-of-life']:
        sc3.append([cat, random.randint(1200, 38_000),
                    random.choice(['Spend-based', 'Activity-based', 'Hybrid']),
                    random.choice(['High', 'Medium', 'Estimated'])])
    _xlsx(path, [
        ('Scope 1', 'Scope 1 - Direct GHG Emissions (FY2025)', sc1[0], sc1[1:]),
        ('Scope 2', 'Scope 2 - Purchased Electricity (FY2025)', sc2[0], sc2[1:]),
        ('Scope 3', 'Scope 3 - Value Chain (FY2025)', sc3[0], sc3[1:]),
    ])


@reg('ESG_Data_Pack.xlsx')
def _b_esg_pack(path):
    env = [['Metric', 'FY2023', 'FY2024', 'FY2025', 'FY2026 Target'],
           ['Scope 1+2 GHG (tCO2e)', '142,800', '136,400', '128,900', '118,000'],
           ['Scope 3 GHG (tCO2e)', '438,200', '462,500', '484,300', '440,000'],
           ['Energy intensity (MWh/MYR M revenue)', '18.4', '17.2', '16.1', '14.5'],
           ['Water withdrawal (m3)', '1,820,000', '1,742,000', '1,684,000', '1,550,000'],
           ['Waste diverted from landfill %', '64.2%', '68.8%', '72.3%', '80%'],
           ['Renewable energy share %', '8.4%', '11.6%', '15.2%', '25%']]
    soc = [['Metric', 'FY2023', 'FY2024', 'FY2025', 'FY2026 Target'],
           ['LTIFR (per million hrs)', '2.84', '2.21', '1.78', '1.20'],
           ['Fatalities', '1', '0', '0', '0'],
           ['Training hours per employee', '24.6', '27.8', '31.2', '36.0'],
           ['Female board %', '22%', '28%', '33%', '40%'],
           ['Female senior mgmt %', '24%', '27%', '31%', '38%'],
           ['Employee engagement score', '72', '75', '78', '82']]
    gov = [['Metric', 'FY2023', 'FY2024', 'FY2025', 'Status'],
           ['Code of Conduct attestation %', '94%', '97%', '99%', 'On-track'],
           ['Anti-bribery training completion %', '88%', '95%', '99%', 'On-track'],
           ['Whistleblowing cases substantiated', '4', '6', '3', 'Under-investigation'],
           ['ISO 27001 sites covered', '2', '4', '6', 'Expanding'],
           ['Bursa CG Code compliance', '85%', '92%', '97%', 'On-track']]
    _xlsx(path, [
        ('Environment', 'ESG Pack - Environmental', env[0], env[1:]),
        ('Social', 'ESG Pack - Social', soc[0], soc[1:]),
        ('Governance', 'ESG Pack - Governance', gov[0], gov[1:]),
    ])


@reg('FIN_Trial_Balance.xlsx')
def _b_fin_tb(path):
    accts = [
        ('1000', 'Cash - KL Operating', 'Asset', 28_400_000, 0),
        ('1010', 'Cash - SGD Hedge', 'Asset', 12_800_000, 0),
        ('1100', 'Trade receivables', 'Asset', 184_200_000, 0),
        ('1200', 'Inventory - finished', 'Asset', 67_800_000, 0),
        ('1500', 'PPE - Buildings', 'Asset', 412_300_000, 0),
        ('1510', 'PPE - Plant', 'Asset', 278_900_000, 0),
        ('2000', 'Trade payables', 'Liability', 0, 142_500_000),
        ('2100', 'Short-term debt', 'Liability', 0, 88_400_000),
        ('2200', 'Long-term debt', 'Liability', 0, 412_700_000),
        ('3000', 'Share capital', 'Equity', 0, 250_000_000),
        ('3100', 'Retained earnings', 'Equity', 0, 224_300_000),
        ('4000', 'Revenue - Banking', 'Revenue', 0, 1_842_000_000),
        ('4010', 'Revenue - Insurance', 'Revenue', 0, 928_400_000),
        ('4020', 'Revenue - Property', 'Revenue', 0, 612_800_000),
        ('5000', 'COGS', 'Expense', 1_424_300_000, 0),
        ('6000', 'Opex - Personnel', 'Expense', 312_400_000, 0),
        ('6100', 'Opex - Selling', 'Expense', 142_800_000, 0),
        ('6200', 'Opex - G&A', 'Expense', 88_300_000, 0),
        ('7000', 'Finance cost', 'Expense', 84_700_000, 0),
        ('8000', 'Tax expense', 'Expense', 42_300_000, 0),
    ]
    rows = [[a, n, c, f'{d:,}' if d else '', f'{cr:,}' if cr else ''] for a, n, c, d, cr in accts]
    _xlsx(path, [('Trial Balance', 'Zava Group - Trial Balance as at 30-Apr-2026 (MYR)',
                  ['Account', 'Description', 'Category', 'Debit', 'Credit'], rows)])


@reg('GENINS_Open_Claims_T0_to_T7.xlsx')
def _b_genins_claims(path):
    lobs = ['Motor', 'Fire', 'Marine', 'Liability', 'PA', 'Medical']
    rows = []
    for i in range(60):
        rows.append([f'CL-2026-{40000 + i:05d}',
                     random.choice(lobs),
                     _d(-random.randint(0, 7)),
                     random.choice(['Apex', 'Pacific', 'Asianova', 'Cendana', 'Mawar']) + ' Bhd',
                     f'MYR {random.randint(8, 480) * 1000:,}',
                     random.choice(['T+0', 'T+1', 'T+2', 'T+3', 'T+5', 'T+7']),
                     random.choice(['FNOL', 'Investigating', 'Reserved', 'Negotiating', 'Quantum-agreed']),
                     random.choice(PEOPLE)])
    _xlsx(path, [('Open Claims', 'General Insurance - Open Claims (T+0 to T+7)',
                  ['Claim Ref', 'LOB', 'Date', 'Insured', 'Reserve',
                   'Aging', 'Status', 'Adjuster'], rows)])


@reg('GLC_Portfolio_Master.xlsx')
def _b_glc_portfolio(path):
    rows = []
    investees = [
        ('Cendana Power Bhd', 'Power', 38.4, 'Listed', 6.2),
        ('Apex Banking Group', 'Banking', 18.7, 'Listed', 12.4),
        ('Pacific Communications', 'Telco', 24.5, 'Listed', 4.8),
        ('Asianova Plantation Bhd', 'Plantation', 56.2, 'Listed', 3.1),
        ('Sapphire Property Trust', 'REIT', 42.0, 'Listed', 2.4),
        ('Tropika Logistics Sdn Bhd', 'Logistics', 100.0, 'Private', 1.8),
        ('Bumi Construction Bhd', 'Construction', 51.3, 'Listed', 2.2),
        ('Mawar Industries', 'Manufacturing', 67.5, 'Listed', 1.6),
        ('Cendana Aviation', 'Aviation', 49.0, 'Private', 3.4),
        ('Permata Retail Holdings', 'Retail', 33.0, 'Listed', 2.8),
        ('Hijau Renewables', 'Renewables', 100.0, 'Private', 0.9),
        ('Pertiwi Healthcare', 'Healthcare', 28.0, 'Listed', 4.6),
    ]
    for nm, sect, stake, listed, val in investees:
        rows.append([nm, sect, f'{stake:.1f}%', listed, f'MYR {val:.1f}B',
                     random.choice(['Hold', 'Strengthen', 'Trim', 'Strategic']),
                     random.choice(PEOPLE)])
    _xlsx(path, [('Portfolio Master', 'GLC Portfolio Master - 12 Investees',
                  ['Investee', 'Sector', 'Stake %', 'Listing', 'Carrying value',
                   'Stewardship stance', 'Portfolio mgr'], rows)])


@reg('GOV_PQ_Tracker.xlsx')
def _b_gov_pq(path):
    rows = []
    topics = ['Highway tolls', 'Flood mitigation', 'Education subsidies', 'GLC dividends',
             'Solar capacity targets', 'Subsidy rationalisation', 'Cost of living', 'Healthcare access']
    for i in range(30):
        rows.append([f'PQ-26-{i + 1:03d}', random.choice(topics),
                     random.choice(['Lisan', 'Bertulis']),
                     f'YB {random.choice(["Ahmad", "Lim", "Kumar", "Siti", "Aisyah"])} bin {random.choice(["Hassan", "Wei", "Raja"])}',
                     _d(-random.randint(0, 14)),
                     random.choice(PEOPLE) + ' (Ministry)',
                     random.choice(['Drafted', 'Reviewed', 'Cleared', 'Tabled']),
                     _d(random.randint(1, 21))])
    _xlsx(path, [('PQ Tracker', 'Parliamentary Question Tracker - Dewan Rakyat',
                  ['PQ Ref', 'Topic', 'Type', 'MP', 'Date received',
                   'Drafter', 'Status', 'Due date'], rows)])


@reg('GRP_Capex_Pipeline_Master.xlsx')
def _b_grp_capex(path):
    rows = []
    divs = ['Banking', 'Insurance', 'Property', 'Plantation', 'Power', 'Retail',
            'Manufacturing', 'Telco', 'Logistics']
    for d in divs:
        for j in range(random.randint(2, 4)):
            rows.append([f'{d}-CX-{j + 1}', d,
                         f'{d} Project {chr(65 + j)}',
                         f'MYR {random.randint(20, 480) * 1_000_000:,}',
                         random.choice(['FY26', 'FY27', 'FY28']),
                         random.choice(['Approved', 'In FID', 'Awaiting board', 'Deferred', 'On-hold']),
                         f'{random.uniform(7.4, 21.8):.1f}%',
                         random.choice(PEOPLE)])
    _xlsx(path, [('Group Capex', 'Group Capex Pipeline - FY2026-2028 Master',
                  ['ID', 'Division', 'Project', 'Total cost', 'FID year',
                   'Gate', 'IRR', 'Sponsor'], rows)])


@reg('HOTEL_Booking_Pace.xlsx')
def _b_hotel_pace(path):
    properties = ['Zava KL Sentral', 'Zava Penang Beach', 'Zava KK Resort',
                  'Zava Cyberjaya Conference', 'Zava JB Marina']
    rows = []
    for p in properties:
        for w in range(8):
            on_books = random.randint(280, 1240)
            ly_pickup = random.randint(220, 1180)
            adr = random.randint(380, 920)
            rows.append([p, f'W+{w}', f'{on_books:,}', f'{ly_pickup:,}',
                         f'{(on_books - ly_pickup) / max(ly_pickup, 1) * 100:+.1f}%',
                         f'MYR {adr}', f'MYR {adr * on_books // 1000:,}K',
                         f'{random.randint(48, 92)}%'])
    _xlsx(path, [('Booking Pace', 'Hotel Group Booking Pace - 8-week forward',
                  ['Property', 'Week', 'On the books', 'LY pickup', 'YoY pickup %',
                   'ADR', 'Revenue', 'Occupancy %'], rows)])


@reg('HR_Performance_Ratings.xlsx')
def _b_hr_perf(path):
    bands = ['Group CEO Office', 'Banking', 'Insurance', 'Property', 'Plantation',
             'Power', 'Retail', 'Manufacturing', 'Telco', 'Logistics', 'Shared Services']
    rows = []
    for b in bands:
        headcount = random.randint(120, 4200)
        rows.append([b, headcount,
                     f'{random.randint(8, 16)}%', f'{random.randint(20, 30)}%',
                     f'{random.randint(45, 58)}%', f'{random.randint(4, 12)}%',
                     f'{random.randint(1, 5)}%',
                     f'{random.uniform(3.3, 4.1):.2f}',
                     random.choice(PEOPLE)])
    _xlsx(path, [('Perf Distribution', 'FY2025 Performance Rating Distribution',
                  ['Business unit', 'Headcount', 'Outstanding %', 'Exceeds %',
                   'Meets %', 'Below %', 'Unsatisfactory %', 'Avg rating', 'BU HRBP'], rows)])


@reg('IB_Target_3yr_Financials.xlsx')
def _b_ib_target(path):
    rows = [
        ['Revenue', 412_500_000, 458_300_000, 484_700_000, 9.4],
        ['Cost of sales', 308_400_000, 339_200_000, 354_600_000, 7.3],
        ['Gross profit', 104_100_000, 119_100_000, 130_100_000, 11.8],
        ['Opex', 48_300_000, 53_700_000, 58_400_000, 9.9],
        ['EBITDA', 55_800_000, 65_400_000, 71_700_000, 13.3],
        ['D&A', 11_200_000, 12_400_000, 13_300_000, 8.9],
        ['EBIT', 44_600_000, 53_000_000, 58_400_000, 14.4],
        ['Net interest', 8_400_000, 8_800_000, 7_900_000, -3.0],
        ['PBT', 36_200_000, 44_200_000, 50_500_000, 18.1],
        ['Tax', 9_100_000, 11_000_000, 12_600_000, 17.6],
        ['PAT', 27_100_000, 33_200_000, 37_900_000, 18.3],
    ]
    rows_str = [[r[0], f'{r[1]:,}', f'{r[2]:,}', f'{r[3]:,}', f'{r[4]:+.1f}%'] for r in rows]
    _xlsx(path, [('Target Financials', 'Acquisition Target - 3yr Financials (MYR)',
                  ['Line', 'FY2024A', 'FY2025A', 'FY2026B', 'CAGR %'], rows_str)])


@reg('LIFE_Policy_Master_5yr.xlsx')
def _b_life_master(path):
    products = ['Endowment 20yr', 'Endowment 30yr', 'Whole-life', 'Investment-linked',
                'Term-10', 'Term-20', 'Annuity', 'Critical-Illness rider']
    rows = []
    for pr in products:
        iforce = random.randint(38_000, 218_000)
        nb = random.randint(2_400, 18_400)
        lap = random.uniform(2.4, 11.8)
        ape = random.randint(48, 380)
        rows.append([pr, iforce, nb, f'{lap:.2f}%',
                     f'MYR {ape}M', f'{random.uniform(38, 92):.1f}%',
                     f'{random.uniform(2.8, 5.4):.2f}%',
                     f'MYR {random.randint(180, 980)}M'])
    _xlsx(path, [('Policy Master', 'Life Insurance - 5yr Policy Master Snapshot',
                  ['Product', 'In-force count', 'NB-26', 'Lapse %',
                   'APE FY25', '13M persistency', 'Loss ratio', 'Reserves'], rows)])


@reg('LOG_Network_Capacity.xlsx')
def _b_log_network(path):
    lanes = [
        ('KUL -> SIN', 1480, 1320, 'Daily'),
        ('KUL -> PEN', 980, 720, 'Daily'),
        ('KUL -> JHB', 1240, 980, 'Daily'),
        ('PEN -> BKK', 620, 540, '5/wk'),
        ('JHB -> CGK', 720, 680, 'Daily'),
        ('SIN -> HKG', 1380, 1180, '6/wk'),
        ('KUL -> NRT', 480, 420, '4/wk'),
        ('KCH -> KUL', 420, 320, 'Daily'),
        ('BKI -> KUL', 580, 480, 'Daily'),
    ]
    rows = [[ln, f'{cap:,}', f'{ut:,}', f'{ut / cap * 100:.1f}%', fq,
             f'{random.uniform(0.92, 0.99):.3f}',
             f'{random.uniform(0.84, 0.97):.3f}',
             random.choice(PEOPLE)]
            for ln, cap, ut, fq in lanes]
    _xlsx(path, [('Lane Capacity', 'Logistics Network Capacity Utilisation',
                  ['Lane', 'Capacity (TEU/wk)', 'Utilised', 'Util %',
                   'Frequency', 'On-time arrival', 'On-time departure', 'Lane mgr'], rows)])


@reg('MTG_Past_Due_30_90.xlsx')
def _b_mtg_past_due(path):
    rows = []
    for i in range(45):
        bal = random.randint(180_000, 1_480_000)
        rows.append([f'MTG-2026-{50000 + i:05d}',
                     random.choice(PEOPLE),
                     random.choice(['Klang Valley', 'Penang', 'JB', 'KK', 'Kuching']),
                     f'MYR {bal:,}',
                     f'MYR {random.randint(1800, 12400):,}',
                     random.choice([30, 45, 60, 75, 90]),
                     random.choice(['Stage 2', 'Stage 2', 'Stage 3']),
                     random.choice(['Pre-legal letter', 'Restructure offered', 'Foreclosure-Q', 'Repossessed', 'Cured'])])
    _xlsx(path, [('Past Due 30-90', 'Mortgage - Past Due 30-90 DPD',
                  ['Loan ID', 'Borrower', 'Region', 'Balance',
                   'Monthly instalment', 'DPD', 'Stage', 'Disposition'], rows)])


@reg('OEE_30D_Log_Line5.xlsx')
def _b_oee_line5(path):
    rows = []
    for d in range(30):
        for shift in ['A', 'B', 'C']:
            avail = random.uniform(0.86, 0.98)
            perf = random.uniform(0.82, 0.97)
            qual = random.uniform(0.92, 0.998)
            oee = avail * perf * qual
            rows.append([_d(-30 + d), shift, f'{avail:.3f}', f'{perf:.3f}',
                         f'{qual:.3f}', f'{oee * 100:.1f}%',
                         random.randint(0, 4),
                         random.choice(['', 'Changeover overrun', 'Quality reject', 'Material starvation', 'Tooling change'])])
    _xlsx(path, [('Line 5 - 30 day log', 'OEE - Line 5 - 30 Day Shift Log',
                  ['Date', 'Shift', 'Availability', 'Performance', 'Quality',
                   'OEE', 'Stoppages', 'Top loss reason'], rows)])


@reg('OGU_Field_Production_Forecast.xlsx')
def _b_ogu_forecast(path):
    fields = ['Bunga Kekwa', 'Tapis-K', 'Resak North', 'Permai East',
              'Anggerik Deepwater', 'Tembikai-2']
    rows = []
    for f in fields:
        for q in ['Q3-26', 'Q4-26', 'Q1-27', 'Q2-27']:
            rows.append([f, q, f'{random.randint(8200, 48_400):,} bopd',
                         f'{random.randint(12, 280):,} mmscfd',
                         f'{random.uniform(0.4, 18.4):.1f}%',
                         f'USD {random.randint(72, 92)}/bbl',
                         f'{random.randint(1800, 9800):,}'])
    _xlsx(path, [('Field Production', 'OG Upstream - Field Production Forecast',
                  ['Field', 'Quarter', 'Oil', 'Gas', 'Decline',
                   'Realised price', 'Net daily revenue (USD K)'], rows)])


@reg('PAY_T_minus_1_Transactions.xlsx')
def _b_pay_t_minus_1(path):
    rows = []
    for i in range(80):
        amt = random.randint(120, 4_800_000)
        rows.append([f'TXN-{20260514}-{i + 1:06d}',
                     _d(-1) + f' {random.randint(0, 23):02d}:{random.randint(0, 59):02d}',
                     random.choice(['DuitNow', 'IBG', 'RENTAS', 'SWIFT', 'Card', 'eWallet']),
                     random.choice(['MYR', 'USD', 'SGD', 'IDR']),
                     f'{amt:,}',
                     random.choice(CUSTOMERS),
                     random.choice(['Settled', 'Pending', 'Reversed', 'Investigating']),
                     random.choice(['', 'AML-flag', 'Fraud-rule', 'Sanctions-hit', 'Limit-breach'])])
    _xlsx(path, [('T-1 Transactions', 'Payments - T-1 Transaction Log',
                  ['Txn ID', 'Timestamp', 'Rail', 'Currency', 'Amount',
                   'Counterparty', 'Status', 'Flag'], rows)])


@reg('PD_Project_Feasibility.xlsx')
def _b_pd_feasibility(path):
    rows = [
        ['Land cost (MYR M)', 92.4],
        ['Construction cost (MYR M)', 218.6],
        ['Soft cost (MYR M)', 42.8],
        ['Total dev cost (MYR M)', 353.8],
        ['Saleable area (sq ft)', 482_000],
        ['Avg selling price (MYR/sq ft)', 1080],
        ['GDV (MYR M)', 520.6],
        ['Gross development profit (MYR M)', 166.8],
        ['GDP margin %', '32.0%'],
        ['IRR (5-yr)', '21.4%'],
        ['NPV @ 10% (MYR M)', 78.2],
        ['Take-up 6 months (%)', '38%'],
        ['Take-up 12 months (%)', '64%'],
        ['Take-up 24 months (%)', '92%'],
        ['Sensitivity: -10% ASP NPV', 22.4],
        ['Sensitivity: +10% cost NPV', 36.8],
    ]
    _xlsx(path, [('Feasibility', 'Bandar Zava Phase 2A - Project Feasibility',
                  ['Line item', 'Value'], rows)])


@reg('PROC_Bid_Master.xlsx')
def _b_proc_bid(path):
    rows = []
    rfps = ['Cyberjaya Tower D MEP', 'Iskandar Hub Civil Works',
            'KK Hospital Wing C Fit-out', 'Solar Farm EPC 90MW',
            'Penang DC Cooling Upgrade', 'Group ERP Implementation']
    for r in rfps:
        for v in ['Apex Engineering', 'Pacific Builders', 'Asianova Infra', 'Mawar Civil']:
            rows.append([r, v, f'MYR {random.randint(28, 480) * 1_000_000:,}',
                         random.choice(['Compliant', 'Conditional', 'Non-compliant']),
                         random.uniform(3.4, 4.8),
                         random.choice(['Shortlisted', 'Eliminated', 'Held']),
                         random.choice(PEOPLE)])
    _xlsx(path, [('Bid Master', 'Procurement Bid Master - Open RFPs',
                  ['RFP', 'Vendor', 'Bid value', 'Technical compliance',
                   'Tech score (5)', 'Status', 'Lead'], rows)])


@reg('REF_Crude_Slate_T0.xlsx')
def _b_ref_crude(path):
    rows = [
        ['Tapis (MY)', 38, 45.3, 0.04, 28.4, 'Domestic'],
        ['Bintulu Light', 22, 44.1, 0.06, 26.8, 'Domestic'],
        ['Murban (UAE)', 18, 40.4, 0.78, 22.6, 'Term'],
        ['Arab Light', 12, 33.2, 1.84, 18.4, 'Term'],
        ['Brent (NWE)', 6, 38.7, 0.40, 24.2, 'Spot'],
        ['Forties', 4, 39.2, 0.42, 23.1, 'Spot'],
    ]
    rows_str = [[r[0], f'{r[1]:.0f}%', f'{r[2]}', f'{r[3]}%',
                 f'{r[4]:.1f}', r[5]] for r in rows]
    _xlsx(path, [('Crude Slate T0', 'Refinery Crude Slate - Today',
                  ['Crude grade', 'Slate %', 'API', 'Sulphur %', 'Gasoline yield %', 'Source'], rows_str)])


@reg('REIT_Tenant_Master.xlsx')
def _b_reit_tenant(path):
    rows = []
    tenants = [
        ('Apex Bank Bhd', 'Banking', 'KL Tower A', 18_200, 11.50),
        ('Pacific Investments', 'Financial Services', 'KL Tower A', 12_400, 11.20),
        ('Asianova Tech Sdn Bhd', 'Tech', 'Cyberjaya Tech Park', 28_400, 8.40),
        ('Cendana Law Chambers', 'Legal', 'KL Tower A', 8_200, 11.80),
        ('Mawar Retail (Padini Concept)', 'Retail', 'Bandar Mall', 4_200, 16.50),
        ('Tropika F&B (3 outlets)', 'F&B', 'Bandar Mall', 6_800, 22.40),
        ('Hijau Pharmacy', 'Retail', 'Bandar Mall', 1_800, 18.20),
        ('Sapphire Logistics', 'Logistics', 'Shah Alam Warehouse', 84_000, 2.80),
        ('Bumi MedCentre', 'Healthcare', 'Penang Tower B', 12_800, 9.40),
        ('Permata Coworking', 'Office', 'Penang Tower B', 18_400, 8.20),
    ]
    for nm, sect, prop, sqft, rate in tenants:
        rows.append([nm, sect, prop, f'{sqft:,}',
                     f'MYR {rate:.2f}',
                     f'MYR {sqft * rate / 1000:.1f}K',
                     random.choice(['2026', '2027', '2028', '2029']),
                     f'{random.randint(78, 99)}%'])
    _xlsx(path, [('Tenant Master', 'REIT Tenant Master - Rent Roll',
                  ['Tenant', 'Sector', 'Property', 'Sq ft', 'Rate/sq ft',
                   'Monthly rent', 'Lease expiry', 'Collection rate'], rows)])


@reg('REM_Corridor_Volume_90D.xlsx')
def _b_rem_corridor(path):
    corridors = ['MY-PH', 'MY-ID', 'MY-BD', 'MY-NP', 'MY-IN', 'MY-MM',
                 'MY-VN', 'SG-PH', 'SG-ID', 'SG-IN']
    rows = []
    for c in corridors:
        vol = random.randint(180_000, 2_840_000)
        avg = random.randint(280, 1240)
        rows.append([c, f'{vol:,}', f'USD {avg}',
                     f'USD {vol * avg / 1000:,.0f}',
                     f'{random.uniform(2.4, 4.8):.2f}%',
                     f'{random.uniform(0.32, 0.98):.2f}%',
                     random.choice(['Cash', 'Bank a/c', 'Wallet'])])
    _xlsx(path, [('Corridor 90D', 'Remittance Corridor Volumes - last 90 days',
                  ['Corridor', 'Txn count', 'Avg ticket', 'GMV (USD K)',
                   'FX margin', 'Fee/txn', 'Dominant payout'], rows)])


@reg('REN_Resource_Assessment.xlsx')
def _b_ren_resource(path):
    sites = [
        ('Solar - Sik Kedah 90MW', 'Solar', 1820, 18.4, 'Operational'),
        ('Solar - Pekan 60MW', 'Solar', 1740, 17.2, 'Construction'),
        ('Wind - Mersing 40MW', 'Wind', 7.2, 32.4, 'FID'),
        ('Hydro - Sg Perak 12MW (mini)', 'Hydro', 4.2, 48.6, 'Operational'),
        ('Biomass - Lahad Datu 10MW', 'Biomass', 8400, 78.4, 'Operational'),
        ('Solar - Sandakan 50MW', 'Solar', 1880, 19.2, 'Feasibility'),
    ]
    rows = []
    for nm, kind, base, cf, stat in sites:
        rows.append([nm, kind,
                     f'{base} kWh/m2/yr' if kind == 'Solar' else f'{base} m/s' if kind == 'Wind' else f'{base} kg/hr',
                     f'{cf:.1f}%',
                     f'MYR {random.randint(40, 380)}M',
                     f'{random.uniform(8.4, 14.8):.1f}%',
                     stat,
                     random.choice(PEOPLE)])
    _xlsx(path, [('Resource Assessment', 'Renewables - Resource Assessment Pipeline',
                  ['Project', 'Type', 'Resource', 'Capacity factor',
                   'Capex', 'Project IRR', 'Stage', 'Sponsor'], rows)])


@reg('RTL_Store_PnL.xlsx')
def _b_rtl_store(path):
    stores = []
    for loc in LOCATIONS:
        stores.append([f'Zava Mart {loc}',
                       f'MYR {random.randint(180, 980) * 1000:,}',
                       f'MYR {random.randint(48, 320) * 1000:,}',
                       f'{random.uniform(22, 38):.1f}%',
                       f'MYR {random.randint(28, 180) * 1000:,}',
                       f'{random.uniform(8, 18):.1f}%',
                       f'{random.randint(8200, 38_400):,}',
                       f'MYR {random.randint(38, 82)}'])
    _xlsx(path, [('Store P&L', 'Retail Store P&L - last month',
                  ['Store', 'Revenue', 'Gross profit', 'GP margin',
                   'EBITDA', 'EBITDA margin', 'Footfall', 'Avg ticket'], stores)])


@reg('SC_Backlog_By_Customer_Node.xlsx')
def _b_sc_backlog(path):
    rows = []
    customers = ['Apex Tech Bhd', 'Pacific Auto Sdn Bhd', 'Asianova Electronics',
                 'Cendana Industrial', 'Mawar Engineering', 'Sapphire Aerospace']
    nodes = ['KL DC', 'Penang DC', 'JB DC', 'Iskandar Hub']
    for c in customers:
        for n in nodes:
            rows.append([c, n,
                         random.randint(120, 4800),
                         f'MYR {random.randint(38, 480) * 1000:,}',
                         random.randint(3, 28),
                         random.choice(['On-plan', 'At-risk', 'Overdue']),
                         random.choice(PEOPLE)])
    _xlsx(path, [('Backlog', 'Supply Chain Backlog - by Customer + Node',
                  ['Customer', 'Node', 'Open lines', 'Value',
                   'Avg age (days)', 'Status', 'Planner'], rows)])


@reg('SHIP_Fleet_Schedule.xlsx')
def _b_ship_fleet(path):
    rows = []
    vessels = ['MV Zava Pacific', 'MV Zava Atlantic', 'MV Zava Indian Ocean',
               'MV Zava Java', 'MV Zava Borneo', 'MV Zava Malacca',
               'MV Zava Selat', 'MV Zava Andaman']
    for v in vessels:
        rows.append([v,
                     random.choice(['Tanker-Suezmax', 'Bulk-Capesize', 'Container-Feeder', 'LNG-Carrier']),
                     random.choice(['SG', 'PEN', 'TPP', 'JKT', 'YGN', 'HKG']),
                     random.choice(['ROT', 'FUJ', 'NSA', 'YOK', 'SHA', 'LAX']),
                     _d(random.randint(1, 14)),
                     _d(random.randint(15, 35)),
                     f'{random.randint(38, 92)}%',
                     f'USD {random.randint(12_000, 78_000):,}/day'])
    _xlsx(path, [('Fleet Schedule', 'Fleet Schedule - next 30 days',
                  ['Vessel', 'Class', 'Loading port', 'Discharge port',
                   'ETD', 'ETA', 'Utilisation', 'TC rate'], rows)])


@reg('Site_Performance_Dashboard.xlsx')
def _b_site_perf(path):
    rows = []
    sites = ['Sg Buloh Plant 1', 'Sg Buloh Plant 2', 'Pasir Gudang',
             'Bayan Lepas', 'Iskandar Plant', 'Kuantan', 'Rawang']
    for s in sites:
        rows.append([s,
                     f'{random.uniform(78, 94):.1f}%',
                     f'{random.uniform(86, 98):.1f}%',
                     f'{random.uniform(90, 99.8):.2f}%',
                     f'{random.uniform(0.4, 3.2):.2f}',
                     random.randint(0, 6),
                     f'{random.uniform(82, 98):.1f}%',
                     random.choice(PEOPLE)])
    _xlsx(path, [('Site Dashboard', 'Site Performance Dashboard - April 2026',
                  ['Site', 'OEE', 'Schedule attainment', 'First-pass yield',
                   'LTIFR', 'NCRs', 'Energy intensity %', 'Site GM'], rows)])


@reg('Target_Financials_3Y.xlsx')
def _b_target_fin_3y(path):
    rows = [
        ['Revenue', 318_200_000, 358_400_000, 392_800_000, '11.0%'],
        ['Gross profit', 86_400_000, 99_800_000, 112_400_000, '14.0%'],
        ['Gross margin', '27.2%', '27.8%', '28.6%', ''],
        ['Opex', 42_800_000, 47_200_000, 51_300_000, '9.5%'],
        ['EBITDA', 43_600_000, 52_600_000, 61_100_000, '18.4%'],
        ['EBITDA margin', '13.7%', '14.7%', '15.6%', ''],
        ['D&A', 9_800_000, 11_200_000, 12_400_000, ''],
        ['EBIT', 33_800_000, 41_400_000, 48_700_000, '20.0%'],
        ['Net interest', 4_800_000, 5_200_000, 4_800_000, ''],
        ['PBT', 29_000_000, 36_200_000, 43_900_000, '23.0%'],
        ['Tax', 7_300_000, 9_100_000, 11_000_000, ''],
        ['PAT', 21_700_000, 27_100_000, 32_900_000, '23.1%'],
    ]
    rows_str = []
    for r in rows:
        rows_str.append([r[0],
                         f'{r[1]:,}' if isinstance(r[1], int) else r[1],
                         f'{r[2]:,}' if isinstance(r[2], int) else r[2],
                         f'{r[3]:,}' if isinstance(r[3], int) else r[3],
                         r[4]])
    _xlsx(path, [('3Y Financials', 'M&A Target - 3 Year Financials (MYR)',
                  ['Line', 'FY2024A', 'FY2025A', 'FY2026E', 'CAGR'], rows_str)])


@reg('Tax_Position_Summary.xlsx')
def _b_tax_position(path):
    rows = [
        ['Current tax payable - Malaysia (LHDN)', 28_400_000, 'Filed', '30-Jun-2026'],
        ['Current tax payable - Indonesia (DGT)', 18_200_000, 'In-progress', '31-Mar-2027'],
        ['Current tax payable - Singapore (IRAS)', 8_400_000, 'Filed', '30-Nov-2026'],
        ['Deferred tax asset - PPE timing', 12_800_000, 'Provisioned', 'n/a'],
        ['Deferred tax liability - Fair value', 18_400_000, 'Provisioned', 'n/a'],
        ['WHT receivable - Cross-border services', 4_200_000, 'Open', 'On-going'],
        ['SST payable - Q1 2026', 9_800_000, 'Filed', '30-Apr-2026'],
        ['RPGT - Property disposal Phase 1A', 6_400_000, 'Filed', '15-May-2026'],
        ['Pillar Two top-up estimate (jurisdictional)', 14_200_000, 'Modelling', '31-Dec-2026'],
        ['Tax loss carryforward - Zava Logistics', 22_800_000, 'Asset', 'Recoverable'],
    ]
    rows_str = [[r[0], f'MYR {r[1]:,}', r[2], r[3]] for r in rows]
    _xlsx(path, [('Tax Position', 'Group Tax Position Summary - 30 Apr 2026',
                  ['Tax position', 'Amount (MYR)', 'Status', 'Filing/Due'], rows_str)])


@reg('TKF_Tabarru_Position_QE.xlsx')
def _b_tkf_tabarru(path):
    rows = [
        ['Tabarru contributions received', 142_800_000],
        ['Less: Wakalah fee (20%)', -28_560_000],
        ['Net to Risk Fund', 114_240_000],
        ['Add: Investment income', 6_200_000],
        ['Less: Claims paid', -38_400_000],
        ['Less: Re-takaful ceded', -22_800_000],
        ['Less: Reserves movement', -8_400_000],
        ['Surplus / (deficit) available', 50_840_000],
        ['Less: Shareholder share (50%)', -25_420_000],
        ['Less: Reserve build-up', -10_000_000],
        ['Surplus distributable to participants', 15_420_000],
    ]
    rows_str = [[r[0], f'MYR {r[1]:,}'] for r in rows]
    _xlsx(path, [('Tabarru Q1 FY26', 'Takaful Operator - Tabarru Fund Position Q1 FY2026',
                  ['Item', 'Amount (MYR)'], rows_str)])


@reg('TYR_Trial_Data_Run_42.xlsx')
def _b_tyr_trial(path):
    rows = []
    for sample in range(40):
        rows.append([f'S-{sample + 1:03d}',
                     random.choice(['175/65 R14', '185/65 R15', '205/55 R16', '225/45 R17', '235/40 R18']),
                     f'{random.uniform(38_000, 62_000):.0f} km',
                     f'{random.uniform(1.4, 3.2):.2f} mm',
                     f'{random.uniform(40, 78):.1f} dB',
                     f'{random.uniform(0.92, 1.08):.3f}',
                     random.choice(['Pass', 'Pass', 'Pass', 'Marginal', 'Fail']),
                     random.choice(PEOPLE) + ' - QA'])
    _xlsx(path, [('Trial Run 42', 'Tyre R&D - Trial Data Run 42',
                  ['Sample', 'Spec', 'Endurance', 'Tread wear (mm)',
                   'Noise', 'Rolling resistance ratio', 'Result', 'QA engineer'], rows)])


@reg('UTIL_Outage_Telemetry.xlsx')
def _b_util_outage(path):
    rows = []
    feeders = ['F-101 KL Central', 'F-118 PJ South', 'F-204 Klang Port',
               'F-301 Penang Mainland', 'F-401 JB Tebrau', 'F-501 KK Inanam',
               'F-601 Kuching Padungan', 'F-701 Ipoh Tasek']
    for f in feeders:
        rows.append([f,
                     random.choice(['Voltage sag', 'Trip - over-current', 'Cable fault', 'Equipment failure', 'Lightning']),
                     _d(-random.randint(0, 14)) + f' {random.randint(0, 23):02d}:{random.randint(0, 59):02d}',
                     f'{random.randint(8, 240)} min',
                     random.randint(420, 18_400),
                     random.choice(['Auto-restored', 'Manual switch', 'Crew dispatched', 'Investigating']),
                     f'{random.uniform(0.04, 1.84):.3f}',
                     random.choice(['Yes', 'No'])])
    _xlsx(path, [('Outage Telemetry', 'Utility - Outage Telemetry (last 14 days)',
                  ['Feeder', 'Cause', 'Start time', 'Duration',
                   'Customers affected', 'Status', 'SAIDI minutes', 'Reportable to ST'], rows)])

# ============================================================================
# Notebook Library XLSX archetype builders (23)
# ============================================================================

@reg('Scoring_Rubric_and_Requirements.xlsx')
def _b_scoring_rubric(path):
    criteria = [
        ('Mandatory', 'M-01', 'ISO 27001 certification (current, valid > 12 months)', 8),
        ('Mandatory', 'M-02', 'Local data residency in Malaysia + Singapore DR', 8),
        ('Mandatory', 'M-03', 'PDPA 2010 + Singapore PDPA Data Processor compliance', 8),
        ('Mandatory', 'M-04', 'Financial - audited statements 3 most recent FY', 6),
        ('Mandatory', 'M-05', 'Bumiputera CIDB G7 or equivalent partner', 6),
        ('Technical', 'T-01', 'SAP S/4HANA 2024 certified delivery practice', 7),
        ('Technical', 'T-02', 'Local consultants - minimum 40 SAP-certified onshore', 7),
        ('Technical', 'T-03', 'Multi-country localisation (MY, ID, SG, VN)', 7),
        ('Technical', 'T-04', 'Integration to Group EAI bus (MuleSoft)', 6),
        ('Technical', 'T-05', 'Change management methodology', 5),
        ('Commercial', 'C-01', 'Total programme cost - lump sum + T&E cap', 8),
        ('Commercial', 'C-02', 'Payment milestones vs deliverables', 6),
        ('Commercial', 'C-03', 'Performance bond - on-demand bank guarantee', 6),
        ('Commercial', 'C-04', 'IP ownership and source-code escrow', 5),
        ('Commercial', 'C-05', 'Warranty period - minimum 12 months post go-live', 5),
        ('Service', 'S-01', 'Hyper-care - 90 days post go-live, on-site', 5),
        ('Service', 'S-02', 'Knowledge transfer plan + acceptance criteria', 5),
        ('Service', 'S-03', 'L3 support SLA 4h response, 24h resolution', 6),
    ]
    rows_c = [[cat, code, desc, w] for cat, code, desc, w in criteria]
    weights = [['Mandatory', sum(w for c, _, _, w in criteria if c == 'Mandatory')],
               ['Technical', sum(w for c, _, _, w in criteria if c == 'Technical')],
               ['Commercial', sum(w for c, _, _, w in criteria if c == 'Commercial')],
               ['Service', sum(w for c, _, _, w in criteria if c == 'Service')]]
    scale = [
        [5, 'Fully addressed with measurable commitment and evidence'],
        [4, 'Fully addressed; some quantitative evidence missing'],
        [3, 'Addressed; partial evidence'],
        [2, 'Partially addressed; gaps'],
        [1, 'Acknowledged but not addressed'],
        [0, 'Silent or non-responsive'],
    ]
    _xlsx(path, [
        ('Rubric', 'RFP Evaluation Rubric - Group ERP Programme',
         ['Category', 'Code', 'Criterion', 'Weight'], rows_c),
        ('Category Weights', 'Weights by Category',
         ['Category', 'Total weight'], weights),
        ('Scoring Scale', 'Scoring Scale 0-5',
         ['Score', 'Definition'], scale),
    ])


@reg('Lab_Results_Panel.xlsx')
def _b_lab_results(path):
    panel = [
        ('Haemoglobin', 'g/dL', '13.2', '13.0 - 17.0', 'Normal'),
        ('White cell count', 'x10^9/L', '12.8', '4.0 - 11.0', 'High'),
        ('Platelets', 'x10^9/L', '328', '150 - 400', 'Normal'),
        ('Sodium', 'mmol/L', '138', '135 - 145', 'Normal'),
        ('Potassium', 'mmol/L', '4.4', '3.5 - 5.0', 'Normal'),
        ('Urea', 'mmol/L', '7.2', '2.5 - 7.0', 'High'),
        ('Creatinine', 'umol/L', '124', '60 - 110', 'High'),
        ('eGFR', 'mL/min/1.73', '58', '> 60', 'Low'),
        ('ALT', 'U/L', '48', '< 41', 'High'),
        ('AST', 'U/L', '54', '< 35', 'High'),
        ('GGT', 'U/L', '102', '< 60', 'High'),
        ('Alk phos', 'U/L', '142', '40 - 130', 'High'),
        ('Total bilirubin', 'umol/L', '18', '< 21', 'Normal'),
        ('Glucose (fasting)', 'mmol/L', '9.8', '3.9 - 5.5', 'High'),
        ('HbA1c', '%', '8.2', '< 6.5', 'High'),
        ('Total cholesterol', 'mmol/L', '6.2', '< 5.2', 'High'),
        ('LDL-C', 'mmol/L', '4.1', '< 2.6', 'High'),
        ('HDL-C', 'mmol/L', '0.9', '> 1.0', 'Low'),
        ('Triglycerides', 'mmol/L', '3.6', '< 1.7', 'High'),
        ('Troponin I', 'ng/mL', '0.02', '< 0.04', 'Normal'),
        ('CRP', 'mg/L', '24', '< 5', 'High'),
        ('TSH', 'mIU/L', '2.4', '0.4 - 4.0', 'Normal'),
    ]
    serial = [
        ['Day 1', 4.82, 32, 8.4, 1.2],
        ['Day 2', 18.40, 38, 7.8, 1.4],
        ['Day 3', 22.10, 28, 7.2, 1.6],
        ['Day 5', 14.20, 18, 6.4, 1.8],
        ['Day 7', 6.40, 12, 5.8, 2.1],
        ['Day 10', 1.80, 8, 5.4, 2.4],
    ]
    flags = [
        ['Creatinine 124', 'Renal function reduced - consider contrast safety + ACE-I dose review', 'Amber'],
        ['HbA1c 8.2%', 'Glycaemic control sub-optimal - intensify per pathway', 'Red'],
        ['LDL-C 4.1', 'Lipid target not met - high-intensity statin', 'Red'],
        ['ALT/AST elevated', 'Transaminitis - rule out statin intolerance vs hepatic congestion', 'Amber'],
        ['CRP 24', 'Inflammation - corroborate with WCC; consider infection screen', 'Amber'],
    ]
    _xlsx(path, [
        ('Panel', 'Lab Results - Patient P-44188 - 28 April 2026',
         ['Analyte', 'Unit', 'Result', 'Reference', 'Flag'],
         [list(r) for r in panel]),
        ('Serial Troponin', 'Serial Cardiac Markers - 0h to Day-10',
         ['Time point', 'Troponin I (ng/mL)', 'CRP (mg/L)',
          'Creatinine trend (umol/L /10)', 'Glucose trend (mmol/L)'], serial),
        ('Clinical Flags', 'Out-of-range values with clinical commentary',
         ['Finding', 'Commentary', 'Severity'], flags),
    ])


@reg('Audit_Findings_Last_Cycle.xlsx')
def _b_audit_findings(path):
    findings = []
    cats = ['ITGC', 'AML KYC', 'Credit Underwriting', 'Market Risk Model',
            'Branch Operations', 'Treasury', 'Reg Reporting', 'Procurement', 'HR Payroll']
    sev = ['High', 'Medium', 'Low']
    statuses = ['Open', 'In Progress', 'Closed - Verified', 'Closed - Pending Verification']
    owners = ['Lim Mei Ling', 'Vijay Subramaniam', 'Ahmad bin Hassan',
              'Kumar Subramaniam', 'Daichi Maruyama', 'Hadar Caspit', 'Siti Aishah']
    for i in range(36):
        cat = random.choice(cats)
        s = random.choice(sev)
        st = random.choice(statuses)
        findings.append([
            f'GIA-2025-{1000 + i:04d}',
            cat,
            s,
            f'{cat} - finding {i + 1}',
            random.choice(owners),
            _d(-365 + i * 8),
            _d(-90 + random.randint(-30, 90)),
            st,
            f'{random.randint(0, 8)} of {random.randint(2, 6)} actions complete',
        ])
    summary = [
        ['High', 8, 3, 5, 0],
        ['Medium', 16, 6, 7, 3],
        ['Low', 12, 4, 5, 3],
    ]
    overdue = []
    for f in findings:
        if f[7] == 'Open' and f[2] in ('High', 'Medium'):
            overdue.append([f[0], f[1], f[2], f[3], f[4], f[5]])
            if len(overdue) >= 12:
                break
    _xlsx(path, [
        ('Findings Register', 'Group Internal Audit - Findings Register (FY2025)',
         ['Audit ID', 'Area', 'Severity', 'Finding', 'Owner',
          'Raised on', 'Target close', 'Status', 'Action progress'], findings),
        ('Summary by Severity', 'Summary by Severity',
         ['Severity', 'Total', 'Open', 'In Progress', 'Closed'], summary),
        ('Overdue Watchlist', 'Overdue Watchlist (open, target date passed)',
         ['Audit ID', 'Area', 'Severity', 'Finding', 'Owner', 'Raised on'], overdue),
    ])


@reg('Compliance_Register_Live.xlsx')
def _b_compliance_register(path):
    items = []
    regs = [
        ('BNM/RH/PD 029', 'Climate Risk Management', 'Risk', 'Daichi Maruyama'),
        ('BNM/RH/PD 028', 'Operational Risk', 'Risk', 'Hadar Caspit'),
        ('BNM/RH/PD 036', 'Cyber Resilience', 'IT Risk', 'Mod Admin'),
        ('AMLA 2001', 'AML KYC Refresh', 'Compliance', 'Lim Mei Ling'),
        ('MFRS 9', 'Expected Credit Loss', 'Finance', 'Catherine Wong'),
        ('Bursa MMLR', 'Continuous Disclosure', 'CompSec', 'Sasha Ouellet'),
        ('Companies Act 2016', 'Annual Return Filing', 'CompSec', 'Sasha Ouellet'),
        ('PDPA 2010', 'Data Protection - Customer', 'Privacy', 'Vijay Subramaniam'),
        ('FSA 2013 S.214', 'Banking License Conditions', 'Reg Affairs', 'Daichi Maruyama'),
        ('IFSA 2013 S.91', 'Shariah Governance', 'Shariah', 'Dr Siti Aishah'),
        ('Income Tax Act', 'CbCR Filing', 'Tax', 'Tan Wei Ming'),
        ('GST/SST Act', 'Quarterly SST Return', 'Tax', 'Tan Wei Ming'),
        ('Employment Act', 'Working Hours and Overtime', 'HR', 'Nurul Huda'),
        ('OSHA 1994', 'HSE Notification', 'HSE', 'Kumar Subramaniam'),
        ('CMSA 2007', 'Insider Dealing Register', 'CompSec', 'Sasha Ouellet'),
    ]
    statuses = ['Compliant', 'Compliant - watch', 'Action required', 'Breach - reported']
    for i, (code, title, fn, owner) in enumerate(regs * 2):
        items.append([
            f'CR-{2026000 + i + 1}',
            code,
            f'{title} ({i % 2 + 1}/2)',
            fn,
            owner,
            _d(-180 + i * 4),
            _d(45 + i * 6),
            random.choice(statuses),
            'Quarterly' if i % 2 else 'Monthly',
        ])
    summary = [
        ['Compliant', 22],
        ['Compliant - watch', 5],
        ['Action required', 2],
        ['Breach - reported', 1],
    ]
    _xlsx(path, [
        ('Live Register', 'Group Compliance Register - Live (May 2026)',
         ['Register ID', 'Regulation', 'Obligation', 'Function', 'Owner',
          'Last reviewed', 'Next review', 'Status', 'Cadence'], items),
        ('Status Summary', 'Status Distribution',
         ['Status', 'Count'], summary),
    ])


@reg('Target_Cap_Table.xlsx')
def _b_target_captable(path):
    shareholders = [
        ['Founder - Datuk Hadar Caspit', 'Ordinary', 6800000, 0.34, 'Yes', '14-Mar-2015'],
        ['Co-Founder - Catherine Wong', 'Ordinary', 4200000, 0.21, 'Yes', '14-Mar-2015'],
        ['Series A - Apex Capital Ventures', 'Preference', 3200000, 0.16, 'No', '22-Aug-2018'],
        ['Series B - Sapphire Growth Fund', 'Preference', 2800000, 0.14, 'No', '14-Sep-2021'],
        ['Series C - Cendana Capital', 'Preference', 1800000, 0.09, 'No', '04-Apr-2024'],
        ['ESOP Pool (vested)', 'Options', 900000, 0.045, 'No', '14-Mar-2015'],
        ['Friends and Family', 'Ordinary', 200000, 0.01, 'No', '14-Mar-2015'],
        ['Strategic - Apex Banking Group', 'Preference', 100000, 0.005, 'No', '02-May-2026'],
    ]
    waterfall = [
        ['Liquidation preference 1x - Apex Capital Ventures (Series A)', 28000000],
        ['Liquidation preference 1.2x - Sapphire (Series B)', 36000000],
        ['Liquidation preference 1.5x - Cendana (Series C)', 42000000],
        ['Strategic preferred - Apex Banking Group', 6000000],
        ['Remaining waterfall to ordinary + ESOP pro-rata', 0],
    ]
    options_grants = []
    for i in range(14):
        options_grants.append([
            f'OPT-2024-{40 + i:03d}',
            random.choice(['Tan Wei Ming', 'Lim Mei Ling', 'Vijay Raja',
                           'Farah Idris', 'Chong Kah Wai', 'Daichi Maruyama']),
            random.choice([20000, 25000, 30000, 40000, 60000, 100000]),
            0.42,
            _d(-720 + i * 30),
            _d(720 + i * 30),
            '4-year cliff 1yr',
            random.choice([0.0, 0.25, 0.5, 0.75, 1.0]),
        ])
    _xlsx(path, [
        ('Cap Table', 'Target Company - Capitalisation Table (Fully Diluted)',
         ['Shareholder', 'Class', 'Shares', '% FD', 'Board Seat', 'Date'],
         shareholders),
        ('Waterfall', 'Liquidation Waterfall (illustrative MYR 200M sale)',
         ['Tranche', 'Cash Out (MYR)'], waterfall),
        ('Options Grants', 'Outstanding Options Grants',
         ['Grant ID', 'Grantee', 'Shares', 'Strike (MYR)', 'Grant date',
          'Vest end', 'Schedule', '% vested'], options_grants),
    ])


@reg('Litigation_Register.xlsx')
def _b_litigation(path):
    cases = []
    statuses = ['Active', 'Settled', 'Dismissed', 'Pending Filing', 'Under Appeal']
    types = ['Contract dispute', 'Employment', 'IP infringement',
             'Regulatory enforcement', 'Tax assessment', 'Tort', 'Shareholder action']
    courts = ['KL High Court', 'Shah Alam Sessions', 'Penang High Court',
              'Court of Appeal', 'Industrial Court', 'SC Adjudication']
    for i in range(24):
        t = random.choice(types)
        st = random.choice(statuses)
        cases.append([
            f'LIT-2025-{200 + i:04d}',
            t,
            random.choice([
                'Cendana Engineering Sdn Bhd', 'Mawar Trading',
                'Former Employee - Mr A', 'BNM Enforcement',
                'IRB Tax Assessment 2022', 'Apex Software Ltd (UK)',
                'Minority shareholder group',
            ]),
            random.choice(courts),
            st,
            random.choice([1200000, 2800000, 4400000, 8800000, 12400000, 24000000]),
            random.choice([0.2, 0.3, 0.4, 0.5, 0.6, 0.7]),
            random.choice([0, 1200000, 800000, 2400000, 3600000]),
            _d(-720 + i * 16),
            random.choice(['Tan Wei Ming', 'Lim Mei Ling', 'Sasha Ouellet']),
        ])
    provisions = [
        ['Total claim exposure', sum(c[5] for c in cases)],
        ['Expected loss (weighted)', round(sum(c[5] * c[6] for c in cases))],
        ['Provision booked at FY2025', sum(c[7] for c in cases)],
        ['Provisioning gap', round(sum(c[5] * c[6] for c in cases) - sum(c[7] for c in cases))],
    ]
    _xlsx(path, [
        ('Cases', 'Litigation Register - Target Company',
         ['Case ID', 'Type', 'Counterparty', 'Court', 'Status',
          'Claim (MYR)', 'P(loss)', 'Provision (MYR)', 'Filed', 'Internal owner'],
         cases),
        ('Provisioning', 'Provisioning Summary',
         ['Item', 'Value (MYR)'], provisions),
    ])


@reg('CFO_Financial_Pack.xlsx')
def _b_cfo_pack(path):
    pl = [
        ['Revenue', 11200, 10580, 5.9, 12100, 92.6],
        ['Cost of sales', -7240, -6820, 6.2, -7820, 92.6],
        ['Gross profit', 3960, 3760, 5.3, 4280, 92.5],
        ['Operating expenses', -2120, -2060, 2.9, -2280, 93.0],
        ['EBITDA', 1840, 1700, 8.2, 2000, 92.0],
        ['Depreciation and Amortisation', -360, -340, 5.9, -380, 94.7],
        ['EBIT', 1480, 1360, 8.8, 1620, 91.4],
        ['Net finance costs', -180, -160, 12.5, -200, 90.0],
        ['Profit before tax', 1300, 1200, 8.3, 1420, 91.5],
        ['Tax', -312, -288, 8.3, -344, 90.7],
        ['Profit after tax', 988, 912, 8.3, 1076, 91.8],
    ]
    divisions = [
        ['Banking', 928, 298, 32.1, 'On plan'],
        ['Insurance', 482, 88, 18.3, 'Below plan'],
        ['Property', 412, 72, 17.5, 'Below plan'],
        ['Plantation', 318, 38, 11.9, 'Below plan'],
        ['Power', 418, 128, 30.6, 'Above plan'],
        ['Retail', 512, 32, 6.3, 'Below plan'],
        ['Manufacturing', 380, -8, -2.1, 'Below plan'],
        ['Telco', 328, 78, 23.8, 'On plan'],
        ['Logistics', 62, 16, 25.8, 'Above plan'],
    ]
    cash = [
        ['Cash and equivalents (open)', 1620],
        ['Operating cash flow', 1840],
        ['Capex', -640],
        ['Acquisitions / Disposals', -180],
        ['Dividends paid', -420],
        ['Net financing', 80],
        ['FX translation', 28],
        ['Cash and equivalents (close)', 2328],
    ]
    kpis = [
        ['Net debt / EBITDA', '2.18x', '<= 2.50x', 'Green'],
        ['Group CET-1 (Banking)', '13.84%', '>= 14.0%', 'Amber'],
        ['Group LCR (Banking)', '142%', '>= 130%', 'Green'],
        ['ROE - Group', '12.4%', '>= 12.0%', 'Green'],
        ['Operating Cost Ratio', '53.6%', '<= 54.0%', 'Green'],
        ['Provision Coverage Ratio', '108%', '>= 100%', 'Green'],
        ['Days Sales Outstanding', '46 days', '<= 45 days', 'Amber'],
    ]
    _xlsx(path, [
        ('Group P&L', 'Group P&L - Q1 FY2026 (MYR million)',
         ['Line item', 'Q1 FY2026', 'Q1 FY2025', 'YoY %', 'Plan', '% of plan'], pl),
        ('Divisional', 'Q1 Revenue and EBITDA by Division',
         ['Division', 'Revenue (MYR M)', 'EBITDA (MYR M)', 'EBITDA Margin %', 'vs Plan'], divisions),
        ('Cash Flow', 'Q1 Cash Flow Summary (MYR M)',
         ['Item', 'Q1 FY2026'], cash),
        ('KPI Dashboard', 'Group KPI Dashboard',
         ['Metric', 'Q1 FY2026', 'Target / Threshold', 'RAG'], kpis),
    ])


@reg('Risk_Heatmap.xlsx')
def _b_risk_heatmap(path):
    risks = [
        ['R-001', 'Credit risk - SME concentration', 'Banking', 4, 4, 16, 'Amber',
         'Hadar Caspit', 'Quarterly portfolio review + sector cap'],
        ['R-002', 'IT operational - core banking patching', 'Banking', 4, 5, 20, 'Red',
         'Mod Admin', 'New change-management policy + 90-day patch SLA'],
        ['R-003', 'Cyber - ransomware', 'Group', 3, 5, 15, 'Amber',
         'Mod Admin', 'Air-gapped backup + tabletop drill quarterly'],
        ['R-004', 'Conduct risk - mis-selling', 'Insurance', 2, 4, 8, 'Amber',
         'Lim Mei Ling', 'Mystery shopping + customer call-back QA'],
        ['R-005', 'Climate physical - flooding', 'Property', 3, 4, 12, 'Amber',
         'Catherine Wong', 'Flood elevation in new projects + insurance review'],
        ['R-006', 'CPO price volatility', 'Plantation', 4, 3, 12, 'Amber',
         'Hadar Caspit', 'Hedging programme + cost discipline'],
        ['R-007', 'Talent attrition - engineering', 'Group', 3, 3, 9, 'Amber',
         'Nurul Huda', 'Engineering refresh + new comp band'],
        ['R-008', 'Regulatory - climate stress test (BNM)', 'Banking', 3, 3, 9, 'Amber',
         'Daichi Maruyama', 'Joint working group with BNM Q2/Q3'],
        ['R-009', 'Supply chain - chip shortage', 'Manufacturing', 3, 4, 12, 'Amber',
         'Ahmad bin Hassan', '6-month inventory buffer + dual-source'],
        ['R-010', 'FX exposure - USD short', 'Group Treasury', 3, 3, 9, 'Amber',
         'Tan Wei Ming', 'Hedge 80% rolling 12-month'],
        ['R-011', 'AML / sanctions screening false-neg', 'Banking', 2, 5, 10, 'Amber',
         'Vijay Subramaniam', 'Quarterly model validation + audit'],
        ['R-012', 'Project execution - Bandar Zava Phase 2', 'Property', 3, 4, 12, 'Amber',
         'Catherine Wong', 'Independent monitor + PRC governance'],
        ['R-013', 'Cyber - 3rd party SaaS breach', 'Group', 3, 4, 12, 'Amber',
         'Mod Admin', '3rd-party assurance programme + SOC2 mandatory'],
        ['R-014', 'Reputation - regulator media', 'Group', 2, 4, 8, 'Amber',
         'Sasha Ouellet', 'Crisis playbook + spokesperson roster'],
        ['R-015', 'Climate transition - stranded assets', 'Energy', 3, 4, 12, 'Amber',
         'Hadar Caspit', 'Forward stranded-asset review + IFRS S2 disclosure'],
    ]
    movements = [
        ['R-002', '12', '20', 'Up', 'IT patch backlog discovered post audit'],
        ['R-006', '16', '12', 'Down', 'CPO price stabilising + hedging in place'],
        ['R-008', '12', '9', 'Down', 'Climate WG progress with regulator'],
        ['R-013', '8', '12', 'Up', 'New SaaS breach in sister industry - exposure recheck'],
    ]
    appetite = [
        ['Operational loss > MYR 5M events', 4, 8, 'Below threshold'],
        ['Single-name concentration', '<= 8% Tier-1', '11.2% Tier-1', 'BREACH'],
        ['VaR 99% daily', '<= MYR 28M', 'MYR 24M', 'Within'],
        ['Group LCR', '>= 130%', '142%', 'Within'],
        ['Group CET-1', '>= 14.0%', '13.84%', 'BREACH (-16 bps)'],
    ]
    _xlsx(path, [
        ('Heatmap', 'Group Risk Heatmap - Q1 FY2026',
         ['Risk ID', 'Risk', 'Owner Division', 'Likelihood (1-5)',
          'Impact (1-5)', 'Score', 'RAG', 'Owner', 'Mitigation'], risks),
        ('QoQ Movement', 'Movements vs Prior Quarter',
         ['Risk ID', 'Prior score', 'Current score', 'Direction', 'Driver'], movements),
        ('Appetite Status', 'Risk Appetite Indicators',
         ['Indicator', 'Threshold', 'Current', 'Status'], appetite),
    ])


@reg('Usage_Telemetry_Last_12M.xlsx')
def _b_usage_telem(path):
    months = ['May-25', 'Jun-25', 'Jul-25', 'Aug-25', 'Sep-25', 'Oct-25',
              'Nov-25', 'Dec-25', 'Jan-26', 'Feb-26', 'Mar-26', 'Apr-26']
    rows = []
    base_dau = 4200
    base_wau = 14800
    base_mau = 32400
    for i, m in enumerate(months):
        trend = 1 + i * 0.018 - (0.04 if i == 7 else 0) + random.uniform(-0.02, 0.02)
        rows.append([
            m,
            int(base_dau * trend),
            int(base_wau * trend),
            int(base_mau * trend),
            round(28.4 + i * 0.4 + random.uniform(-0.6, 0.6), 1),
            round(0.62 + i * 0.006 + random.uniform(-0.012, 0.012), 3),
            int(48400 * trend),
        ])
    features = [
        ['Dashboard - Cashflow Forecast', 0.84, 12.4, 'Power user'],
        ['Workflow - Approval Routing', 0.72, 8.2, 'Heavy'],
        ['Mobile - Card Controls', 0.68, 6.8, 'Heavy'],
        ['Reporting - Custom Builder', 0.42, 14.4, 'Medium'],
        ['Notification - Smart Alerts', 0.36, 2.4, 'Medium'],
        ['Open Banking - Aggregator', 0.28, 3.2, 'Low'],
        ['Admin - SSO Setup', 0.22, 0.4, 'Setup'],
        ['Beta - AI Insights', 0.18, 4.8, 'Beta'],
    ]
    cohorts = [
        ['New (< 90 days)', 1284, 412, 0.32],
        ['Active core (90-365 days)', 18400, 14820, 0.81],
        ['Long-tenure (1+ yrs)', 22480, 19980, 0.89],
        ['Dormant (> 60 days inactive)', 4480, 0, 0.00],
    ]
    _xlsx(path, [
        ('Monthly', 'Usage Telemetry - Last 12 Months',
         ['Month', 'DAU', 'WAU', 'MAU', 'Avg session (min)',
          'WAU/MAU', 'Active in last 30d'], rows),
        ('Feature Adoption', 'Feature Adoption Rates',
         ['Feature', 'Adoption %', 'Median uses/wk per user', 'Tier'], features),
        ('Cohorts', 'Activity by Cohort',
         ['Cohort', 'Users', 'Active (30d)', 'Active rate'], cohorts),
    ])


@reg('Support_Tickets_Log.xlsx')
def _b_support_log(path):
    cats = ['Login / SSO', 'Card decline', 'Payment failed', 'App crash',
            'Statement download', 'Notification missing', 'Feature request',
            'Account closure', 'Fraud report', 'Limit increase']
    sev = ['P1', 'P2', 'P3', 'P4']
    states = ['Open', 'In progress', 'Awaiting customer', 'Resolved', 'Closed']
    rows = []
    for i in range(60):
        cat = random.choice(cats)
        s = random.choice(sev)
        st = random.choice(states)
        rows.append([
            f'TKT-2026-{40000 + i}',
            _d(-30 + i // 2),
            cat,
            s,
            st,
            random.choice(['App', 'Web', 'Call', 'Chat', 'Branch']),
            random.choice(['Tier-1', 'Tier-2', 'Tier-3', 'Engineering']),
            random.randint(2, 480),
            random.choice([1, 2, 3, 4, 5, None]),
        ])
    by_cat = {}
    for r in rows:
        by_cat[r[2]] = by_cat.get(r[2], 0) + 1
    cat_summary = [[k, v] for k, v in sorted(by_cat.items(), key=lambda x: -x[1])]
    sla = [
        ['P1', '15 min', '12 min', 0.98, 'Green'],
        ['P2', '1 hour', '52 min', 0.96, 'Green'],
        ['P3', '4 hours', '4.2 hours', 0.88, 'Amber'],
        ['P4', '24 hours', '18 hours', 0.94, 'Green'],
    ]
    _xlsx(path, [
        ('Tickets', 'Support Tickets - Last 30 Days',
         ['Ticket ID', 'Opened', 'Category', 'Severity', 'State',
          'Channel', 'Routed to', 'Resolution (min)', 'CSAT (1-5)'], rows),
        ('By Category', 'Volume by Category',
         ['Category', 'Tickets'], cat_summary),
        ('SLA Performance', 'SLA Performance by Severity',
         ['Severity', 'SLA Target', 'Avg Time to Resolve',
          'Within SLA %', 'RAG'], sla),
    ])


@reg('Pager_Alerts_Log.xlsx')
def _b_pager_alerts(path):
    services = ['payments-orchestrator', 'auth-service', 'cards-issuer',
                'core-banking-east', 'core-banking-west', 'sso-gateway',
                'mobile-bff', 'web-portal', 'reporting-warehouse',
                'fraud-engine', 'kyc-screening']
    runbooks = ['RB-001', 'RB-014', 'RB-022', 'RB-038', 'RB-044', 'RB-052', 'RB-063']
    rows = []
    for i in range(80):
        svc = random.choice(services)
        sev = random.choice(['P1', 'P2', 'P3'])
        ack_min = random.randint(1, 12) if sev == 'P1' else random.randint(2, 30)
        rows.append([
            f'PD-2026-{50000 + i}',
            _d(-10 + i // 12),
            f'{12 + i % 24:02d}:{random.randint(0, 59):02d}',
            svc,
            sev,
            random.choice(['Error rate', 'Latency p99', 'CPU', 'Memory',
                           'Queue depth', 'Disk', 'Hearbeat lost']),
            random.choice(['Vijay Raja', 'Mod Admin', 'Kumar Subramaniam',
                           'Chong Kah Wai', 'Tan Wei Ming']),
            ack_min,
            random.choice(['Self-healed', 'Mitigated', 'Escalated', 'Open']),
            random.choice(runbooks),
        ])
    by_svc = {}
    for r in rows:
        by_svc[r[3]] = by_svc.get(r[3], 0) + 1
    svc_top = [[k, v] for k, v in sorted(by_svc.items(), key=lambda x: -x[1])]
    p1 = [r for r in rows if r[4] == 'P1'][:14]
    _xlsx(path, [
        ('Alerts', 'PagerDuty Alerts - Last 10 Days',
         ['Alert ID', 'Date', 'Time', 'Service', 'Severity',
          'Symptom', 'On-call', 'Ack (min)', 'Outcome', 'Runbook'], rows),
        ('Top Services', 'Alert Volume by Service',
         ['Service', 'Alerts'], svc_top),
        ('P1 Alerts', 'P1 Alerts - Detail',
         ['Alert ID', 'Date', 'Time', 'Service', 'Severity', 'Symptom',
          'On-call', 'Ack (min)', 'Outcome', 'Runbook'], p1),
    ])


@reg('Past_Similar_Incidents.xlsx')
def _b_past_incidents(path):
    incidents = [
        ['INC-2024-038', '14-Mar-2024', 'P1', 'payments-orchestrator',
         'Silent failover - hbeat lost', '3h 14m',
         'Heartbeat alarm missing - only queue-depth alarm', 'Added hbeat alarm', 'Closed'],
        ['INC-2024-104', '22-Jul-2024', 'P1', 'auth-service',
         'Cert rotation failure',
         '1h 42m', 'Cert expiry not in maintenance window',
         'Cert auto-renew runbook published', 'Closed'],
        ['INC-2024-218', '14-Oct-2024', 'P2', 'core-banking-east',
         'EOD batch slowdown', '6h 28m',
         'Index missing on collateral table', 'Index added + nightly index health check',
         'Closed'],
        ['INC-2025-014', '08-Jan-2025', 'P1', 'fraud-engine',
         'Model service throttled', '2h 04m',
         'Vendor rate limit lower than load test',
         'Negotiated higher rate limit + autoscale ahead', 'Closed'],
        ['INC-2025-088', '14-Apr-2025', 'P2', 'payments-orchestrator',
         'Backlog drift', '3h 38m',
         'Drift detection missing on queue depth metric',
         'Drift detection added in Datadog', 'Closed'],
        ['INC-2025-128', '02-Jul-2025', 'P1', 'sso-gateway',
         'Identity provider 5xx', '0h 48m',
         'IdP regional outage',
         'Multi-region IdP failover designed (in flight)', 'In progress'],
        ['INC-2025-198', '14-Oct-2025', 'P2', 'mobile-bff',
         'Increased error rate after deploy',
         '1h 18m', 'Deploy validation gate too lenient',
         'Stricter canary gates', 'Closed'],
        ['INC-2025-244', '18-Dec-2025', 'P2', 'core-banking-west',
         'Connection pool exhaustion',
         '2h 22m', 'New microservice opened too many connections',
         'Connection pool limits + chaos testing', 'Closed'],
        ['INC-2026-104', '14-Feb-2026', 'P2', 'cards-issuer',
         'Token decryption latency spike',
         '1h 06m', 'HSM throttling under load',
         'Burst capacity uplift + queue tier 2', 'Closed'],
        ['INC-2026-188', '11-May-2026', 'P2', 'payments-orchestrator',
         'IBG queue backlog 18.4k',
         '4h 30m', 'Silent failure - heartbeat alarm STILL not configured',
         'Heartbeat alarm + canary 25%', 'Action open'],
    ]
    similar = [
        ['INC-2024-038', 'IDENTICAL pattern - same service - SAME root cause family',
         'Action from 2024 NOT applied to all environments'],
        ['INC-2025-088', 'Drift detection added in Datadog but heartbeat alarm not added',
         'Two parallel controls - one added, one missed'],
        ['INC-2025-244', 'Connection pool fix was scoped to west - east not addressed',
         'Cross-environment scope is recurring weakness'],
    ]
    _xlsx(path, [
        ('History', 'Past Similar Incidents - Payments Orchestrator family',
         ['Incident ID', 'Date', 'Severity', 'Service', 'Symptom',
          'Duration', 'Root Cause', 'Corrective Action', 'Status'], incidents),
        ('Pattern Match', 'Incidents Matching the Current Pattern',
         ['Reference', 'Match Reason', 'Note'], similar),
    ])


@reg('Sensitivity_Scenarios.xlsx')
def _b_sensitivity(path):
    scenarios = [
        ['Base', 1.00, 1.00, 1.00, 1.00, 18.4, 218, 4.8],
        ['Upside +20% revenue', 1.20, 1.00, 1.00, 1.00, 26.8, 312, 3.8],
        ['Downside -10% revenue', 0.90, 1.00, 1.00, 1.00, 13.8, 144, 6.4],
        ['Capex +15%', 1.00, 1.00, 1.15, 1.00, 15.6, 178, 5.4],
        ['Opex +5%', 1.00, 1.05, 1.00, 1.00, 17.2, 198, 5.0],
        ['Cost of debt +200bps', 1.00, 1.00, 1.00, 1.02, 16.8, 188, 5.2],
        ['Stress - rev -20% + capex +10%', 0.80, 1.00, 1.10, 1.00, 9.8, 82, 8.4],
        ['Optimised - rev +10% + opex -3%', 1.10, 0.97, 1.00, 1.00, 22.4, 268, 4.2],
        ['Delay launch 6 months', 0.92, 1.00, 1.00, 1.00, 14.2, 158, 5.8],
        ['Commodity hit - input +18%', 1.00, 1.09, 1.00, 1.00, 14.6, 168, 5.6],
    ]
    tornado = [
        ['Revenue', -8.6, 8.4],
        ['Capex', -2.8, 2.8],
        ['Opex', -1.2, 1.6],
        ['Cost of debt', -1.6, 1.4],
        ['Forex (USD/MYR)', -2.4, 2.2],
        ['Commodity input', -3.8, 1.8],
    ]
    breakeven = [
        ['Min revenue for IRR > 12%', '-7.2% from base'],
        ['Min revenue for NPV > 0', '-22.4% from base'],
        ['Max capex for IRR > 12%', '+18.4% from base'],
        ['Max input cost for margin > 14%', '+12.8% from base'],
    ]
    _xlsx(path, [
        ('Scenarios', 'Sensitivity Scenarios - Project Bandar Zava Ph2',
         ['Scenario', 'Rev x', 'Opex x', 'Capex x', 'Cost of debt x',
          'IRR %', 'NPV (MYR M)', 'Payback (yrs)'], scenarios),
        ('Tornado', 'Tornado - NPV Impact (MYR M)',
         ['Driver', 'Downside', 'Upside'], tornado),
        ('Breakeven', 'Breakeven Analysis',
         ['Question', 'Answer'], breakeven),
    ])


@reg('Channel_Spend_Performance.xlsx')
def _b_channel_spend(path):
    rows = [
        ['Search (Google + Bing)', 320000, 18400, 4200, 0.228, 17.40, 76.20, 4.20],
        ['Meta (FB + IG)', 1536000, 84200, 12800, 0.152, 18.24, 120.00, 2.84],
        ['TikTok', 768000, 48400, 6200, 0.128, 15.87, 123.87, 1.42],
        ['YouTube', 1920000, 128400, 10200, 0.079, 14.95, 188.24, 0.84],
        ['CTV / OTT', 512000, 32400, 2400, 0.074, 15.80, 213.33, 0.42],
        ['Programmatic display', 640000, 24800, 1800, 0.073, 25.81, 355.56, 0.18],
        ['LinkedIn (B2B)', 384000, 6200, 980, 0.158, 61.94, 391.84, 0.62],
        ['Radio', 384000, 0, 0, 0.0, 0.0, 0.0, 0.0],
        ['OOH', 320000, 0, 0, 0.0, 0.0, 0.0, 0.0],
    ]
    weekly = []
    for i in range(13):
        weekly.append([
            f'W{i + 1}',
            random.randint(420000, 580000),
            random.randint(28000, 48000),
            random.randint(2400, 6800),
            round(random.uniform(2.4, 4.2), 2),
        ])
    quality = [
        ['Search', 0.84, 0.62, 0.41, 'Highest converting'],
        ['Meta', 0.62, 0.48, 0.22, 'Mid'],
        ['TikTok', 0.42, 0.28, 0.14, 'Top-funnel'],
        ['YouTube', 0.38, 0.22, 0.09, 'Awareness'],
        ['LinkedIn', 0.74, 0.48, 0.28, 'B2B-strong'],
    ]
    _xlsx(path, [
        ('Spend by Channel', 'Channel Spend and Performance - Q1 FY2026',
         ['Channel', 'Spend (MYR)', 'Clicks/Plays', 'Funded accounts',
          'Funded rate', 'CPA (MYR per click)', 'Cost / funded',
          'Funded per MYR1K'], rows),
        ('Weekly Trend', 'Total Spend and Outcomes by Week',
         ['Week', 'Spend (MYR)', 'Clicks/Plays', 'Funded', 'ROAS'], weekly),
        ('Lead Quality', 'Lead Quality by Channel (MQL/SQL/Funded conversion)',
         ['Channel', 'MQL rate', 'SQL rate', 'Funded rate', 'Note'], quality),
    ])


@reg('Survey_and_Sentiment_Pulse.xlsx')
def _b_survey_sentiment(path):
    themes = [
        ['Pricing relative to competitors', 'Negative', 0.38, 'High', 'Tactical'],
        ['App speed and reliability', 'Negative', 0.32, 'High', 'Strategic'],
        ['Customer service responsiveness', 'Mixed', 0.28, 'High', 'Operational'],
        ['Branch closure communication', 'Negative', 0.24, 'Medium', 'Comms'],
        ['Rewards programme attractiveness', 'Mixed', 0.22, 'Medium', 'Tactical'],
        ['Mobile app new design', 'Positive', 0.18, 'Medium', 'Strategic'],
        ['ATM availability in Sabah/Sarawak', 'Negative', 0.18, 'Medium', 'Operational'],
        ['Onboarding ease', 'Positive', 0.14, 'Medium', 'Strategic'],
        ['Card design and benefits', 'Positive', 0.12, 'Low', 'Tactical'],
        ['Sustainability messaging', 'Mixed', 0.08, 'Low', 'Comms'],
    ]
    nps = []
    months = ['May-25', 'Jun-25', 'Jul-25', 'Aug-25', 'Sep-25', 'Oct-25',
              'Nov-25', 'Dec-25', 'Jan-26', 'Feb-26', 'Mar-26', 'Apr-26']
    base = 32
    for i, m in enumerate(months):
        base += random.uniform(-2, 3)
        nps.append([m, round(base, 1), random.randint(2400, 4200),
                    round(random.uniform(0.18, 0.32), 3)])
    quotes = [
        ['Pricing', 'Negative',
         '"Your spread on FX is wider than the digital banks. Why am I still here?"'],
        ['App reliability', 'Negative',
         '"App froze again during payday transfer - third time this quarter."'],
        ['Service', 'Positive',
         '"The chat agent solved my issue in 4 minutes - first time that ever happened."'],
        ['Branch closure', 'Negative',
         '"You closed my branch and the next one is 28km. No notice. Awful."'],
        ['Mobile redesign', 'Positive',
         '"The new dashboard is much cleaner - I can see my cashflow at a glance now."'],
    ]
    _xlsx(path, [
        ('Themes', 'Customer Themes - Q1 FY2026',
         ['Theme', 'Sentiment', 'Share of mentions', 'Priority', 'Type'], themes),
        ('NPS Trend', 'NPS Trend - Last 12 Months',
         ['Month', 'NPS', 'Respondents', 'Response rate'], nps),
        ('Verbatims', 'Sample Verbatim Quotes',
         ['Theme', 'Sentiment', 'Quote'], quotes),
    ])


@reg('Sales_Pipeline_Influence.xlsx')
def _b_pipeline_influence(path):
    opps = []
    stages = ['MQL', 'SQL', 'Qualified', 'Proposal', 'Negotiation', 'Closed Won', 'Closed Lost']
    sources = ['Search', 'Meta', 'TikTok', 'YouTube', 'LinkedIn', 'CTV', 'Direct', 'Referral']
    for i in range(40):
        st = random.choice(stages)
        amt = random.choice([4800, 12400, 28000, 44000, 88000, 184000, 320000])
        opps.append([
            f'OPP-2026-{2000 + i}',
            random.choice(['Bumi Konsortium Sdn Bhd', 'Sapphire Holdings Bhd',
                           'Apex Banking Group', 'Tropika Sdn Bhd', 'Cendana Properties',
                           'Hijau Industries', 'Mawar Group', 'Tanjung Resources',
                           'Bayu Logistics', 'Sinar Capital']),
            random.choice(['Lim Mei Ling', 'Hadar Caspit', 'Vijay Raja',
                           'Sasha Ouellet', 'Daichi Maruyama']),
            st,
            amt,
            random.choice(sources),
            _d(-90 + i * 2),
            _d(random.choice([14, 21, 28, 45, 60, 90])),
        ])
    by_source = {}
    by_source_won = {}
    for o in opps:
        s = o[5]
        by_source[s] = by_source.get(s, 0) + o[4]
        if o[3] == 'Closed Won':
            by_source_won[s] = by_source_won.get(s, 0) + o[4]
    src_rows = []
    for s in sources:
        src_rows.append([s, by_source.get(s, 0), by_source_won.get(s, 0),
                         round(100 * by_source_won.get(s, 0) / max(by_source.get(s, 1), 1), 1)])
    funnel = [
        ['MQL', 412, 1.00],
        ['SQL', 184, 0.45],
        ['Qualified', 92, 0.50],
        ['Proposal', 48, 0.52],
        ['Negotiation', 26, 0.54],
        ['Closed Won', 14, 0.54],
    ]
    _xlsx(path, [
        ('Opportunities', 'Sales Pipeline - Q1 FY2026',
         ['Opp ID', 'Customer', 'AE', 'Stage', 'Amount (MYR)',
          'Source', 'Created', 'Close target'], opps),
        ('By Source', 'Pipeline and Won Revenue by Source',
         ['Source', 'Pipeline (MYR)', 'Closed Won (MYR)', 'Win %'], src_rows),
        ('Funnel', 'Funnel Conversion - Last 90 Days',
         ['Stage', 'Volume', 'Stage conversion'], funnel),
    ])


@reg('Treasury_Cash_Positions.xlsx')
def _b_treasury(path):
    positions = []
    banks = ['Apex Bank', 'CIMB', 'Maybank', 'HSBC', 'Standard Chartered', 'DBS', 'BNP Paribas']
    ccys = [('MYR', 1.0), ('USD', 4.42), ('SGD', 3.28), ('EUR', 4.74),
            ('JPY', 0.029), ('IDR', 0.000284), ('VND', 0.000182)]
    accts = ['Operating', 'Tax', 'Dividend', 'Capex Reserve', 'Working Capital', 'CSR']
    for i in range(28):
        ccy, fx = random.choice(ccys)
        amt = random.choice([1240000, 4800000, 12400000, 24800000, 48400000, 84200000])
        positions.append([
            f'ACC-{700000 + i}',
            random.choice(banks),
            ccy,
            random.choice(accts),
            amt,
            round(amt * fx, 2),
            random.choice(['Current', 'Money Market', 'Fixed Deposit']),
            random.choice(['Z+1', 'Z+3', 'Z+7']),
        ])
    by_ccy = {}
    for p in positions:
        by_ccy[p[2]] = by_ccy.get(p[2], 0) + p[5]
    ccy_rows = sorted([[k, round(v, 2)] for k, v in by_ccy.items()],
                      key=lambda x: -x[1])
    forecast = []
    for i in range(14):
        forecast.append([_d(i),
                         round(184000000 + i * 1200000 + random.uniform(-2400000, 2400000), 2),
                         round(176000000 + i * 1800000, 2),
                         round((184000000 + i * 1200000) - (176000000 + i * 1800000), 2)])
    _xlsx(path, [
        ('Positions', 'Treasury Cash Positions - 12 May 2026',
         ['Account', 'Bank', 'CCY', 'Purpose', 'Balance (local)',
          'Balance (MYR equiv)', 'Type', 'Liquidity'], positions),
        ('By Currency', 'Cash by Currency (MYR equivalent)',
         ['CCY', 'Total (MYR equiv)'], ccy_rows),
        ('14-Day Forecast', '14-Day Cash Forecast',
         ['Date', 'Forecast inflow (MYR)', 'Forecast outflow (MYR)',
          'Net (MYR)'], forecast),
    ])


@reg('FX_Hedging_Register.xlsx')
def _b_fx_hedging(path):
    rows = []
    pairs = ['USD/MYR', 'SGD/MYR', 'EUR/MYR', 'JPY/MYR', 'IDR/MYR', 'VND/MYR']
    instruments = ['Forward', 'NDF', 'Cross-Currency Swap', 'Option', 'Vanilla Collar']
    for i in range(24):
        pair = random.choice(pairs)
        instr = random.choice(instruments)
        rows.append([
            f'HDG-2026-{3000 + i}',
            pair,
            instr,
            random.choice(['Buy', 'Sell']),
            random.choice([1, 2, 4, 8, 12, 18, 24]),
            random.choice([4.20, 4.32, 4.42, 4.48]),
            random.choice([4.42, 4.46, 4.52]),
            random.choice([1200000, 4800000, 12400000, 24800000]),
            random.choice(['Apex Bank', 'CIMB', 'HSBC', 'Standard Chartered', 'DBS']),
            _d(-60 + i),
            _d(60 + i * 18),
        ])
    coverage = [
        ['USD - 12-month forecast exposure', 'USD 412M', 'USD 332M', 0.806],
        ['SGD - 12-month forecast exposure', 'SGD 84M', 'SGD 76M', 0.905],
        ['EUR - 12-month forecast exposure', 'EUR 42M', 'EUR 34M', 0.810],
        ['JPY - 12-month forecast exposure', 'JPY 4.8B', 'JPY 3.6B', 0.750],
        ['IDR - 12-month forecast exposure', 'IDR 248B', 'IDR 188B', 0.758],
        ['VND - 12-month forecast exposure', 'VND 1.2T', 'VND 0.9T', 0.750],
    ]
    pnl = []
    months = ['Nov-25', 'Dec-25', 'Jan-26', 'Feb-26', 'Mar-26', 'Apr-26']
    for m in months:
        pnl.append([m,
                    round(random.uniform(-2400000, 2400000), 2),
                    round(random.uniform(-1800000, 1800000), 2),
                    round(random.uniform(-600000, 1200000), 2)])
    _xlsx(path, [
        ('Open Hedges', 'FX Hedging Register - Open Trades',
         ['Hedge ID', 'Pair', 'Instrument', 'Direction', 'Tenor (months)',
          'Strike', 'Spot at trade', 'Notional', 'Counterparty',
          'Trade date', 'Maturity'], rows),
        ('Coverage', 'Hedge Coverage Ratios by Currency',
         ['Exposure', 'Total', 'Hedged', 'Coverage'], coverage),
        ('Realised P&L', 'Realised Hedge P&L (Last 6 Months)',
         ['Month', 'Hedge realised P&L (MYR)', 'FX mark-to-market (MYR)',
          'Net (MYR)'], pnl),
    ])


@reg('Capex_and_Maintenance_Plan.xlsx')
def _b_capex_maint(path):
    capex = []
    sites = ['Klang Plant', 'Cyberjaya', 'Iskandar Refinery', 'Penang',
             'Kota Kinabalu', 'Kuching', 'Shah Alam', 'Bangi']
    cats = ['Reliability', 'Capacity', 'Compliance', 'HSE', 'Digitalisation', 'Sustainability']
    for i in range(28):
        site = random.choice(sites)
        cat = random.choice(cats)
        capex.append([
            f'CAPEX-2026-{500 + i}',
            site,
            cat,
            f'{cat} - {site} - Project {i + 1}',
            random.choice([180000, 480000, 1240000, 2800000, 4800000, 12400000]),
            random.choice(['Q2 26', 'Q3 26', 'Q4 26', 'Q1 27']),
            random.choice(['Approved', 'In Approval', 'Deferred', 'Cancelled']),
            random.choice(['Ahmad bin Hassan', 'Kumar Subramaniam', 'Lim Mei Ling']),
            random.choice([0.14, 0.18, 0.22, 0.26, 0.32]),
        ])
    maint = []
    eqs = ['CT-103', 'P-218', 'FCC-101', 'CDU-101', 'K-201 Granulator',
           'K-202 Press', 'Line BR-2 Filler', 'HVAC AHU-04']
    for i in range(20):
        maint.append([
            f'WO-{42000 + i}',
            random.choice(eqs),
            random.choice(['Preventive', 'Predictive', 'Corrective', 'Inspection']),
            _d(i * 4),
            random.choice([4, 8, 16, 24, 48]),
            random.choice(sites),
            random.choice(['Planned', 'Open', 'Completed']),
            random.choice(['Vijay Subramaniam', 'Kumar Subramaniam', 'Tan Wei Ming']),
        ])
    summary = [
        ['Reliability', 28400000],
        ['Capacity', 64200000],
        ['Compliance', 14800000],
        ['HSE', 8400000],
        ['Digitalisation', 24800000],
        ['Sustainability', 18200000],
    ]
    _xlsx(path, [
        ('Capex', 'FY2026 Capex Plan',
         ['ID', 'Site', 'Category', 'Project', 'Value (MYR)',
          'Target', 'Status', 'Sponsor', 'IRR'], capex),
        ('Maintenance', 'Planned Maintenance Schedule (next 90 days)',
         ['WO', 'Equipment', 'Type', 'Date', 'Duration (hrs)',
          'Site', 'Status', 'Owner'], maint),
        ('Category Summary', 'Capex by Category (MYR)',
         ['Category', 'Total (MYR)'], summary),
    ])


@reg('Supplier_Performance_Scorecard.xlsx')
def _b_supplier(path):
    suppliers = []
    names = ['Bumi Konsortium Sdn Bhd', 'Mawar Trading', 'Cendana Logistics',
             'Sapphire Industrial Supplies', 'Permata Components',
             'Tropika Chemicals', 'Hijau Renewables', 'Anggerik Bearings',
             'Pertiwi Engineering', 'Gemilang Castings', 'Sinar Plastics',
             'Bayu Couriers', 'Tanjung Tubes']
    cats = ['Critical', 'Strategic', 'Operational', 'Tactical']
    for i, n in enumerate(names):
        suppliers.append([
            f'SUP-{8000 + i}',
            n,
            random.choice(cats),
            round(random.uniform(78, 98), 1),
            round(random.uniform(82, 99), 1),
            round(random.uniform(0.5, 4.8), 1),
            round(random.uniform(84, 99), 1),
            round(random.uniform(60, 95), 1),
            random.choice(['Green', 'Green', 'Amber', 'Amber', 'Red']),
        ])
    issues = [
        ['SUP-8001', 'Bumi Konsortium', 'Q1 delay - 2 deliveries 6 days late',
         'Buffer stock + alternate sourcing identified'],
        ['SUP-8005', 'Permata Components', '3 quality rejects on incoming inspection',
         'On supplier improvement plan - 90 day'],
        ['SUP-8010', 'Gemilang Castings', 'Audit non-conformance (ISO 9001)',
         'Re-audit scheduled August 2026'],
        ['SUP-8012', 'Bayu Couriers', 'Cyber incident at courier - data exposed',
         'Vendor risk re-rating + replacement RFP started'],
    ]
    spend = [
        ['Q1 FY2026', 84200000],
        ['Q4 FY2025', 92800000],
        ['Q3 FY2025', 88400000],
        ['Q2 FY2025', 80200000],
    ]
    _xlsx(path, [
        ('Scorecard', 'Supplier Performance Scorecard - Q1 FY2026',
         ['Supplier ID', 'Name', 'Category', 'OTD %', 'Quality %',
          'Defect rate %', 'Compliance %', 'Cost competitiveness',
          'RAG'], suppliers),
        ('Issues', 'Open Issues',
         ['Supplier ID', 'Name', 'Issue', 'Action'], issues),
        ('Spend Trend', 'Total Spend Trend',
         ['Quarter', 'Spend (MYR)'], spend),
    ])


@reg('Worker_Safety_Incidents.xlsx')
def _b_worker_safety(path):
    rows = []
    cats = ['Near-miss', 'First-aid', 'Medical-treatment', 'Lost-time', 'Restricted-duty']
    types = ['Slip/Trip', 'Manual handling', 'Caught-between', 'Struck-by',
             'Chemical exposure', 'Heat stress', 'Working at height', 'Electrical']
    sites = ['Klang Plant', 'Cyberjaya', 'Iskandar Refinery', 'Penang',
             'Kota Kinabalu', 'Bangi']
    for i in range(30):
        sev = random.choice(cats)
        rows.append([
            f'INC-SAFE-2026-{i + 1:03d}',
            _d(-90 + i * 3),
            random.choice(sites),
            random.choice(types),
            sev,
            random.choice(['Day shift', 'Night shift']),
            random.choice([0, 0, 0, 1, 3, 7, 14]) if sev == 'Lost-time' else 0,
            random.choice(['Ahmad bin Hassan', 'Kumar Subramaniam',
                           'Siti Aishah', 'Hadar Caspit']),
            random.choice(['Closed', 'Closed', 'Closed', 'In RCA', 'Open']),
        ])
    rates = []
    months = ['Nov-25', 'Dec-25', 'Jan-26', 'Feb-26', 'Mar-26', 'Apr-26']
    for m in months:
        rates.append([m, round(random.uniform(0.3, 0.9), 2),
                      round(random.uniform(1.4, 2.4), 2),
                      random.randint(2, 6)])
    by_type = {}
    for r in rows:
        by_type[r[3]] = by_type.get(r[3], 0) + 1
    type_rows = sorted([[k, v] for k, v in by_type.items()], key=lambda x: -x[1])
    _xlsx(path, [
        ('Incidents', 'Worker Safety Incidents - Last 90 Days',
         ['Incident ID', 'Date', 'Site', 'Type', 'Severity', 'Shift',
          'Days lost', 'Investigation lead', 'Status'], rows),
        ('Monthly Rates', 'LTIFR + TRIFR by Month',
         ['Month', 'LTIFR', 'TRIFR', 'Near-misses logged'], rates),
        ('By Type', 'Incident Count by Type',
         ['Type', 'Count'], type_rows),
    ])


@reg('Disclosure_Framework_Mapping.xlsx')
def _b_disclosure_map(path):
    rows = []
    frameworks = ['TCFD', 'IFRS S1', 'IFRS S2', 'GRI', 'CDP Climate',
                  'CDP Water', 'Bursa Sust', 'SDGs', 'SASB Banking']
    disclosures = [
        ('Governance', 'Board oversight of climate'),
        ('Governance', 'Management role'),
        ('Strategy', 'Climate risks and opportunities'),
        ('Strategy', 'Scenario analysis'),
        ('Strategy', 'Resilience of strategy'),
        ('Risk Management', 'Processes for ID + assessment'),
        ('Risk Management', 'Integration into ERM'),
        ('Metrics and Targets', 'Scope 1 emissions'),
        ('Metrics and Targets', 'Scope 2 emissions'),
        ('Metrics and Targets', 'Material Scope 3'),
        ('Metrics and Targets', 'Climate-related targets'),
        ('Metrics and Targets', 'Capex aligned to transition'),
        ('Metrics and Targets', 'Water withdrawal'),
        ('Metrics and Targets', 'Waste and recycling'),
        ('Social', 'Workforce diversity'),
        ('Social', 'Health and safety LTIFR'),
        ('Social', 'Community investment'),
        ('Governance', 'Anti-bribery training'),
    ]
    for i, (cat, d) in enumerate(disclosures):
        rows.append([
            f'DISC-{100 + i}',
            cat, d,
            random.choice(['Yes', 'Yes', 'Yes', 'No']),
            random.choice(['Yes', 'Yes', 'No']),
            random.choice(['Yes', 'Yes', 'No']),
            random.choice(['Yes', 'No']),
            random.choice(['Yes', 'Yes']),
            random.choice(['Yes', 'Yes', 'No']),
            random.choice(['Yes', 'Yes']),
            random.choice(['Yes', 'No']),
            random.choice(['Yes', 'Yes']),
            random.choice(['Daichi Maruyama', 'Lim Mei Ling',
                           'Sasha Ouellet', 'Kumar Subramaniam', 'Catherine Wong']),
            random.choice(['Annual Report', 'Sustainability Report',
                           'CDP filing', 'Bursa Sust filing', 'TCFD section']),
        ])
    gaps = [
        ['Strategy - Resilience of strategy', 'IFRS S2 + TCFD',
         'No quantified resilience disclosure under 1.5C scenario',
         'Add resilience commentary in FY2026 report'],
        ['Metrics - Material Scope 3', 'IFRS S2 + TCFD + CDP',
         'Financed emissions Banking partially disclosed',
         'Complete financed emissions methodology + disclose FY2026'],
        ['Metrics - Climate targets', 'IFRS S2',
         'No interim 2030 targets disclosed for Property and Manufacturing',
         '2030 interim targets to be set by Sustainability Council Q3 FY2026'],
    ]
    timeline = [
        ['Initial gap close - data availability', 'Q3 FY2026'],
        ['Internal assurance walkthrough', 'Q4 FY2026'],
        ['Final disclosure draft', 'Q1 FY2027'],
        ['External assurance (limited)', 'Q2 FY2027'],
        ['Bursa filing', 'Q3 FY2027'],
    ]
    _xlsx(path, [
        ('Framework Mapping', 'Disclosure Framework Mapping - FY2026',
         ['Disclosure ID', 'Category', 'Required disclosure', 'TCFD', 'IFRS S1',
          'IFRS S2', 'GRI', 'CDP Climate', 'CDP Water', 'Bursa Sust',
          'SDGs', 'SASB Banking', 'Owner', 'Location'], rows),
        ('Gap Register', 'Disclosure Gaps Identified',
         ['Disclosure', 'Frameworks affected', 'Gap', 'Closure plan'], gaps),
        ('Closure Timeline', 'Gap Closure Timeline',
         ['Milestone', 'Target'], timeline),
    ])


@reg('Compliance_Training_Catalog.xlsx')
def _b_training_catalog(path):
    catalog = [
        ['CT-001', 'Anti-Bribery and Corruption', 'All staff', 'Annual', 45, 'Required'],
        ['CT-002', 'AML / CFT - Foundations', 'All staff', 'Annual', 60, 'Required'],
        ['CT-003', 'AML / CFT - Advanced', 'Banking + Insurance staff', 'Annual', 90, 'Required'],
        ['CT-004', 'PDPA 2010 - Data Privacy', 'All staff', 'Annual', 40, 'Required'],
        ['CT-005', 'Cyber Awareness', 'All staff', 'Bi-annual', 30, 'Required'],
        ['CT-006', 'Cyber Awareness - Privileged Users', 'Privileged users', 'Bi-annual', 60, 'Required'],
        ['CT-007', 'Sanctions Screening', 'Banking + Trade staff', 'Annual', 60, 'Required'],
        ['CT-008', 'Insider Dealing', 'Listed-entity insiders', 'Annual', 40, 'Required'],
        ['CT-009', 'Conduct Risk + Mis-selling', 'Frontline staff', 'Annual', 45, 'Required'],
        ['CT-010', 'Health and Safety Foundations', 'Plant + Field', 'Annual', 30, 'Required'],
        ['CT-011', 'Working at Height', 'Plant + Field', 'Annual', 45, 'Role-based'],
        ['CT-012', 'Confined Space Entry', 'Refinery + Plant', 'Annual', 60, 'Role-based'],
        ['CT-013', 'Whistleblowing Channels', 'All staff', 'Bi-annual', 20, 'Required'],
        ['CT-014', 'Code of Conduct', 'All staff', 'Annual', 40, 'Required'],
        ['CT-015', 'Shariah Governance', 'Islamic Banking staff', 'Annual', 60, 'Role-based'],
        ['CT-016', 'Sustainability Foundations', 'All staff', 'One-off', 30, 'Recommended'],
        ['CT-017', 'TCFD + Climate Risk', 'Risk + Finance staff', 'Bi-annual', 60, 'Role-based'],
        ['CT-018', 'IFRS 9 ECL Methodology', 'Banking Finance + Risk', 'Annual', 90, 'Role-based'],
    ]
    onboarding = [
        ['New Hire - Day 1', ['CT-014', 'CT-001', 'CT-013']],
        ['New Hire - Day 30', ['CT-002', 'CT-004', 'CT-005']],
        ['New Hire - Day 60', ['CT-016']],
        ['Role-based - Banker', ['CT-003', 'CT-007', 'CT-009']],
        ['Role-based - Field', ['CT-010', 'CT-011']],
        ['Role-based - Refinery', ['CT-010', 'CT-011', 'CT-012']],
        ['Role-based - Privileged IT', ['CT-006']],
        ['Role-based - Risk + Finance', ['CT-017', 'CT-018']],
    ]
    onb_rows = [[label, ', '.join(courses)] for label, courses in onboarding]
    compliance = [
        ['All staff (Group)', 28400, 26840, 0.945],
        ['Banking', 8240, 7980, 0.969],
        ['Insurance', 2480, 2380, 0.960],
        ['Plant + Field', 4820, 4520, 0.938],
        ['Privileged IT', 218, 218, 1.000],
    ]
    _xlsx(path, [
        ('Course Catalog', 'Compliance Training Catalog FY2026',
         ['Course ID', 'Title', 'Audience', 'Cadence', 'Duration (min)', 'Status'],
         catalog),
        ('Onboarding Plan', 'Standard Onboarding Plan by Role',
         ['Milestone / Role', 'Courses'], onb_rows),
        ('Compliance Status', 'Compliance Status - Q1 FY2026',
         ['Cohort', 'Required', 'Completed', 'Compliance %'], compliance),
    ])


# ============================================================================
# DOCX archetype builders (26)
# ============================================================================

@reg('Account_Plan.docx')
def _b_account_plan(path):
    _docx(path, 'Strategic Account Plan - Apex Banking Group FY2026', [
        ('p', 'Account: Apex Banking Group | Account Manager: Hadar Caspit | Coverage tier: Tier-1 Strategic | Plan period: FY2026 (Jul 2026 - Jun 2027)'),
        ('h1', '1. Account snapshot'),
        ('p', 'Apex Banking Group is the second-largest commercial bank in Malaysia with consolidated assets of MYR 412B, deposit franchise of MYR 318B, and a retail customer base of 4.8M. The relationship has been classified Tier-1 Strategic since FY2023. FY2025 wallet share is estimated at 26%.'),
        ('h1', '2. Revenue history and FY2026 target'),
        ('table', [
            ['Metric', 'FY2023', 'FY2024', 'FY2025', 'FY2026 Target'],
            ['Revenue (MYR M)', '38.4', '46.2', '52.8', '62.5'],
            ['Gross margin', '42%', '44%', '45%', '47%'],
            ['Wallet share', '22%', '24%', '26%', '32%'],
            ['Net Promoter Score', '+38', '+44', '+51', '+55'],
        ]),
        ('h1', '3. Buying centres and stakeholder map'),
        ('bullets', [
            'Group CEO - Datuk Azman bin Ibrahim - sponsor (warm)',
            'Group CFO - Catherine Wong - economic buyer (warm)',
            'Group CIO - Vijay Pillai - technical buyer (neutral, leans competitor X)',
            'Head of Retail Banking - Lim Mei Ling - end-user champion (champion)',
            'Head of Procurement - Ahmad Faizal - process gatekeeper (neutral)',
        ]),
        ('h1', '4. FY2026 plays'),
        ('numbered', [
            'Cloud migration extension - 36-month renewal + 4 additional workloads (MYR 8.4M ACV)',
            'AI Copilot rollout to 12,000 retail bankers (MYR 14.2M ACV, M365 E5 base)',
            'Risk modelling co-innovation - shared IP, 2 PoCs already running',
            'Data platform consolidation - replace 3 legacy warehouses (MYR 6.8M services)',
        ]),
        ('h1', '5. Competitive landscape'),
        ('p', 'Incumbent: Microsoft (us). Active threats: Google Cloud (CIO sponsor), AWS (CTO-level relationship via prior employer). Defensive moves: Q1 executive sponsorship visit; quarterly innovation council; co-funded AI Centre of Excellence.'),
        ('h1', '6. Risks and watch-outs'),
        ('bullets', [
            'CIO turnover risk - 6-month tenure remaining per market intel',
            'Tier-1 capital pressure may compress FY2026 IT capex',
            'Regulator scrutiny on AI usage (BNM/RH/STD/036)',
            'Talent shortage in their AI team - we need to over-rotate on enablement',
        ]),
    ])


@reg('Campaign_Brief.docx')
def _b_campaign_brief(path):
    _docx(path, 'Campaign Brief - Bandar Zava Phase 2A Launch', [
        ('p', 'Brand: Zava Properties | Campaign: Bandar Zava Phase 2A Launch | In-market: 15 Jun - 30 Sep 2026 | Budget: MYR 4.8M | Owner: Sasha Ouellet (Marketing)'),
        ('h1', '1. Business objective'),
        ('p', 'Sell-through 320 units (out of 540 launched) within first 6 months at average price of MYR 1,080 per sq ft. Sustain GDV pipeline of MYR 520M.'),
        ('h1', '2. Audience'),
        ('bullets', [
            'Primary: 32-45 year old upgrader couples, household income MYR 18K-32K/month, Klang Valley',
            'Secondary: First-time buyers with parental co-signers, 28-35, MYR 12K+ household income',
            'Tertiary: Investor buyers, Singapore-based Malaysians, second-home seekers',
        ]),
        ('h1', '3. Single-minded proposition'),
        ('p', 'Bandar Zava Phase 2A is the only freehold, 35-minute-to-KL township that bundles a 5-acre community park, an international primary school within the master plan, and a 24-hour smart-home concierge - at a sub-MYR 1,200 psf entry price.'),
        ('h1', '4. Channel mix and budget'),
        ('table', [
            ['Channel', 'Budget (MYR)', 'Share', 'Primary KPI'],
            ['Digital (Meta + Google + property portals)', '1,920,000', '40%', 'Cost per qualified lead'],
            ['OOH (Federal Highway + LDP + NPE)', '720,000', '15%', 'Awareness lift'],
            ['Sales gallery activation', '1,080,000', '22.5%', 'Walk-in count'],
            ['Influencer + PR', '480,000', '10%', 'Earned reach'],
            ['Print + Radio', '300,000', '6%', 'Reach uplift in BM-Chinese segments'],
            ['Contingency', '300,000', '6%', '-'],
        ]),
        ('h1', '5. Creative direction'),
        ('p', 'Two creative streams: (a) aspirational family-life storytelling for upgraders; (b) yield-and-capital-gain rational frame for investor segment. Tone: warm-confident-grounded. Avoid the cliched "luxury skyline" visual language used by competitor X.'),
        ('h1', '6. Risks'),
        ('bullets', [
            'OPR upcycle could compress affordability by 60-80 bps',
            'Competitor X launching adjacent project 4 weeks later at lower psf',
            'Material price index up 8.4% YoY - margin pressure if discounting kicks in',
        ]),
    ])


@reg('CEO_Report.docx')
def _b_ceo_report(path):
    _docx(path, 'CEO Monthly Report to Board - April 2026', [
        ('p', 'To: Board of Directors | From: Group CEO | Period: April 2026 | Classification: Board-Confidential'),
        ('h1', '1. Executive summary'),
        ('p', 'April closed at MYR 3.84B revenue (98.4% of plan, +6.2% YoY) and MYR 642M EBITDA (94.1% of plan, +3.1% YoY). The 5.9% EBITDA miss is concentrated in two divisions: Plantation (FFB price compression) and Manufacturing (input cost spike). Liquidity is strong at MYR 1.84B cash plus MYR 2.6B undrawn RCF.'),
        ('h1', '2. By division'),
        ('table', [
            ['Division', 'Revenue (MYR M)', 'vs Plan', 'EBITDA (MYR M)', 'vs Plan'],
            ['Banking', '928', '+3.2%', '298', '+1.8%'],
            ['Insurance', '482', '+1.4%', '88', '-2.4%'],
            ['Property', '412', '-6.8%', '72', '-12.4%'],
            ['Plantation', '318', '-14.2%', '38', '-32.6%'],
            ['Power', '418', '+2.6%', '128', '+4.2%'],
            ['Retail', '512', '+1.8%', '32', '-8.4%'],
            ['Manufacturing', '380', '-4.6%', '-8', 'NM'],
            ['Telco', '328', '+2.2%', '78', '+0.8%'],
            ['Logistics', '62', '+8.4%', '16', '+12.4%'],
        ]),
        ('h1', '3. Capital and treasury'),
        ('p', 'Net debt / EBITDA closed at 2.18x (covenant 3.50x). Average cost of debt 4.62%. FX exposure: USD 412M net short, hedged 78%. Refinancing schedule for FY2026 is fully covered; FY2027 has MYR 1.2B maturity to address by H1.'),
        ('h1', '4. Strategic agenda - April update'),
        ('bullets', [
            'ZAVA FORWARD 2030 - Pillar 2 (digital) is on track; Pillar 4 (energy transition) slipped 6 weeks',
            'Bandar Zava Phase 2 launch is confirmed for 15 June',
            'Cendana Power acquisition - SPA signing targeted for end-May',
            'Group AI office: 14 deployed agents, 38 in pipeline',
        ]),
        ('h1', '5. Risks raised at GMC this month'),
        ('bullets', [
            'CPO price spike risk to Plantation margins (downside MYR 28M Q4)',
            'BNM stress test results expected end-May - Banking liaison ready',
            'Cyber-security: 1 ransomware attempt foiled by SOC; tabletop scheduled',
        ]),
    ])


@reg('CSEC_Last_AGM_Minutes.docx')
def _b_csec_agm(path):
    _docx(path, 'Minutes of the 28th Annual General Meeting - Zava Group Berhad', [
        ('p', 'Date: 24 April 2026 | Venue: Grand Ballroom, Connexion Conference Centre, KL Sentral | Time commenced: 10:00 hours | Time concluded: 13:42 hours'),
        ('h1', '1. Chairman and quorum'),
        ('p', 'The Chairman, Tan Sri Dr Lim Cheng Yong, opened the meeting at 10:00 hours. The Group Company Secretary, Sasha Ouellet, confirmed a quorum of 472 shareholders representing 67.4% of issued share capital was present or represented by proxy, satisfying the requirement under Article 73 of the Constitution.'),
        ('h1', '2. Resolutions tabled'),
        ('table', [
            ['No.', 'Resolution', 'Voted For', 'Voted Against', 'Outcome'],
            ['1', 'Adoption of Audited Financial Statements FY2025', '99.84%', '0.16%', 'Carried'],
            ['2', 'Declaration of Final Single-Tier Dividend of 24.0 sen', '99.78%', '0.22%', 'Carried'],
            ['3', 'Re-election of Datuk Wong Lai Ling as Director', '98.42%', '1.58%', 'Carried'],
            ['4', 'Re-election of Mr Vijay Subramaniam as Director', '97.84%', '2.16%', 'Carried'],
            ['5', 'Approval of Directors\' Fees of MYR 4.8M', '94.62%', '5.38%', 'Carried'],
            ['6', 'Re-appointment of KPMG as Auditors', '99.91%', '0.09%', 'Carried'],
            ['7', 'Authority to issue shares - Section 75', '88.42%', '11.58%', 'Carried'],
            ['8', 'Proposed share buyback', '92.18%', '7.82%', 'Carried'],
        ]),
        ('h1', '3. Q&A highlights from the floor'),
        ('bullets', [
            'MSWG raised dividend payout ratio - Chairman confirmed 50-60% policy maintained',
            'Minority Shareholder questioned Plantation EBITDA fall - CFO Hadar Caspit provided written follow-up',
            'Sustainability rating queries addressed by Head of IR Daichi Maruyama',
        ]),
        ('h1', '4. Closing remarks and adjournment'),
        ('p', 'The Chairman thanked shareholders for their continued support and confirmed the next AGM is targeted for end-April 2027. The meeting was adjourned at 13:42 hours.'),
        ('p', 'Certified true minutes - Sasha Ouellet, Group Company Secretary | Approved by the Board on 12 May 2026.'),
    ])


@reg('ECOM_Last_Sale_Postmortem.docx')
def _b_ecom_postmortem(path):
    _docx(path, 'E-Commerce Sale Post-Mortem - 4.4 Mega Sale 2026', [
        ('p', 'Event: 4.4 Mega Sale | Dates: 1-4 April 2026 | Reviewer: Hadar Caspit (e-Commerce GM) | Distributed to: Trading Council, Marketing, Operations, Tech'),
        ('h1', '1. Outcome vs target'),
        ('table', [
            ['Metric', 'Target', 'Actual', 'Variance'],
            ['GMV (MYR M)', '142.0', '128.4', '-9.6%'],
            ['Orders (k)', '420', '388', '-7.6%'],
            ['AOV (MYR)', '338', '331', '-2.1%'],
            ['Conversion rate', '3.4%', '3.1%', '-30 bps'],
            ['Return rate', '8.0%', '11.4%', '+340 bps'],
            ['Customer acquisition', '38,000', '42,800', '+12.6%'],
            ['Repeat-buyer share', '54%', '49%', '-5 pp'],
        ]),
        ('h1', '2. What worked'),
        ('bullets', [
            'New-user acquisition exceeded plan by 12.6% via TikTok creator program',
            'Voucher stacking strategy lifted AOV by MYR 18 vs control cell',
            'Live-commerce stream peak concurrent viewers 84,000 (vs target 60,000)',
        ]),
        ('h1', '3. What broke'),
        ('numbered', [
            'Checkout outage on Day 1 22:00-22:48 - estimated MYR 4.8M GMV loss; root cause: cart-service auto-scaling lag',
            'Payment gateway A timeouts on Day 2 between 14:00-15:30 - declined 2,800 orders worth MYR 920K',
            'Fashion category over-discounted by 4-6 pp vs plan due to inconsistent supplier matching - margin damage MYR 1.2M',
            'Returns surge driven by 3 sellers with poor quality control - now under suspension review',
        ]),
        ('h1', '4. Lessons for 6.6 Mid-Year Sale'),
        ('bullets', [
            'Pre-scale cart and checkout services 96 hours ahead - confirmed with SRE',
            'Dual-rail payment gateway with hot failover (gateway A primary, gateway B 25% mirror)',
            'Discount governance - margin guardrails enforced in PIM at submission time',
            'Seller QA gate - any seller with > 12% return rate auto-removed from campaign',
        ]),
    ])


@reg('FMCG_Brand_Strategy.docx')
def _b_fmcg_brand(path):
    _docx(path, 'Brand Strategy - Mawar Personal Care FY2026', [
        ('p', 'Brand: Mawar | Category: Personal Care - Body Wash, Shampoo, Soap | Owner: Lim Mei Ling (Brand Director) | Period: FY2026'),
        ('h1', '1. Where the brand stands today'),
        ('p', 'Mawar is the #3 mass personal-care brand in West Malaysia with 14.2% volume share and 11.8% value share in Body Wash, 12.4% / 10.6% in Shampoo. Brand health scores: Aided awareness 92%, Consideration 38%, P3M purchase 18%. Penetration grew 1.8 points YoY to 41% in target households.'),
        ('h1', '2. Strategic question'),
        ('p', 'How do we shift Mawar from a price-fighter brand to a premium-mass brand and lift gross margin from 28% to 34% over 3 years - without losing the loyal MYR 8-12 SKU buyer?'),
        ('h1', '3. Brand purpose and proposition'),
        ('p', 'Purpose: Mawar believes everyday self-care is a small act of dignity. Proposition: Salon-grade ingredients, supermarket prices, designed in Malaysia. Three pillars: (1) Botanical efficacy (Centella, Tualang Honey, Pegaga); (2) Halal + dermatologically tested; (3) Refill economy.'),
        ('h1', '4. Portfolio reshape FY2026'),
        ('table', [
            ['SKU tier', 'FY2025 share of revenue', 'FY2026 target', 'Move'],
            ['Core - MYR 8-12 entry', '58%', '46%', 'Sustain'],
            ['Premium-mass - MYR 14-22', '34%', '42%', 'Grow'],
            ['Hero - MYR 28-38 Tualang Honey', '8%', '12%', 'Launch'],
        ]),
        ('h1', '5. Communications spine'),
        ('bullets', [
            'Hero film: "Lebih daripada wangian" - 60s + 30s + 6s cutdowns',
            'Always-on: 5-week sprint cadence per pillar - TVC + digital + retail media',
            'Influencer roster: 3 Tier-1 lifestyle, 12 Tier-2 dermatology + hair, 80 nano-influencers',
            'In-store: Re-launched gondola with refill station in 480 hypermarket doors by Q2',
        ]),
        ('h1', '6. KPIs'),
        ('bullets', [
            'Value share: 11.8% to 13.6% by end-FY2026',
            'Gross margin: 28% to 31% by end-FY2026, on track to 34% by FY2028',
            'Brand power score (Kantar): 32 to 38',
            'Refill SKU share of total volume: 4% to 12%',
        ]),
    ])


@reg('GLV_Change_Control.docx')
def _b_glv_change(path):
    _docx(path, 'Change Control Record - GLV-2026-117 - Cooling Tower 3 Cell Replacement', [
        ('p', 'CC ID: GLV-2026-117 | Asset: Cooling Tower 3 | Plant: Iskandar Refinery | Initiated by: Mod Admin (Maintenance Eng) | Date raised: 09-May-2026'),
        ('h1', '1. Change summary'),
        ('p', 'Replace deteriorated drift eliminator and fill media in Cell 3 of CT-103. Original installation 2014; ultrasonic inspection 12-Apr-2026 identified 38% loss of fill efficiency.'),
        ('h1', '2. Classification'),
        ('table', [
            ['Field', 'Value'],
            ['Change category', 'Like-for-Like with material upgrade'],
            ['HSE classification', 'Medium'],
            ['Production impact', 'Cell 3 isolated; other 3 cells continue (75% capacity)'],
            ['Estimated outage', '36 hours (Cell only)'],
            ['Budget', 'MYR 280,000'],
            ['Risk score (P x I)', '8 / 25 (Medium)'],
            ['Target start', '02-Jun-2026 08:00'],
            ['Target end', '03-Jun-2026 20:00'],
        ]),
        ('h1', '3. Approvals required'),
        ('bullets', [
            'Plant Manager - Lim Mei Ling - Approved 11-May-2026',
            'Process Safety Lead - Kumar Subramaniam - Approved 12-May-2026',
            'Operations Superintendent - Ahmad bin Hassan - Approved 12-May-2026',
            'HSE Manager - Siti Aishah - Conditional approval pending JHA review',
        ]),
        ('h1', '4. JHA outline'),
        ('numbered', [
            'Confined-space entry permit required - certified attendant + gas-test every 15 min',
            'Hot-work permit required for any cutting - fire-watch + extinguisher kit on standby',
            'Working at height >2m - full harness + double-lanyard, anchor inspected',
            'Chemical hazard - residual chlorine in basin water - SCBA available, eye-wash within 10m',
        ]),
        ('h1', '5. Rollback plan'),
        ('p', 'If new fill media fails performance test (Approach > 5 degF, Range < 22 degF), restore old fill from staging and re-test within 8 hours. Decision authority: Lim Mei Ling (Plant Mgr).'),
        ('h1', '6. Post-implementation review'),
        ('p', 'PIR scheduled for 17-Jun-2026 (10 working days after restart). Owner: Mod Admin.'),
    ])


@reg('HC_Patient_EMR_Excerpt.docx')
def _b_hc_emr(path):
    _docx(path, 'Patient EMR Excerpt (de-identified)', [
        ('p', 'MRN: MRN-2026-44188 | Patient: P-44188 (de-identified) | Age: 58 | Sex: M | Attending: Dr Nurul Huda | Ward: 6B-12 | Admission date: 28-Apr-2026'),
        ('h1', 'Presenting complaint'),
        ('p', 'Patient presented to the Emergency Department on 28-Apr-2026 at 04:12 with 3-hour history of central chest tightness radiating to left arm, associated with diaphoresis and nausea. No syncope. Pain was 7 / 10 at rest, worse on exertion.'),
        ('h1', 'Past medical history'),
        ('bullets', [
            'Type 2 Diabetes Mellitus (diagnosed 2014) - HbA1c 8.2% (Mar 2026)',
            'Hypertension (diagnosed 2016) - on Amlodipine 10mg + Perindopril 5mg',
            'Dyslipidaemia - LDL-C 3.8 mmol/L (Mar 2026), on Atorvastatin 20mg',
            'Ex-smoker (quit 2018, prior 20 pack-years)',
            'Family Hx of premature CAD (father MI age 51)',
        ]),
        ('h1', 'Examination on admission'),
        ('p', 'BP 158/96, HR 102, RR 22, SpO2 96% RA, Temp 36.8 degC. Heart sounds dual, no added sounds. Lungs clear. No peripheral oedema. ECG: ST-elevation 2-3mm in V2-V5 with reciprocal changes II, III, aVF.'),
        ('h1', 'Investigations'),
        ('table', [
            ['Test', 'Result', 'Reference'],
            ['Troponin I (0h)', '4.82 ng/mL', '< 0.04'],
            ['Troponin I (3h)', '18.40 ng/mL', '< 0.04'],
            ['Creatinine', '108 umol/L', '60-110'],
            ['Potassium', '4.4 mmol/L', '3.5-5.0'],
            ['HbA1c', '8.2%', '< 6.5%'],
            ['LDL-C', '3.8 mmol/L', '< 1.8'],
        ]),
        ('h1', 'Management to date'),
        ('numbered', [
            'STEMI pathway activated 04:24 - door-to-cathlab 38 min',
            'Primary PCI - LAD proximal 95% lesion - DES deployed - TIMI 3 flow restored',
            'DAPT: Aspirin 100mg + Ticagrelor 90mg BD',
            'Beta-blocker initiated: Bisoprolol 2.5mg OD',
            'High-intensity statin: Atorvastatin uptitrated to 80mg ON',
            'ACE-I continued; metformin held pending creatinine review',
        ]),
        ('h1', 'Plan'),
        ('p', 'Transfer to CCU; echo within 24 hrs; cardiac rehab referral; smoking-cessation reinforcement; HbA1c target < 7.0%; LDL-C target < 1.4 mmol/L; outpatient follow-up Cardiology Clinic 2 weeks.'),
    ])


@reg('HR_Role_JD.docx')
def _b_hr_jd(path):
    _docx(path, 'Job Description - Senior Manager, Group Internal Audit', [
        ('p', 'Position: Senior Manager, Group Internal Audit | Reports to: Chief Internal Auditor | Function: Internal Audit | Grade: M5 | Location: KL Sentral'),
        ('h1', '1. Purpose of the role'),
        ('p', 'Lead end-to-end internal audit engagements across the Group, with primary coverage of Financial Services and Investment Holdings divisions, providing independent assurance to the Board Audit Committee on risk management, control adequacy, and governance.'),
        ('h1', '2. Key responsibilities'),
        ('bullets', [
            'Plan and execute 8-10 risk-based audit engagements per year across assigned coverage',
            'Lead an audit team of 3-4 auditors; coach and develop junior team members',
            'Present findings and recommendations to senior management and the Audit Committee',
            'Follow-up on management action plans; escalate overdue items per CAE protocol',
            'Contribute to the annual risk-based audit plan and continuous risk monitoring',
            'Maintain proficiency in BNM, SC, Bursa, and Sarbanes-Oxley equivalent regulations',
        ]),
        ('h1', '3. Required qualifications'),
        ('bullets', [
            'Recognised degree in Accounting, Finance, or related discipline',
            'CIA, CISA, CA(M), CPA, ACCA or equivalent professional qualification',
            'Minimum 8 years total experience with at least 4 in Internal Audit or Big 4 Risk Advisory',
            'Demonstrated track record in Financial Services audit (Banking, Insurance, AMC)',
            'IIA International Standards proficiency mandatory',
        ]),
        ('h1', '4. Preferred qualifications'),
        ('bullets', [
            'Data analytics tooling (ACL, IDEA, Power BI, Python)',
            'IT general controls audit experience',
            'Bahasa Malaysia proficiency for working with regulator and management interviews',
        ]),
        ('h1', '5. Behavioural competencies'),
        ('bullets', [
            'Stakeholder management - C-suite presence',
            'Critical thinking - challenges status quo respectfully',
            'Drive for results - delivers within tight timelines',
            'Integrity and independence - non-negotiable',
        ]),
        ('h1', '6. Compensation band'),
        ('p', 'Grade M5 | Base salary band: MYR 18,000 - 26,000 per month | Variable bonus target: 25% of base | Group long-term incentive eligibility: Yes (after 12 months)'),
    ])


@reg('HSE_Incident_Wartime_Log.docx')
def _b_hse_war(path):
    _docx(path, 'HSE Incident Wartime Log - Iskandar Refinery FCC Trip 06-May-2026', [
        ('p', 'Incident Ref: HSE-2026-104 | Site: Iskandar Refinery | Unit: FCC-101 | Severity (provisional): Level 3 (Significant) | Incident commander: Kumar Subramaniam'),
        ('h1', '1. Timeline (real-time log)'),
        ('table', [
            ['Time', 'Event', 'Action / Owner'],
            ['14:42', 'FCC-101 high differential pressure alarm (P-114) - operator response timer started', 'Console - Lim'],
            ['14:46', 'Slurry circulation pump P-218 high vibration; aux pump P-218B failed to start', 'Field - Ahmad'],
            ['14:51', 'Catalyst hopper level dropped below 22%; FCC began controlled trip', 'DCS'],
            ['14:54', 'FCC tripped; flare load increased by 38 t/hr', 'Flare KO drum - Siti'],
            ['14:58', 'Site Emergency Response activated - Tier 2', 'IC - Kumar'],
            ['15:08', 'Adjacent CDU-101 reduced rate to 78% to balance hydrogen demand', 'Console'],
            ['15:24', 'Pump P-218B isolated; mechanical inspection commenced', 'Maintenance - Vijay'],
            ['16:12', 'Root cause provisional: bearing failure on P-218 (8.4 yr in service)', 'Reliability - Hadar'],
            ['18:48', 'FCC restart preparation - feedstock blend approved by tech', 'Tech - Mod Admin'],
            ['22:14', 'FCC restart - first feed at 18% rate', 'Console'],
            ['07-May 04:32', 'Returned to 92% rate; flare load normalised', 'Console'],
        ]),
        ('h1', '2. Personnel impact'),
        ('bullets', [
            'Zero personal injuries',
            'No environmental release beyond permitted flare envelope',
            'Community impact assessment - nil complaints received via hotline',
        ]),
        ('h1', '3. Production and financial impact'),
        ('p', 'Estimated production loss 38,400 bbl gasoline + 12,800 bbl middle distillate. Realised loss MYR 8.2M before insurance recovery. Reportable to MIDA under Industrial Hazard Reporting Order (within 24 hrs).'),
        ('h1', '4. Immediate corrective actions'),
        ('numbered', [
            'All slurry pumps (P-218A through P-218E) scheduled for vibration check by 12-May-2026',
            'Reliability standard tightened: replace bearings at 7-year mark (was 10-year run-to-failure)',
            'P-218B impeller replacement; full overhaul by 31-May-2026',
            'Tabletop review with shift teams 09-May-2026 - lessons cascaded',
        ]),
        ('h1', '5. Reporting obligations'),
        ('p', 'DOSH notified 06-May-2026 16:24 (within 4 hr SLA). Internal HSE Committee scheduled 13-May-2026. Final RCA report due to DOSH by 06-Jun-2026.'),
    ])


@reg('INC_Wartime_Channel_Log.docx')
def _b_inc_war(path):
    _docx(path, 'Wartime Channel Log - INC-2026-204 - Group Online Banking Degradation', [
        ('p', 'Incident: INC-2026-204 | Severity: P1 | Started: 12-May-2026 09:42 MYT | Resolved: 12-May-2026 13:28 MYT | Incident Commander: Mod Admin | Comms Lead: Sasha Ouellet'),
        ('h1', '1. Wartime channel - Microsoft Teams "P1-INC-204"'),
        ('table', [
            ['Time', 'Speaker', 'Statement'],
            ['09:42', 'Synthetic monitor', 'Login success rate dropped to 62% nationwide'],
            ['09:43', 'On-call SRE - Vijay', 'Confirming - PagerDuty also firing'],
            ['09:44', 'IC - Mod Admin', 'Declaring P1. Bridge open. Comms paging Daichi'],
            ['09:48', 'Network - Kumar', 'DC1 primary load balancer showing 38% packet loss on internal segment'],
            ['09:52', 'Daichi (Head of IR)', 'On bridge. BNM notification window starts now per RH/STD/008'],
            ['09:58', 'IC', 'Failover to DC2 LB pair authorised. Risk Council notified'],
            ['10:14', 'SRE', 'Failover complete. Login success climbing - now 92%'],
            ['10:22', 'Customer Care', 'Call volume 4.2x baseline. IVR message updated'],
            ['11:00', 'IC', 'Sustained > 99% for 38 minutes - declaring stable'],
            ['11:30', 'BNM', 'Acknowledged regulator email - awaiting written sitrep within 24 hrs'],
            ['13:28', 'IC', 'Incident closed. Wartime channel archived. Post-mortem 15-May 10:00'],
        ]),
        ('h1', '2. Customer impact'),
        ('bullets', [
            'Approximately 184,000 unique customers affected over 3h 46m window',
            'Failed login attempts: 412,000 (vs baseline 18,000 over same period)',
            'Failed transactions: 12,400 worth approx MYR 28.4M (auto-retried successfully post-recovery)',
            'Compensation policy: MYR 50 e-voucher to customers with > 3 failed attempts',
        ]),
        ('h1', '3. Notifications'),
        ('numbered', [
            'BNM Cyber Risk Unit - email sent 10:04 MYT - acknowledged 10:34',
            'PIDM - notified via watchlist channel 10:24',
            'Group Risk Council - emergency meeting 11:30',
            'External advisory issued via app + SMS 12:14',
        ]),
        ('h1', '4. Provisional root cause'),
        ('p', 'Stale ARP table on DC1 primary load balancer following firmware push 11-May-2026 23:14. Failed to detect via canary because canary traffic bypassed the affected interface. Vendor TAC ticket #VND-2026-44218 raised. Patch validation rollback procedure being revised by Network Engineering.'),
    ])


@reg('Incident_Timeline.docx')
def _b_incident_timeline(path):
    _docx(path, 'Incident Timeline Report - INC-2026-188 - Payments Queue Backlog', [
        ('p', 'Incident: INC-2026-188 | Severity: P2 | Detected: 11-May-2026 02:18 | Resolved: 11-May-2026 06:48 | IC: Hadar Caspit'),
        ('h1', '1. Summary'),
        ('p', 'IBG outbound queue accumulated a backlog of 18,400 transactions due to silent failure on the payment-orchestrator-east cluster. No customer-facing failure; settlement delayed by up to 4 hours within cut-off window.'),
        ('h1', '2. Detailed timeline'),
        ('table', [
            ['Time', 'Event', 'Source', 'Action'],
            ['02:18', 'Queue depth alert > 2,400 (warning)', 'Datadog', 'Auto-page on-call'],
            ['02:24', 'Queue depth crossed 4,800 (critical)', 'Datadog', 'PagerDuty SRE'],
            ['02:32', 'On-call Vijay acknowledged', 'PagerDuty', 'Bridge opened'],
            ['02:48', 'Identified orchestrator-east not heart-beating', 'Investigation', 'Failover prepared'],
            ['03:14', 'Failover to orchestrator-west', 'Manual', 'Backlog drain started'],
            ['03:42', 'Backlog reduced to 8,200', 'Datadog', 'Continued drain'],
            ['04:28', 'Backlog reduced to 1,400', 'Datadog', 'Within SLA'],
            ['05:18', 'Backlog cleared to 240 (steady state)', 'Datadog', 'Monitoring'],
            ['06:48', 'Sustained stable - incident closed', 'IC', 'Post-mortem scheduled'],
        ]),
        ('h1', '3. Customer impact'),
        ('bullets', [
            'Settlement delay 1-4 hours for 12,400 IBG transactions',
            'All transactions settled within next-day cut-off',
            'Zero failed transactions; zero reversal',
            'BNM threshold not breached (P2 - no mandatory notification)',
        ]),
        ('h1', '4. Action items'),
        ('numbered', [
            'Add heart-beat-failure alarm on orchestrator instances (currently only queue-depth alarm)',
            'Increase canary coverage from 10% to 25% to catch silent failures faster',
            'Quarterly DR drill includes orchestrator failover (currently annual)',
            'Document orchestrator failover runbook + add to wartime kit',
        ]),
    ])


@reg('IR_Group_Strategy.docx')
def _b_ir_group(path):
    _docx(path, 'Investor Relations Briefing Pack - Group Strategy Update Q1 FY2026', [
        ('p', 'Author: Daichi Maruyama (Head of IR) | Audience: Sell-side analysts, top-30 institutional holders | Distribution date: 15-May-2026 | Classification: For-disclosure (post Bursa filing)'),
        ('h1', '1. Strategic framework - ZAVA FORWARD 2030'),
        ('p', 'Our 5-pillar strategy remains unchanged. We are 22 months into the 60-month plan. Pillar 2 (Digital and Customer Experience) and Pillar 3 (Operating Model Transformation) are tracking ahead of plan; Pillar 4 (Energy Transition) is 6 weeks behind original schedule due to renewable project FID delay.'),
        ('h1', '2. Pillar-level FY2026 milestones'),
        ('table', [
            ['Pillar', 'FY2026 Milestone', 'Status'],
            ['1. Portfolio reshape', 'Divest non-core Cendana Aviation (12%)', 'SPA signed - completion Q3'],
            ['2. Digital and CX', 'Group AI Copilot rollout to 28,000 users', 'On-track (14,200 live)'],
            ['3. Operating Model', 'Shared Services consolidation - 6 GBS centres into 3', 'On-track'],
            ['4. Energy Transition', '500MW renewables operational, 50% reduction in Scope 1+2 GHG vs 2020', '6-week slip'],
            ['5. Talent and Culture', '40% female senior management by FY2027', 'On-track (33%)'],
        ]),
        ('h1', '3. Capital allocation'),
        ('bullets', [
            'Targeted dividend payout 50-60% of recurring PAT - unchanged',
            'Capex envelope MYR 4.2B for FY2026 - flat YoY',
            'Net debt / EBITDA target band 1.8x - 2.5x - currently 2.18x',
            'Share buyback authorisation MYR 600M - 38% deployed YTD'
        ]),
        ('h1', '4. FY2026 guidance reiteration'),
        ('p', 'Group revenue guidance MYR 44.0 - 46.5B (unchanged). Group EBITDA margin 15.0 - 16.0% (unchanged). PAT growth 8-12% YoY (unchanged). FX assumption: USD/MYR 4.42, JPY/MYR 0.029. Commodity: CPO MYR 4,100/t, Brent USD 82/bbl.'),
        ('h1', '5. Key risks to monitor'),
        ('bullets', [
            'OPR upcycle and impact on retail credit quality',
            'CPO price volatility and impact on Plantation',
            'Renewable project FID schedule',
            'Cyber-security maturity of mid-tier subsidiaries',
        ]),
    ])


@reg('IT_Incident_Timeline.docx')
def _b_it_incident(path):
    _docx(path, 'IT Incident Timeline - INC-2026-194 - SharePoint Online Slow Performance', [
        ('p', 'Incident: INC-2026-194 | Severity: P3 (productivity degradation) | Detected: 13-May-2026 09:18 | Closed: 13-May-2026 14:32 | IC: Vijay Subramaniam (Service Mgr - Modern Work)'),
        ('h1', '1. Symptoms'),
        ('bullets', [
            'File-open latency 8-14 seconds (baseline 1-2 seconds)',
            'Affecting approximately 28,000 KL HQ + KL Sentral users',
            'Mobile users on cellular networks unaffected',
            'Microsoft 365 service health dashboard: orange status',
        ]),
        ('h1', '2. Timeline'),
        ('table', [
            ['Time', 'Event', 'Owner'],
            ['09:18', 'User reports surge - 240 tickets in 12 min', 'Service Desk'],
            ['09:24', 'Synthetic monitor confirms latency spike', 'NOC'],
            ['09:32', 'IC engaged. Bridge opened. M365 health checked', 'Vijay'],
            ['09:48', 'Network team confirms uplink saturation on KL DC ExpressRoute primary', 'Network'],
            ['10:12', 'Confirmed root cause: ExpressRoute primary BGP flapping - secondary healthy', 'Network'],
            ['10:18', 'Failover to ExpressRoute secondary initiated', 'Network'],
            ['10:36', 'Failover complete. Latency normalising', 'Synthetic'],
            ['11:04', 'User-facing latency within 1-2s baseline', 'NOC'],
            ['14:32', 'Sustained stable for 3.5 hrs - incident closed', 'IC'],
        ]),
        ('h1', '3. Root cause'),
        ('p', 'Primary ExpressRoute circuit experienced BGP route flapping after carrier-side maintenance ran outside the announced window. Telco carrier RFO received. Mitigation: failover to secondary held for 7 days while primary is hard-failed and re-provisioned.'),
        ('h1', '4. Action items'),
        ('numbered', [
            'Schedule carrier review meeting - 19 May - regarding maintenance window discipline',
            'Stand up 3rd ExpressRoute circuit for true diversity by Q3 FY2026',
            'Update DR runbook to reflect new failover steps',
            'Communicate to user community on Yammer + Viva Engage',
        ]),
    ])


@reg('LEG_Draft_Contract.docx')
def _b_leg_draft(path):
    _docx(path, 'Master Services Agreement (Draft v2.3) - Zava Group and Apex Technology Bhd', [
        ('p', 'Parties: Zava Holdings Berhad (Customer) and Apex Technology Bhd (Supplier) | Document version: v2.3 | Drafter: Lim Mei Ling (Group Legal) | Date: 12-May-2026'),
        ('h1', '1. Definitions'),
        ('p', 'In this Agreement, unless the context otherwise requires: (a) "Services" means the managed cloud services described in Schedule 1; (b) "Effective Date" means the date of the last signature on the signature page; (c) "Term" means the period commencing on the Effective Date and ending 36 months thereafter, unless terminated earlier in accordance with Clause 14; (d) "Confidential Information" means any information disclosed by one Party to the other that is marked confidential or would reasonably be understood to be confidential.'),
        ('h1', '2. Services and SLAs'),
        ('bullets', [
            'Service Catalogue per Schedule 1 - 14 service items',
            'Availability SLA: 99.95% per calendar month - measured per Schedule 2',
            'Service credits cap at 10% of monthly recurring charge',
            'Right of termination if SLA breached in 3 consecutive months > 1% below target',
        ]),
        ('h1', '3. Commercial terms (commercially-sensitive)'),
        ('table', [
            ['Item', 'Value'],
            ['Initial Term', '36 months from Effective Date'],
            ['Annual Service Fee', 'MYR 12,400,000 indexed to CPI annually (cap 4%)'],
            ['Payment terms', '30 days from valid e-invoice'],
            ['Late-payment interest', '3-month KLIBOR + 1.5%'],
            ['Renewal', 'Automatic renewal 12 months unless 6 months prior notice'],
        ]),
        ('h1', '4. Confidentiality'),
        ('p', 'Each Party shall keep confidential all Confidential Information of the other Party for the duration of the Term and for 5 years after termination. Exceptions: information that is (a) publicly available without breach; (b) independently developed; (c) required to be disclosed by law or regulator.'),
        ('h1', '5. Data protection and security'),
        ('bullets', [
            'PDPA 2010 (Malaysia) compliance - Supplier acts as Data Processor',
            'Security: ISO 27001 certification required; annual re-cert evidence to be provided',
            'Data location: Malaysia primary, Singapore DR; written approval for any change',
            'Breach notification: within 24 hours of awareness to the Customer DPO',
        ]),
        ('h1', '6. Limitation of liability'),
        ('p', 'Each Party\'s aggregate liability under this Agreement (whether in contract, tort, or otherwise) shall be capped at 100% of the fees paid or payable in the 12 months immediately preceding the event giving rise to the claim. The cap shall not apply to: (a) fraud or wilful misconduct; (b) breach of confidentiality; (c) infringement of intellectual property; (d) breach of data protection obligations; (e) personal injury or death caused by negligence.'),
        ('h1', '7. Open items for negotiation'),
        ('bullets', [
            'Clause 6 cap - we are seeking 200% of trailing 12-month fees',
            'Clause 4 - mutual NDA period - Supplier wants 3 years; we want 5',
            'Schedule 2 service credits - Supplier wants cap at 5%; we are pushing for 10%',
        ]),
    ])


@reg('MEDIA_Campaign_Brief.docx')
def _b_media_brief(path):
    _docx(path, 'Media Campaign Brief - Zava Bank "Save Smarter" Q3 FY2026', [
        ('p', 'Campaign: Save Smarter | In-market: 1 July - 30 September 2026 | Budget: MYR 6.4M | Marketing Lead: Lim Mei Ling | Media agency: Apex Media (record retained)'),
        ('h1', '1. Business problem'),
        ('p', 'Zava Bank\'s retail savings deposit base grew 3.2% YoY in FY2025, lagging market growth of 5.8%. We are losing share to digital-only challengers among 28-38 year-old segment. Q3 FY2026 needs to deliver +MYR 2.4B incremental retail deposits and rebuild momentum for FY2027 plan.'),
        ('h1', '2. Audience'),
        ('bullets', [
            'Primary: 28-38, household income MYR 12-28K/month, urban, dual-income, raising young children',
            'Secondary: 38-48, household income MYR 24-48K/month, mortgage holders, saving for child education',
            'Both audiences are heavy YouTube + Facebook + TikTok consumers with > 4 hours daily digital time',
        ]),
        ('h1', '3. Key message'),
        ('p', 'Save Smarter is more than an interest rate. Earn up to 4.20% p.a. on your everyday savings with our new Smart Save account - no minimum balance, no lock-in, weekly interest pay-out, and a real human banker on WhatsApp seven days a week.'),
        ('h1', '4. Media channel mix and budget'),
        ('table', [
            ['Channel', 'Budget (MYR)', 'Reach Target', 'Frequency'],
            ['YouTube (preroll + bumpers)', '1,920,000', '6.8M unique', '4.2x'],
            ['Meta (Facebook + Instagram)', '1,536,000', '5.4M unique', '5.8x'],
            ['TikTok (creator + spark ads)', '768,000', '2.8M unique', '3.2x'],
            ['Programmatic display + native', '640,000', '8.4M unique', '2.6x'],
            ['CTV / OTT (Astro + Tonton + Apple TV)', '512,000', '1.8M unique', '3.8x'],
            ['Radio - Hot FM, Era FM, MIX FM', '384,000', '4.2M unique', '8.4x'],
            ['Search (Google + Bing)', '320,000', 'n/a', 'n/a'],
            ['OOH - Federal Highway + LDP', '320,000', '12.8M impressions', 'n/a'],
        ]),
        ('h1', '5. KPIs'),
        ('bullets', [
            'Account opening applications: 84,000 over 13 weeks',
            'Funded accounts: 52,000 (62% conversion)',
            'Average funded balance: MYR 12,400',
            'Total incremental deposits: MYR 644M',
            'Brand consideration lift: +6 points (Kantar tracking)',
        ]),
    ])


@reg('MKT_Brand_Strategy.docx')
def _b_mkt_brand(path):
    _docx(path, 'Group Brand Strategy - Zava Group FY2026-2028', [
        ('p', 'Strategy owner: Lim Mei Ling (Group Brand Director) | Sponsor: Sasha Ouellet (Group CoS) | Horizon: 3-year (FY2026-FY2028) | Status: Endorsed by ExCo 28-Apr-2026'),
        ('h1', '1. Brand context'),
        ('p', 'Zava Group operates 9 divisions under 14 customer-facing brands. Brand fragmentation has historically been a strength (channel specialisation) but is now creating friction in (a) digital go-to-market, (b) talent attraction, (c) Group cross-sell. Brand Stack 2030 is the first systemic re-architecture in 12 years.'),
        ('h1', '2. Brand architecture - target end-state'),
        ('table', [
            ['Tier', 'Brand', 'Role', 'Relationship to Zava'],
            ['Master', 'Zava', 'Endorser brand for the Group, employer brand, corporate', 'Master brand'],
            ['Strategic', 'Zava Bank', 'Banking division consumer brand', 'Endorsed - "Zava Bank, a Zava Group company"'],
            ['Strategic', 'Zava Insure', 'Insurance and Takaful brand', 'Endorsed'],
            ['Strategic', 'Zava Properties', 'Property dev + REIT', 'Endorsed'],
            ['Strategic', 'Zava Energy', 'Power, renewables, plantation under one umbrella', 'Endorsed (new in FY2026)'],
            ['Strategic', 'Zava Mart', 'Retail grocery + hypermart', 'Endorsed'],
            ['Tactical', 'Mawar, Cendana, Permata (FMCG)', 'Product brands', 'Silent - parent disclosed only on pack'],
        ]),
        ('h1', '3. Strategic intent'),
        ('bullets', [
            'Consolidate 14 customer-facing brands into 6 strategic + 8 tactical product brands over 3 years',
            'Build Zava as a Top-3 Most Loved Employer brand in ASEAN by FY2028',
            'Lift cross-sell index across Group divisions by 18% by FY2028',
            'Standardise brand expression - 1 master design system across digital + brick-and-mortar',
        ]),
        ('h1', '4. FY2026 priorities'),
        ('numbered', [
            'Brand identity refresh - new master logo, design system, voice and tone',
            'Launch Zava Energy as the consolidated brand (Apex Power + Hijau Renewables + Gemilang Plantations)',
            'Refresh Zava Bank visual identity - phase out legacy "Cendana Bank" sub-branding',
            'Employer brand campaign - global launch with Q4 priority on Malaysia + Indonesia',
        ]),
        ('h1', '5. Investment envelope'),
        ('p', 'Total 3-year brand investment: MYR 84M. FY2026: MYR 38M (identity refresh, system roll-out, Zava Energy launch, employer brand). FY2027: MYR 28M. FY2028: MYR 18M (steady-state).'),
    ])


@reg('OPS_BU1_Procedure.docx')
def _b_ops_proc(path):
    _docx(path, 'Operating Procedure BU1-OPS-014 - Daily Production Shift Handover', [
        ('p', 'Document: BU1-OPS-014 | Version: 4.2 | Effective: 01-May-2026 | Process owner: Ahmad bin Hassan (Plant Operations Manager) | Review cycle: Annual'),
        ('h1', '1. Purpose'),
        ('p', 'Defines the mandatory procedure for handover between Production shifts at BU1 manufacturing plants to ensure operational continuity, equipment status awareness, and safe transition between shift personnel.'),
        ('h1', '2. Scope'),
        ('bullets', [
            'All production lines L1 through L7 at BU1 Sg Buloh and BU1 Bayan Lepas',
            'Three shifts daily: A (06:00-14:00), B (14:00-22:00), C (22:00-06:00)',
            'Applies to Production Supervisors, Line Leaders, Maintenance Technicians, and QA Inspectors',
        ]),
        ('h1', '3. Procedure - step by step'),
        ('numbered', [
            'Outgoing supervisor completes Shift Handover Log (Form BU1-OPS-014-F1) within 30 minutes of shift end',
            'Log captures: line-by-line production status, OEE, downtime events, quality holds, safety observations, open work-orders, and pending decisions',
            'Outgoing supervisor walks the line with incoming supervisor for minimum 20 minutes prior to shift change',
            'Joint inspection of high-risk equipment (Line 5 hopper, Line 3 furnace, all CIP stations)',
            'Both supervisors sign the handover log; copy filed to plant SharePoint',
            'Incoming supervisor briefs incoming Line Leaders within first 15 minutes of new shift',
            'Any high-severity open items (P1/P2) escalated to Plant Manager via WhatsApp within 5 minutes',
        ]),
        ('h1', '4. Records and retention'),
        ('table', [
            ['Record', 'Owner', 'Retention period', 'Location'],
            ['Shift Handover Log', 'Production Supervisor', '3 years', 'SharePoint Plant Operations'],
            ['Pre-shift Toolbox Talk', 'HSE Officer', '5 years', 'EHS Portal'],
            ['Equipment Status Sheet', 'Maintenance Tech Lead', '2 years', 'CMMS'],
        ]),
        ('h1', '5. Performance indicators'),
        ('bullets', [
            'Handover compliance rate >= 98% monthly',
            'OEE delta between shifts <= 4 percentage points',
            'Zero unreported incidents discovered in subsequent shift',
        ]),
    ])


@reg('PHA_Change_Control_Package.docx')
def _b_pha_change(path):
    _docx(path, 'Change Control Package PHA-CC-2026-088 - Manufacturing Site Transfer (Product P-2418)', [
        ('p', 'Change ID: PHA-CC-2026-088 | Product: P-2418 (Active: Metformin HCl 500mg sustained-release tablets) | Initiating site: Apex Pharma Cyberjaya | Receiving site: Apex Pharma Klang'),
        ('h1', '1. Description of change'),
        ('p', 'Transfer of commercial manufacturing of P-2418 from Cyberjaya facility to Klang facility. Reason: Cyberjaya capacity constraint; Klang has spare capacity on Line K-202 (similar fluid-bed granulation, similar tablet press). NPRA submission scheduled August 2026.'),
        ('h1', '2. Regulatory classification'),
        ('table', [
            ['Authority', 'Variation Category', 'Filing Type', 'Target submission'],
            ['NPRA Malaysia', 'Major Variation (MaV-1)', 'Prior-approval supplement', '15-Aug-2026'],
            ['HSA Singapore (exports)', 'Major Variation - Manufacturing site', 'Prior-approval', '01-Sep-2026'],
            ['BPOM Indonesia (exports)', 'Major Variation', 'Prior-approval', '15-Sep-2026'],
        ]),
        ('h1', '3. Risk assessment summary'),
        ('p', 'Cross-functional risk assessment held 28-Apr-2026 chaired by QA. Highest risks: (a) granulation process robustness on K-202 (Risk 16/25 - High), (b) tablet hardness consistency (Risk 12/25 - Medium), (c) packaging line speed mismatch (Risk 8/25 - Medium). Mitigations defined for each.'),
        ('h1', '4. Process validation strategy'),
        ('numbered', [
            'Engineering batches: 3 batches at full commercial scale to confirm process parameters',
            'Process Performance Qualification (PPQ): 3 consecutive successful batches per ICH Q7',
            'Continued Process Verification (CPV) for 12 months post-launch with monthly review',
            'In-process critical quality attributes monitored: tablet weight, hardness, friability, disintegration, content uniformity',
        ]),
        ('h1', '5. Stability commitment'),
        ('p', 'Three PPQ batches enrolled into stability per ICH Q1A. Long-term 25C/60%RH and accelerated 40C/75%RH. Stability data filed as commitment with NPRA submission; final 12-month data available March 2027.'),
        ('h1', '6. Cross-functional approvals'),
        ('bullets', [
            'Quality Assurance Head - Lim Mei Ling - Approved 02-May-2026',
            'Manufacturing Director - Ahmad bin Hassan - Approved 03-May-2026',
            'Regulatory Affairs Director - Vijay Subramaniam - Approved 05-May-2026',
            'Site Quality Heads (Cyberjaya and Klang) - Approved 05-May-2026',
            'Qualified Person (QP) - Dr Siti Aishah - Approved subject to PPQ completion',
        ]),
    ])


@reg('RE_Export_Application_Form.docx')
def _b_re_export(path):
    _docx(path, 'Application for Export Permit - Renewable Energy Equipment (Form ST-RE-EX-2026)', [
        ('p', 'Applicant: Hijau Renewables Sdn Bhd | Permit Authority: Suruhanjaya Tenaga (Energy Commission of Malaysia) | Form: ST-RE-EX-2026 | Submitted: 12-May-2026 | Reference: ST-RE-EX-2026-04488'),
        ('h1', 'Section A - Applicant Details'),
        ('table', [
            ['Field', 'Value'],
            ['Legal Name', 'Hijau Renewables Sdn Bhd'],
            ['Company Number', '202301044218'],
            ['Registered Address', 'Level 18, Menara Hijau, Jalan Bangsar Utama, 59000 Kuala Lumpur'],
            ['Contact Person', 'Hadar Caspit - General Manager, Trade'],
            ['Email', 'hadar.caspit@hijaure.com.my'],
            ['Telephone', '+603-2284 0428'],
            ['ST License No.', 'ST-RE-2024-2218 (issued 14-Mar-2024, valid through 13-Mar-2029)'],
        ]),
        ('h1', 'Section B - Export Consignment Details'),
        ('table', [
            ['Field', 'Value'],
            ['Equipment Description', 'Solar PV inverters - string-inverter 100kW units (commercial-scale)'],
            ['HS Tariff Code', '8504.40.30'],
            ['Quantity', '48 units'],
            ['Unit Value (USD)', '8,400'],
            ['Total FOB Value (USD)', '403,200'],
            ['Destination Country', 'Vietnam'],
            ['Importing Party', 'Apex Energy Vietnam JSC'],
            ['Port of Loading', 'Port Klang (PKLG)'],
            ['Port of Discharge', 'Ho Chi Minh City (VNSGN)'],
            ['Vessel / Voyage', 'MV Pacific Dawn / V-2026-208'],
            ['ETD / ETA', '12-Jun-2026 / 19-Jun-2026'],
        ]),
        ('h1', 'Section C - End-Use Declaration'),
        ('p', 'The applicant hereby declares that the export consignment is for commercial sale to Apex Energy Vietnam JSC for installation at a 4.8MW rooftop solar farm at Long An Province. The equipment will not be re-exported, used for military purposes, or diverted to any sanctioned destination.'),
        ('h1', 'Section D - Supporting Documents Checklist'),
        ('bullets', [
            'Commercial Invoice (signed)',
            'Packing List',
            'Bill of Lading (draft)',
            'Certificate of Origin (preferential, under ATIGA)',
            'Test certificate - IEC 62109',
            'ST License (copy)',
        ]),
        ('h1', 'Section E - Declaration and Signature'),
        ('p', 'I, Hadar Caspit, in my capacity as General Manager - Trade for Hijau Renewables Sdn Bhd, declare that the information provided in this application is true and accurate to the best of my knowledge. Date: 12-May-2026. Signed: ______________'),
    ])


@reg('REG_Entity_Risk_Profile.docx')
def _b_reg_risk(path):
    _docx(path, 'Regulated Entity Risk Profile - Apex Banking Group Bhd', [
        ('p', 'Entity: Apex Banking Group Bhd | Regulator file: BNM-CLR-2026-44218 | Reviewing officer: Daichi Maruyama (BNM Banking Supervision Department) | Period: FY2025'),
        ('h1', '1. Entity overview'),
        ('table', [
            ['Field', 'Value'],
            ['Banking License', 'Commercial Banking License under FSA 2013'],
            ['Total Assets', 'MYR 412 billion'],
            ['Tier-1 Capital', 'MYR 38.4 billion'],
            ['CET-1 Ratio', '13.84%'],
            ['Risk-Weighted Assets', 'MYR 282 billion'],
            ['Branches (MY)', '218 nationwide'],
            ['Employees', '24,800'],
            ['Latest On-site Examination', 'Q4 FY2025 (concluded 18-Feb-2026)'],
        ]),
        ('h1', '2. Risk dimensions - heat map'),
        ('table', [
            ['Dimension', 'Inherent', 'Control', 'Residual', 'Trend'],
            ['Credit Risk', 'High', 'Strong', 'Medium', 'Stable'],
            ['Market Risk', 'Medium', 'Strong', 'Low', 'Stable'],
            ['Liquidity Risk', 'Medium', 'Strong', 'Low', 'Stable'],
            ['Operational Risk', 'High', 'Adequate', 'Medium', 'Increasing'],
            ['IT and Cyber Risk', 'High', 'Adequate', 'Medium', 'Increasing'],
            ['Conduct Risk', 'Medium', 'Strong', 'Low', 'Stable'],
            ['AML / CFT', 'High', 'Strong', 'Medium', 'Stable'],
            ['Climate and ESG Risk', 'Medium', 'Adequate', 'Medium', 'Increasing'],
        ]),
        ('h1', '3. Findings from latest on-site examination'),
        ('numbered', [
            'IT Operations - 3 high-severity observations on change management; rectification commitment 30-Jun-2026',
            'Credit Underwriting - SME portfolio concentration above internal threshold; remediation plan accepted',
            'AML KYC refresh - 8.4% backlog in high-risk customer reviews; remediation 31-Dec-2026',
            'Climate stress test - methodology gap on transition risk for Plantation lending; closure by Q4 FY2026',
        ]),
        ('h1', '4. Supervisory rating and stance'),
        ('p', 'Composite Supervisory Rating: 2 (Satisfactory) - unchanged from previous cycle. Supervisory stance for FY2026: monitor IT operational risk closely; require quarterly progress on AML KYC backlog; escalate to MOU if IT change management findings are not closed by 30-Sep-2026.'),
    ])


@reg('RSK_Current_RAS.docx')
def _b_rsk_ras(path):
    _docx(path, 'Group Risk Appetite Statement - FY2026 (Approved)', [
        ('p', 'Document: Group Risk Appetite Statement | Version: 5.0 | Approval: Board Risk Committee 28-Apr-2026; Board ratification 12-May-2026 | Effective: 01-May-2026 through 30-Apr-2027'),
        ('h1', '1. Strategic intent'),
        ('p', 'Zava Group will pursue diversified growth across financial services, real economy, and consumer divisions while maintaining a Group Common Equity Tier-1 ratio comfortably above regulatory minimum, investment-grade credit profile, and zero tolerance for financial crime breaches or fatal incidents.'),
        ('h1', '2. Quantitative limits'),
        ('table', [
            ['Risk Type', 'Metric', 'Green', 'Amber', 'Red'],
            ['Capital', 'Group CET-1 ratio', '>= 14.0%', '12.5% - 13.99%', '< 12.5%'],
            ['Capital', 'Net Debt / EBITDA', '<= 2.5x', '2.51x - 3.25x', '> 3.25x'],
            ['Liquidity', 'Group LCR', '>= 130%', '110% - 129%', '< 110%'],
            ['Credit', 'NPL ratio (Banking)', '<= 1.8%', '1.81% - 2.50%', '> 2.50%'],
            ['Market', 'VaR at 99% over book', '<= MYR 28M/day', '28-42M/day', '> 42M/day'],
            ['Operational', 'Operational loss events > MYR 5M', '<= 4 / yr', '5 - 8 / yr', '> 8 / yr'],
            ['Concentration', 'Single name exposure', '<= 8% Tier-1', '8 - 10%', '> 10%'],
            ['ESG', 'Scope 1+2 GHG vs glide path', 'on-track or better', 'within 5% slip', '> 5% slip'],
        ]),
        ('h1', '3. Qualitative statements'),
        ('bullets', [
            'Zero tolerance for breaches of AML, sanctions, and anti-bribery regulations',
            'Zero tolerance for fatalities and serious injuries',
            'Low tolerance for cyber breaches involving customer PII',
            'Low tolerance for material misstatement of financial reporting',
            'Moderate appetite for innovation and new-business launch risk',
        ]),
        ('h1', '4. Escalation triggers'),
        ('numbered', [
            'Any indicator entering Amber - escalation to Group Risk Committee within 10 working days',
            'Any indicator entering Red - immediate escalation to GCEO and Chair of Board Risk Committee within 24 hours',
            'Cumulative breach of 3+ amber indicators - Board notification within 5 working days',
        ]),
    ])


@reg('SHC_Audit_Findings_FY.docx')
def _b_shc_audit(path):
    _docx(path, 'External Audit Findings Report - Apex Hospital Chain Bhd FY2025', [
        ('p', 'Auditor: KPMG PLT | Engagement Partner: Vijay Subramaniam | Period: FYE 31 December 2025 | Date issued: 12 May 2026 | Recipients: Audit Committee + Group CFO'),
        ('h1', '1. Audit opinion'),
        ('p', 'In our opinion, the consolidated financial statements of Apex Hospital Chain Bhd and its subsidiaries give a true and fair view of the consolidated financial position as at 31 December 2025, and of their consolidated financial performance and cash flows for the year then ended, in accordance with Malaysian Financial Reporting Standards (MFRS) and Companies Act 2016. Unqualified opinion.'),
        ('h1', '2. Key audit matters'),
        ('bullets', [
            'KAM-1 - Revenue recognition for managed care contracts (significant judgement)',
            'KAM-2 - Impairment assessment of goodwill from FY2022 acquisition of Cendana Medical Centre',
            'KAM-3 - Provisions for medical malpractice claims (estimation uncertainty)',
        ]),
        ('h1', '3. Findings classified by severity'),
        ('table', [
            ['Severity', 'Findings', 'Owner', 'Target close'],
            ['Significant Deficiency', '2', 'CFO Office', '30-Jun-2026'],
            ['Control Deficiency', '7', 'Various Heads', '31-Dec-2026'],
            ['Recommendation', '14', 'Various Heads', 'FY2026 cycle'],
        ]),
        ('h1', '4. Significant deficiencies - detail'),
        ('numbered', [
            'SD-1 - Inadequate IT general controls in revenue cycle management application (RCM-2) at 4 hospitals - manual reconciliation does not fully compensate. Management response: ITGC remediation plan engaged with vendor; close by 30-Jun-2026.',
            'SD-2 - Delayed reconciliation of intercompany balances between holding and 3 subsidiaries (4-week delay vs 2-week policy). Management response: process redesign and SLA enforcement; close by 31-May-2026.',
        ]),
        ('h1', '5. Audit fees and non-audit services'),
        ('p', 'Audit fees for FY2025: MYR 2,840,000 (FY2024: MYR 2,620,000). Non-audit services pre-approved by Audit Committee: MYR 480,000 (transfer-pricing advisory and ICAP review). Non-audit fee ratio: 17%, within Bursa MMLR 15.18 threshold.'),
    ])


@reg('STR_Macro_Inputs.docx')
def _b_str_macro(path):
    _docx(path, 'Strategy Macro Input Pack - Quarterly Scan Q2 FY2026', [
        ('p', 'Owner: Mod Admin (Group Strategy Director) | Distribution: ExCo, Board Strategy Committee | Cut-off: 30 April 2026 | Cycle: Quarterly'),
        ('h1', '1. Global outlook in one paragraph'),
        ('p', 'Global growth tracking 2.7% in 2026 (IMF April update, -0.2 vs Jan). US labour market resilient but disinflation slow. ECB beginning rate-cut cycle. China property drag persisting but stimulus offsetting. Geopolitical risk premium on energy elevated.'),
        ('h1', '2. Malaysia outlook'),
        ('table', [
            ['Indicator', '2025 Actual', '2026 Forecast', 'Source'],
            ['Real GDP growth', '4.6%', '4.8% (range 4.5-5.2%)', 'BNM Annual Report'],
            ['Headline CPI', '2.1%', '2.4-3.0% (subsidy rationalisation)', 'BNM'],
            ['OPR', '3.00%', '3.00% (hold)', 'BNM Statement Mar-26'],
            ['USD/MYR average', '4.62', '4.42 (range 4.35-4.55)', 'Consensus'],
            ['Brent average (USD/bbl)', '82', '82 (range 75-88)', 'EIA STEO'],
            ['CPO average (MYR/t)', '4,180', '4,100 (range 3,800-4,400)', 'MPOB'],
        ]),
        ('h1', '3. Indonesia outlook'),
        ('bullets', [
            'GDP growth 5.1% in 2026 vs 5.0% in 2025 (Bappenas + consensus)',
            'BI Rate held at 6.00%; cuts expected H2 if Fed eases',
            'IDR/USD average 15,800; trend depreciation pressure',
            'Coal price USD 118/t (down 8% YoY)',
        ]),
        ('h1', '4. Sector-level signals'),
        ('numbered', [
            'Banking - NIM pressure persists; CET-1 strong; loan growth 4-5% range',
            'Property - residential transactions +2.8% YoY; commercial soft',
            'Plantation - CPO range-bound MYR 3,800-4,400; FFB yield improving',
            'Power - peninsular reserve margin tight in Q3; gas prices volatile',
            'Telco - ARPU stable; 5G capex steady; Pillar 2 tower-sharing accelerating',
            'Retail - K-shaped recovery; premium-mass growing, mass-market flat',
        ]),
        ('h1', '5. Strategic implications for ExCo discussion'),
        ('bullets', [
            'FX hedging - dollar strength assumption to be re-stressed at next FX Committee',
            'Plantation - revisit FY2026 CPO assumption (currently MYR 4,100)',
            'Banking - prepare for higher provision overlay if subsidy rationalisation impacts B40 credit quality',
            'Energy Transition - accelerate renewables FID despite project delay (window narrowing)',
        ]),
    ])


@reg('TELCO_Outage_Timeline.docx')
def _b_telco_outage(path):
    _docx(path, 'Telco Outage Timeline - INC-TELCO-2026-088 - LTE Core Disruption Klang Valley', [
        ('p', 'Incident: INC-TELCO-2026-088 | Severity: P1 | Service affected: 4G LTE Voice and Data | Region: Klang Valley + Selangor | Start: 10-May-2026 02:18 | End: 10-May-2026 04:48 | IC: Kumar Subramaniam'),
        ('h1', '1. Impact assessment'),
        ('bullets', [
            'Approximately 1.84 million subscribers affected with intermittent voice and data',
            '4G LTE data success rate dropped from 99.6% baseline to 38% during peak impact',
            'Voice call setup success rate dropped from 98.4% baseline to 42% during peak impact',
            '5G NSA subscribers also affected (anchor on LTE)',
        ]),
        ('h1', '2. Timeline'),
        ('table', [
            ['Time', 'Event', 'Owner'],
            ['02:18', 'MME-1 cluster CPU spike to 98% (baseline 38%)', 'NOC'],
            ['02:24', 'Subscriber attach failures begin - alarm storm', 'NOC'],
            ['02:30', 'P1 declared; bridge open; vendor TAC engaged', 'IC'],
            ['02:48', 'Root cause hypothesis: corrupted subscriber-profile cache after software patch 09-May 22:00', 'Vendor TAC'],
            ['03:12', 'MME-1 cluster restart authorised', 'IC + Network VP'],
            ['03:18', 'MME-1 restart in progress; traffic shifted to MME-2 and MME-3', 'Network'],
            ['03:42', 'MME-1 back online; cache rebuilt', 'Network'],
            ['04:14', 'Subscriber attach success rate climbing - now 92%', 'NOC'],
            ['04:48', 'Sustained 99%+ for 30 minutes - incident closed', 'IC'],
        ]),
        ('h1', '3. Customer-facing comms timeline'),
        ('numbered', [
            'App push notification at 02:48 acknowledging service impact',
            'SMS broadcast at 03:24 with apology + restoration ETA',
            'Twitter and Facebook status update every 30 minutes',
            'Service-restored notification 05:00',
            'Goodwill data top-up 5GB to all affected subscribers - issued 12 May',
        ]),
        ('h1', '4. Regulator notification'),
        ('p', 'MCMC notified at 02:54 (within mandatory 30-minute window for P1 incidents affecting > 1M subscribers per Standards SBKK-001). Written sitrep submitted by 10 May 18:00. Final root cause analysis report due to MCMC by 09-Jun-2026.'),
        ('h1', '5. Vendor accountability'),
        ('p', 'Patch validation gap on vendor side. SLA breach: USD 480,000 penalty being calculated per contract terms. Vendor required to re-baseline patch validation procedure within 30 days.'),
    ])


@reg('TH_CEO_Script_Outline.docx')
def _b_th_ceo_script(path):
    _docx(path, 'Town Hall - CEO Script Outline - All-Hands May 2026', [
        ('p', 'Event: Group Town Hall May 2026 | Date: 22 May 2026 | Time: 15:00-16:30 MYT | Format: Hybrid (KL HQ Auditorium + Teams Live) | Speaker: Group CEO | Script owner: Sasha Ouellet'),
        ('h1', '1. Opening - 4 minutes'),
        ('bullets', [
            'Welcome and thank you - acknowledge the 28,000 people on the call',
            'Lead-in: "It has been four months since our last town hall - lots to share"',
            'Frame the agenda: Q1 results, strategy update, people-first changes, Q&A',
        ]),
        ('h1', '2. Q1 FY2026 results - 12 minutes'),
        ('bullets', [
            'Headline: Q1 revenue MYR 11.2B (+5.8% YoY), EBITDA MYR 1.84B (+3.4% YoY)',
            'Wins: Banking strong NIM defence, Telco subscriber adds best in 5 years',
            'Challenges: Plantation CPO compression, Property soft pipeline, Manufacturing input cost spike',
            'Liquidity strong - MYR 1.84B cash + MYR 2.6B undrawn RCF - dividend on track',
        ]),
        ('h1', '3. Strategy update - 18 minutes'),
        ('numbered', [
            'ZAVA FORWARD 2030 - we are 22 months in. Pillars 2 and 3 ahead of plan',
            'Energy Transition - Cendana Power acquisition signing 28 May - transforms our renewables footprint',
            'Group AI Office - 14 agents in production, 38 in pipeline; we will train 28,000 employees on Copilot by end-FY2027',
            'Bandar Zava Phase 2 launch 15 June - this is the biggest property launch in our history',
            'Talent: we are at 33% female senior management - we will hit 40% by end-FY2027',
        ]),
        ('h1', '4. People announcements - 8 minutes'),
        ('bullets', [
            'New Group Chief People Officer joining 1 July - Datin Aishah Mokhtar (ex-Apex Bank)',
            'Career Marketplace - launching internal mobility platform 1 August',
            'Wellbeing budget - increasing wellness allowance by 18% effective 1 July',
            'New "Innovation Days" - 2 days per quarter for cross-team innovation projects',
        ]),
        ('h1', '5. Q&A - 25 minutes'),
        ('bullets', [
            'Pre-submitted top 5 questions answered first (from Yammer poll, ranked by votes)',
            'Live questions from auditorium + Teams chat - moderated by Sasha',
            'Anticipated topics: AI impact on jobs, dividend outlook, return-to-office stance',
        ]),
        ('h1', '6. Closing - 3 minutes'),
        ('p', 'Recap three takeaways: (1) Q1 is solid despite divisional pressure; (2) strategy is on track and accelerating in digital and energy transition; (3) people-first changes coming in H2. Thank you. Next town hall: 28 August 2026.'),
    ])

# ============================================================================
# PDF archetype builders (7)
# ============================================================================

@reg('CON_Contract_FIDIC_Yellow.pdf')
def _b_con_fidic(path):
    _pdf(path, 'FIDIC Yellow Book Contract - East Klang Valley Highway Extension Package 4', [
        ('p', 'Contract Reference: ZAVA-INFRA-2026-PKG4 | Form: FIDIC Yellow Book 2017 (Plant and Design-Build) | Employer: Zava Infrastructure Sdn Bhd | Contractor: Apex Engineering Bhd | Engineer: Asianova Consulting'),
        ('h1', 'Article 1 - Contract Particulars'),
        ('table', [
            ['Item', 'Detail'],
            ['Project', 'East Klang Valley Highway Extension Package 4 (8.4km)'],
            ['Contract Sum', 'MYR 482,400,000 (lump sum)'],
            ['Commencement Date', '01 July 2026'],
            ['Time for Completion', '36 months from Commencement Date'],
            ['Defects Notification Period', '24 months from Taking-Over'],
            ['Currency', 'Malaysian Ringgit (MYR)'],
            ['Governing Law', 'Laws of Malaysia'],
            ['Language', 'English (BM authoritative for statutory)'],
            ['Dispute Resolution', 'KLRCA arbitration seat KL, single arbitrator'],
        ]),
        ('h1', 'Article 2 - Scope and Specifications'),
        ('p', 'The Contractor shall execute the Works comprising design, supply, construction, testing, and commissioning of an 8.4 kilometre dual-carriageway highway extension including 2 grade-separated interchanges, 1 viaduct (240m), 14 culverts, drainage, road furniture, lighting, ITS systems, and toll plaza ancillaries. Reference design provided in Appendix A. Employer requirements in Appendix B.'),
        ('h1', 'Article 3 - Payment Terms'),
        ('table', [
            ['Milestone', 'Trigger', 'Value (% of Contract Sum)'],
            ['Mobilisation Advance', 'Performance bond + advance guarantee in place', '10%'],
            ['Monthly Interim Payment', 'Engineer-certified IPC', 'against value'],
            ['Substantial Completion', 'Taking-Over Certificate issued', '5%'],
            ['Final Account', 'DNP expiry + final certificate', '5%'],
        ]),
        ('h1', 'Article 4 - Performance Security'),
        ('p', 'Contractor shall provide a Performance Bond in the form of an On-Demand Bank Guarantee from an approved Malaysian licensed bank valued at 10% of the Contract Sum, issued within 28 days of Letter of Acceptance, valid until expiry of the Defects Notification Period plus 90 days.'),
        ('h1', 'Article 5 - Liquidated Damages'),
        ('p', 'In the event of late completion, the Contractor shall pay Liquidated Damages of MYR 240,000 per day of delay, capped at 10% of the Contract Sum, in accordance with FIDIC Sub-Clause 8.8.'),
        ('h1', 'Article 6 - Variation Mechanism'),
        ('bullets', [
            'Variations administered per FIDIC Yellow Book Clause 13',
            'Quotation submission within 28 days of Variation Order',
            'Engineer determination per Sub-Clause 3.7',
            'Maximum Variation Limit: 15% of Contract Sum without Employer Board re-approval',
        ]),
        ('pagebreak', None),
        ('h1', 'Article 7 - Special Conditions of Contract'),
        ('h2', '7.1 - Local Content Requirement'),
        ('p', 'Minimum 70% local content (by value) for materials, plant, and services. Quarterly reporting to the Engineer; non-compliance triggers a 0.25% Contract Sum deduction per percentage-point shortfall.'),
        ('h2', '7.2 - Bumiputera Sub-Contracting'),
        ('p', 'Minimum 30% (by value) of sub-contracted scope must be awarded to Bumiputera Class A/B/C contractors registered with CIDB. Compliance audited quarterly.'),
        ('h2', '7.3 - Health, Safety, and Environment'),
        ('p', 'Contractor shall maintain a CIDB Green Card site, OHSAS 18001 certification, and DOSH-compliant SH&E plan. Lost-Time Injury Frequency Rate target: <= 0.5 per 200,000 man-hours. Fatal incident triggers immediate suspension and root-cause investigation.'),
        ('h2', '7.4 - Sustainability'),
        ('p', 'Minimum 35% recycled content in pavement base course. Construction waste diversion >= 80%. Carbon footprint reporting per ISO 14064-1 quarterly. Tree replacement ratio 1:3 for site clearance.'),
        ('h1', 'Article 8 - Signatures'),
        ('table', [
            ['Party', 'Signatory', 'Position', 'Date'],
            ['Employer', '________________', 'Group MD - Zava Infrastructure', '___________'],
            ['Contractor', '________________', 'CEO - Apex Engineering Bhd', '___________'],
            ['Witness', '________________', 'Engineer Representative', '___________'],
        ]),
    ])


@reg('Employee_Handbook.pdf')
def _b_emp_handbook(path):
    _pdf(path, 'Zava Group Employee Handbook (Edition 12 - 2026)', [
        ('p', 'Approved by: Group HR Council, March 2026 | Effective: 1 April 2026 | Applies to: All Zava Group employees worldwide | Owner: Group Chief People Officer'),
        ('h1', '1. Welcome to Zava Group'),
        ('p', 'Welcome to Zava Group. As one of 28,000 colleagues across 9 divisions and 11 countries, you are part of a 46-year journey to build a diversified ASEAN champion. This handbook codifies the policies that protect you and that govern how we work together. It does not form part of your employment contract; for contractual terms refer to your Letter of Appointment.'),
        ('h1', '2. Our Values and Code of Conduct'),
        ('bullets', [
            'Integrity above all - we act with honesty even when no one is watching',
            'Bold and pragmatic - we make decisions and own outcomes',
            'One Zava - we win as a group, not as silos',
            'Customer at the centre - we earn trust through service',
            'Sustainability mindset - we make decisions a 30-year owner would make',
        ]),
        ('h1', '3. Working Hours and Leave'),
        ('table', [
            ['Leave Type', 'Entitlement', 'Carry Forward'],
            ['Annual Leave (5+ years service)', '21 working days', 'Up to 7 days'],
            ['Annual Leave (2-5 years)', '18 working days', 'Up to 5 days'],
            ['Annual Leave (< 2 years)', '14 working days', 'Up to 3 days'],
            ['Medical Leave (with cert)', '14 - 22 days (per MA1955)', 'No'],
            ['Maternity Leave', '98 consecutive days (per MA1955)', 'N/A'],
            ['Paternity Leave', '7 consecutive days', 'N/A'],
            ['Marriage Leave', '5 working days', 'N/A'],
            ['Compassionate Leave', '3 working days per occurrence', 'N/A'],
            ['Examination Leave', '5 working days per year', 'N/A'],
        ]),
        ('h1', '4. Performance and Compensation'),
        ('p', 'Performance is measured against a balanced scorecard with four dimensions: Financial, Customer, People, Operational Excellence. The annual cycle runs Jan-Dec with mid-year check-in in July. Annual increments are differentiated and tied to performance rating. Variable bonus is funded by company performance and distributed against individual performance.'),
        ('h1', '5. Remote and Flexible Work'),
        ('bullets', [
            'Hybrid working: minimum 3 days in office per week (Tue-Wed-Thu fixed core days)',
            'Flexible start: 07:30 to 10:00 with 8-hour working window',
            'Workcation: up to 4 weeks per year working from approved overseas location',
            'Compressed work week available - manager + HR business partner approval',
        ]),
        ('pagebreak', None),
        ('h1', '6. Wellbeing'),
        ('bullets', [
            'Annual medical screening - fully paid',
            'Mental health hotline - 24/7 confidential',
            'Wellness allowance - MYR 1,800 per year (gym, fitness, dental)',
            'Onsite clinic at KL HQ + KL Sentral (weekdays 09:00-17:00)',
            'EAP coverage - 12 free counselling sessions per year for self + family',
        ]),
        ('h1', '7. Code of Conduct'),
        ('h2', '7.1 Anti-Bribery and Corruption'),
        ('p', 'Zero tolerance. No employee may give or receive gifts, hospitality, or any form of advantage that may influence or be perceived to influence business decisions. Gifts received in good faith must be declared via the Gifts Register if their value exceeds MYR 250. Refer to the Group Anti-Bribery Policy (PL-COMP-014) for detail.'),
        ('h2', '7.2 Conflict of Interest'),
        ('p', 'Employees must declare actual, potential, or perceived conflicts of interest to their line manager and HR business partner. Outside directorships, employment, and material shareholdings (>5%) require prior written approval. Annual COI declaration is mandatory.'),
        ('h2', '7.3 Confidentiality'),
        ('p', 'Confidentiality obligations survive your employment. Disclosure of confidential information without authorisation is a serious misconduct subject to summary termination and legal action under contract.'),
        ('h1', '8. Grievance and Whistleblowing'),
        ('p', 'Speak Up Hotline (anonymous): 1-800-ZAVA-HELP or speakup.zava.com. Managed by an independent third-party (Deloitte). Investigations are conducted by Group Internal Investigation. Retaliation against whistleblowers is strictly prohibited and subject to disciplinary action.'),
        ('h1', '9. Departures and Final Settlements'),
        ('bullets', [
            'Notice periods per Letter of Appointment - typically 1-3 months',
            'Pro-rata bonus payable if employed for full performance year',
            'Annual Leave encashment - up to 7 days payable',
            'Exit interview with HR business partner mandatory',
            'Return of Group assets (laptop, phone, badge) within last working day',
        ]),
    ])


@reg('KYC_Corporate_Registry_Extract.pdf')
def _b_kyc_extract(path):
    _pdf(path, 'Suruhanjaya Syarikat Malaysia (SSM) - Company Information Extract', [
        ('p', 'This extract is generated by SSM e-Search facility on 12 May 2026 14:38 MYT. The information contained reflects the records of SSM as at the date and time of extraction.'),
        ('h1', '1. Company Particulars'),
        ('table', [
            ['Field', 'Detail'],
            ['Company Name', 'Apex Energy Holdings Sdn Bhd'],
            ['Old Name (if any)', 'Hijau Power Investments Sdn Bhd (changed 14-Mar-2022)'],
            ['Company Number', '202101044218'],
            ['Date of Incorporation', '18 February 2021'],
            ['Country of Incorporation', 'Malaysia'],
            ['Company Type', 'Private Limited Company (Sdn Bhd)'],
            ['Status', 'Existing'],
            ['Last Annual Return Filed', '30 April 2025'],
            ['Last Financial Statements Filed', 'FYE 31-Dec-2024'],
        ]),
        ('h1', '2. Registered Office'),
        ('p', 'Level 18, Menara Hijau, No. 3 Jalan Bangsar Utama 1, 59000 Kuala Lumpur, Wilayah Persekutuan'),
        ('h1', '3. Business Address'),
        ('p', 'Suite 14-2, Block A, Mid Valley City, Lingkaran Syed Putra, 59200 Kuala Lumpur'),
        ('h1', '4. Principal Activities'),
        ('bullets', [
            '(35101) Generation of electricity',
            '(35130) Distribution of electricity',
            '(64200) Activities of holding companies',
            '(70200) Management consultancy activities',
        ]),
        ('h1', '5. Authorised and Paid-up Capital'),
        ('table', [
            ['Field', 'Detail'],
            ['Authorised Share Capital', 'MYR 50,000,000.00 (50,000,000 ordinary shares of MYR 1.00 each)'],
            ['Issued and Paid-up Capital', 'MYR 24,800,000.00'],
            ['Number of Issued Shares', '24,800,000'],
        ]),
        ('h1', '6. Shareholders'),
        ('table', [
            ['Name', 'No. of Shares', '%'],
            ['Zava Holdings Berhad', '14,880,000', '60.00%'],
            ['Cendana Infrastructure Capital Sdn Bhd', '6,200,000', '25.00%'],
            ['EPF Investment (Nominee)', '2,480,000', '10.00%'],
            ['Public (Sundry)', '1,240,000', '5.00%'],
        ]),
        ('h1', '7. Directors'),
        ('table', [
            ['Name', 'NRIC / Passport', 'Date Appointed', 'Status'],
            ['Datuk Hadar Caspit', '740812-14-5318', '18-Feb-2021', 'Active'],
            ['Tan Sri Lim Cheng Yong', '610524-08-6217', '18-Feb-2021', 'Active'],
            ['Mod Admin', '780318-10-4421', '14-Mar-2022', 'Active'],
            ['Sasha Ouellet', 'A8412884 (USA)', '14-Mar-2022', 'Active'],
            ['Dr Siti Aishah Mohd Razali', '690914-12-5408', '01-Jul-2023', 'Active'],
        ]),
        ('h1', '8. Company Secretary'),
        ('p', 'Lim Mei Ling (MAICSA 7044218), Appointed 18-Feb-2021, Practicing Certificate valid through 30-Jun-2026'),
        ('h1', '9. Charges Registered'),
        ('table', [
            ['Charge No.', 'Type', 'Amount', 'Chargee', 'Date'],
            ['CH-2022-04428', 'Fixed and Floating', 'MYR 180,000,000', 'CIMB Bank Berhad', '14-Sep-2022'],
            ['CH-2024-08841', 'Fixed', 'MYR 60,000,000', 'Maybank Investment Bank', '22-Aug-2024'],
        ]),
        ('h1', '10. Certification'),
        ('p', 'This is a true extract of SSM records as at the time of extraction. Generated by user reference: BNK-CLR-2026-44218.'),
    ])


@reg('PLT_Audit_NCRs.pdf')
def _b_plt_ncrs(path):
    _pdf(path, 'Plant Audit Non-Conformance Report - Klang Manufacturing Facility', [
        ('p', 'Audit Date: 28-Apr to 02-May 2026 | Audit Type: Internal QMS Audit (ISO 9001:2015) | Lead Auditor: Vijay Subramaniam | Auditee: Klang Plant Management | Report Date: 12-May-2026'),
        ('h1', '1. Audit Scope and Methodology'),
        ('p', 'Five-day internal audit covering: production planning, manufacturing operations, quality control, calibration, document control, training records, customer complaint handling, and management review. Audit conducted in accordance with ISO 19011:2018 guidelines. Audit team: 4 auditors, total 80 audit hours.'),
        ('h1', '2. Summary of Findings'),
        ('table', [
            ['Category', 'Count', 'Examples'],
            ['Major Non-Conformance', '2', 'NCR-2026-014, NCR-2026-017'],
            ['Minor Non-Conformance', '8', 'NCR-2026-015 through NCR-2026-022'],
            ['Observation', '14', 'OBS-2026-018 through OBS-2026-031'],
            ['Opportunity for Improvement', '6', 'OFI-2026-008 through OFI-2026-013'],
        ]),
        ('h1', '3. Major Non-Conformance Details'),
        ('h2', 'NCR-2026-014 (Major) - Calibration of critical measurement equipment'),
        ('p', 'Standard: ISO 9001:2015 Clause 7.1.5.2. Three out of fourteen critical pressure gauges in Line K-201 had expired calibration certificates (expired 14-Jan-2026, used through 28-Apr-2026). Root cause: calibration schedule alert mechanism failed silently after CMMS upgrade in December 2025.'),
        ('p', 'Required Corrective Action: (a) immediate withdrawal and recalibration of affected gauges; (b) impact assessment on production from 14-Jan-2026 onward; (c) review of all calibration certs across all lines; (d) CMMS alert mechanism fix. Target close date: 30-May-2026.'),
        ('h2', 'NCR-2026-017 (Major) - Customer complaint handling timeliness'),
        ('p', 'Standard: ISO 9001:2015 Clause 9.1.2 + internal SOP. SLA for customer complaint acknowledgement is 24 hours and resolution within 14 days. Audit sample of 50 complaints from Q1 2026 showed: 18% breached acknowledgement SLA, 28% breached 14-day resolution SLA. Root cause: complaint routing tool migration in February 2026 caused mis-routing for 6 weeks.'),
        ('p', 'Required Corrective Action: (a) root-cause complaint routing fix verified; (b) overdue complaints triaged and closed by 31-May; (c) SLA reporting dashboard restored; (d) cross-check sample to be re-audited in October 2026. Target close date: 30-Jun-2026.'),
        ('pagebreak', None),
        ('h1', '4. Minor Non-Conformance Summary'),
        ('table', [
            ['NCR No.', 'Area', 'Brief Description', 'Owner', 'Target Close'],
            ['NCR-2026-015', 'Training Records', '4 operators missing annual refresher', 'HR Manager', '30-Jun-2026'],
            ['NCR-2026-016', 'Document Control', 'Obsolete SOP revision in use at WS-K204', 'QA Manager', '14-May-2026'],
            ['NCR-2026-018', 'Internal Audit', 'Q4-2025 audit overdue (planned Nov-25)', 'QA Manager', '31-May-2026'],
            ['NCR-2026-019', 'Mgmt Review', 'Mgmt review minutes incomplete - no action tracker', 'Plant Mgr', '14-May-2026'],
            ['NCR-2026-020', 'Risk Register', 'Top-10 risks not reviewed in last 9 months', 'Plant Mgr', '30-Jun-2026'],
            ['NCR-2026-021', 'Supplier Eval', '3 of 12 strategic suppliers no annual eval FY2025', 'Procurement', '30-Jun-2026'],
            ['NCR-2026-022', 'PPE Compliance', 'Hearing protection not consistently used in Line K-203', 'Shift Sup', '21-May-2026'],
        ]),
        ('h1', '5. Auditor Recommendation'),
        ('p', 'Despite the two Major NCRs noted, the QMS is generally operating effectively. We recommend continued certification subject to satisfactory closure of all Major NCRs by 30-Jun-2026. Re-audit of NCR-2026-014 and NCR-2026-017 to be performed in July 2026.'),
        ('h1', '6. Lead Auditor Signature'),
        ('p', 'Vijay Subramaniam (Lead Auditor) | Date: 12-May-2026 | Distribution: Plant GM, Group QA Director, Audit Committee'),
    ])


@reg('QA_Failed_Test_Result.pdf')
def _b_qa_failed(path):
    _pdf(path, 'QA Failed Test Result Report - Batch ZA-2026-2188-A (Mawar Body Wash 250ml)', [
        ('p', 'Batch: ZA-2026-2188-A | Product: Mawar Body Wash 250ml (SKU MW-BW-250-OR) | Production Date: 08-May-2026 | Lab Reference: QA-2026-4488 | Test Date: 09-May-2026'),
        ('h1', '1. Batch Information'),
        ('table', [
            ['Field', 'Value'],
            ['Batch Quantity', '14,840 bottles'],
            ['Production Line', 'Line BR-2 at Klang Plant'],
            ['Shift', 'Shift B (14:00-22:00)'],
            ['Shift Supervisor', 'Ahmad bin Hassan'],
            ['Hold Status', 'Quality Hold - DO NOT release'],
            ['Hold Location', 'Quarantine Bay Q-08'],
        ]),
        ('h1', '2. Failed Test Parameters'),
        ('table', [
            ['Parameter', 'Specification', 'Result', 'Status'],
            ['pH', '6.0 - 7.0', '5.6', 'FAIL (out of spec, below LSL)'],
            ['Viscosity (cP)', '8,000 - 12,000', '7,200', 'FAIL (out of spec, below LSL)'],
            ['Active surfactant', '12.0 - 14.0%', '12.8%', 'PASS'],
            ['Preservative', '0.30 - 0.50%', '0.42%', 'PASS'],
            ['Total Plate Count (CFU/g)', '< 100', '< 10', 'PASS'],
            ['Yeast and Mould (CFU/g)', '< 10', '< 1', 'PASS'],
            ['Appearance', 'Clear orange liquid', 'Clear pale orange', 'PASS'],
            ['Fragrance', 'Match standard', 'Match standard', 'PASS'],
        ]),
        ('h1', '3. Investigation Findings'),
        ('p', 'Initial investigation indicates pH and viscosity deviation likely due to: (a) Citric Acid charging weighed 2.4kg over target due to mis-calibrated load cell on Tank LT-203 (calibration check found offset of +1.8kg); (b) Carbomer addition rate slower than spec - 14 minute hold rather than 8-minute hold caused under-thickening. Both contributory factors confirmed via batch sheet and DCS trend review.'),
        ('h1', '4. Disposition Decision'),
        ('p', 'Batch CANNOT be released as-is. Two reprocessing options under evaluation: (a) Reprocess with Caustic Soda to lift pH + Carbomer slurry to recover viscosity - estimated cost MYR 4,800 + 6 hr line time; (b) Scrap and reissue raw material credit - estimated value loss MYR 18,400. Final decision authority: QA Manager Lim Mei Ling, awaiting customer impact review by 14-May-2026 13:00.'),
        ('h1', '5. Corrective Actions'),
        ('bullets', [
            'Load cell LT-203 recalibration completed 10-May-2026 06:30',
            'All load cells across the line scheduled for early recalibration by 16-May-2026',
            'Process timer SOP refresher session for all shift operators by 21-May-2026',
            'Failure mode added to FMEA review - update by 31-May-2026',
            'Root-cause investigation report due by 23-May-2026 (Owner: Vijay Subramaniam)',
        ]),
        ('h1', '6. Sign-off'),
        ('p', 'QA Analyst: Siti Aishah (10-May-2026) | QA Manager: Lim Mei Ling (pending) | Production Manager: Ahmad bin Hassan (pending) | Distribution: Plant Manager, Group QA Director, Quality Council'),
    ])


@reg('Regulator_Circular_2026.pdf')
def _b_reg_circular(path):
    _pdf(path, 'Bank Negara Malaysia - Policy Document', [
        ('p', 'Circular: BNM/RH/PD 036-44 | Title: Climate Risk Management and Scenario Analysis (Revised) | Issued by: Financial Stability Department | Issue date: 12 May 2026 | Effective date: 01 January 2027'),
        ('h1', 'Part A - Overview'),
        ('h2', 'A.1 Background'),
        ('p', 'This policy document supersedes BNM/RH/PD 029-7 issued in 2022. It sets out Bank Negara Malaysia\'s requirements on the management of climate-related risks (physical and transition risks) and on the conduct of climate scenario analysis by financial institutions, in alignment with the Joint Committee on Climate Change Phase 2 commitments.'),
        ('h2', 'A.2 Applicability'),
        ('bullets', [
            'Locally-incorporated commercial banks (S 213/214 of FSA 2013)',
            'Locally-incorporated Islamic banks (S 8/9 of IFSA 2013)',
            'Insurers and Takaful operators',
            'Investment banks',
            'Development Financial Institutions under DFIA 2002',
        ]),
        ('h1', 'Part B - Key Requirements'),
        ('h2', 'B.1 Governance'),
        ('p', 'Boards must establish board-level oversight of climate risks, with at least one director having documented climate-risk competency. Senior management must designate a Climate Risk Owner of Senior Vice President grade or above with direct reporting line to the CRO.'),
        ('h2', 'B.2 Risk Identification and Materiality Assessment'),
        ('p', 'Institutions must perform an annual climate-risk materiality assessment covering: (i) Physical risks - acute (flood, storm, heatwave) and chronic (sea-level rise); (ii) Transition risks - policy, technology, market, reputation, liability. The assessment must be approved by the Board Risk Committee and incorporated into the Group Risk Appetite Statement.'),
        ('h2', 'B.3 Scenario Analysis'),
        ('bullets', [
            'Minimum 3 scenarios annually: Disorderly Transition, Net Zero by 2050, Hot House World',
            'Time horizon: short (3 yr), medium (10 yr), long (30 yr)',
            'Methodology aligned with NGFS reference scenarios v4.2',
            'Material exposures: Plantation, Property, Power, Manufacturing portfolios',
            'Results presented to Board Risk Committee + filed with BNM within 30 days of approval',
        ]),
        ('h2', 'B.4 Disclosure'),
        ('p', 'Annual climate-related disclosure aligned with TCFD framework, integrated within the Annual Report. Disclosure must cover Governance, Strategy, Risk Management, Metrics and Targets. Scope 1, 2, and material Scope 3 financed emissions to be disclosed. Independent assurance required from FY2028 onward.'),
        ('pagebreak', None),
        ('h1', 'Part C - Implementation Timeline'),
        ('table', [
            ['Milestone', 'Target Date', 'Action'],
            ['Internal gap assessment complete', '30-Sep-2026', 'Institutions'],
            ['Updated policies and procedures', '31-Dec-2026', 'Institutions'],
            ['Effective date', '01-Jan-2027', 'BNM'],
            ['First materiality assessment under PD036', '31-Mar-2027', 'Institutions'],
            ['First scenario analysis under PD036', '30-Jun-2027', 'Institutions'],
            ['Independent assurance start', '01-Jan-2028', 'Institutions'],
        ]),
        ('h1', 'Part D - Supervisory Approach'),
        ('p', 'BNM will engage with institutions on a thematic-review basis in 2027 and a full-cycle review in 2028. Institutions that fail to meet the policy requirements will be subject to capital add-ons under Pillar 2 of the regulatory capital framework, ranging from 25 bps to 200 bps of risk-weighted assets, until remediation is satisfactory.'),
        ('h1', 'Part E - Contacts'),
        ('p', 'For queries on this policy document, contact: Financial Stability Department, Bank Negara Malaysia. Email: climate_supervision@bnm.gov.my. Telephone: +603-2698 8044. Office hours: 09:00 - 17:30 MYT, Monday to Friday.'),
        ('h1', 'Part F - Annexes'),
        ('bullets', [
            'Annex 1 - NGFS scenario specifications (v4.2)',
            'Annex 2 - TCFD-aligned disclosure template',
            'Annex 3 - Materiality assessment methodology',
            'Annex 4 - Financed emissions calculation guidelines',
            'Annex 5 - Frequently Asked Questions',
        ]),
    ])


@reg('Vendor_A_RFP_Response.pdf')
def _b_vendor_rfp(path):
    _pdf(path, 'Vendor RFP Response - Group ERP Implementation Programme', [
        ('p', 'RFP Reference: ZAVA-IT-RFP-2026-008 | Bidder: Apex Technology Bhd | Submitted: 12-May-2026 | Validity: 90 days | Bidder Contact: Vijay Subramaniam (Account Director)'),
        ('h1', '1. Executive Summary'),
        ('p', 'Apex Technology Bhd is pleased to submit our response to Zava Group\'s RFP for the Group ERP Implementation Programme. We bring 14 years of S/4HANA implementation experience across 22 ASEAN conglomerates, a dedicated 320-person regional practice, and demonstrated ability to deliver multi-country ERP transformations on time and on budget. Our proposed solution implements SAP S/4HANA 2024 across all 9 divisions in 28 months with a phased rollout aligned to your risk appetite.'),
        ('h1', '2. Proposed Solution Architecture'),
        ('bullets', [
            'Core: SAP S/4HANA 2024 (Public Cloud Edition where suitable, Private Cloud for regulated entities)',
            'Add-ons: SAP Analytics Cloud, SAP IBP, SAP Concur, SAP SuccessFactors',
            'Integration: SAP Integration Suite + Group EAI bus',
            'Localisation: Malaysia (e-Invoicing LHDN), Indonesia (e-Faktur), Singapore (IRAS), Vietnam (e-Invoicing GDT)',
            'Reporting: SAP Group Reporting with statutory + management views',
        ]),
        ('h1', '3. Implementation Approach'),
        ('table', [
            ['Wave', 'Scope', 'Go-Live', 'FTE (Apex)', 'FTE (Zava)'],
            ['Wave 1', 'Holding Co + Treasury + Group Procurement', 'Q1 2027', '32', '14'],
            ['Wave 2', 'Banking + Insurance Division', 'Q3 2027', '48', '22'],
            ['Wave 3', 'Property + Retail + Telco Division', 'Q1 2028', '52', '24'],
            ['Wave 4', 'Manufacturing + Plantation + Power Division', 'Q3 2028', '46', '20'],
            ['Wave 5', 'Hyper-care + Optimisation', 'Q4 2028', '18', '8'],
        ]),
        ('h1', '4. Commercial Proposal'),
        ('table', [
            ['Cost Category', 'Value (MYR)'],
            ['Software Licenses (3-year subscription)', '38,400,000'],
            ['Implementation Services (Apex)', '128,400,000'],
            ['Third-Party Services (Hardware partner)', '24,800,000'],
            ['Training and Change Management', '18,200,000'],
            ['Contingency (10% of services)', '17,140,000'],
            ['Total Programme Cost', '226,940,000'],
        ]),
        ('h1', '5. Team and Key Personnel'),
        ('bullets', [
            'Programme Director: Catherine Wong (12 years S/4HANA, 4 prior ASEAN conglomerate programmes)',
            'Solution Architect: Dr Vikram Patel (SAP Certified, 18 years)',
            'Change Management Lead: Lim Mei Ling (10 years, prior Zava Group engagement on HRIS)',
            'Quality Assurance: PwC (third-party QA per Group governance standard)',
            'Bench Strength: 320 SAP-certified consultants in regional practice',
        ]),
        ('h1', '6. References'),
        ('bullets', [
            'Apex Banking Group - S/4HANA Wave 1 Holding Co (FY2024) - on time and on budget',
            'Pacific Plantation Bhd - S/4HANA Group rollout (FY2022-FY2024) - 8 country, 14 entities',
            'Asianova Hospitality - S/4HANA + SAP Concur (FY2023-FY2025)',
        ]),
        ('h1', '7. Risks and Assumptions'),
        ('numbered', [
            'Assumed: Wave 1 design freeze by 30-Sep-2026; slip in design freeze pushes Wave 1 go-live by 1 month per month of slip',
            'Risk: Multi-country data residency and tax localisation complexity; mitigated by dedicated localisation team',
            'Risk: Change management bandwidth in Operating Divisions; mitigated by joint Apex+Zava CM team',
            'Assumed: Zava IT provides infrastructure and L1/L2 support post hyper-care; Apex provides L3 only',
        ]),
        ('h1', '8. Validity and Signature'),
        ('p', 'This proposal is valid for 90 days from submission. Apex Technology Bhd is committed to engaging on negotiations and to refining the proposal in dialogue with Zava Group. Signed: Hadar Caspit (Country MD - Apex Technology Bhd) | Date: 12-May-2026.'),
    ])

# ============================================================================
# PNG archetype builders (1)
# ============================================================================

@reg('DBS_Sample_E-Stmt.png')
def _b_dbs_estmt(path):
    W, H = 1100, 1400
    img = Image.new('RGB', (W, H), 'white')
    drw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype('arialbd.ttf', 28)
        font_h = ImageFont.truetype('arialbd.ttf', 14)
        font_b = ImageFont.truetype('arial.ttf', 12)
        font_s = ImageFont.truetype('arial.ttf', 10)
    except Exception:
        font_title = ImageFont.load_default()
        font_h = ImageFont.load_default()
        font_b = ImageFont.load_default()
        font_s = ImageFont.load_default()

    drw.rectangle([(0, 0), (W, 90)], fill='#A4051F')
    drw.text((40, 28), 'APEX BANK', fill='white', font=font_title)
    drw.text((W - 360, 36), 'CONSOLIDATED STATEMENT', fill='white', font=font_h)

    drw.text((40, 110), 'Account Holder', fill='#666666', font=font_s)
    drw.text((40, 125), 'MR HADAR CASPIT', fill='black', font=font_b)
    drw.text((40, 145), 'Level 18, Menara Hijau, 59000 Kuala Lumpur', fill='black', font=font_s)
    drw.text((40, 160), 'MALAYSIA', fill='black', font=font_s)

    drw.text((640, 110), 'Statement Period', fill='#666666', font=font_s)
    drw.text((640, 125), '01 APR 2026 - 30 APR 2026', fill='black', font=font_b)
    drw.text((640, 145), 'Statement Date: 02 MAY 2026', fill='black', font=font_s)
    drw.text((640, 160), 'Statement Ref: 2026-04-44218', fill='black', font=font_s)

    drw.rectangle([(40, 200), (W - 40, 320)], outline='#1F4E78', width=1)
    drw.rectangle([(40, 200), (W - 40, 230)], fill='#1F4E78')
    drw.text((50, 207), 'ACCOUNT SUMMARY', fill='white', font=font_h)
    drw.text((50, 248), 'Account Type', fill='#666666', font=font_s)
    drw.text((50, 268), 'Multi-Currency Savings Account', fill='black', font=font_b)
    drw.text((50, 290), 'Account Number: 7204-4218-4488', fill='black', font=font_s)
    drw.text((420, 248), 'Opening Balance (01 Apr 2026)', fill='#666666', font=font_s)
    drw.text((420, 268), 'SGD 84,218.42', fill='black', font=font_b)
    drw.text((420, 290), 'Currency: Singapore Dollar', fill='black', font=font_s)
    drw.text((780, 248), 'Closing Balance (30 Apr 2026)', fill='#666666', font=font_s)
    drw.text((780, 268), 'SGD 102,448.18', fill='#1F4E78', font=font_b)
    drw.text((780, 290), 'Net Change: +SGD 18,229.76', fill='black', font=font_s)

    drw.text((40, 350), 'TRANSACTION HISTORY', fill='#1F4E78', font=font_h)

    headers = [('Date', 40), ('Description', 130), ('Reference', 460),
               ('Debit (SGD)', 620), ('Credit (SGD)', 770), ('Balance (SGD)', 920)]
    drw.rectangle([(40, 380), (W - 40, 410)], fill='#1F4E78')
    for label, x in headers:
        drw.text((x, 388), label, fill='white', font=font_h)

    transactions = [
        ('02 Apr', 'Salary Credit - Zava Holdings', 'SAL-04488', '', '24,800.00', '109,018.42'),
        ('03 Apr', 'GIRO - Tenaga Nasional', 'BIL-08842', '848.40', '', '108,170.02'),
        ('05 Apr', 'Card Purchase - Aerodrome Hotels', 'POS-14488', '1,248.00', '', '106,922.02'),
        ('07 Apr', 'IBG Transfer - to Sasha O', 'IBG-21488', '4,400.00', '', '102,522.02'),
        ('09 Apr', 'Interest Credit', 'INT-04488', '', '128.42', '102,650.44'),
        ('11 Apr', 'Card Purchase - Sushi Kanzo', 'POS-24818', '284.00', '', '102,366.44'),
        ('12 Apr', 'FX Conversion MYR 8,400 to SGD', 'FX-44218', '', '2,418.40', '104,784.84'),
        ('14 Apr', 'GIRO - Indah Water', 'BIL-08418', '88.00', '', '104,696.84'),
        ('16 Apr', 'Card Purchase - Cendana Mart', 'POS-44248', '418.20', '', '104,278.64'),
        ('17 Apr', 'GIRO - Astro', 'BIL-14488', '184.40', '', '104,094.24'),
        ('19 Apr', 'ATM Withdrawal', 'ATM-04488', '800.00', '', '103,294.24'),
        ('21 Apr', 'Card Purchase - Eslite Spectrum', 'POS-04488', '418.18', '', '102,876.06'),
        ('22 Apr', 'IBG In - Sapphire Investments', 'IBG-44818', '', '3,200.00', '106,076.06'),
        ('24 Apr', 'Card Purchase - Mawar Salon', 'POS-44218', '184.20', '', '105,891.86'),
        ('26 Apr', 'Standing Instr - Insurance Premium', 'STD-04488', '1,288.00', '', '104,603.86'),
        ('27 Apr', 'Card Purchase - Bumi Konsortium F&B', 'POS-44488', '484.40', '', '104,119.46'),
        ('28 Apr', 'IBG Transfer - to Daichi M', 'IBG-44848', '2,400.00', '', '101,719.46'),
        ('29 Apr', 'Card Purchase - Tropika Spa', 'POS-44918', '212.40', '', '101,507.06'),
        ('30 Apr', 'Service Charge', 'CHG-04488', '8.50', '', '101,498.56'),
        ('30 Apr', 'Interest Credit - End of Cycle', 'INT-44218', '', '949.62', '102,448.18'),
    ]
    y = 420
    stripe = False
    for row in transactions:
        if stripe:
            drw.rectangle([(40, y), (W - 40, y + 26)], fill='#F2F2F2')
        for (val, (_, x)) in zip(row, headers):
            drw.text((x, y + 8), val, fill='black', font=font_b)
        y += 26
        stripe = not stripe

    drw.rectangle([(40, y + 12), (W - 40, y + 64)], outline='#1F4E78', width=1)
    drw.text((50, y + 22), 'Period Totals', fill='#1F4E78', font=font_h)
    drw.text((50, y + 42), 'Total Debits: SGD 13,266.68', fill='black', font=font_s)
    drw.text((420, y + 42), 'Total Credits: SGD 31,496.44', fill='black', font=font_s)
    drw.text((780, y + 42), 'Interest Earned: SGD 1,078.04', fill='black', font=font_s)

    drw.text((40, H - 60), 'Apex Bank Berhad (Company No. 198101044218) | Customer Service: 1-800-APEX-BNK | apexbank.com.my', fill='#666666', font=font_s)
    drw.text((40, H - 42), 'This statement is generated electronically and does not require a signature.', fill='#666666', font=font_s)
    drw.text((40, H - 24), 'Please report discrepancies within 14 days of statement date.', fill='#666666', font=font_s)

    img.save(str(path), 'PNG')
# ============================================================================
# Main loop
# ============================================================================
if __name__ == '__main__':
    OUT.mkdir(exist_ok=True)
    print(f'Will write {len(FILE_HANDLERS)} files to {OUT}')
    for fname, builder in FILE_HANDLERS.items():
        try:
            builder(OUT / fname)
            print(f'  ok   {fname}')
        except Exception as e:
            print(f'  FAIL {fname}: {e}')
    print(f'Done. {len(FILE_HANDLERS)} files processed.')
