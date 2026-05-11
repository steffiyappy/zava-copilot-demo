"""Generate Notebook source data files for all 13 archetypes (65 files total).

Each archetype's files are placed under:
    notebook-files/<archetype-id>/<filename>

Content is realistic enough for the Notebook prompts to produce sensible
grounded answers while staying generic (Zava Conglomerate context) so the
files work across all entries that share the archetype.
"""
from __future__ import annotations

import os
from datetime import datetime, timedelta
from pathlib import Path

# Document libs
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

REPO = Path(__file__).resolve().parent
OUT = REPO / 'notebook-files'
OUT.mkdir(exist_ok=True)

# ─── PDF helpers ───
_styles = getSampleStyleSheet()
_styleN = _styles['Normal']
_styleN.fontSize = 10
_styleN.leading = 13
_styleH1 = ParagraphStyle('H1', parent=_styles['Heading1'], fontSize=16, textColor=colors.HexColor('#0F4C81'), spaceAfter=10)
_styleH2 = ParagraphStyle('H2', parent=_styles['Heading2'], fontSize=12, textColor=colors.HexColor('#1F2937'), spaceAfter=6, spaceBefore=10)
_styleSmall = ParagraphStyle('Small', parent=_styleN, fontSize=8, textColor=colors.HexColor('#6B7280'))
_styleQuote = ParagraphStyle('Quote', parent=_styleN, leftIndent=12, rightIndent=12, fontSize=9, textColor=colors.HexColor('#374151'), backColor=colors.HexColor('#F3F4F6'))

def _pdf(path: Path, title: str, blocks):
    """blocks = list of ('h1'|'h2'|'p'|'small'|'quote'|'spacer'|'table', value)"""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm,
                            title=title)
    flow = []
    for kind, val in blocks:
        if kind == 'h1':
            flow.append(Paragraph(val, _styleH1))
        elif kind == 'h2':
            flow.append(Paragraph(val, _styleH2))
        elif kind == 'p':
            flow.append(Paragraph(val, _styleN))
            flow.append(Spacer(1, 4))
        elif kind == 'small':
            flow.append(Paragraph(val, _styleSmall))
        elif kind == 'quote':
            flow.append(Paragraph(val, _styleQuote))
            flow.append(Spacer(1, 4))
        elif kind == 'spacer':
            flow.append(Spacer(1, val or 8))
        elif kind == 'pagebreak':
            flow.append(PageBreak())
        elif kind == 'table':
            t = Table(val, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F4C81')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 8),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F9FAFB'), colors.white]),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D1D5DB')),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 4),
                ('RIGHTPADDING', (0,0), (-1,-1), 4),
                ('TOPPADDING', (0,0), (-1,-1), 3),
                ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ]))
            flow.append(t)
            flow.append(Spacer(1, 6))
    doc.build(flow)

# ─── DOCX helpers ───
def _docx_init(title: str, subtitle: str = ''):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    h = d.add_heading(title, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    if subtitle:
        p = d.add_paragraph(subtitle)
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    return d

def _docx_h2(d, text):
    h = d.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def _docx_p(d, text, bold=False, italic=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def _docx_bullets(d, items):
    for it in items:
        d.add_paragraph(it, style='List Bullet')

def _docx_table(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)

# ─── XLSX helpers ───
_HEAD_FONT = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
_HEAD_FILL = PatternFill('solid', fgColor='0F4C81')
_HEAD_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)
_BORDER = Border(left=Side(style='thin', color='D1D5DB'),
                 right=Side(style='thin', color='D1D5DB'),
                 top=Side(style='thin', color='D1D5DB'),
                 bottom=Side(style='thin', color='D1D5DB'))
_RAG_RED  = PatternFill('solid', fgColor='FEE2E2')
_RAG_AMB  = PatternFill('solid', fgColor='FEF3C7')
_RAG_GRN  = PatternFill('solid', fgColor='D1FAE5')

def _xl_write_sheet(ws, headers, rows, widths=None, rag_col=None, rag_map=None):
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = _HEAD_FONT; c.fill = _HEAD_FILL; c.alignment = _HEAD_ALIGN; c.border = _BORDER
    for ri, row in enumerate(rows, 2):
        for ci, v in enumerate(row, 1):
            c = ws.cell(row=ri, column=ci, value=v)
            c.border = _BORDER
            c.alignment = Alignment(vertical='top', wrap_text=True)
            if rag_col is not None and ci == rag_col and rag_map:
                rag = rag_map(row)
                if rag == 'R': c.fill = _RAG_RED
                elif rag == 'A': c.fill = _RAG_AMB
                elif rag == 'G': c.fill = _RAG_GRN
    ws.freeze_panes = 'A2'
    if widths:
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w
    ws.auto_filter.ref = ws.dimensions

def _xl_new(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    return wb

def _xl_save(wb, path: Path):
    wb.save(str(path))

# ─── Archetype 1: RFP procurement ───
def gen_rfp_procurement():
    arch_dir = OUT / 'nb-rfp-procurement'
    arch_dir.mkdir(parents=True, exist_ok=True)

    vendors = [
        ('Vendor A', 'Apex Cloud Solutions',  'Kuala Lumpur', 'MYR', 4_850_000, 4.6, 99.95, '24x7 follow-the-sun', 'AWS, Azure', 'ISO 27001, SOC 2 Type II'),
        ('Vendor B', 'Pacific Digital Group', 'Singapore',    'MYR', 5_120_000, 4.2, 99.90, '12x5 + on-call',     'Azure',       'ISO 27001'),
        ('Vendor C', 'Asianova Systems',      'Jakarta',      'MYR', 4_320_000, 3.8, 99.50, '8x5 business hours', 'GCP, AWS',    'ISO 27001'),
        ('Vendor D', 'Northwind Integrators', 'Manila',       'MYR', 5_680_000, 4.8, 99.99, '24x7 follow-the-sun', 'AWS, Azure, GCP', 'ISO 27001, SOC 2, FedRAMP'),
    ]

    # Vendor PDFs
    for vid, vname, vloc, vccy, vprice, vrating, vsla, vsupport, vcloud, vcerts in vendors:
        blocks = [
            ('h1', f'{vname} — RFP Response'),
            ('small', f'Submission for: Zava Conglomerate — Group-Wide Cloud Modernisation RFP-2026-04'),
            ('small', f'Submitted: 2026-04-22 · Vendor HQ: {vloc} · Reference: ZAVA/{vid.replace(" ","")}/0426'),
            ('spacer', 8),
            ('h2', '1. Executive Summary'),
            ('p', f'{vname} is pleased to respond to Zava Conglomerate\'s Group-Wide Cloud Modernisation RFP. We propose a 36-month engagement covering migration assessment, landing-zone build, application re-platforming for 24 priority workloads, and run-state managed services across {vcloud}.'),
            ('p', f'Our delivery is anchored in {vsla}% availability commitment, {vsupport} support model, and certifications: {vcerts}. Total Contract Value (TCV) over 36 months: {vccy} {vprice:,.0f} (inclusive of professional services, licence pass-through, and run-state).'),
            ('h2', '2. Technical Approach'),
            ('p', f'2.1 Landing-Zone — We will deliver a hub-and-spoke landing zone on {vcloud} aligned to the {vcloud.split(",")[0].strip()} Well-Architected Framework. Identity federates through Zava\'s Entra ID tenant. Networking uses dedicated ExpressRoute / Direct Connect circuits with 99.95% SLA.'),
            ('p', f'2.2 Migration Wave Plan — 24 workloads grouped into 4 waves of 6 each, sequenced by business criticality. Wave 1 (low-risk): file shares, dev/test, internal portals. Wave 2: data warehouse, analytics. Wave 3: ERP + payroll. Wave 4: customer-facing transactional systems.'),
            ('p', f'2.3 Re-platforming Strategy — 6 lift-and-shift, 10 re-platform (containerise), 6 refactor (cloud-native), 2 retire. Each workload has its own runbook, rollback plan, and DR pair.'),
            ('p', f'2.4 Security — Zero-Trust model, MFA on all admin accounts, JIT access via PIM, KMS-managed keys (BYOK option), CSPM tooling, and continuous compliance scanning against {vcerts}.'),
            ('pagebreak', None),
            ('h2', '3. Commercial'),
            ('table', [
                ['Line Item', 'Year 1', 'Year 2', 'Year 3', 'Total (MYR)'],
                ['Professional services (migration)', f'{vprice*0.45/1000:.0f}k', f'{vprice*0.15/1000:.0f}k', '—', f'{vprice*0.60/1000:.0f}k'],
                ['Run-state managed services',         f'{vprice*0.05/1000:.0f}k', f'{vprice*0.15/1000:.0f}k', f'{vprice*0.20/1000:.0f}k', f'{vprice*0.40/1000:.0f}k'],
                ['Licence pass-through (capped)',      'as consumed', 'as consumed', 'as consumed', 'metered'],
                ['Total commitment',                   f'{vprice*0.50/1000:.0f}k', f'{vprice*0.30/1000:.0f}k', f'{vprice*0.20/1000:.0f}k', f'{vprice/1000:.0f}k'],
            ]),
            ('p', f'3.1 Payment Terms — Net 45 from invoice. 10% retention on Year 1 professional services, released on go-live sign-off per wave.'),
            ('p', f'3.2 Pricing Validity — 90 days from RFP submission date.'),
            ('p', f'3.3 Price Escalation — CPI+1% annual cap on run-state from Year 2 onwards. Professional services rate card fixed for Year 1.'),
            ('h2', '4. SLA & Support'),
            ('p', f'4.1 Availability SLA: {vsla}% measured per workload per calendar month.'),
            ('p', f'4.2 Support model: {vsupport}. P1 response 15 min, P2 30 min, P3 4 hours.'),
            ('p', f'4.3 Service credits: 10% of monthly run-state fees per 0.1% below SLA, capped at 25%.'),
            ('h2', '5. Compliance & Certifications'),
            ('p', f'{vname} holds the following certifications relevant to this engagement: {vcerts}.'),
            ('p', f'Data residency: all customer data remains within Malaysia / ASEAN region as required by PDPA 2010. Government and regulated-workload data may be ring-fenced to a sovereign-cloud landing zone on request.'),
            ('h2', '6. References'),
            ('p', f'6.1 Public-listed financial-services conglomerate — 18-month engagement, 32 workloads migrated, completed 2025.'),
            ('p', f'6.2 ASEAN-headquartered energy group — landing-zone + 14 workloads, completed 2024.'),
            ('p', f'6.3 Government-linked transportation operator — sovereign-cloud landing-zone, ongoing.'),
            ('p', 'Detailed reference contact details available on request under NDA.'),
            ('h2', '7. Assumptions, Exclusions & Open Items'),
            ('p', '7.1 Assumption — Zava provides existing application architecture documentation within 14 days of award.'),
            ('p', '7.2 Exclusion — Application code remediation for non-cloud-ready workloads is out of scope and would be a change order.'),
            ('p', f'7.3 Open item — Sustainability / Scope-3 emissions reporting commitment NOT included in this response. {vname} can provide a separate proposal on request.'),
        ]
        _pdf(arch_dir / f'{vid.replace(" ", "_")}_RFP_Response.pdf', f'{vname} RFP Response', blocks)

    # Scoring rubric XLSX
    wb = _xl_new(arch_dir / 'Scoring_Rubric_and_Requirements.xlsx')
    ws = wb.active; ws.title = 'Rubric'
    _xl_write_sheet(ws,
        ['Criterion ID', 'Category', 'Requirement', 'Weight %', 'Mandatory?', 'Scoring Scale (1-5)', 'Evidence Required'],
        [
            ['C-01','Technical','Landing-zone aligned to Well-Architected Framework on AWS or Azure', 12, 'Y', '1=absent .. 5=fully described with diagrams', 'Architecture diagram + framework mapping'],
            ['C-02','Technical','Migration wave plan with rollback per wave',                          8,  'Y', '1=high-level only .. 5=runbook per workload',  'Wave plan + sample runbook'],
            ['C-03','Technical','Re-platforming strategy specifies disposition per workload',         6,  'Y', '1=generic .. 5=disposition + rationale per app','Workload disposition table'],
            ['C-04','Security','Zero-Trust model with MFA, JIT, KMS, CSPM',                            10, 'Y', '1=partial .. 5=all 4 controls + tooling listed','Security architecture'],
            ['C-05','Security','Data residency commitment for ASEAN/Malaysia',                        6,  'Y', '1=silent .. 5=explicit + sovereign option',    'Data residency clause'],
            ['C-06','Commercial','TCV ≤ MYR 5.5M across 36 months',                                    12, 'Y', '1=>6M .. 5=<=4.5M',                            'Pricing schedule'],
            ['C-07','Commercial','Run-state escalation capped at CPI+1%',                              5,  'N', '1=>CPI+3% .. 5=CPI+0%',                       'Pricing clause'],
            ['C-08','Commercial','Payment terms net 45 or longer',                                     4,  'N', '1=net 15 .. 5=net 60+',                       'Payment terms clause'],
            ['C-09','SLA','Workload availability ≥ 99.9%',                                            8,  'Y', '1=<99.5 .. 5=>=99.99',                        'SLA schedule'],
            ['C-10','SLA','24x7 support model',                                                       5,  'N', '1=8x5 .. 5=24x7 follow-the-sun',              'Support model schedule'],
            ['C-11','Compliance','ISO 27001 + SOC 2 certifications',                                  7,  'Y', '1=neither .. 5=both + FedRAMP/equivalent',    'Certificate copies'],
            ['C-12','Compliance','PDPA 2010 (MY) compliance attested',                                 4,  'Y', '1=silent .. 5=explicit attestation + DPO contact','Compliance attestation'],
            ['C-13','References','3+ relevant ASEAN reference engagements',                            6,  'N', '1=0 references .. 5=3+ with public-listed',  'Reference list (NDA OK)'],
            ['C-14','Sustainability','Scope-3 emissions reporting capability',                         4,  'N', '1=silent .. 5=quarterly Scope-3 reporting included','Sustainability schedule'],
            ['C-15','Risk','All assumptions/exclusions explicit',                                     3,  'N', '1=hidden .. 5=fully documented + priced',     'Assumptions section'],
        ],
        widths=[12,14,55,9,11,40,30]
    )
    ws.cell(row=18, column=1, value='Total Weight').font = Font(bold=True)
    ws.cell(row=18, column=4, value=100).font = Font(bold=True)

    ws2 = wb.create_sheet('Scoring Method')
    _xl_write_sheet(ws2,
        ['Step', 'Description'],
        [
            ['1','Evaluators score each criterion 1-5 independently using the scale in column F.'],
            ['2','Where the response is silent on a MANDATORY criterion, score = 1 and flag for clarification.'],
            ['3','Weighted score = score x (weight %); maximum possible = 500.'],
            ['4','Shortlist threshold = 350 weighted points + no mandatory criterion below 3.'],
            ['5','Discrepancies > 1 point between evaluators reconciled in calibration session.'],
            ['6','BAFO requested from the top 2 weighted scorers before final award.'],
        ],
        widths=[6, 110]
    )

    ws3 = wb.create_sheet('Evaluation Team')
    _xl_write_sheet(ws3,
        ['Name', 'Role', 'Function', 'Vote Weight'],
        [
            ['Hadar Caspit',   'Group CFO',                 'Commercial / Risk',  '20%'],
            ['Sasha Ouellet',  'Group Chief of Staff',      'Programme governance', '15%'],
            ['Daichi Maruyama','Head of Investor Relations','Sustainability / reporting','10%'],
            ['Mod Admin',      'Group Strategy Director',   'Strategic fit',      '15%'],
            ['CIO (named)',    'Group CIO',                 'Technical lead',     '25%'],
            ['CISO (named)',   'Group CISO',                'Security lead',      '15%'],
        ],
        widths=[20, 28, 28, 14]
    )
    _xl_save(wb, arch_dir / 'Scoring_Rubric_and_Requirements.xlsx')

# ─── Archetype 2: Clinical case ───
def gen_clinical_case():
    d = OUT / 'nb-clinical-case'
    d.mkdir(parents=True, exist_ok=True)

    # Patient_History_Summary.docx
    doc = _docx_init('Patient History Summary', 'Apex Specialist Centre — Confidential · MRN: APX-2026-04421')
    _docx_h2(doc, '1. Demographics')
    _docx_table(doc,
        ['Field','Value'],
        [['Name (pseudonymised)','Patient X'],['Age','58 years'],['Sex','Female'],
         ['Occupation','Retired teacher'],['Residence','Kuala Lumpur'],['BMI','29.4 (overweight)'],
         ['Smoking history','Never'],['Alcohol','Occasional'],['MRN','APX-2026-04421'],
         ['Primary care GP','Dr. Tan, Apex Family Clinic']])
    _docx_h2(doc, '2. Presenting Complaint')
    _docx_p(doc, 'Patient X presented to the outpatient clinic on 2026-04-03 with a 6-week history of progressive lower-back pain radiating to the left leg, associated with intermittent numbness in the lateral foot. Pain is worse on standing and walking, partially relieved by rest. No bowel or bladder symptoms. No saddle anaesthesia.')
    _docx_h2(doc, '3. History of Presenting Illness')
    _docx_p(doc, 'Pain commenced 2026-02-18, initially attributed to gardening. Initial GP review on 2026-02-25 prescribed paracetamol 1g QDS and naproxen 500mg BD for 2 weeks. Pain improved 30% but did not resolve. Re-presented 2026-03-15 with worsening leg pain. GP added pregabalin 75mg BD and ordered MRI lumbar spine.')
    _docx_h2(doc, '4. Past Medical History')
    _docx_bullets(doc, [
        'Type 2 Diabetes Mellitus (diagnosed 2018; HbA1c 7.4% on metformin 1g BD)',
        'Hypertension (diagnosed 2015; controlled on perindopril 8mg OD)',
        'Hyperlipidaemia (on atorvastatin 20mg ON)',
        'Right total knee replacement 2022 — uncomplicated recovery',
        'No previous spinal surgery',
        'No known drug allergies',
    ])
    _docx_h2(doc, '5. Medications (Current)')
    _docx_table(doc,
        ['Drug','Dose','Frequency','Indication'],
        [['Metformin','1g','BD','T2DM'],['Perindopril','8mg','OD','HTN'],
         ['Atorvastatin','20mg','ON','Hyperlipidaemia'],
         ['Paracetamol','1g','QDS PRN','Back pain'],['Pregabalin','75mg','BD','Neuropathic pain']])
    _docx_h2(doc, '6. Examination Findings (2026-04-03)')
    _docx_p(doc, 'Vitals: BP 138/82, HR 78, RR 16, T 36.7°C, SpO₂ 98% RA.')
    _docx_p(doc, 'Neurological: power 4/5 left ankle dorsiflexion, 5/5 elsewhere. Reduced sensation L5 dermatome left lower limb. Straight-leg-raise positive at 40° on left, negative on right. Reflexes preserved and symmetrical. No saddle anaesthesia. No long-tract signs.')
    _docx_h2(doc, '7. Working Diagnosis')
    _docx_p(doc, 'Left L5 radiculopathy, most likely secondary to lumbar disc protrusion. Awaiting MRI for confirmation. Differential: spinal stenosis, facet joint pathology, foraminal narrowing.')
    _docx_h2(doc, '8. Plan')
    _docx_bullets(doc, [
        'MRI lumbar spine — completed 2026-04-08 (report attached)',
        'Refer to Spinal MDT once imaging available',
        'Continue conservative management: physiotherapy, analgesia ladder',
        'Reinforce diabetic control given delayed wound healing risk if surgery indicated',
        'Review at MDT decision point',
    ])
    _docx_p(doc, 'Dictated by: Dr. K Wong, Specialist Orthopaedic Surgeon · Apex Specialist Centre · 2026-04-03 14:25', italic=True)
    doc.save(str(d / 'Patient_History_Summary.docx'))

    # Imaging_Report_MRI.pdf
    _pdf(d / 'Imaging_Report_MRI.pdf', 'MRI Lumbar Spine Report', [
        ('h1','MRI Lumbar Spine — Report'),
        ('small','Apex Specialist Centre · Radiology · Confidential'),
        ('small','Patient: Patient X (MRN: APX-2026-04421) · Date of scan: 2026-04-08 · Reporting radiologist: Dr. L. Chen, FRCR'),
        ('spacer', 6),
        ('h2','Clinical Indication'),
        ('p','58 F with 6-week history of progressive lower back pain radiating to the left leg, L5 dermatomal distribution, positive SLR. Rule out lumbar disc pathology.'),
        ('h2','Technique'),
        ('p','Sagittal T1, T2, STIR; axial T2 from L1 to S1. No contrast administered.'),
        ('h2','Findings'),
        ('p','<b>Alignment:</b> Lumbar lordosis preserved. No spondylolisthesis.'),
        ('p','<b>Vertebral bodies:</b> Normal height and marrow signal throughout L1-S1. No fractures, no marrow infiltration, no enhancing lesion.'),
        ('p','<b>L4/L5 level:</b> <i>Posterolateral disc protrusion measuring 6mm extending into the left lateral recess, contacting and displacing the traversing left L5 nerve root. Mild facet joint arthropathy bilaterally. Ligamentum flavum thickened to 4mm.</i> Moderate left foraminal narrowing. Right foramen patent.'),
        ('p','<b>L5/S1 level:</b> Disc desiccation with 2mm broad-based bulge. No focal protrusion. Both neural foramina patent.'),
        ('p','<b>L3/L4 level:</b> Disc desiccation only. No protrusion, no significant facet arthropathy. Central canal and foramina patent.'),
        ('p','<b>Upper levels (L1/L2, L2/L3):</b> Unremarkable.'),
        ('p','<b>Conus medullaris:</b> Terminates at L1, normal signal.'),
        ('p','<b>Paraspinal soft tissues:</b> No mass, no abscess, no significant abnormality.'),
        ('h2','Impression'),
        ('quote','1. Left posterolateral L4/L5 disc protrusion (6mm) compressing the traversing left L5 nerve root, with moderate left foraminal narrowing. Correlates with clinical L5 radiculopathy.'),
        ('quote','2. Mild multi-level degenerative change. No central canal stenosis. No red-flag features (no mass, no infection, no fracture).'),
        ('h2','Recommendation'),
        ('p','Findings consistent with mechanical L5 radiculopathy. Clinical correlation. Consider conservative management trial; if no improvement at 6 weeks or if neurological deficit progresses, escalate per pathway.'),
        ('spacer', 8),
        ('small','Report electronically signed by Dr. L. Chen, FRCR on 2026-04-08 17:42.'),
    ])

    # Lab_Results_Panel.xlsx
    wb = _xl_new(d / 'Lab_Results_Panel.xlsx')
    ws = wb.active; ws.title = 'Lab Panel'
    rows = [
        ['Haemoglobin','12.8','11.5–16.5','g/dL','2026-04-02','G'],
        ['WBC','7.4','4.0–11.0','x10⁹/L','2026-04-02','G'],
        ['Platelets','268','150–450','x10⁹/L','2026-04-02','G'],
        ['Neutrophils','4.6','2.0–7.5','x10⁹/L','2026-04-02','G'],
        ['CRP','12','<5','mg/L','2026-04-02','A'],
        ['ESR','38','<20','mm/hr','2026-04-02','A'],
        ['Urea','5.8','2.5–7.5','mmol/L','2026-04-02','G'],
        ['Creatinine','78','50–110','µmol/L','2026-04-02','G'],
        ['eGFR','82','>60','mL/min/1.73m²','2026-04-02','G'],
        ['Na+','139','135–145','mmol/L','2026-04-02','G'],
        ['K+','4.4','3.5–5.0','mmol/L','2026-04-02','G'],
        ['Fasting glucose','7.8','3.9–5.6','mmol/L','2026-04-02','R'],
        ['HbA1c','7.4','<6.5','%','2026-04-02','R'],
        ['LDL-C','2.4','<3.0','mmol/L','2026-04-02','G'],
        ['HDL-C','1.3','>1.0','mmol/L','2026-04-02','G'],
        ['Triglycerides','1.8','<1.7','mmol/L','2026-04-02','A'],
        ['Calcium (corrected)','2.35','2.20–2.60','mmol/L','2026-04-02','G'],
        ['Phosphate','1.1','0.8–1.5','mmol/L','2026-04-02','G'],
        ['ALT','24','<40','U/L','2026-04-02','G'],
        ['TSH','2.1','0.4–4.5','mIU/L','2026-04-02','G'],
        ['Vitamin D','42','>50 optimal','nmol/L','2026-04-02','A'],
    ]
    _xl_write_sheet(ws,
        ['Test','Result','Reference Range','Units','Sample Date','RAG'],
        rows, widths=[24,12,18,14,14,8],
        rag_col=6, rag_map=lambda r: r[5])
    ws2 = wb.create_sheet('Notes')
    _xl_write_sheet(ws2,
        ['Flag','Interpretation'],
        [['CRP / ESR (Amber)','Mildly elevated — consistent with low-grade systemic inflammation. Not specific. Trend over 4 weeks.'],
         ['HbA1c / Fasting glucose (Red)','Diabetic control suboptimal. Pre-operatively HbA1c should be <8.0%. Endocrinology consult if surgery considered.'],
         ['Triglycerides (Amber)','Borderline elevated. Lifestyle counselling.'],
         ['Vitamin D (Amber)','Insufficient. Supplement 1000 IU/day for 12 weeks then retest.']],
        widths=[28, 100])
    _xl_save(wb, d / 'Lab_Results_Panel.xlsx')

    # Clinical_Guidelines_Pathway.pdf
    _pdf(d / 'Clinical_Guidelines_Pathway.pdf', 'Lumbar Radiculopathy Pathway', [
        ('h1','Lumbar Radiculopathy — Clinical Pathway'),
        ('small','Apex Specialist Centre · Pathway Reference: ORTHO-SPINE-04 · Version 4.2 · Effective 2026-01-01'),
        ('spacer', 8),
        ('h2','1. Scope'),
        ('p','This pathway applies to adult patients (≥18 years) presenting with unilateral lumbar radiculopathy. It does NOT apply to: cauda equina syndrome (emergency referral), progressive neurological deficit (urgent surgical review), suspected malignancy, infection, or trauma.'),
        ('h2','2. Red-Flag Triage (do these BEFORE entering the pathway)'),
        ('p','• Saddle anaesthesia / bladder or bowel dysfunction → ED, exclude cauda equina'),
        ('p','• Bilateral leg weakness or rapidly progressive deficit → urgent surgical review'),
        ('p','• Weight loss, night pain, history of malignancy → urgent imaging + oncology workup'),
        ('p','• Fever, recent procedure, IVDU → suspect discitis / epidural abscess → urgent MRI + ID'),
        ('h2','3. Conservative Phase (Weeks 0–6)'),
        ('p','First-line: analgesic ladder (paracetamol → NSAIDs → neuropathic agents), structured physiotherapy 2x/week, activity modification. No bed rest. Re-assess at 6 weeks.'),
        ('h2','4. Imaging'),
        ('p','MRI lumbar spine is the imaging of choice if: symptoms persist >6 weeks, neurological deficit present, OR red flags identified.'),
        ('h2','5. Treatment Options (after imaging confirms disc pathology)'),
        ('p','<b>Option A — Continued Conservative:</b> for patients with stable or improving symptoms, no progressive deficit. Continue physio + analgesia for further 6 weeks. Indicated when imaging shows protrusion ≤7mm and SLR positive but no motor weakness >grade 1.'),
        ('p','<b>Option B — Targeted Epidural Steroid Injection (ESI):</b> for patients with significant radicular pain unresponsive to oral medication, imaging-confirmed nerve root contact, NO motor weakness progression. Contraindicated in: uncontrolled diabetes (HbA1c >8.5%), anticoagulation, active infection.'),
        ('p','<b>Option C — Surgical Decompression (microdiscectomy):</b> indicated when (1) progressive neurological deficit, OR (2) failure of 6-12 weeks of conservative + ESI, AND (3) imaging-clinical correlation clear. Contraindications: HbA1c >9.0% (delayed wound healing), active systemic infection, severe cardiac comorbidity.'),
        ('p','<b>Option D — Multidisciplinary Pain Programme:</b> for chronic radiculopathy >12 weeks not amenable to surgery. CBT + physio + pain medicine.'),
        ('h2','6. MDT Decision'),
        ('p','Patients NOT meeting clear surgical criteria must be presented at the Spinal MDT (Wednesdays 14:00) before progressing to any invasive treatment. Required at MDT: history summary, imaging report, recent labs (HbA1c, FBC, U&E, CRP), patient preference / consent discussion.'),
        ('h2','7. Outcomes & Follow-Up'),
        ('p','• Conservative: review week 6 and 12, discharge if pain VAS ≤2 sustained'),
        ('p','• ESI: review 4 weeks post-injection'),
        ('p','• Surgical: review 2 weeks, 6 weeks, 3 months post-op'),
        ('p','• All patients: lifestyle counselling, weight management, glycaemic optimisation as relevant'),
    ])

    # Multidisciplinary_Team_Notes.docx
    doc = _docx_init('Multidisciplinary Team — Previous Meeting Notes', 'Apex Specialist Centre · Spinal MDT · 2026-03-27 (most recent meeting prior to this case)')
    _docx_h2(doc, 'Attendees')
    _docx_bullets(doc, [
        'Mr K. Wong — Specialist Orthopaedic Surgeon (Chair)',
        'Dr L. Chen — Consultant Radiologist',
        'Dr A. Lim — Pain Medicine Consultant',
        'Ms S. Tan — Senior Physiotherapist',
        'Ms R. Cheong — Specialist Nurse, Spine',
        'Dr H. Ng — Endocrinology (corresponding member, diabetic cases)',
    ])
    _docx_h2(doc, 'Cases Discussed (3 of 6 relevant to current case)')
    _docx_h2(doc, 'Case 1 — L5/S1 disc protrusion, female 55y, diabetic (HbA1c 8.1%)')
    _docx_p(doc, 'Outcome: Deferred surgery. Diabetic control optimisation referral. Trial ESI scheduled. Decision rationale: HbA1c above pathway threshold of 8.0% for elective surgery, conservative options not exhausted, no progressive deficit. Review in 6 weeks.')
    _docx_h2(doc, 'Case 2 — L4/L5 lateral recess stenosis, male 62y, progressive foot drop')
    _docx_p(doc, 'Outcome: Listed for urgent decompression within 2 weeks. Progressive motor deficit grade 3/5 dorsiflexion satisfies surgical criterion regardless of conservative duration.')
    _docx_h2(doc, 'Case 3 — Chronic radiculopathy 14 months, multiple failed treatments')
    _docx_p(doc, 'Outcome: Referred to multidisciplinary pain programme. Surgery offered no incremental benefit per evidence review.')
    _docx_h2(doc, 'Recurring themes / governance')
    _docx_bullets(doc, [
        'Diabetic-control threshold (HbA1c <8.0% for elective spinal surgery) is being applied consistently; 2 cases this cycle deferred for endocrine optimisation.',
        'Pre-MDT checklist completion improved from 67% (prior quarter) to 91% (this quarter). Continued focus on imaging-clinical correlation.',
        'Patient-preference documentation gap: only 4/6 cases had explicit consent discussion summary. Action: nursing team to update template.',
    ])
    _docx_p(doc, 'Minutes prepared by: Ms R. Cheong, Specialist Nurse · Apex Specialist Centre · 2026-03-27', italic=True)
    doc.save(str(d / 'Multidisciplinary_Team_Notes.docx'))


# ─── Archetype 3: Regulator cross-ref ───
def gen_regulator_crossref():
    d = OUT / 'nb-regulator-crossref'
    d.mkdir(parents=True, exist_ok=True)

    _pdf(d / 'Regulator_Circular_2026.pdf', 'Regulator Circular 2026', [
        ('h1','Regulator Circular 2026/04 — Strengthening Operational Resilience'),
        ('small','Issued: 2026-03-15 · Effective: 2026-10-01 · Supersedes: 2024 Circular on Outsourcing'),
        ('h2','Section 1 — Scope'),
        ('p','Applicable to all licensed entities under the regulator\'s purview, including subsidiaries and material outsourcing arrangements. Compliance attestation required by Group Risk Officer or equivalent.'),
        ('h2','Section 2 — Governance Obligations'),
        ('p','2.1 — The Board must approve a Group Operational Resilience Framework no later than 2026-09-30.'),
        ('p','2.2 — The framework must identify Important Business Services (IBS) using customer-impact, financial-stability, and market-conduct criteria.'),
        ('p','2.3 — For each IBS, the entity must set Impact Tolerances stated in maximum tolerable disruption duration and maximum tolerable data loss.'),
        ('h2','Section 3 — Mapping & Testing'),
        ('p','3.1 — Map each IBS end-to-end including people, processes, technology, facilities, and third parties.'),
        ('p','3.2 — Conduct annual scenario testing using severe but plausible scenarios; report results to the Board within 60 days of test completion.'),
        ('p','3.3 — Identify and remediate vulnerabilities preventing the entity from remaining within Impact Tolerances.'),
        ('h2','Section 4 — Third-Party Risk'),
        ('p','4.1 — Material outsourcing must include contractual rights to audit, data residency clauses, exit / step-in provisions, and continuous service-level monitoring.'),
        ('p','4.2 — Material technology service providers concentration: no single provider may support more than 40% of IBS by criticality weight unless mitigated.'),
        ('h2','Section 5 — Incident Reporting'),
        ('p','5.1 — Notify the regulator within 4 hours of confirmed P1 operational incident affecting any IBS.'),
        ('p','5.2 — Submit root-cause analysis within 30 calendar days. Customer redress plan within 60 days where customer impact occurred.'),
        ('h2','Section 6 — Attestation'),
        ('p','6.1 — Group Risk Officer (or equivalent) must attest annually by 31 March that the Operational Resilience Framework is operating effectively.'),
        ('p','6.2 — First attestation due: 2027-03-31 covering FY2026.'),
        ('h2','Section 7 — Penalties'),
        ('p','Failure to comply may attract supervisory action up to MYR 25 million per breach and, in severe cases, restriction of licence. Individual director accountability under the Senior Manager Regime applies.'),
        ('h2','Annexes'),
        ('p','Annex A — IBS taxonomy (15 service categories)'),
        ('p','Annex B — Sample Impact Tolerance metrics'),
        ('p','Annex C — Reporting templates (Form OR-01 to OR-05)'),
    ])

    # Internal_Policy_Manual.docx
    doc = _docx_init('Zava Conglomerate — Internal Operational Resilience Policy', 'Document ID: ZAVA-POL-RES-01 · Version 2.1 · Effective 2025-04-01 · Owner: Group Risk Officer')
    _docx_h2(doc, '1. Purpose & Scope')
    _docx_p(doc, 'This policy sets out Zava Conglomerate\'s operational resilience framework. It applies to all subsidiaries and joint ventures controlled by Zava Conglomerate.')
    _docx_h2(doc, '2. Governance')
    _docx_p(doc, '2.1 — The Group Risk Committee reviews the operational resilience framework annually.')
    _docx_p(doc, '2.2 — The Board approves the Important Business Services (IBS) list and Impact Tolerances every 24 months.')
    _docx_p(doc, '(Note: this policy currently states 24-month Board review cycle. Regulator Circular 2026 requires annual Board attestation. POLICY GAP — to be remediated.)')
    _docx_h2(doc, '3. Important Business Services')
    _docx_p(doc, 'Current IBS catalogue (12 services):')
    _docx_bullets(doc, [
        'Customer-facing digital banking',
        'Payment processing (domestic + cross-border)',
        'Card authorisation',
        'Retail branch banking',
        'Treasury operations',
        'Investor portal and disclosures',
        'Insurance claims processing',
        'Healthcare patient records system',
        'Plantation logistics dispatch',
        'Group payroll',
        'Group financial close',
        'Cybersecurity operations centre',
    ])
    _docx_h2(doc, '4. Impact Tolerances')
    _docx_table(doc,
        ['IBS','Max Disruption','Max Data Loss','Last Test','Result'],
        [['Digital banking','2 hrs','0 transactions','2025-11','Pass'],
         ['Payment processing','1 hr','0 transactions','2025-10','Pass'],
         ['Card authorisation','30 min','0 transactions','2025-11','Pass'],
         ['Branch banking','4 hrs','15 min','2025-09','Pass'],
         ['Treasury','2 hrs','0 transactions','2025-08','Partial'],
         ['Insurance claims','24 hrs','24 hrs','Not tested','—'],
         ['Healthcare records','1 hr','15 min','Not tested','—']])
    _docx_h2(doc, '5. Third-Party Risk')
    _docx_p(doc, 'All material outsourcing must include audit rights, data residency, and exit clauses. Concentration assessment performed annually.')
    _docx_p(doc, '(Note: 2 of our 7 material technology providers exceed 40% concentration per regulator threshold. Remediation in progress; see Audit Findings register.)')
    _docx_h2(doc, '6. Incident Reporting')
    _docx_p(doc, 'P1 operational incidents must be reported to Group Risk within 1 hour. Regulator notification handled by Compliance team — current internal SLA is 8 hours, which exceeds the regulator\'s 4-hour requirement under Circular 2026. POLICY GAP — to be remediated.')
    _docx_h2(doc, '7. Review')
    _docx_p(doc, 'This policy is reviewed every 24 months or sooner if material regulatory change occurs. Next scheduled review: 2027-04-01.')
    doc.save(str(d / 'Internal_Policy_Manual.docx'))

    # Audit_Findings_Last_Cycle.xlsx
    wb = _xl_new(d / 'Audit_Findings_Last_Cycle.xlsx')
    ws = wb.active; ws.title = 'Findings'
    rows = [
        ['AF-2025-001','Operational Resilience','Board attestation cycle (24m) inconsistent with regulator expectation (12m)','Group Risk Officer','High','Open','2026-09-30'],
        ['AF-2025-002','Third-Party Risk','Concentration: 2 of 7 material TSPs exceed 40% threshold','CIO + CISO','High','Open','2026-12-31'],
        ['AF-2025-003','Incident Reporting','Internal SLA 8h vs regulator 4h notification window','Group Compliance','High','In Progress','2026-06-30'],
        ['AF-2025-004','IBS Mapping','Insurance claims + Healthcare records not yet scenario-tested','COO','Medium','Open','2026-09-30'],
        ['AF-2025-005','Treasury Resilience','Last test only partial pass — recovery exceeded tolerance once','Group Treasurer','Medium','In Progress','2026-08-31'],
        ['AF-2025-006','Documentation','15 of 132 outsourcing contracts missing audit-rights clause','Procurement','Medium','In Progress','2026-07-31'],
        ['AF-2025-007','Data Residency','3 cloud services storing customer data outside ASEAN region','CISO','High','Remediated','2025-12-15'],
        ['AF-2025-008','Cybersecurity','MFA not enforced on 12% of privileged accounts','CISO','High','Remediated','2025-11-30'],
        ['AF-2025-009','Vendor Exit','No tested exit plan for 4 of 7 material TSPs','CIO','Medium','Open','2026-12-31'],
        ['AF-2025-010','Disclosure','ESG Scope-3 measurement methodology not aligned with framework','Head of IR','Low','In Progress','2026-09-30'],
        ['AF-2025-011','HR','Annual code-of-conduct attestation completion at 87% (target 100%)','CHRO','Low','In Progress','2026-06-30'],
        ['AF-2025-012','Internal Audit','Coverage of regulated subsidiaries below 3-year cycle target','Head of Internal Audit','Medium','In Progress','2027-03-31'],
    ]
    _xl_write_sheet(ws,
        ['Finding ID','Category','Description','Owner','Severity','Status','Target Closure'],
        rows, widths=[14,22,55,22,10,14,16],
        rag_col=5, rag_map=lambda r: 'R' if r[4]=='High' and r[5] not in ('Remediated','Closed') else 'A' if r[4]=='Medium' else 'G')
    _xl_save(wb, d / 'Audit_Findings_Last_Cycle.xlsx')

    # Compliance_Register_Live.xlsx
    wb = _xl_new(d / 'Compliance_Register_Live.xlsx')
    ws = wb.active; ws.title = 'Live Register'
    rows = [
        ['OB-001','Operational Resilience','Board approval of Group Op Resilience Framework','Annual','Group Risk Officer','2026-09-30','Open','Policy 2.2 currently 24m cycle'],
        ['OB-002','Operational Resilience','Identify all IBS with customer-impact criteria','One-off','Group Risk Officer','2026-09-30','In Progress','12 of 15 IBS categories mapped'],
        ['OB-003','Operational Resilience','Set Impact Tolerances per IBS','Annual','Group Risk Officer','2026-09-30','In Progress','7 of 12 tested'],
        ['OB-004','Operational Resilience','Annual scenario testing per IBS','Annual','COO','2026-09-30','In Progress','Insurance + Healthcare not yet tested'],
        ['OB-005','Third-Party Risk','Material outsourcing audit rights','Continuous','Procurement','Live','In Progress','15 contracts missing clause'],
        ['OB-006','Third-Party Risk','Concentration ≤ 40% per TSP','Continuous','CIO','Live','Open','2 TSPs exceed threshold'],
        ['OB-007','Incident Reporting','Notify regulator within 4 hours of P1','Continuous','Compliance','Live','Open','Internal SLA currently 8h'],
        ['OB-008','Incident Reporting','RCA within 30 days; redress plan within 60','Per incident','Compliance','Live','In Progress','RCA timeline aligned, redress under review'],
        ['OB-009','Governance','Group Risk Officer annual attestation','Annual','GRO','2027-03-31','Not Started','First filing FY2026'],
        ['OB-010','Data Residency','Customer data within ASEAN','Continuous','CISO','Live','Remediated','3 services migrated 2025-12'],
        ['OB-011','PDPA','Annual DPO report','Annual','DPO','2026-06-30','In Progress','Drafting'],
        ['OB-012','PDPA','Breach notification 72h','Continuous','DPO','Live','In Progress','No breaches reported in last cycle'],
    ]
    _xl_write_sheet(ws,
        ['Obligation ID','Domain','Description','Frequency','Owner','Due','Status','Notes'],
        rows, widths=[14,20,50,12,22,14,14,40])
    _xl_save(wb, d / 'Compliance_Register_Live.xlsx')

    # Board_Risk_Appetite_Statement.docx
    doc = _docx_init('Group Risk Appetite Statement', 'Zava Conglomerate — Board approved 2025-06-30 · Document ID: ZAVA-GOV-RAS-02')
    _docx_h2(doc, '1. Strategic Risk Appetite')
    _docx_p(doc, 'Zava Conglomerate pursues sustainable long-term value creation. We accept moderate strategic risk where the upside is commensurate with disciplined capital allocation and clear stress-testing.')
    _docx_h2(doc, '2. Operational Risk Appetite')
    _docx_p(doc, 'Zero tolerance for: regulatory non-compliance that materially threatens our licences, customer harm that breaches fiduciary duty, and material data breaches affecting PII or customer assets.')
    _docx_p(doc, 'Low tolerance for: operational incidents disrupting any Important Business Service beyond its Impact Tolerance, third-party concentration above policy thresholds.')
    _docx_h2(doc, '3. Compliance Risk Appetite')
    _docx_p(doc, 'Zero tolerance for: wilful regulatory breach, market abuse, financial-crime control failure, breach of sanctions regime.')
    _docx_p(doc, 'Compliance with all binding regulatory circulars — including but not limited to those addressing operational resilience, capital adequacy, conduct of business, data protection — is non-negotiable.')
    _docx_h2(doc, '4. Financial Risk Appetite')
    _docx_p(doc, 'Group Net Debt / EBITDA ceiling: 3.0x. Investment-grade credit rating maintained. Liquidity coverage ratio ≥ 130% at all times.')
    _docx_h2(doc, '5. Reputational Risk Appetite')
    _docx_p(doc, 'Low tolerance — we will exit any client, market, or supplier relationship that materially damages stakeholder trust.')
    _docx_h2(doc, '6. People & Conduct Appetite')
    _docx_p(doc, 'Zero tolerance for breaches of the Code of Conduct, harassment, or workplace safety failures resulting in major injury or fatality.')
    _docx_h2(doc, '7. Review')
    _docx_p(doc, 'This statement is reviewed annually by the Board Risk Committee. Material breaches of stated appetite are escalated to the Board within 5 business days.')
    doc.save(str(d / 'Board_Risk_Appetite_Statement.docx'))


# ─── Archetype 4: M&A due diligence ───
def gen_ma_due_diligence():
    d = OUT / 'nb-ma-due-diligence'
    d.mkdir(parents=True, exist_ok=True)

    # Target_Financials_3Y.xlsx
    wb = _xl_new(d / 'Target_Financials_3Y.xlsx')

    # P&L
    ws = wb.active; ws.title = 'P&L'
    pnl = [
        ['Revenue',                   '482.4','531.0','602.8','670.2','744.6','Growth +10% CAGR'],
        ['Cost of sales',             '(289.4)','(312.8)','(355.6)','(394.1)','(434.6)','GM stable ~40%'],
        ['Gross profit',              '193.0','218.2','247.2','276.1','310.0','—'],
        ['Operating expenses',        '(138.5)','(146.2)','(164.1)','(180.5)','(199.6)','—'],
        ['EBITDA',                    '54.5','72.0','83.1','95.6','110.4','Margin trending 14-16%'],
        ['Depreciation & amortisation','(18.4)','(20.1)','(22.8)','(25.6)','(28.0)','—'],
        ['EBIT',                      '36.1','51.9','60.3','70.0','82.4','—'],
        ['Net interest',              '(7.2)','(8.5)','(9.8)','(10.5)','(11.2)','—'],
        ['Profit before tax',         '28.9','43.4','50.5','59.5','71.2','—'],
        ['Income tax',                '(6.9)','(10.4)','(12.1)','(14.3)','(17.1)','24% effective rate'],
        ['Profit after tax',          '22.0','33.0','38.4','45.2','54.1','—'],
    ]
    _xl_write_sheet(ws,
        ['MYR M','FY2023A','FY2024A','FY2025A','FY2026F (Mgmt)','FY2027F (Mgmt)','Auditor Note'],
        pnl, widths=[30,12,12,12,16,16,40])

    # Balance Sheet
    ws2 = wb.create_sheet('Balance Sheet')
    bs = [
        ['Cash & equivalents','48.2','62.4','71.8','—'],
        ['Trade receivables','71.5','78.6','89.1','DSO 54 days — trending up'],
        ['Inventory','38.4','42.1','47.6','—'],
        ['Other current assets','12.4','14.2','16.0','—'],
        ['Total current assets','170.5','197.3','224.5','—'],
        ['PP&E (net)','148.6','167.4','184.2','Capex ~7% of revenue'],
        ['Goodwill & intangibles','62.1','62.1','62.1','From 2018 acquisition'],
        ['Other non-current','14.2','15.8','17.4','—'],
        ['Total non-current','224.9','245.3','263.7','—'],
        ['TOTAL ASSETS','395.4','442.6','488.2','—'],
        ['Trade payables','58.4','63.9','71.2','—'],
        ['Short-term debt','34.8','38.2','41.0','—'],
        ['Other current liabilities','22.6','25.1','28.4','—'],
        ['Total current liabilities','115.8','127.2','140.6','—'],
        ['Long-term debt','78.4','82.1','89.5','Lender covenants — see register'],
        ['Other non-current','18.6','20.1','22.0','—'],
        ['Total non-current liabilities','97.0','102.2','111.5','—'],
        ['Total equity','182.6','213.2','236.1','—'],
        ['TOTAL EQUITY + LIABILITIES','395.4','442.6','488.2','—'],
    ]
    _xl_write_sheet(ws2,
        ['MYR M','FY2023A','FY2024A','FY2025A','Auditor Note'],
        bs, widths=[32,12,12,12,42])

    # Cash flow
    ws3 = wb.create_sheet('Cash Flow')
    cf = [
        ['EBITDA','54.5','72.0','83.1'],
        ['Changes in working capital','(8.2)','(11.4)','(14.6)'],
        ['Cash from operations','46.3','60.6','68.5'],
        ['Capex (maintenance)','(18.4)','(22.1)','(24.8)'],
        ['Capex (growth)','(14.2)','(18.5)','(22.0)'],
        ['Free cash flow','13.7','20.0','21.7'],
        ['Net interest paid','(7.2)','(8.5)','(9.8)'],
        ['Tax paid','(6.9)','(10.4)','(12.1)'],
        ['Cash flow before financing','(0.4)','1.1','(0.2)'],
        ['Net debt drawdown / (repayment)','5.4','3.7','7.4'],
        ['Dividends paid','(4.8)','(6.2)','(7.2)'],
        ['Net change in cash','0.2','(1.4)','—'],
    ]
    _xl_write_sheet(ws3,
        ['MYR M','FY2023A','FY2024A','FY2025A'],
        cf, widths=[34,12,12,12])

    # KPIs / Quality of earnings
    ws4 = wb.create_sheet('KPIs')
    kpis = [
        ['Revenue growth %','—','10.1%','13.5%'],
        ['Gross margin %','40.0%','41.1%','41.0%'],
        ['EBITDA margin %','11.3%','13.6%','13.8%'],
        ['EBIT margin %','7.5%','9.8%','10.0%'],
        ['Net margin %','4.6%','6.2%','6.4%'],
        ['DSO (days)','54','54','54'],
        ['DIO (days)','48','49','49'],
        ['DPO (days)','74','75','73'],
        ['Cash conversion cycle (days)','28','28','30'],
        ['Net debt / EBITDA','1.20x','0.80x','0.71x'],
        ['Net debt (MYR M)','65.0','57.9','58.7'],
        ['ROIC','9.8%','13.5%','14.2%'],
        ['Customer concentration: top-3','22%','27%','31%','Trending — top customer now 14%'],
        ['Customer concentration: top-10','58%','61%','66%','Trending — risk increasing'],
    ]
    _xl_write_sheet(ws4,
        ['Metric','FY2023A','FY2024A','FY2025A','Comment'],
        kpis, widths=[34,12,12,12,40])

    # Quality of Earnings adjustments
    ws5 = wb.create_sheet('QoE Adjustments')
    qoe = [
        ['FY2025 reported EBITDA','83.1','—','Reported'],
        ['Adj: One-time legal settlement','+2.4','Add-back','Single litigation matter (settled 2025)'],
        ['Adj: COVID-recovery grants','-1.8','Remove','Non-recurring government support'],
        ['Adj: Owner remuneration above market','+3.2','Add-back','To be normalised at market rate post-deal'],
        ['Adj: Related-party rent below market','-1.4','Deduct','Lease to be revised at market post-deal'],
        ['Adj: R&D capitalisation review','-0.9','Deduct','Software development costs over-capitalised'],
        ['Adj: Inventory write-down (one-off)','+1.1','Add-back','Discontinued SKU clear-out'],
        ['Adjusted FY2025 EBITDA','85.7','—','Pro-forma normalised'],
        ['EV / EBITDA at offer price (8.5x)','728.5','—','Implied enterprise value at offer'],
    ]
    _xl_write_sheet(ws5,
        ['Item','MYR M','Direction','Auditor Note'],
        qoe, widths=[44,10,12,40])

    _xl_save(wb, d / 'Target_Financials_3Y.xlsx')

    # Target_Cap_Table.xlsx
    wb = _xl_new(d / 'Target_Cap_Table.xlsx')
    ws = wb.active; ws.title = 'Cap Table'
    cap = [
        ['Founder (CEO)','3,420,000','34.2%','Ordinary','—','Class A — 1 vote per share'],
        ['Founder (CTO)','2,180,000','21.8%','Ordinary','—','Class A — 1 vote per share'],
        ['Co-founder #3','1,460,000','14.6%','Ordinary','—','Class A — 1 vote per share'],
        ['Pacific Growth Fund','1,200,000','12.0%','Series A Pref','2021','Liquidation preference 1x non-participating'],
        ['Asianova Ventures','800,000','8.0%','Series A Pref','2021','Liquidation preference 1x non-participating'],
        ['Apex Strategic Partners','480,000','4.8%','Series B Pref','2023','Liquidation preference 1x participating, 2x cap'],
        ['ESOP pool (allocated)','280,000','2.8%','ESOP (Vested)','—','Vesting 4-yr, 1-yr cliff'],
        ['ESOP pool (unallocated)','120,000','1.2%','ESOP (Unvested)','—','Available for grant'],
        ['Angel investors (12 holders)','60,000','0.6%','Ordinary','2019-2020','—'],
        ['Total','10,000,000','100.0%','—','—','—'],
    ]
    _xl_write_sheet(ws,
        ['Holder','Shares','% Fully Diluted','Class','Vintage','Notes / Preference Terms'],
        cap, widths=[34,14,16,20,10,52])

    ws2 = wb.create_sheet('Preferred Terms')
    _xl_write_sheet(ws2,
        ['Class','Holders','Liquidation Preference','Conversion','Anti-Dilution','Board Seats','Drag/Tag','Information Rights'],
        [
            ['Series A','Pacific Growth + Asianova','1x non-participating','1:1','Weighted average','1','Drag if 66.7% Series A','Quarterly + audited annual'],
            ['Series B','Apex Strategic','1x participating capped at 2x return','1:1','Full ratchet','1','Drag if Series B + majority','Monthly + board observer rights'],
        ],
        widths=[14,30,38,14,18,12,28,38])

    ws3 = wb.create_sheet('Notes')
    _xl_write_sheet(ws3,
        ['Note','Detail'],
        [
            ['Outstanding warrants','None'],
            ['Convertible notes','None'],
            ['Promised but unissued options','30,000 to incoming CFO (offer outstanding)'],
            ['Right of first refusal','All shareholders subject to ROFR on transfer'],
            ['Tag-along rights','All shareholders'],
            ['Drag-along rights','Triggered by 66.7% Series A + Founders majority'],
            ['Pending matters','Series A consent required for sale above MYR 500M'],
        ],
        widths=[36,80])
    _xl_save(wb, d / 'Target_Cap_Table.xlsx')

    # Material_Contracts_Index.pdf
    _pdf(d / 'Material_Contracts_Index.pdf', 'Material Contracts Index', [
        ('h1','Material Contracts Index — Project Hawthorn (Target)'),
        ('small','Confidential — prepared under NDA · Data room: Section 3 · Date: 2026-04-15'),
        ('h2','1. Customer Contracts (Top 10 by revenue, FY2025)'),
        ('table', [
            ['#','Customer (anonymised)','Revenue MYR M FY25','Tenor','Auto-renewal','Change-of-control'],
            ['1','Customer #1 (Telco)','86.4','3 years to 2027','Yes (24m notice)','Termination right'],
            ['2','Customer #2 (Banking)','62.1','5 years to 2029','No','Notification only'],
            ['3','Customer #3 (Government)','48.9','3 years to 2026','No (re-tender)','Termination right'],
            ['4','Customer #4 (Healthcare)','40.2','2 years rolling','Yes (12m notice)','Notification only'],
            ['5','Customer #5 (Retail)','38.6','2 years to 2026','Yes (6m notice)','Termination right'],
            ['6','Customer #6 (FMCG)','32.4','3 years to 2027','Yes (12m notice)','Notification only'],
            ['7','Customer #7 (Utilities)','30.1','5 years to 2028','No','Termination right'],
            ['8','Customer #8 (Manufacturing)','28.0','2 years to 2026','Yes (6m notice)','Notification only'],
            ['9','Customer #9 (O&G)','25.4','3 years to 2026','No','Termination right'],
            ['10','Customer #10 (Insurance)','24.1','3 years to 2027','Yes (12m notice)','Notification only'],
        ]),
        ('p','Customer #1 termination-on-change-of-control is a material risk in deal context. Customer #3 has re-tender notice due 2026-08; pipeline currently rated Amber by Sales.'),
        ('h2','2. Supplier Contracts (key)'),
        ('p','Cloud infrastructure (Pacific Digital Group, 5-year contract to 2027) — most-favoured-nation pricing clause, audit rights, exit period 12 months.'),
        ('p','Software licences (4 enterprise vendors) — change-of-control consent required for 2 of 4. Estimated re-licence fees if denied: MYR 4-6M.'),
        ('p','Outsourced support services (2-year contract to 2026) — auto-renewal unless 6m notice.'),
        ('h2','3. Leases'),
        ('p','HQ lease (KL) — 12 years to 2034, 5-yearly rent reviews. Subletting permitted with consent. No change-of-control clause.'),
        ('p','Operations site (Penang) — 8 years to 2030. Termination on change-of-control without compensation.'),
        ('p','Office leases (3 secondary sites) — 3-year rolling, easy to exit.'),
        ('h2','4. Banking / Financing'),
        ('p','RM 90M senior facility — Net Debt/EBITDA covenant 3.0x; Interest Cover 4.0x; Material Adverse Change clause; change-of-control triggers immediate repayment unless lender consents.'),
        ('p','RM 25M working capital revolver — change-of-control consent required.'),
        ('h2','5. Litigation & Disputes'),
        ('p','See separate Litigation Register. One material matter outstanding (commercial claim from former customer, quantum MYR 8.4M, defended by management as without merit). All other matters immaterial.'),
        ('h2','6. Key Personnel Contracts'),
        ('p','CEO + CTO — change-of-control "double trigger" entitlement (12-month salary + accelerated vesting if terminated within 12 months post-deal).'),
        ('p','5 other senior executives — single-trigger 6-month salary; deal team to assess retention package strategy.'),
        ('p','Non-compete: 12 months post-employment, customary geography.'),
        ('h2','7. Insurance'),
        ('p','D&O — USD 25M aggregate; runs through 2026-11. Buyer should consider tail coverage.'),
        ('p','Cyber — MYR 10M aggregate; current.'),
        ('p','Professional indemnity — MYR 20M aggregate; current.'),
        ('h2','8. Regulatory Approvals'),
        ('p','Operating licences in 4 jurisdictions. Change of beneficial ownership above 30% requires regulator notification in 3 of 4; only one jurisdiction requires prior approval (timeline 90-120 days).'),
    ])

    # Litigation_Register.xlsx
    wb = _xl_new(d / 'Litigation_Register.xlsx')
    ws = wb.active; ws.title = 'Litigation'
    lit = [
        ['LIT-2024-001','Commercial','Former customer claim for contract termination damages','MYR 8.4M','High','Active — defence filed','External counsel: Tier-1 firm','Trial 2026-Q4'],
        ['LIT-2023-014','Employment','Former employee unfair dismissal','MYR 240k','Low','Settled 2024-09','Settled out of court','Closed'],
        ['LIT-2023-008','IP','Software copyright (defendant)','MYR 1.2M','Low','Dismissed 2024-04','—','Closed'],
        ['LIT-2025-003','Tax','Transfer pricing review, FY2022','MYR 1.8M (assessed)','Medium','Under appeal','Tax advisor engaged','Decision Q3 2026'],
        ['LIT-2025-007','Regulatory','Data-protection complaint by customer','None monetary','Low','Closed by regulator','No fine imposed','Closed 2025-11'],
        ['LIT-2026-001','Commercial','Supplier counterclaim re payment terms','MYR 380k','Low','Active — mediation','—','Mediation 2026-05'],
    ]
    _xl_write_sheet(ws,
        ['Case ID','Type','Description','Quantum','Severity','Status','Counsel','Milestone'],
        lit, widths=[14,14,46,18,12,22,22,18])

    ws2 = wb.create_sheet('Assessment')
    _xl_write_sheet(ws2,
        ['Risk Item','Likely Impact','Comment'],
        [
            ['LIT-2024-001 — commercial claim','MYR 5-8M downside','Defended robustly. Insurer engaged. Trial Q4 2026.'],
            ['LIT-2025-003 — transfer pricing','MYR 1.8M + penalties','Possible warranty / indemnity item in SPA.'],
            ['Cumulative reserves on books','MYR 2.4M','Adequate per auditor.'],
            ['Insurance coverage','D&O + Cyber + PI active','Tail coverage recommended post-deal.'],
        ],
        widths=[44,22,80])
    _xl_save(wb, d / 'Litigation_Register.xlsx')

    # Management_Presentation.pdf
    _pdf(d / 'Management_Presentation.pdf', 'Management Presentation', [
        ('h1','Project Hawthorn — Management Presentation'),
        ('small','Prepared by Target Management · For Confidential Discussion Only · 2026-04 · Buyer side: Zava Conglomerate Strategy Office'),
        ('h2','1. Executive Summary'),
        ('p','Target is a leading provider of industrial digital solutions across ASEAN. Revenue MYR 602M (FY2025), EBITDA MYR 83M (13.8% margin), 1,420 employees, presence in 4 countries. Founded 2014, profitable since FY2018, two priced equity rounds (Series A 2021, Series B 2023).'),
        ('p','Management is seeking a strategic partner to (a) accelerate cross-ASEAN expansion, (b) deepen enterprise customer penetration, and (c) provide founder liquidity. Open to majority or 100% buyout structures.'),
        ('h2','2. Business Overview'),
        ('p','Three business lines — Industrial Automation Solutions (54% rev), Digital Operations Platform (28%), Managed Services (18%). All three positive contribution; Platform fastest-growing (+38% YoY).'),
        ('p','Customer base: 86 enterprise accounts. Top-3 concentration 31% (up from 22% in FY2023 — a trend the buyer should evaluate).'),
        ('h2','3. Market & Strategy'),
        ('p','ASEAN industrial digitalisation TAM estimated USD 4.2bn (2025), growing 11% CAGR through 2030. Target current market share ~3%, fragmented competitive landscape.'),
        ('p','Strategic priorities (next 36 months): (1) win 4 Top-30 enterprise logos per year, (2) expand into Vietnam + Thailand, (3) deepen platform R&D investment, (4) build verticalised offerings for banking, healthcare, plantation.'),
        ('h2','4. Financial Highlights'),
        ('p','Revenue CAGR FY23-25: 11.8%. EBITDA margin expanded from 11.3% to 13.8% over 2 years. ROIC FY25 14.2%. Strong cash conversion (FCF / EBITDA ~26%).'),
        ('p','FY2026 management plan: Revenue MYR 670M (+11%), EBITDA MYR 96M (14.3% margin), based on signed pipeline of MYR 142M and historical conversion rates.'),
        ('h2','5. Synergies (Indicative, Buyer to Validate)'),
        ('p','Revenue: cross-sell to Zava\'s banking + healthcare divisions — estimated MYR 30-50M incremental revenue over 3 years.'),
        ('p','Cost: shared services consolidation (HR, Finance, IT) — MYR 4-6M run-rate savings.'),
        ('p','Capital: leverage Zava\'s balance sheet to accelerate inorganic moves.'),
        ('h2','6. Deal Considerations'),
        ('p','Valuation indication: EV MYR 720-780M implying 8.5-9.0x FY25 adjusted EBITDA. Founder rollover available up to 25%. Indemnity package: 12 months G&W, 18 months tax, 60 months fundamentals.'),
        ('p','Conditions precedent expected: regulator approval (3 of 4 jurisdictions), customer #1 consent (change-of-control), lender consents (senior facility), waiver of preferred share liquidation preference.'),
        ('h2','7. Risks (Management View)'),
        ('p','Customer #1 termination-on-CoC clause. Top-3 concentration trend. Series B preference (full ratchet anti-dilution). Senior facility CoC clause. CEO + CTO retention post-deal.'),
        ('h2','8. Suggested Timeline'),
        ('p','Q2 2026 — Confirmatory DD; Q3 2026 — Negotiation + SPA; Q4 2026 — Regulator filings; Q1 2027 — Closing.'),
    ])


# ─── Archetype 5: Board pre-read ───
def gen_board_prereread():
    d = OUT / 'nb-board-prereread'
    d.mkdir(parents=True, exist_ok=True)

    # CEO_Report.docx
    doc = _docx_init('Group CEO Report to the Board', 'Zava Conglomerate · Q1 FY2026 (Jan-Mar 2026) · Board Meeting 2026-05-22')
    _docx_h2(doc, '1. Headline Performance')
    _docx_p(doc, 'Group revenue MYR 11.2bn (+8% YoY, +2% vs plan). Group EBITDA MYR 1.74bn (15.5% margin, -40bps vs plan). PATAMI MYR 612m. Free cash flow MYR 480m. Net debt / EBITDA stable at 2.1x.')
    _docx_p(doc, 'Three of eleven divisions are tracking below plan (Banking — credit cycle headwind; Plantation — fresh-fruit-bunch price softness; Property — slow take-up). Other eight divisions on or ahead of plan.')
    _docx_h2(doc, '2. Strategic Highlights')
    _docx_bullets(doc, [
        'Project Hawthorn (M&A target — industrial digital solutions) progressed to confirmatory DD. Indicative EV MYR 720-780M.',
        'Group Operational Resilience Framework draft tabled to Risk Committee on 2026-04-18; Board approval requested at this meeting (Item 7).',
        'Sustainability — Scope-1 + Scope-2 emissions reduced 4.2% YoY; Scope-3 measurement scope expanded to 11 of 11 divisions.',
        'Digital — 24-workload cloud RFP shortlist tabled (Item 9). Decision target Q3 2026.',
        '3 senior leadership appointments made — Group CTO, Head of Cyber, Head of Cross-Border M&A.',
    ])
    _docx_h2(doc, '3. Material Risks')
    _docx_bullets(doc, [
        'Regulatory — Op Resilience Circular 2026 effective Oct 2026; remediation plan being executed (3 of 12 obligations not yet started).',
        'Geo-political — heightened cross-border tax scrutiny; transfer-pricing documentation refresh underway.',
        'People — voluntary attrition at 14.2% Group-wide (above target 12%); retention programme refreshed.',
        'Climate — physical-risk assessment for 7 plantation estates flagged Amber for flood/extreme-weather; mitigation plan in train.',
    ])
    _docx_h2(doc, '4. Asks of the Board')
    _docx_bullets(doc, [
        'Approval — Group Operational Resilience Framework (Item 7)',
        'Approval — Cloud Modernisation vendor selection (Item 9)',
        'Approval — FY2026 Half-Year Dividend (Item 12)',
        'Discussion — Project Hawthorn Q3/Q4 decision points',
    ])
    _docx_h2(doc, '5. Q1 KPI Scorecard')
    _docx_table(doc,
        ['Metric','Actual','Plan','Status','Comment'],
        [['Revenue (MYR bn)','11.2','11.0','Green','+2%'],
         ['EBITDA (MYR bn)','1.74','1.81','Amber','-4%, margin compression'],
         ['PATAMI (MYR m)','612','640','Amber','-4%'],
         ['Net debt / EBITDA','2.1x','2.0x','Amber','Within appetite'],
         ['Voluntary attrition','14.2%','12.0%','Red','—'],
         ['LTIFR','1.18','1.50','Green','Below target — improving'],
         ['NPS','+38','+35','Green','Improving'],
         ['Scope 1+2 emissions ('+'tCO2e'+')','-4.2% YoY','-5% YoY','Amber','On track'],
         ['Cyber maturity score','3.4 / 5','3.5 / 5','Amber','Close to target']])
    _docx_p(doc, 'Submitted by: Hadar Caspit, Group CFO (on behalf of Group CEO) · 2026-05-15', italic=True)
    doc.save(str(d / 'CEO_Report.docx'))

    # CFO_Financial_Pack.xlsx — multi-tab
    wb = _xl_new(d / 'CFO_Financial_Pack.xlsx')
    ws = wb.active; ws.title = 'Group P&L'
    _xl_write_sheet(ws,
        ['MYR M','Q1 FY26 Actual','Q1 FY26 Plan','Var %','FY26 YTD','FY25 YTD','YoY %'],
        [
            ['Revenue','11,238','11,012','2.1%','11,238','10,402','8.0%'],
            ['Cost of sales','(7,948)','(7,712)','3.1%','(7,948)','(7,310)','8.7%'],
            ['Gross profit','3,290','3,300','-0.3%','3,290','3,092','6.4%'],
            ['SG&A','(1,552)','(1,490)','4.2%','(1,552)','(1,438)','7.9%'],
            ['EBITDA','1,738','1,810','-4.0%','1,738','1,654','5.1%'],
            ['D&A','(412)','(420)','-1.9%','(412)','(388)','6.2%'],
            ['EBIT','1,326','1,390','-4.6%','1,326','1,266','4.7%'],
            ['Net interest','(298)','(286)','4.2%','(298)','(272)','9.6%'],
            ['Tax','(248)','(265)','-6.4%','(248)','(238)','4.2%'],
            ['PATAMI','612','645','-5.1%','612','584','4.8%'],
            ['EBITDA margin %','15.5%','16.4%','-90bps','15.5%','15.9%','-40bps'],
        ],
        widths=[28,16,16,10,14,14,10])

    # Division-level
    ws2 = wb.create_sheet('By Division')
    divs = [
        ['Banking','3,420','3,580','-4.5%','520','590','-12%','Credit cycle'],
        ['Insurance','1,820','1,800','+1.1%','264','252','+5%','Higher claims under control'],
        ['Healthcare','1,140','1,120','+1.8%','188','176','+7%','Outpatient growth'],
        ['Plantation','842','910','-7.5%','142','188','-24%','FFB price down 18%'],
        ['Property','620','680','-8.8%','82','118','-31%','Take-up slow'],
        ['O&G Downstream','1,180','1,140','+3.5%','164','144','+14%','Refining margin up'],
        ['Utilities','920','905','+1.7%','198','188','+5%','Stable'],
        ['Telco','580','578','+0.3%','98','94','+4%','In line'],
        ['Retail','410','395','+3.8%','58','54','+7%','Festive period strong'],
        ['Industrial Manufacturing','220','218','+0.9%','38','34','+12%','Order book healthy'],
        ['Hospitality','86','86','flat','-14','-8','—','Seasonal pre-opening costs'],
    ]
    _xl_write_sheet(ws2,
        ['Division','Revenue Q1','Plan Q1','Rev Var','EBITDA Q1','EBITDA Q1 Plan','EBITDA Var','Commentary'],
        divs, widths=[28,12,12,10,12,16,14,38])

    # Variance
    ws3 = wb.create_sheet('Variance Drivers')
    _xl_write_sheet(ws3,
        ['Driver','Impact MYR M','Direction','Affected Divisions','Action / Owner'],
        [
            ['FFB price softness','-46','Negative','Plantation','Hedge programme — Treasury'],
            ['Banking credit provisions','-70','Negative','Banking','Risk appetite review — Group Risk'],
            ['Property handover delay','-36','Negative','Property','Customer comms + handover ramp — COO'],
            ['Refining margin upside','+20','Positive','O&G Downstream','Lock-in via term — Strategy'],
            ['Healthcare outpatient volume','+12','Positive','Healthcare','Capacity expansion — COO'],
            ['Telco postpaid mix','+4','Positive','Telco','Continue — Sales'],
            ['Working capital improvement','+18','Positive','Group','Sustain — Group Treasurer'],
            ['Energy cost inflation','-22','Negative','Utilities + Manufacturing','Pass-through review — Strategy'],
        ],
        widths=[34,14,12,28,30])

    # Balance Sheet
    ws4 = wb.create_sheet('Balance Sheet')
    _xl_write_sheet(ws4,
        ['MYR M','Q1 FY26','FY25 Close','Change'],
        [['Cash & equivalents','3,820','3,540','+280'],
         ['Net working capital','5,460','5,148','+312'],
         ['PP&E (net)','24,628','24,182','+446'],
         ['Goodwill & intangibles','3,840','3,840','—'],
         ['Total assets','38,946','37,998','+948'],
         ['Total debt','11,420','11,180','+240'],
         ['Net debt','7,600','7,640','-40'],
         ['Total equity','19,840','19,228','+612'],
         ['Net debt / EBITDA','2.1x','2.1x','flat']],
        widths=[28,12,12,12])

    # Liquidity
    ws5 = wb.create_sheet('Liquidity & Covenants')
    _xl_write_sheet(ws5,
        ['Facility','Limit MYR M','Drawn MYR M','Headroom MYR M','Covenant','Limit','Current','Status'],
        [['RCF Tranche A','3,000','1,680','1,320','Net Debt/EBITDA','3.0x','2.1x','Pass'],
         ['RCF Tranche B','2,000','920','1,080','Interest Cover','4.0x','5.8x','Pass'],
         ['Working Capital','1,500','680','820','Min Liquidity','MYR 2,000M','MYR 3,820M','Pass'],
         ['Term Loan A','4,000','4,000','—','Min Equity','MYR 15,000M','MYR 19,840M','Pass'],
         ['Term Loan B','2,500','2,500','—','—','—','—','—']],
        widths=[24,12,12,14,22,12,18,10])
    _xl_save(wb, d / 'CFO_Financial_Pack.xlsx')

    # Risk_Heatmap.xlsx
    wb = _xl_new(d / 'Risk_Heatmap.xlsx')
    ws = wb.active; ws.title = 'Heat Map'
    risks = [
        ['R-01','Operational Resilience','Failure to comply with Circular 2026 by Oct','GRO','High','4','5','20','Mitigating','Aligned remediation plan; Board approval requested'],
        ['R-02','Credit','Banking division — corporate NPL uptick','Banking CEO','High','4','5','20','Mitigating','Loss given default review; appetite refresh'],
        ['R-03','Cyber','Ransomware attack via third-party','CISO','Medium','3','5','15','Mitigating','MFA enforcement complete; vendor consolidation in progress'],
        ['R-04','Climate','Physical risk to plantation estates','Plantation CEO','Medium','3','4','12','Mitigating','Flood-proofing capex; insurance renewed'],
        ['R-05','Talent','Attrition above appetite','CHRO','Medium','4','3','12','Mitigating','Retention programme rolled out'],
        ['R-06','Regulatory','Cross-border tax review','CFO','Medium','3','4','12','Mitigating','TP documentation refresh underway'],
        ['R-07','Strategic','Customer concentration in target (M&A)','Strategy','Low','2','4','8','Watch','DD addressing'],
        ['R-08','Market','Commodity price (FFB, oil)','CFO','Medium','5','2','10','Watch','Hedging programme active'],
        ['R-09','Reputational','ESG disclosure quality','Head of IR','Low','2','3','6','Watch','Framework alignment in train'],
        ['R-10','Liquidity','Refinancing of term loan B (2027)','Treasurer','Low','2','3','6','Watch','Refi plan started'],
    ]
    _xl_write_sheet(ws,
        ['Risk ID','Category','Description','Owner','Severity','Likelihood','Impact','Score','Status','Mitigation'],
        risks, widths=[10,18,40,18,10,12,10,8,12,44],
        rag_col=8, rag_map=lambda r: 'R' if int(r[7])>=15 else 'A' if int(r[7])>=10 else 'G')

    ws2 = wb.create_sheet('Top-5 Movement')
    _xl_write_sheet(ws2,
        ['Risk','Prior Quarter Score','Current Score','Direction','Trigger'],
        [
            ['R-01 Op Resilience','16','20','↑','Circular issued — clock started'],
            ['R-02 Banking NPL','12','20','↑','Q1 NPL formation higher than expected'],
            ['R-03 Cyber','15','15','→','Stable — controls tightening'],
            ['R-05 Attrition','10','12','↑','New cycle of voluntary exits in Q1'],
            ['R-06 Tax review','9','12','↑','Receipt of inquiry letter'],
        ],
        widths=[28,18,16,12,40])
    _xl_save(wb, d / 'Risk_Heatmap.xlsx')

    # Strategy_Update_Deck.pdf
    _pdf(d / 'Strategy_Update_Deck.pdf', 'Strategy Update Deck', [
        ('h1','Strategy Update — Q1 FY2026'),
        ('small','Group Strategy Office · Zava Conglomerate · For Board Discussion'),
        ('h2','1. Zava Forward 2030 — Progress'),
        ('p','Pillar 1 — Customer-Obsessed Growth: NPS +38 (vs +35 target); Top-30 enterprise win-rate 28% (vs 25% target).'),
        ('p','Pillar 2 — Capital Productivity: ROIC 12.4% (vs 12% target); WACC 8.4%; ROIC-WACC spread 400bps.'),
        ('p','Pillar 3 — Operational Excellence: 4.2% emissions reduction; LTIFR 1.18; cost-to-income improved 60bps.'),
        ('p','Pillar 4 — Digital & Data: 12 of 24 priority workloads migrated to cloud; data fabric live in 4 of 11 divisions.'),
        ('p','Pillar 5 — Sustainability & Trust: TCFD reporting expanded; Scope-3 measurement now 11 of 11 divisions.'),
        ('h2','2. Capital Allocation Update'),
        ('p','FY2026 capex envelope: MYR 4.8bn. Committed YTD: MYR 1.1bn. M&A pipeline value: MYR 1.4bn (Hawthorn + 2 smaller). Dividend MYR 1.6bn declared; payout ratio 38%.'),
        ('h2','3. Project Hawthorn (M&A)'),
        ('p','Industrial digital solutions target. Indicative EV MYR 720-780M (8.5-9.0x adj EBITDA). Strategic fit High — cross-sell into Banking + Healthcare + Plantation digital programmes.'),
        ('p','Synergies: revenue MYR 30-50M / 3yr; cost MYR 4-6M run-rate. Risk: customer concentration trend, change-of-control clauses, founder retention.'),
        ('h2','4. Operating Plan Adjustments'),
        ('p','Plantation guidance revised — FFB headwind to persist into H2. Property guidance held — H2 launches expected to recover.'),
        ('p','Banking provisioning recalibrated upward MYR 70M; capital remains comfortably above regulatory minima.'),
        ('h2','5. Recommendations'),
        ('p','(1) Approve Op Resilience Framework. (2) Approve cloud vendor recommendation. (3) Endorse Hawthorn progression to definitive agreement workstream. (4) Note guidance reset.'),
    ])

    # Previous_Board_Minutes.docx
    doc = _docx_init('Previous Board Meeting — Minutes', 'Zava Conglomerate · Q4 FY2025 Board Meeting · 2026-02-20 · CONFIDENTIAL')
    _docx_h2(doc, 'Attendees')
    _docx_bullets(doc, ['Chair: Independent Director','CEO: Group CEO','CFO: Hadar Caspit','Group Chief of Staff: Sasha Ouellet',
                        '7 other Independent Directors','Group Risk Officer (item 6 only)','Head of Investor Relations: Daichi Maruyama'])
    _docx_h2(doc, 'Decisions Carried')
    _docx_table(doc,
        ['Item','Decision','Status as at Q1 FY26'],
        [['FY25 audited accounts','Approved','Filed'],
         ['Final dividend declaration','Approved MYR 1.6bn','Paid'],
         ['Internal audit plan FY26','Approved','In execution'],
         ['Risk appetite — refresh','Approved','In force'],
         ['Op Resilience Framework v0.9','Tabled, asked for management to return with v1.0','Returning at this meeting (Item 7)'],
         ['Cloud Modernisation procurement','Tabled — asked for shortlist','Returning at this meeting (Item 9)']])
    _docx_h2(doc, 'Key Discussion Points')
    _docx_p(doc, 'Banking — Board sought reassurance on credit cycle. CRO walked through stress test results; Board comfortable.')
    _docx_p(doc, 'M&A pipeline — Board requested clearer synergy quantification on Hawthorn before progressing.')
    _docx_p(doc, 'Cybersecurity — Board endorsed accelerated investment plan; CISO to present roadmap at Q3 meeting.')
    _docx_p(doc, 'Climate — Board requested next pre-read include physical-risk assessment summary for all divisions.')
    _docx_h2(doc, 'Items Outstanding for This Meeting')
    _docx_bullets(doc, [
        'Op Resilience Framework — Board approval requested',
        'Cloud vendor selection — Board approval requested',
        'Half-year dividend — Board approval requested',
        'Hawthorn synergy quantification — addressed in Strategy Update Deck',
        'Physical-risk assessment summary — included in CEO Report Section 3',
    ])
    doc.save(str(d / 'Previous_Board_Minutes.docx'))


# ─── Archetype 6: Renewal book ───
def gen_renewal_book():
    d = OUT / 'nb-renewal-book'
    d.mkdir(parents=True, exist_ok=True)

    # Account_Plan.docx
    doc = _docx_init('Account Plan — Apex Industries', 'Zava Conglomerate · Renewal cycle FY2026 · Customer Success · CONFIDENTIAL')
    _docx_h2(doc, '1. Account Profile')
    _docx_table(doc, ['Field','Value'],
        [['Customer','Apex Industries Sdn Bhd'],['Industry','Industrial manufacturing'],
         ['HQ','Selangor, Malaysia'],['Annual revenue (customer)','~MYR 1.2bn'],
         ['Employees','3,400'],['Contract value (ours, FY25)','MYR 4.6M'],
         ['Renewal date','2026-09-30'],['Account tenure','4 years (since 2022)'],
         ['Sponsor','VP Operations'],['Day-to-day owner','IT Director'],
         ['Champion','Head of Plant Engineering']])
    _docx_h2(doc, '2. Subscription & Services')
    _docx_bullets(doc, [
        'Core Platform — 1,200 seats (currently using 1,054)',
        'Analytics Add-on — 1,200 seats (currently using 642)',
        'Premium Support — 24x7',
        'Professional Services — 240 hours / year retainer',
        'Implementation services completed 2022',
    ])
    _docx_h2(doc, '3. Health Signals')
    _docx_bullets(doc, [
        'Adoption — core platform usage healthy (88%); analytics add-on low (54%)',
        'Stickiness — 3 of 7 key workflows automated end-to-end',
        'Support — 14 tickets last 12m; 1 P1 (resolved within SLA)',
        'Champion — Head of Plant Engineering is renewing internal advocacy',
        'Sponsor risk — VP Operations role under change; new VP joining 2026-06',
    ])
    _docx_h2(doc, '4. Renewal Strategy')
    _docx_p(doc, '4.1 Target — 3-year renewal, 8% net price increase, expand analytics add-on to 800 active users (currently 642).')
    _docx_p(doc, '4.2 Sales motion — VP-to-VP sponsor reset with new VP Operations in July; champion-led usage review; ROI memo from analytics rollout.')
    _docx_p(doc, '4.3 Concessions framework — accept 2-year renewal if customer pushes back on 3-year; accept 4% net increase if multi-year secured; accept seat true-up at renewal not at signing.')
    _docx_p(doc, '4.4 Walk-away — annual roll-over at 4% net increase; refuse below 0% net increase regardless of term.')
    _docx_h2(doc, '5. Risks')
    _docx_bullets(doc, [
        'New VP unknown — may want to evaluate alternatives',
        'Competitor X is actively targeting; placed pitch June 2025 (we won then)',
        'Analytics add-on underutilised — value justification harder',
        'Industry capex cycle softening — customer may push for price cut',
    ])
    _docx_h2(doc, '6. Action Plan (next 90 days)')
    _docx_bullets(doc, [
        'Week 1-2: champion alignment meeting; usage review prepared',
        'Week 3-4: ROI memo drafted with quantified business outcomes',
        'Week 5-6: sponsor reset meeting with new VP (target before 2026-07-15)',
        'Week 7-8: term-sheet exchange',
        'Week 9-12: legal + commercial negotiation, signed by 2026-09-15',
    ])
    doc.save(str(d / 'Account_Plan.docx'))

    # Usage_Telemetry_Last_12M.xlsx
    wb = _xl_new(d / 'Usage_Telemetry_Last_12M.xlsx')
    ws = wb.active; ws.title = 'Monthly Usage'
    months = ['2025-05','2025-06','2025-07','2025-08','2025-09','2025-10','2025-11','2025-12','2026-01','2026-02','2026-03','2026-04']
    cu = [820, 860, 890, 910, 940, 980, 1010, 1024, 1018, 1042, 1054, 1054]
    au = [380, 420, 460, 520, 560, 580, 610, 622, 618, 632, 642, 642]
    rows = []
    for i, m in enumerate(months):
        rows.append([m, cu[i], 1200, f'{cu[i]/1200:.0%}', au[i], 1200, f'{au[i]/1200:.0%}'])
    _xl_write_sheet(ws,
        ['Month','Core Active Users','Core Seats Licensed','Core Adoption %','Analytics Active','Analytics Seats','Analytics Adoption %'],
        rows, widths=[14,18,20,16,16,16,20])

    ws2 = wb.create_sheet('Feature Usage')
    _xl_write_sheet(ws2,
        ['Feature','Monthly Active (last 30d)','Usage Trend','Note'],
        [['Dashboard',968,'↑ 12% YoY','Workhorse — heavy use'],
         ['Workflow Builder',542,'↑ 28% YoY','Strong expansion'],
         ['Reporting Library',418,'→ flat','Stable'],
         ['Alerts & Triggers',312,'↑ 18% YoY','Growing'],
         ['Mobile App',186,'↑ 8% YoY','Field operations'],
         ['API Integrations',128,'↑ 45% YoY','Power-user growth'],
         ['ML Forecasting',64,'→ flat','Analytics add-on — under-used'],
         ['Anomaly Detection',82,'↑ 22% YoY','Analytics add-on — growing'],
         ['Custom Dashboards',218,'↑ 14% YoY','—'],
         ['Audit Trail',180,'→ flat','Compliance use']],
        widths=[28,22,18,40])

    ws3 = wb.create_sheet('Workflow Automation')
    _xl_write_sheet(ws3,
        ['Workflow','Status','Volume / month','Time Saved (est, hrs/month)','Value'],
        [['Plant downtime alerts','Live','1,420','62','High'],
         ['Maintenance work-order auto-create','Live','840','38','High'],
         ['Quality NCR routing','Live','312','24','Medium'],
         ['Supplier delivery alerts','Pilot','—','—','Pending'],
         ['Energy cost variance','Pilot','—','—','Pending'],
         ['Inventory replenishment','Not started','—','—','Opportunity'],
         ['HSE incident triage','Not started','—','—','Opportunity']],
        widths=[34,12,18,22,12])
    _xl_save(wb, d / 'Usage_Telemetry_Last_12M.xlsx')

    # Support_Tickets_Log.xlsx
    wb = _xl_new(d / 'Support_Tickets_Log.xlsx')
    ws = wb.active; ws.title = 'Tickets'
    tix = [
        ['T-2025-301','2025-06-12','P3','Performance','Dashboard slow under load','Resolved','3 days','Optimisation applied'],
        ['T-2025-318','2025-07-04','P2','Bug','Mobile app crash on iOS 17','Resolved','2 days','Patch released'],
        ['T-2025-322','2025-07-21','P3','How-to','Custom dashboard config','Resolved','1 day','Training delivered'],
        ['T-2025-359','2025-09-08','P3','Feature request','Add SAP integration','In backlog','—','Roadmap Q3 FY26'],
        ['T-2025-401','2025-10-14','P2','Bug','Alert delivery delay','Resolved','4 days','Queue scaling fix'],
        ['T-2025-428','2025-11-22','P3','How-to','API rate limit increase','Resolved','1 day','Configured'],
        ['T-2025-444','2025-12-03','P2','Bug','Workflow Builder save error','Resolved','3 days','Patch released'],
        ['T-2026-008','2026-01-12','P1','Outage','Dashboard unavailable 38 min','Resolved','same day','RCA delivered; service credit applied'],
        ['T-2026-024','2026-02-04','P3','How-to','SSO config refresh','Resolved','1 day','—'],
        ['T-2026-035','2026-02-21','P3','Feature request','SSO with custom IdP','In backlog','—','Roadmap Q4 FY26'],
        ['T-2026-052','2026-03-14','P2','Bug','Alert false-positive rate','In progress','—','Investigating'],
        ['T-2026-068','2026-04-02','P3','How-to','Audit-log export format','Resolved','1 day','—'],
        ['T-2026-081','2026-04-22','P3','Feature request','Anomaly detection sensitivity','In backlog','—','Roadmap Q3 FY26'],
        ['T-2026-094','2026-05-04','P3','How-to','ML model retraining cadence','Resolved','1 day','Documented'],
    ]
    _xl_write_sheet(ws,
        ['Ticket','Opened','Priority','Category','Summary','Status','TTR','Notes'],
        tix, widths=[12,12,8,16,40,14,8,38])

    ws2 = wb.create_sheet('Trends')
    _xl_write_sheet(ws2,
        ['Metric','Last 12m','Prior 12m','Trend'],
        [['Total tickets','14','17','↓ Healthy'],
         ['P1 incidents','1','0','↑ One outage'],
         ['P2 incidents','4','5','↓ Improving'],
         ['Avg time to resolve P1+P2','2.4 days','2.8 days','↓ Improving'],
         ['Feature requests','3','2','↑ Healthy backlog growth'],
         ['SLA compliance','100%','94%','↑ Improved']],
        widths=[34,12,12,18])
    _xl_save(wb, d / 'Support_Tickets_Log.xlsx')

    # Renewal_Contract_Draft.pdf
    _pdf(d / 'Renewal_Contract_Draft.pdf', 'Renewal Contract Draft', [
        ('h1','Renewal Master Subscription Agreement — DRAFT'),
        ('small','Customer: Apex Industries Sdn Bhd · Effective Date: 2026-10-01 · Term: 36 months · DRAFT FOR DISCUSSION'),
        ('h2','1. Scope of Service'),
        ('p','Vendor shall provide the Subscription Services as set out in Schedule 1 (Service Description) and the Professional Services Hours as set out in Schedule 2. Software-as-a-Service delivery model on Vendor\'s multi-tenant cloud.'),
        ('h2','2. Subscription Fees'),
        ('p','2.1 Annual Subscription Fee — Year 1: MYR 4,968,000. Year 2: MYR 5,366,000 (8% increase). Year 3: MYR 5,795,000 (8% increase).'),
        ('p','2.2 Seat counts as per Schedule 1. Customer may add seats mid-term at the same per-seat price (true-up at renewal). Mid-term seat reductions not permitted.'),
        ('p','2.3 Invoicing — annually in advance. Payment terms net 45 days. Currency MYR.'),
        ('h2','3. Service Levels'),
        ('p','3.1 Availability — 99.95% per calendar month, measured per workload.'),
        ('p','3.2 Support — Premium 24x7. P1 response 15 min, P2 30 min, P3 4 hours.'),
        ('p','3.3 Service credits — 10% of monthly fees per 0.1% below SLA, capped at 25% of monthly fees per month.'),
        ('h2','4. Term & Termination'),
        ('p','4.1 Initial term 36 months from Effective Date.'),
        ('p','4.2 Termination for convenience — customer may terminate with 12 months\' written notice; no refund of pre-paid annual fees.'),
        ('p','4.3 Termination for cause — material breach not cured within 30 days; refund pro-rata of pre-paid unused term.'),
        ('h2','5. Data, Privacy & Security'),
        ('p','5.1 Customer data remains the property of Customer. Vendor processes data only as Customer\'s data processor.'),
        ('p','5.2 Data residency — ASEAN region. Vendor warrants compliance with PDPA 2010 (Malaysia).'),
        ('p','5.3 Security certifications — ISO 27001, SOC 2 Type II maintained throughout term.'),
        ('h2','6. Outstanding / To-Be-Discussed'),
        ('p','• Customer requesting 2-year initial term (instead of 3) — DISCUSSION POINT'),
        ('p','• Customer requesting 4% (instead of 8%) annual escalator — DISCUSSION POINT'),
        ('p','• Customer requesting mid-term seat REDUCTIONS allowed — DISCUSSION POINT'),
        ('p','• Vendor proposing seat true-up at renewal — pending Customer acceptance'),
        ('p','• Auto-renewal clause — to be added if 3-year term accepted'),
        ('h2','7. Signatures (placeholder)'),
        ('p','Customer: VP Operations, Apex Industries Sdn Bhd · Vendor: Account Executive, Zava Group · Witnessed by Legal'),
    ])

    # Competitor_Win_Loss_Notes.docx
    doc = _docx_init('Competitor Win/Loss Notes', 'Zava Conglomerate · Sales Intelligence · Last 12 months')
    _docx_h2(doc, '1. Competitor X — Recent Activity at Apex')
    _docx_p(doc, 'June 2025: Competitor X formal pitch at Apex. We won; primary differentiator was depth of plant-engineering workflow templates and 4-year operational history together.')
    _docx_p(doc, 'October 2025: Competitor X re-engaged on analytics add-on (low usage was their angle of attack). We retained via expansion ROI conversation.')
    _docx_p(doc, 'February 2026: Competitor X published case study with a similar industrial manufacturer. Champion forwarded it to us for our response.')
    _docx_h2(doc, '2. Competitor X — Strengths to Counter')
    _docx_bullets(doc, [
        'Aggressive 3-year price commitments with 0% escalators',
        'Industry-vertical solution suites with packaged workflows',
        'Strong analytics ML capability — they have outpaced us in anomaly detection',
        'Reference customers in same industry segment in Indonesia and Thailand',
    ])
    _docx_h2(doc, '3. Competitor X — Weaknesses to Press')
    _docx_bullets(doc, [
        'No native mobile app — only web responsive',
        'Implementation timelines longer (typically 9-12 months vs our 4-6)',
        'Support model only 12x5 business hours; no 24x7 Premium tier',
        'No PDPA-Malaysia data-residency commitment in their standard contract',
        'Two recent ASEAN customer losses to us (2024-Q3 and 2025-Q1)',
    ])
    _docx_h2(doc, '4. Other Competitor Activity')
    _docx_p(doc, 'Competitor Y — niche analytics player. Saw activity at 2 of our other manufacturing accounts. Not yet seen at Apex.')
    _docx_p(doc, 'Competitor Z — large global player. Bundle into ERP. Apex evaluated and rejected in 2022. Not currently active.')
    _docx_h2(doc, '5. Reference Cases to Use')
    _docx_bullets(doc, [
        'Pacific Industrial Co — 18 months, 28% TCO reduction',
        'Asianova Manufacturing — 24 months, 32% downtime reduction',
        'Northwind Production — 12 months, 1,200 hours/year saved across maintenance team',
    ])
    doc.save(str(d / 'Competitor_Win_Loss_Notes.docx'))


# ─── Archetype 7: Incident post-mortem ───
def gen_incident_pmortem():
    d = OUT / 'nb-incident-pmortem'
    d.mkdir(parents=True, exist_ok=True)

    # Incident_Timeline.docx
    doc = _docx_init('Incident Timeline — INC-2026-0418-001', 'Zava Digital Operations · Severity: P1 (Major) · Customer-Impacting · Date: 2026-04-18')
    _docx_h2(doc, 'Incident Overview')
    _docx_p(doc, 'Customer-facing digital banking platform experienced authentication failures and slow transaction confirmations for approximately 42 minutes between 14:18 and 15:00 MYT on 2026-04-18.')
    _docx_p(doc, 'Estimated impact: ~62,000 active users affected; ~14,200 failed authentications; ~3,400 delayed transactions. No data loss. No financial impact to customer accounts. Regulator notified 16:08 (within 4h window).')
    _docx_h2(doc, 'Timeline (all times MYT)')
    _docx_table(doc,
        ['Time','Event','Source','Action'],
        [['13:42','Deployment of release v2026.04.18 to production','Release pipeline','Routine — passed all stage gates'],
         ['14:18','First customer complaint (Twitter)','Social listening','Triaged as anomaly'],
         ['14:20','Authentication service latency alert','Pager','On-call paged'],
         ['14:22','On-call engineer acknowledges page','PagerDuty','—'],
         ['14:25','Incident channel opened — Severity P2','Teams','IM declared'],
         ['14:28','Authentication queue saturation confirmed','Grafana','Dashboard'],
         ['14:31','Customer support tickets spike (38 in 3 min)','Service Desk','Escalation'],
         ['14:32','Severity escalated to P1','Teams','IM upgrades'],
         ['14:35','Customer Comms team activated','Teams','Drafting status page'],
         ['14:36','Decision: rollback v2026.04.18','IC + Tech Lead','Rollback initiated'],
         ['14:38','Status page updated — Investigating','Atlassian Statuspage','—'],
         ['14:42','Rollback initiated','Release pipeline','—'],
         ['14:54','Rollback complete in 3 of 4 regions','Release pipeline','SG, MY, ID rolled back'],
         ['14:58','Fourth region (PH) rolled back','Release pipeline','—'],
         ['15:00','Authentication service nominal','Monitoring','Recovery confirmed'],
         ['15:02','Status page updated — Monitoring','Atlassian Statuspage','—'],
         ['15:38','Status page updated — Resolved','Atlassian Statuspage','—'],
         ['16:08','Regulator notification submitted','Compliance','Filed within 4h window'],
         ['18:00','War room debrief','Teams','Action items captured']])
    _docx_p(doc, 'Incident Commander: Mod Admin (Group Digital Operations) · Tech Lead: Engineering Lead · Communications Lead: Sasha Ouellet (Group CoS)', italic=True)
    doc.save(str(d / 'Incident_Timeline.docx'))

    # Pager_Alerts_Log.xlsx
    wb = _xl_new(d / 'Pager_Alerts_Log.xlsx')
    ws = wb.active; ws.title = 'Alerts'
    al = [
        ['14:20:14','HIGH','auth-service: P95 latency > 2000ms','Critical','On-call paged','Acknowledged 14:22'],
        ['14:21:08','HIGH','auth-service: error rate > 5%','Critical','—','—'],
        ['14:22:31','MEDIUM','session-store: connection pool saturated','Warning','—','—'],
        ['14:24:02','HIGH','transaction-service: P99 latency > 8000ms','Critical','—','—'],
        ['14:24:47','MEDIUM','session-store: queue depth > 8000','Warning','—','—'],
        ['14:28:14','HIGH','auth-service: queue depth > 12000','Critical','—','—'],
        ['14:33:01','MEDIUM','api-gateway: 5xx rate > 2%','Warning','—','—'],
        ['14:36:28','HIGH','release-pipeline: rollback initiated','Critical','—','By Tech Lead'],
        ['14:42:14','MEDIUM','release-pipeline: deploy-job rollback-progress','Warning','—','—'],
        ['14:54:01','LOW','release-pipeline: 3/4 regions complete','Info','—','—'],
        ['14:58:38','LOW','release-pipeline: PH region complete','Info','—','—'],
        ['15:00:12','LOW','auth-service: latency back to nominal','Info','—','Recovery'],
        ['15:01:48','LOW','transaction-service: error rate normal','Info','—','—']
    ]
    _xl_write_sheet(ws,
        ['Time','Severity','Alert','Type','Action','Notes'],
        al, widths=[12,12,52,12,18,20])

    ws2 = wb.create_sheet('Service Health')
    _xl_write_sheet(ws2,
        ['Service','Pre-Incident (13:00-14:15)','During (14:18-15:00)','Post-Recovery (15:00-16:00)'],
        [['Auth — P95 latency','280ms','3,420ms','310ms'],
         ['Auth — error rate','0.02%','12.4%','0.04%'],
         ['Transaction — P99 latency','1,200ms','8,800ms','1,250ms'],
         ['Transaction — success rate','99.96%','86.8%','99.94%'],
         ['Session store — queue depth','120','12,400','180'],
         ['API gateway — 5xx %','0.01%','3.4%','0.02%']],
        widths=[28,30,30,30])

    ws3 = wb.create_sheet('Customer Impact')
    _xl_write_sheet(ws3,
        ['Region','Affected Users','Failed Auths','Delayed Tx','Duration (min)'],
        [['Malaysia',32400,7300,1750,42],
         ['Singapore',12800,3200,720,38],
         ['Indonesia',14200,2900,720,42],
         ['Philippines',2600,800,210,40]],
        widths=[14,18,18,18,18])
    _xl_save(wb, d / 'Pager_Alerts_Log.xlsx')

    # Customer_Comms_Sent.docx
    doc = _docx_init('Customer Communications — Sent', 'Zava Digital Banking · Incident INC-2026-0418-001 · 2026-04-18')
    _docx_h2(doc, '1. Initial Status (14:38 MYT)')
    _docx_p(doc, '"We are aware of an issue affecting authentication on our digital banking platform. Our team is investigating. Affected services: login, transaction confirmation. We will provide an update within 30 minutes." — Status Page')
    _docx_h2(doc, '2. Mid-Incident Update (14:48 MYT)')
    _docx_p(doc, '"Update: we have identified the root cause and rolled back the recent deployment. Services are recovering. Estimated full recovery within 15 minutes." — Status Page + In-App banner')
    _docx_h2(doc, '3. Recovery Notification (15:02 MYT)')
    _docx_p(doc, '"Services have been restored. We are monitoring closely. We sincerely apologise for the inconvenience caused. A detailed post-incident review will be shared on our blog within 5 business days." — Status Page + Push notification')
    _docx_h2(doc, '4. Final Update (15:38 MYT)')
    _docx_p(doc, '"All services confirmed operational. The incident is now resolved. We thank you for your patience. — Customer Care Team"')
    _docx_h2(doc, '5. Customer-Facing Post-Incident Note (draft for blog publication)')
    _docx_p(doc, '"On 2026-04-18 between 14:18 and 15:00 MYT, our digital banking customers experienced authentication failures and slow transaction confirmations. The root cause was an unintended configuration change in a routine release; rolling back the release restored services. No customer data was lost or compromised, and no transactions were processed incorrectly. We are committed to learning from this incident and have implemented additional pre-deployment checks. We are also reviewing our communication speed to make sure customers hear from us within the first 10 minutes of any future incident. We apologise sincerely. — Zava Digital Banking, Chief Customer Officer."')
    doc.save(str(d / 'Customer_Comms_Sent.docx'))

    # System_Architecture_Diagram.pdf
    _pdf(d / 'System_Architecture_Diagram.pdf', 'System Architecture Diagram', [
        ('h1','Digital Banking — System Architecture Overview'),
        ('small','Zava Digital Operations · Architecture Reference Document · Confidential'),
        ('h2','1. Edge Layer'),
        ('p','CloudFront / Front Door edge → WAF → API Gateway. Region routing per geo. Static assets cached at edge.'),
        ('h2','2. API Gateway Layer'),
        ('p','API Gateway (regional). Rate limiting, JWT validation, request routing to downstream services. Tied to upstream timeouts.'),
        ('h2','3. Application Services'),
        ('p','Auth Service (stateful, Redis-backed session store, OIDC provider integration). Transaction Service (stateful, calls into core banking). Notification Service (Kafka-fed). Customer Profile Service. Reporting Service.'),
        ('h2','4. Data Tier'),
        ('p','Primary OLTP — Postgres clusters (regional). Session store — Redis (regional with cross-region async replication). Core banking — mainframe (system-of-record).'),
        ('h2','5. Messaging'),
        ('p','Kafka clusters for event streaming (notifications, audit log, fraud signals).'),
        ('h2','6. Observability'),
        ('p','Metrics: Prometheus → Grafana. Logs: ELK. Traces: OpenTelemetry. Synthetic monitoring + RUM (real user monitoring).'),
        ('h2','7. Critical Dependencies'),
        ('p','Auth Service ← Session Store ← OIDC Provider (3rd-party). Transaction Service ← Core Banking. Notification Service ← SMS gateway (3rd-party).'),
        ('h2','8. Failure Domain Notes'),
        ('p','Session Store is the historical bottleneck. Connection pool saturation has been observed in 2023 and 2024 under similar load patterns. Auto-scaling triggers exist but lag by 90 seconds — within that window, downstream auth queue saturates rapidly.'),
        ('h2','9. Recent Changes (Risk-Relevant, last 30 days)'),
        ('p','2026-04-08 — connection pool default sizing changed from 200 to 80 in a config refactor (intended for non-prod default; mistakenly inherited by prod via the release pipeline) — root cause of the 2026-04-18 incident.'),
        ('h2','10. Resilience Posture'),
        ('p','Multi-region (4 regions). Active-Active for read traffic; Active-Passive for writes. RTO target 15 min; RPO 1 minute.'),
    ])

    # Past_Similar_Incidents.xlsx
    wb = _xl_new(d / 'Past_Similar_Incidents.xlsx')
    ws = wb.active; ws.title = 'Past Incidents'
    pi = [
        ['INC-2024-0218-002','2024-02-18','P1','Auth latency from session-store pool saturation','62 min','Connection pool sized too low after refactor','Hotfix + pool sizing standard'],
        ['INC-2023-0912-004','2023-09-12','P2','Auth degraded during traffic spike','38 min','Slow Redis failover','Cross-region replication tuned'],
        ['INC-2025-0301-001','2025-03-01','P1','Transaction-service timeouts on core-banking call','46 min','Mainframe batch overlap','Batch window rescheduled'],
        ['INC-2025-0712-003','2025-07-12','P2','Notification delivery delay','22 min','Kafka consumer lag','Consumer scaling fixed'],
        ['INC-2026-0103-001','2026-01-03','P3','Status page delay','12 min','Status page CMS slow','Cache layer added'],
        ['INC-2026-0418-001','2026-04-18','P1','Auth queue saturation post-release','42 min','Connection pool config inherited from non-prod','THIS INCIDENT — RCA in progress'],
    ]
    _xl_write_sheet(ws,
        ['Incident ID','Date','Severity','Summary','Duration','Root Cause','Remediation'],
        pi, widths=[24,12,10,40,12,44,30])

    ws2 = wb.create_sheet('Recurring Patterns')
    _xl_write_sheet(ws2,
        ['Pattern','Recurrence','Last Occurrence','Comment'],
        [['Auth service overload via session-store','3 incidents in 24 months','2026-04-18 (latest)','Indicates session-store is a chronic single point of pressure'],
         ['Config drift between non-prod and prod','2 incidents in 24 months','2026-04-18 (latest)','Suggests release-pipeline guard rails need strengthening'],
         ['Customer comms timing >10 min','3 incidents in 24 months','2026-04-18 (latest)','Comms-readiness needs improvement']],
        widths=[40,24,18,52])
    _xl_save(wb, d / 'Past_Similar_Incidents.xlsx')


# ─── Archetype 8: Capex business case ───
def gen_capex_business_case():
    d = OUT / 'nb-capex-business-case'
    d.mkdir(parents=True, exist_ok=True)

    # Capex_Model_Base.xlsx
    wb = _xl_new(d / 'Capex_Model_Base.xlsx')
    ws = wb.active; ws.title = 'Capex Summary'
    _xl_write_sheet(ws,
        ['Category','Year 0','Year 1','Year 2','Year 3','Year 4','Year 5','Total MYR M'],
        [['Land & site works','42.0','—','—','—','—','—','42.0'],
         ['Buildings & civil','—','118.0','42.0','—','—','—','160.0'],
         ['Plant & equipment','—','62.0','148.0','22.0','—','—','232.0'],
         ['IT & automation','—','18.0','22.0','12.0','—','—','52.0'],
         ['Commissioning & ramp','—','—','24.0','18.0','—','—','42.0'],
         ['Working capital','—','12.0','22.0','18.0','—','—','52.0'],
         ['Sustaining capex','—','—','—','8.0','14.0','18.0','40.0'],
         ['TOTAL CAPEX','42.0','210.0','258.0','78.0','14.0','18.0','620.0']],
        widths=[28,10,10,10,10,10,10,14])

    ws2 = wb.create_sheet('Financials')
    _xl_write_sheet(ws2,
        ['MYR M','Y0','Y1','Y2','Y3','Y4','Y5','Y6','Y7','Y8','Y9','Y10'],
        [['Revenue','—','—','82','248','384','468','504','522','540','558','578'],
         ['EBITDA','—','—','12','58','108','138','152','158','164','170','178'],
         ['EBITDA margin','—','—','15%','23%','28%','29%','30%','30%','30%','30%','31%'],
         ['Capex','(42)','(210)','(258)','(78)','(14)','(18)','(8)','(8)','(8)','(8)','(8)'],
         ['Tax','—','—','(2)','(10)','(20)','(26)','(28)','(30)','(31)','(32)','(34)'],
         ['Working capital movement','—','(12)','(22)','(18)','(6)','(4)','(2)','(2)','(2)','(2)','(2)'],
         ['Free cash flow','(42)','(222)','(270)','(48)','+68','+90','+114','+118','+123','+128','+134']],
        widths=[28,8,8,8,8,8,8,8,8,8,8,8])

    ws3 = wb.create_sheet('Returns')
    _xl_write_sheet(ws3,
        ['Metric','Base Case','Hurdle','Status'],
        [['NPV @ WACC 8.4% (MYR M)','+184','>0','Pass'],
         ['IRR','12.8%','>10%','Pass'],
         ['Payback (years)','5.4','<7','Pass'],
         ['ROIC year 5','14.2%','>12%','Pass'],
         ['ROIC year 10','18.6%','>15%','Pass'],
         ['NPV/Capex multiple','0.30x','>0.20','Pass'],
         ['Break-even revenue (year 5)','MYR 320M','—','Comfortably above']],
        widths=[34,18,16,12])

    ws4 = wb.create_sheet('Assumptions')
    _xl_write_sheet(ws4,
        ['Assumption','Base Value','Sensitivity Tested ±%','Source'],
        [['Plant capacity utilisation Y5','82%','±15%','Engineering study'],
         ['Selling price growth','2.5% p.a.','±200bps','Strategy team'],
         ['Variable cost inflation','3.0% p.a.','±200bps','Operations'],
         ['Energy cost / unit','MYR 0.46','±25%','Energy market'],
         ['Maintenance capex Y3-10','4% of revenue','—','Industry benchmark'],
         ['Working capital','15% of revenue','—','Current portfolio'],
         ['Tax rate','24%','—','Statutory'],
         ['WACC','8.4%','±100bps','Group Treasury'],
         ['Terminal value growth','2.0%','—','Conservative']],
        widths=[36,16,22,22])
    _xl_save(wb, d / 'Capex_Model_Base.xlsx')

    # Sensitivity_Scenarios.xlsx
    wb = _xl_new(d / 'Sensitivity_Scenarios.xlsx')
    ws = wb.active; ws.title = 'Scenarios'
    sc = [
        ['Base Case','—','+184','12.8%','5.4'],
        ['Capacity utilisation -15%','Demand softer than plan','+58','9.8%','6.8'],
        ['Capacity utilisation +15%','Strong ramp','+318','15.4%','4.6'],
        ['Selling price -10%','Pricing pressure','+22','9.0%','7.0'],
        ['Selling price +10%','Pricing power','+346','15.8%','4.4'],
        ['Energy +25%','Energy spike','+98','11.0%','6.0'],
        ['Energy -25%','Energy softening','+270','14.5%','5.0'],
        ['Capex overrun +15%','Construction risk','+86','10.5%','6.4'],
        ['Capex overrun +30%','Severe overrun','-12','8.0%','7.6'],
        ['Combined downside','Util -15% + price -10% + capex +15%','-118','5.8%','>10 (no payback)'],
        ['Combined upside','Util +15% + price +10% + energy -25%','+520','19.2%','3.8']
    ]
    _xl_write_sheet(ws,
        ['Scenario','Description','NPV MYR M','IRR','Payback yrs'],
        sc, widths=[28,38,14,10,16],
        rag_col=3, rag_map=lambda r: (lambda v: 'R' if v<0 else 'A' if v<50 else 'G')(int(str(r[2]).replace('+','').split()[0]) if str(r[2]).replace('+','').replace('-','').split()[0].isdigit() else 0))

    ws2 = wb.create_sheet('Tornado')
    _xl_write_sheet(ws2,
        ['Lever','Low Impact on NPV','High Impact on NPV','Range'],
        [['Capacity utilisation ±15%','-126','+134','260'],
         ['Selling price ±10%','-162','+162','324'],
         ['Energy ±25%','-86','+86','172'],
         ['Capex ±15%','-98','+98','196'],
         ['WACC ±100bps','-72','+88','160'],
         ['Tax rate ±200bps','-18','+18','36']],
        widths=[28,18,18,12])
    _xl_save(wb, d / 'Sensitivity_Scenarios.xlsx')

    # Permit_and_EIA_Status.pdf
    _pdf(d / 'Permit_and_EIA_Status.pdf', 'Permits & EIA Status', [
        ('h1','Permits & Environmental Impact Assessment Status'),
        ('small','Zava Capex Project — Site A · Date: 2026-05-08'),
        ('h2','1. Permit Inventory'),
        ('table', [
            ['#','Permit','Authority','Application Date','Status','Expected Decision'],
            ['1','Planning Approval','Local Council','2025-11-12','Approved 2026-02-18','—'],
            ['2','Land Use Conversion','State Authority','2025-09-08','Approved 2026-01-22','—'],
            ['3','EIA — Detailed','Federal Environmental','2025-12-04','Public review complete; decision pending','2026-07-15'],
            ['4','Building Plan','Local Council','2026-03-01','Under review','2026-06-30'],
            ['5','Effluent Discharge','State Environmental','2026-04-12','Under review','2026-07-30'],
            ['6','Construction Permit','Local Council','Pending #4','Not yet filed','Q3 2026'],
            ['7','Hazardous Materials','Department of Safety','Pending #3','Not yet filed','Q3 2026'],
            ['8','Fire Safety','BOMBA','Will follow #4','Not yet filed','Q4 2026'],
            ['9','Power Connection','Utility','2026-04-02','In progress','2026-09-30'],
            ['10','Water Allocation','Water Authority','2026-04-02','In progress','2026-09-30'],
        ]),
        ('h2','2. Detailed EIA — Critical Path'),
        ('p','The Detailed EIA is the gating permit for construction permits #6 and #7. Public review period concluded 2026-04-20 with no objections raised at the public hearing.'),
        ('p','Outstanding requirement: Federal Environmental requires a final mitigation plan covering: (a) stormwater management with retention pond design, (b) noise attenuation barrier specification, (c) emergency response plan with offsite consultation.'),
        ('p','Mitigation plan to be submitted by 2026-06-15. Expected EIA decision 2026-07-15. Risk: 4-6 weeks slippage if Federal Environmental requests additional studies.'),
        ('h2','3. Risk Assessment'),
        ('p','Low risk — Planning Approval, Land Use Conversion (already approved).'),
        ('p','Medium risk — Building Plan, Effluent Discharge (typical processing).'),
        ('p','High risk — Detailed EIA (slippage risk; on critical path). Mitigation: dedicated EIA team, weekly engagement with Federal Environmental.'),
        ('h2','4. Mitigation Costs Already Committed'),
        ('p','MYR 12M provisioned for environmental mitigation infrastructure (already in capex model under "Buildings & civil"). MYR 4M provisioned for community development programmes per CSR commitment.'),
    ])

    # Procurement_Quotes.pdf
    _pdf(d / 'Procurement_Quotes.pdf', 'Procurement Quotes Summary', [
        ('h1','Procurement Quotes — Major Equipment & Civil Works'),
        ('small','Zava Capex Project — Site A · Pre-FID procurement review · 2026-05-10'),
        ('h2','1. Major Plant & Equipment'),
        ('table', [
            ['Package','# Bidders','Lowest MYR M','Highest MYR M','Recommended','Note'],
            ['Process line A','4','78.0','98.0','Bidder 2 @ 84.0','Best balance of price + capability'],
            ['Process line B','3','62.0','78.0','Bidder 1 @ 64.0','Lowest with strong references'],
            ['Utilities & boiler','4','22.0','29.0','Bidder 4 @ 24.0','Energy efficiency premium justified'],
            ['Storage & handling','3','18.0','22.0','Bidder 1 @ 19.0','Adequate'],
            ['Automation / DCS','3','14.0','22.0','Bidder 2 @ 16.0','Open architecture preferred'],
            ['Electrical & instrumentation','5','12.0','18.0','Bidder 3 @ 13.5','Local content qualifying'],
            ['Total plant & equipment','—','206.0','267.0','220.5','—'],
        ]),
        ('h2','2. Civil & Building Works'),
        ('table', [
            ['Package','# Bidders','Lowest MYR M','Highest MYR M','Recommended','Note'],
            ['Earthworks & foundations','5','48.0','62.0','Bidder 2 @ 52.0','Adequate'],
            ['Main building structure','4','58.0','72.0','Bidder 3 @ 64.0','Local content qualifying'],
            ['Mechanical & electrical fit-out','4','22.0','28.0','Bidder 4 @ 24.0','Strong HSE record'],
            ['Roads, fence, landscaping','3','12.0','16.0','Bidder 1 @ 13.0','Adequate'],
            ['Total civil & building','—','140.0','178.0','153.0','—'],
        ]),
        ('h2','3. Total Procurement Position'),
        ('p','Total of recommended packages: MYR 373.5M (vs Capex Model Base allocation of MYR 392M for these line items — MYR 18.5M underrun if all awarded at recommendation).'),
        ('h2','4. Outstanding Procurement'),
        ('p','IT & automation, commissioning services, working-capital items — not yet tendered. Expected total MYR 110-130M.'),
        ('h2','5. Key Commercial Terms (proposed)'),
        ('p','Payment milestones — 10% advance, 60% progress, 20% commissioning, 10% retention.'),
        ('p','Liquidated damages — capped at 10% of contract value for delay.'),
        ('p','Performance bond — 10% of contract value.'),
        ('p','Warranty — 24 months from commissioning for major plant, 12 months for civil.'),
        ('h2','6. Risk Notes'),
        ('p','• 2 process-line bidders have constrained capacity — early award required to lock slot'),
        ('p','• Steel pricing volatility — civil packages contain CPI-linked escalator capped at 8% over the project'),
        ('p','• Local content rules — meeting minimum 30% throughout supply chain'),
    ])

    # Strategic_Rationale_Memo.docx
    doc = _docx_init('Strategic Rationale Memo — Capex Project Site A', 'Zava Conglomerate · Strategy Office · For FID Approval · Date: 2026-05-15')
    _docx_h2(doc, '1. Strategic Context')
    _docx_p(doc, 'Zava\'s industrial-manufacturing division currently runs three plants (Sites B, C, D) at average utilisation 87%. Forecast demand growth of 6-8% p.a. over the next 5 years exceeds available capacity by 2027. Failure to add capacity risks losing 12-18% of incremental demand to competitors.')
    _docx_p(doc, 'Site A capex addresses this gap and adds optionality to serve a higher-margin specialty product line not currently in Zava\'s portfolio. The specialty line has a TAM of MYR 2.4bn and grows at 14% p.a.')
    _docx_h2(doc, '2. Strategic Fit (Zava Forward 2030)')
    _docx_p(doc, 'Aligned to Pillar 1 — Customer-Obsessed Growth (capacity to serve top 5 existing customers + 2 anchor wins).')
    _docx_p(doc, 'Aligned to Pillar 2 — Capital Productivity (ROIC year-5 of 14.2% exceeds Group threshold of 12%).')
    _docx_p(doc, 'Aligned to Pillar 3 — Operational Excellence (BAT plant design, energy efficiency 14% better than current Sites B-D, automated quality lab).')
    _docx_h2(doc, '3. Strategic Risks')
    _docx_bullets(doc, [
        'Demand risk — if industrial cycle softens, capacity sits idle. Mitigated by phased commissioning (Process Line A first, B in Year 3 based on demand).',
        'Capex overrun — Construction permits not yet finalised. Mitigated by independent project owner\'s engineer and 10% contingency.',
        'Operational risk — first plant with this scale automation in our portfolio. Mitigated by partnership with experienced EPC.',
        'Stranded asset risk — if regulatory shift on environmental compliance happens. Mitigated by BAT design + 20% margin above current standards.',
    ])
    _docx_h2(doc, '4. Strategic Alternatives Considered')
    _docx_bullets(doc, [
        'M&A — acquire existing capacity. Rejected — no available target at attractive valuation, integration risk high.',
        'Greenfield at smaller scale — half-size plant. Rejected — unit economics inferior, no specialty line option.',
        'Capacity expansion at Site D — Rejected — physically constrained, environmental headroom exhausted.',
        'Toll manufacturing arrangement — Rejected — strategic dependency on third party, no specialty line capability.',
    ])
    _docx_h2(doc, '5. Recommendation')
    _docx_p(doc, 'Approve FID at base case capex MYR 620M. Authorise Phase 1 commitments (MYR 252M). Conditions: (a) Detailed EIA approval received by 2026-07-31; (b) Process Line A and Civil Earthworks contracts awarded by 2026-08-31; (c) FID stage-gate review at end of 2026 confirming permit status before Phase 2 commitments.')
    _docx_p(doc, 'Submitted by Mod Admin, Group Strategy Director · Endorsed by Hadar Caspit, Group CFO · 2026-05-15', italic=True)
    doc.save(str(d / 'Strategic_Rationale_Memo.docx'))


# ─── Archetype 9: Campaign retro ───
def gen_campaign_retro():
    d = OUT / 'nb-campaign-retro'
    d.mkdir(parents=True, exist_ok=True)

    # Campaign_Brief.docx
    doc = _docx_init('Campaign Brief — "Future Forward" Q1 FY2026', 'Zava Group Marketing · Multi-channel brand + demand · Approved 2026-01-15')
    _docx_h2(doc, '1. Objective')
    _docx_p(doc, 'Drive brand consideration for Zava\'s digital banking proposition across 4 ASEAN markets (Malaysia, Indonesia, Singapore, Philippines) and generate 18,000 qualified leads for the new product line over the 12-week campaign window (2026-02-01 to 2026-04-30).')
    _docx_h2(doc, '2. Target Audience')
    _docx_bullets(doc, [
        'Primary: SME owners, 28-48 years, mobile-first, mid-tier banking customers in target markets',
        'Secondary: aspirational young professionals (25-35) entering business ownership',
        'Audience size: estimated 4.2M reachable',
    ])
    _docx_h2(doc, '3. Channel Mix')
    _docx_table(doc,
        ['Channel','Budget MYR k','Rationale'],
        [['Paid social (Meta + TikTok)','420','Reach + targeted retargeting'],
         ['Paid search (Google)','210','Bottom-funnel intent'],
         ['Programmatic display','120','Top-funnel awareness'],
         ['Influencer partnerships','180','3 hero + 12 mid-tier across markets'],
         ['Out-of-home (KL + Jakarta)','140','Brand-build in core cities'],
         ['Owned: organic social + email','40','Existing audience nurture'],
         ['Creative production','120','Hero film + 15 cutdowns + statics'],
         ['Measurement / brand lift studies','30','Pre/post Nielsen brand-lift'],
         ['Total','1,260','—']])
    _docx_h2(doc, '4. KPI Targets')
    _docx_bullets(doc, [
        'Reach: 22 million (across all paid)',
        'CTR: 1.4% paid average',
        'Lead volume: 18,000 qualified',
        'Cost per qualified lead: <MYR 75',
        'Brand consideration uplift: +6pp post-campaign',
        'Sales pipeline influence: MYR 12M generated',
    ])
    _docx_h2(doc, '5. Creative Strategy')
    _docx_p(doc, 'Hero idea: "Your business, future-forward." Show real ASEAN SME entrepreneurs running their business via mobile with our app as the connective thread. Local-language adaptations per market (Bahasa Malaysia, Bahasa Indonesia, Tagalog, English Singapore).')
    _docx_h2(doc, '6. Measurement Plan')
    _docx_bullets(doc, [
        'Weekly performance dashboard (in-house)',
        'Mid-campaign optimisation review at week 6',
        'Brand lift study at end of campaign (Nielsen)',
        'Sales pipeline attribution via CRM (qualified leads + influence on existing pipeline)',
        'Customer survey to lead-form respondents (sentiment, intent)',
    ])
    doc.save(str(d / 'Campaign_Brief.docx'))

    # Channel_Spend_Performance.xlsx
    wb = _xl_new(d / 'Channel_Spend_Performance.xlsx')
    ws = wb.active; ws.title = 'Channel Summary'
    rows = [
        ['Meta','190','24.8M','3.42M','13.8%','2.1%','71,820','2,140','MYR 88','+8pp'],
        ['TikTok','230','38.4M','4.66M','12.1%','3.8%','177,080','3,810','MYR 60','+9pp'],
        ['Google Search','210','—','1.85M','—','2.4%','44,400','3,420','MYR 61','+2pp'],
        ['Programmatic Display','120','62.8M','3.14M','5.0%','0.6%','18,840','420','MYR 286','+1pp'],
        ['Influencer','180','5.2M','1.18M','22.7%','4.2%','49,560','2,860','MYR 63','+11pp'],
        ['Out-of-Home','140','—','—','—','—','—','—','—','+4pp brand lift'],
        ['Owned Social + Email','40','—','—','—','5.8%','—','1,840','MYR 22','—'],
        ['Production + Measurement','150','—','—','—','—','—','—','—','—'],
        ['TOTAL','1,260','131.2M','14.25M','10.9%','—','361,700','14,490','MYR 76 weighted','+6pp brand'],
    ]
    _xl_write_sheet(ws,
        ['Channel','Spend MYR k','Impressions','Reach','Frequency','CTR','Engagements','Qualified Leads','Cost/Lead','Brand Lift'],
        rows, widths=[20,12,14,12,12,8,16,16,14,14])

    ws2 = wb.create_sheet('Weekly Pacing')
    weeks = [['W1','82','—','7.6%','60','MYR 92'],['W2','98','+18M','8.4%','220','MYR 88'],
             ['W3','110','+22M','9.2%','580','MYR 78'],['W4','124','+24M','9.8%','920','MYR 72'],
             ['W5','128','+24M','10.4%','1,240','MYR 68'],['W6','132','+22M','10.6%','1,490','MYR 65'],
             ['W7','116','+18M','11.2%','1,720','MYR 62'],['W8','118','+18M','11.4%','1,860','MYR 60'],
             ['W9','108','+16M','11.0%','1,720','MYR 62'],['W10','98','+14M','10.8%','1,580','MYR 64'],
             ['W11','86','+12M','10.4%','1,420','MYR 66'],['W12','60','+8M','10.2%','1,280','MYR 70']]
    _xl_write_sheet(ws2,
        ['Week','Spend MYR k','Impressions added','Aggregate CTR','Qualified Leads','Cost/Lead'],
        weeks, widths=[8,14,18,16,16,14])

    ws3 = wb.create_sheet('Creative')
    _xl_write_sheet(ws3,
        ['Asset','Format','Impressions','CTR','Spend','Cost / engagement'],
        [['Hero Film 30s (master)','Video','12.4M','2.8%','MYR 220k','MYR 0.63'],
         ['Cutdown 6s — entrepreneur','Video','22.6M','4.2%','MYR 180k','MYR 0.19'],
         ['Cutdown 15s — testimonial','Video','18.4M','3.4%','MYR 240k','MYR 0.38'],
         ['Static — KL SME','Static','9.6M','1.2%','MYR 80k','MYR 0.69'],
         ['Static — Jakarta SME','Static','11.8M','1.4%','MYR 90k','MYR 0.54'],
         ['Carousel — product features','Carousel','8.2M','2.2%','MYR 110k','MYR 0.61'],
         ['Story — behind-scenes','Story','14.8M','3.6%','MYR 150k','MYR 0.28'],
         ['Reels — partner','Reels','21.4M','5.4%','MYR 190k','MYR 0.16']],
        widths=[34,12,14,8,14,18])
    _xl_save(wb, d / 'Channel_Spend_Performance.xlsx')

    # Creative_Asset_Library.pdf
    _pdf(d / 'Creative_Asset_Library.pdf', 'Creative Asset Library', [
        ('h1','Creative Asset Library — "Future Forward" Campaign'),
        ('small','Zava Group Marketing · Asset Production Log · Campaign Q1 FY2026'),
        ('h2','1. Hero Concepts'),
        ('p','Three hero films produced (30s, 15s, 6s). Same narrative spine, market-localised versions. Total post-production cost MYR 220k including talent, locations, and licensed music.'),
        ('h2','2. Asset Inventory'),
        ('p','3 hero films + 15 cutdowns + 24 static creatives + 8 storyboards + 12 carousel creatives + 18 banner sizes. Total assets produced: 80.'),
        ('h2','3. Market Adaptations'),
        ('p','• Malaysia — Bahasa Malaysia + English versions; KL SME testimonial; outdoor at Bukit Bintang + Mid Valley.'),
        ('p','• Indonesia — Bahasa Indonesia versions; Jakarta + Surabaya SME testimonials; outdoor at Sudirman + Kelapa Gading.'),
        ('p','• Singapore — English version with localised CTAs.'),
        ('p','• Philippines — Tagalog version; Manila SME testimonial.'),
        ('h2','4. Performance Highlights (Asset-Level)'),
        ('p','• Cutdown 6s entrepreneur — highest CTR 4.2%, lowest cost/engagement MYR 0.19. 22.6M impressions.'),
        ('p','• Reels partner — 5.4% CTR; 21.4M impressions; partnered with hero influencer.'),
        ('p','• Hero film 30s — strong brand lift contribution, lower direct CTR (2.8%).'),
        ('p','• Static creatives — lower CTR but useful at top of funnel.'),
        ('h2','5. Brand Compliance'),
        ('p','All assets reviewed against the Zava brand guidelines and Group risk compliance checks before going live. No takedowns or recalls during campaign.'),
        ('h2','6. Assets Archived'),
        ('p','Full asset library archived to brand DAM (digital asset management). Reusable assets tagged for future campaigns. Estimated reuse value MYR 60k.'),
    ])

    # Survey_and_Sentiment_Pulse.xlsx
    wb = _xl_new(d / 'Survey_and_Sentiment_Pulse.xlsx')
    ws = wb.active; ws.title = 'Survey Results'
    sv = [
        ['Pre-campaign brand awareness','42%','—','42%','Baseline'],
        ['Post-campaign brand awareness','58%','+16pp','58%','Strong uplift'],
        ['Pre-campaign brand consideration','24%','—','24%','Baseline'],
        ['Post-campaign brand consideration','30%','+6pp','30%','Met +6pp target'],
        ['Aided ad recall','—','—','38%','Above industry norm (24%)'],
        ['Unaided ad recall','—','—','14%','Within range'],
        ['Message take-out clarity','—','—','64%','Good'],
        ['Brand favourability','—','—','+22 net','Strong'],
        ['Intent to consider','—','—','42%','Strong'],
        ['Likelihood to recommend','—','—','+12 NPS','Improving from -4 pre-campaign']
    ]
    _xl_write_sheet(ws,
        ['Metric','Score','Change','Final %','Comment'],
        sv, widths=[34,12,12,12,30])

    ws2 = wb.create_sheet('Sentiment')
    _xl_write_sheet(ws2,
        ['Channel','Positive','Neutral','Negative','Net Sentiment'],
        [['Meta','62%','28%','10%','+52'],
         ['TikTok','71%','22%','7%','+64'],
         ['Twitter / X','48%','38%','14%','+34'],
         ['YouTube','64%','26%','10%','+54'],
         ['News + PR','58%','38%','4%','+54'],
         ['Reviews + customer forums','54%','38%','8%','+46'],
         ['Aggregate','60%','30%','10%','+50']],
        widths=[24,14,14,14,18])

    ws3 = wb.create_sheet('Open-Ends Summary')
    _xl_write_sheet(ws3,
        ['Theme','Mentions','Tone'],
        [['Mobile-app ease of use','428','Positive'],
         ['Customer service experience','246','Mixed'],
         ['Pricing / fees','198','Slightly negative'],
         ['Trust / reliability','340','Positive'],
         ['Onboarding speed','174','Positive'],
         ['Local language support','118','Positive'],
         ['Branch experience','86','Mixed'],
         ['Fraud / security','42','Slightly negative'],
         ['Loyalty / rewards','64','Neutral']],
        widths=[34,14,16])
    _xl_save(wb, d / 'Survey_and_Sentiment_Pulse.xlsx')

    # Sales_Pipeline_Influence.xlsx
    wb = _xl_new(d / 'Sales_Pipeline_Influence.xlsx')
    ws = wb.active; ws.title = 'Pipeline Influence'
    _xl_write_sheet(ws,
        ['Channel','Direct Leads','Influenced Existing','MYR M Sourced','MYR M Influenced','Conversion %'],
        [['Meta','2,140','480','3.8','1.8','18%'],
         ['TikTok','3,810','620','5.2','2.2','22%'],
         ['Google Search','3,420','120','4.6','0.6','28%'],
         ['Programmatic Display','420','280','0.4','1.2','12%'],
         ['Influencer','2,860','220','2.4','0.8','24%'],
         ['Out-of-Home','—','580','—','2.4','—'],
         ['Owned Social + Email','1,840','340','1.2','1.4','32%'],
         ['Total','14,490','2,640','17.6','10.4','—']],
        widths=[24,14,18,18,18,14])

    ws2 = wb.create_sheet('Funnel')
    _xl_write_sheet(ws2,
        ['Funnel Stage','Volume','Conversion to Next'],
        [['Impressions','131M','—'],
         ['Reach (unique)','14.25M','—'],
         ['Engaged','361,700','2.5%'],
         ['Web visitors','82,400','22.8%'],
         ['Lead form started','21,800','26.5%'],
         ['Qualified Lead','14,490','66.5%'],
         ['Sales-Accepted Lead','7,820','54.0%'],
         ['Opportunity','3,140','40.2%'],
         ['Won deals','940','29.9%']],
        widths=[28,14,18])
    _xl_save(wb, d / 'Sales_Pipeline_Influence.xlsx')


# ─── Archetype 10: Tax & treasury close ───
def gen_tax_treasury_close():
    d = OUT / 'nb-tax-treasury-close'
    d.mkdir(parents=True, exist_ok=True)

    # Tax_Position_Summary.xlsx
    wb = _xl_new(d / 'Tax_Position_Summary.xlsx')
    ws = wb.active; ws.title = 'Tax Position'
    rows = [
        ['Malaysia (HQ)','MYR','2,420','24%','580.8','120.4 (paid)','38.0 (overpayment)','None active','In progress'],
        ['Singapore','SGD (in MYR)','840','17%','142.8','148.0 (paid)','—','None','Complete'],
        ['Indonesia','IDR (in MYR)','1,180','22%','259.6','228.4 (paid)','24.0 (advance refund)','TP review FY22','In progress'],
        ['Philippines','PHP (in MYR)','420','25%','105.0','110.4 (paid)','—','None','Complete'],
        ['Thailand','THB (in MYR)','280','20%','56.0','58.2 (paid)','—','None','Complete'],
        ['Vietnam','VND (in MYR)','180','20%','36.0','32.8 (paid)','3.2','None','In progress'],
        ['Holding / other','—','340','—','38.4','42.0 (paid)','—','None','Complete'],
        ['TOTAL GROUP','—','5,660','—','1,218.6','740.2 (paid)','65.2 (refunds due)','—','—']
    ]
    _xl_write_sheet(ws,
        ['Entity / Jurisdiction','Reporting Currency','Profit Before Tax (MYR M)','Effective Rate','Tax Charge (MYR M)','Tax Paid (MYR M)','Refund / Overpayment','Open Disputes','Filing Status'],
        rows, widths=[24,16,18,12,14,18,18,16,14])

    ws2 = wb.create_sheet('Deferred Tax')
    _xl_write_sheet(ws2,
        ['Source','Asset / Liability','Amount MYR M','Recovery Horizon'],
        [['PP&E timing differences','Liability','148','5-10 years'],
         ['Provisions (legal, restructuring)','Asset','42','2-3 years'],
         ['Pension obligations','Asset','38','>10 years'],
         ['Tax losses carried forward','Asset','62','Within 7 years'],
         ['Goodwill amortisation','Asset','28','Indefinite'],
         ['Net Deferred Tax (Liability)','Liability','-22','—']],
        widths=[34,18,14,18])

    ws3 = wb.create_sheet('Reconciliation')
    _xl_write_sheet(ws3,
        ['Item','Amount MYR M','Source'],
        [['Group PBT','5,660','Consolidated P&L'],
         ['Tax at weighted statutory (22.2%)','1,256','Computed'],
         ['Non-deductible expenses','+48','Provisions, entertainment'],
         ['Tax-exempt income','-32','Capital gains in jurisdictions, dividends from associates'],
         ['Withholding tax adjustments','+22','Cross-border interest, royalties'],
         ['Prior-year over/under-provision','-12','Reversal of prior over-provision'],
         ['Adjusted tax charge','1,282','—'],
         ['Reported tax charge','1,218.6','As per P&L'],
         ['Reconciling difference','-63.4','Tax loss utilisation, R&D incentive (MY), 20% effective rate'],
         ['Effective tax rate','21.5%','—']],
        widths=[34,14,38])

    ws4 = wb.create_sheet('Open Items')
    _xl_write_sheet(ws4,
        ['Item','Jurisdiction','Risk MYR M','Status','Owner','Target'],
        [['Indonesia FY22 TP review','Indonesia','1.8 + penalties','Under appeal','Tax Director','Decision Q3 2026'],
         ['Malaysia GST refund overpayment','Malaysia','12','Filed; awaiting','Tax Advisor','2026-08'],
         ['Singapore transfer-pricing documentation refresh','Singapore','—','In progress','Tax Director','2026-09'],
         ['Vietnam customs reclassification','Vietnam','3.2','Submitted','Tax Advisor','2026-07'],
         ['Cross-border interest WHT (group-internal)','Multiple','None monetary','Documentation refresh','Tax Director','Ongoing']],
        widths=[40,14,12,18,18,16])
    _xl_save(wb, d / 'Tax_Position_Summary.xlsx')

    # Transfer_Pricing_Documentation.pdf
    _pdf(d / 'Transfer_Pricing_Documentation.pdf', 'Transfer Pricing Documentation', [
        ('h1','Group Transfer Pricing Master File — FY2025'),
        ('small','Zava Conglomerate · Prepared by Group Tax · Date: 2026-04 · For internal use + regulator request'),
        ('h2','1. Overview of Group'),
        ('p','Zava Conglomerate operates across 4 ASEAN jurisdictions with 11 business divisions. Cross-border related-party transactions arise from: (i) shared services (HR, Finance, IT), (ii) intercompany financing, (iii) royalties for brand and IP licensing, (iv) supply of goods and services between divisions.'),
        ('h2','2. Functional Analysis'),
        ('p','Each entity\'s functions, assets, and risks are documented. Group HQ in Malaysia performs entrepreneurial risk-taking functions (strategy, treasury, brand). Operating subsidiaries in IDN/SGP/PHL/THA/VNM are characterised as limited-risk distributors or routine-service providers.'),
        ('h2','3. Method Selection'),
        ('p','TNMM (Transactional Net Margin Method) selected as the most appropriate method for intercompany distribution. Net Cost Plus margin range of 4-7% applied based on benchmark studies of 38 comparable independent ASEAN entities.'),
        ('p','Comparable Uncontrolled Price (CUP) applied where third-party comparables exist (e.g., 3rd-party services contracts).'),
        ('p','Cost Plus method applied for shared services (G&A, HR, IT) with markup range 5-8% benchmarked.'),
        ('h2','4. Material Transactions FY2025'),
        ('table', [
            ['Counterparty','Type','MYR M','Method','Margin'],
            ['Indonesia Operations','Distribution','480','TNMM','5.4%'],
            ['Singapore Holding','Financing (loans)','620','CUP (loan rate 5.8%)','—'],
            ['Philippines Operations','Distribution','220','TNMM','4.8%'],
            ['Thailand Operations','Service contract','140','Cost Plus','6.0%'],
            ['Vietnam Operations','Distribution','80','TNMM','5.2%'],
            ['Group HQ — IT Services','Service contract','42','Cost Plus','7.0%'],
            ['Group HQ — Brand Royalty','Royalty','38','CUP (4% of revenue)','—'],
        ]),
        ('h2','5. Benchmark Studies'),
        ('p','Independent benchmark studies refreshed every 3 years. Last refresh 2024-09 covering FY2024 and forward. Sample size 38 ASEAN independent comparables. Interquartile range 3.2% - 6.8% net cost plus. Group entity margins all within range.'),
        ('h2','6. Open Issue — Indonesia FY2022 Review'),
        ('p','Indonesia tax authority issued an assessment letter in 2025-Q3 challenging the TP methodology for FY2022. Position assessed: MYR 1.8M additional tax + penalties. Group position: TNMM remains appropriate; benchmark study supports our margin. Currently under appeal. External tax counsel engaged. Decision expected Q3 2026.'),
        ('h2','7. Controlled-Foreign-Company / Anti-Avoidance'),
        ('p','Group structure assessed for CFC rules across jurisdictions. No CFC charges arise based on current structure (substance + business purpose tests met).'),
        ('h2','8. Country-by-Country Reporting'),
        ('p','CbCR filed in Malaysia for FY2024. Automatic exchange of information operational with all ASEAN jurisdictions plus OECD reporters.'),
    ])

    # Treasury_Cash_Positions.xlsx
    wb = _xl_new(d / 'Treasury_Cash_Positions.xlsx')
    ws = wb.active; ws.title = 'Cash Positions'
    cp = [
        ['Bank A','Operating','MYR','420.0','MYR 420.0','Concentration: Group HQ'],
        ['Bank A','Term','MYR','180.0','MYR 180.0','3-month deposit; 3.8%'],
        ['Bank B','Operating','MYR','280.0','MYR 280.0','Operating in Banking Division'],
        ['Bank C','Operating','SGD','120.0','MYR 384.0','Singapore Holding'],
        ['Bank C','Sweep','SGD','40.0','MYR 128.0','Daily sweep to Tranche A'],
        ['Bank D','Operating','IDR','42,000','MYR 168.0','Indonesia operating'],
        ['Bank D','Term','IDR','18,000','MYR 72.0','Indonesia surplus, 6.2%'],
        ['Bank E','Operating','PHP','38,000','MYR 84.0','Philippines'],
        ['Bank F','Operating','THB','22,000','MYR 64.0','Thailand'],
        ['Bank G','Operating','VND','120,000','MYR 24.0','Vietnam'],
        ['Total','—','—','—','MYR 1,804.0','—']
    ]
    _xl_write_sheet(ws,
        ['Bank','Account Type','Currency','Local Balance','MYR Equivalent','Notes'],
        cp, widths=[12,14,12,14,16,38])

    ws2 = wb.create_sheet('Counterparty Risk')
    _xl_write_sheet(ws2,
        ['Bank','Rating','Group Exposure MYR M','% of Total','Approved Limit MYR M','Headroom'],
        [['Bank A','AA+','600','33.3%','750','150'],
         ['Bank B','AA','280','15.5%','350','70'],
         ['Bank C','AA-','512','28.4%','550','38'],
         ['Bank D','A+','240','13.3%','300','60'],
         ['Bank E','A','84','4.7%','120','36'],
         ['Bank F','A','64','3.5%','100','36'],
         ['Bank G','A-','24','1.3%','50','26']],
        widths=[10,10,18,12,20,12])

    ws3 = wb.create_sheet('30-day Forecast')
    months = []
    base_in = [380, 320, 480, 420, 280, 360]
    base_out = [380, 360, 420, 380, 280, 280]
    cum = 1804
    weekly = []
    for i in range(6):
        cum += base_in[i] - base_out[i]
        weekly.append([f'Week {i+1}', base_in[i], base_out[i], base_in[i]-base_out[i], cum])
    _xl_write_sheet(ws3,
        ['Week','Inflows MYR M','Outflows MYR M','Net MYR M','Cum Cash MYR M'],
        weekly, widths=[10,14,14,12,16])
    _xl_save(wb, d / 'Treasury_Cash_Positions.xlsx')

    # FX_Hedging_Register.xlsx
    wb = _xl_new(d / 'FX_Hedging_Register.xlsx')
    ws = wb.active; ws.title = 'FX Hedges'
    fx = [
        ['FX-2026-001','Forward','USD/MYR','USD 12M sell','3 months','2026-08-22','4.4820','+0.6%','MYR -180k','Group HQ — US receivable'],
        ['FX-2026-002','Forward','USD/MYR','USD 8M sell','6 months','2026-11-12','4.4980','+1.0%','MYR -380k','Banking division — USD bond coupon'],
        ['FX-2026-003','Forward','SGD/MYR','SGD 14M buy','3 months','2026-08-12','3.4920','-0.1%','MYR +18k','Singapore HQ dividend repatriation'],
        ['FX-2026-004','Option','USD/MYR','USD 6M put','12 months','2027-05-12','4.5200 strike','+1.4%','MYR +120k MTM','Insurance USD reserves'],
        ['FX-2026-005','Forward','IDR/MYR','IDR 280bn sell','6 months','2026-11-04','246.20','-0.8%','MYR +180k','Indonesia dividend up-stream'],
        ['FX-2026-006','Forward','PHP/MYR','PHP 1.6bn sell','3 months','2026-08-04','12.40','-0.3%','MYR +42k','Philippines dividend'],
        ['FX-2026-007','Forward','THB/MYR','THB 900M sell','6 months','2026-11-04','7.84','+0.2%','MYR -22k','Thailand operating'],
        ['FX-2026-008','Forward','EUR/MYR','EUR 2M sell','9 months','2027-02-12','4.96','+0.6%','MYR -42k','Equipment vendor payable'],
    ]
    _xl_write_sheet(ws,
        ['Hedge ID','Type','Pair','Notional','Tenor','Maturity','Rate','vs Spot','MTM','Purpose'],
        fx, widths=[14,10,12,18,12,14,16,10,16,30])

    ws2 = wb.create_sheet('Hedge Policy Compliance')
    _xl_write_sheet(ws2,
        ['Policy','Limit','Current','Status'],
        [['Max single counterparty open notional','MYR 500M','MYR 280M','Pass'],
         ['Max tenor (without Board approval)','12 months','9 months','Pass'],
         ['Forecasted exposure hedge ratio','60-90%','78%','Pass'],
         ['Speculative positions','None','None','Pass'],
         ['Mark-to-market reporting','Monthly','Live','Pass'],
         ['Stress testing','Quarterly','Last 2026-04-30','Pass']],
        widths=[40,14,14,12])
    _xl_save(wb, d / 'FX_Hedging_Register.xlsx')

    # Auditor_Open_Queries.docx
    doc = _docx_init('Auditor Open Queries — FY2025 Audit', 'Zava Conglomerate · External Auditor · As at 2026-05-15 · CONFIDENTIAL')
    _docx_h2(doc, '1. Tax')
    _docx_p(doc, 'Q-TAX-01: Provide detailed reconciliation of the deferred tax movement on PP&E timing differences (MYR 22M increase YoY). Status: Drafted, awaiting Group Tax sign-off.')
    _docx_p(doc, 'Q-TAX-02: Confirm position on Indonesia FY22 TP review and adequacy of the provisional provision (MYR 1.8M). Status: External counsel letter received supporting position; appeal scheduled Q3 2026.')
    _docx_p(doc, 'Q-TAX-03: Provide the refreshed transfer pricing benchmark study summary for Singapore Holding intra-group financing rate. Status: Outstanding — due 2026-06-30.')
    _docx_h2(doc, '2. Treasury')
    _docx_p(doc, 'Q-TRE-01: Confirm classification of FX-2026-004 option as hedge accounting (cash flow hedge). Status: Documentation refreshed; pass-through to disclosures.')
    _docx_p(doc, 'Q-TRE-02: Provide counterparty exposure to Bank C (SGD 512M MYR-equivalent — 28.4% of group cash). Concentration above policy threshold of 25%. Status: Treasury proposed action to diversify to 22% by Q3.')
    _docx_h2(doc, '3. Revenue Recognition')
    _docx_p(doc, 'Q-REV-01: Provide aged debtor analysis with credit-loss provisioning for accounts >180 days. Status: Provided; auditor reviewing.')
    _docx_p(doc, 'Q-REV-02: Confirm contract modification accounting for two large telco contracts amended in Q4. Status: Provided; concluded no material impact.')
    _docx_h2(doc, '4. Disclosures')
    _docx_p(doc, 'Q-DISC-01: ESG disclosure — alignment with IFRS S2 climate-related disclosures. Status: In progress (Group Sustainability — see ESG Disclosure pack).')
    _docx_p(doc, 'Q-DISC-02: Material accounting estimates note — refresh required for goodwill impairment testing on Plantation division (FFB price assumption). Status: Updated by Plantation CFO; ready for auditor review.')
    _docx_h2(doc, '5. Provisions & Contingencies')
    _docx_p(doc, 'Q-PROV-01: Update on commercial litigation matter (MYR 8.4M defended claim). Status: Counsel position note attached; provision basis maintained.')
    _docx_p(doc, 'Q-PROV-02: Restructuring provision movement — clarify utilisation of MYR 12M released in Q4. Status: Memo provided.')
    _docx_h2(doc, '6. Severity Summary')
    _docx_table(doc,
        ['Query','Severity','Status','Owner'],
        [['Q-TAX-03 (TP study Singapore)','Medium','Outstanding','Tax Director'],
         ['Q-TRE-02 (Bank C concentration)','Medium','Action plan','Treasurer'],
         ['Q-DISC-01 (IFRS S2)','Medium','In progress','Sustainability'],
         ['Others','Low','In progress / Closed','Various']])
    doc.save(str(d / 'Auditor_Open_Queries.docx'))


# ─── Archetype 11: Ops turnaround ───
def gen_ops_turnaround():
    d = OUT / 'nb-ops-turnaround'
    d.mkdir(parents=True, exist_ok=True)

    # Site_Performance_Dashboard.xlsx
    wb = _xl_new(d / 'Site_Performance_Dashboard.xlsx')
    ws = wb.active; ws.title = 'Site KPIs'
    sites = [
        ['Site Alpha','82%','78%','12.4%','+8%','+12%','1.8','MYR 24.6M','MYR 28.4M','-13.4%','Underperforming'],
        ['Site Beta','94%','92%','3.8%','+2%','+3%','0.9','MYR 18.2M','MYR 18.6M','-2.2%','On plan'],
        ['Site Gamma','88%','85%','7.2%','+5%','+8%','1.2','MYR 14.8M','MYR 15.4M','-3.9%','Slight underperform'],
        ['Site Delta','91%','88%','5.4%','+3%','+6%','1.1','MYR 22.4M','MYR 22.8M','-1.8%','On plan'],
        ['Site Epsilon','78%','72%','15.8%','+12%','+18%','2.4','MYR 12.4M','MYR 16.4M','-24.4%','Critical underperformer'],
        ['Site Zeta','85%','82%','6.4%','+4%','+5%','1.4','MYR 16.8M','MYR 17.6M','-4.5%','Slight underperform'],
        ['Site Eta','92%','90%','4.2%','+2%','+4%','1.0','MYR 19.2M','MYR 19.8M','-3.0%','On plan'],
        ['Group Total','87%','83%','7.8%','+5%','+8%','1.3','MYR 128.4M','MYR 139.0M','-7.6%','Group view']
    ]
    _xl_write_sheet(ws,
        ['Site','OEE Actual','OEE Plan','Defect Rate','Defect vs Plan','Yield Loss vs Plan','LTIFR','Output (MYR Actual)','Output (MYR Plan)','Var %','Status'],
        sites, widths=[14,10,10,12,16,18,10,18,18,10,24])

    ws2 = wb.create_sheet('Downtime Pareto')
    _xl_write_sheet(ws2,
        ['Cause','Hours Lost','Cumulative %','Top-3 Site'],
        [['Equipment breakdown — extruder','420','24%','Alpha + Epsilon'],
         ['Material quality (raw input)','280','40%','Epsilon + Alpha'],
         ['Operator change-over delays','220','52%','Alpha + Gamma'],
         ['Setup / changeover time','198','64%','Multiple'],
         ['Maintenance overruns','164','73%','Epsilon'],
         ['Power interruptions','118','79%','Alpha'],
         ['Quality holds','102','85%','Multiple'],
         ['Tooling / die issues','86','89%','Gamma'],
         ['Software / IT','64','93%','Multiple'],
         ['Other','118','100%','Multiple']],
        widths=[34,14,18,24])
    _xl_save(wb, d / 'Site_Performance_Dashboard.xlsx')

    # Root_Cause_Investigations.docx
    doc = _docx_init('Root-Cause Investigations — Operational Underperformance', 'Zava Industrial Manufacturing · Q1 FY2026 RCA Pack · Operations Office')
    _docx_h2(doc, '1. Site Alpha — Extruder Reliability')
    _docx_p(doc, 'Symptom: OEE 82% vs plan 78%-target-but-on-plan; actually 4pp below 86% benchmark. Extruder #3 unplanned downtime accounted for 38% of total Alpha downtime in Q1.')
    _docx_p(doc, 'Root cause (5-why analysis): bearing failure → bearing not on preventive maintenance schedule → maintenance plan inherited from 2018 has not been updated → equipment was upgraded in 2022 with different bearing spec → no engineering change-management gate caught the omission.')
    _docx_p(doc, 'Corrective action: refresh PM plan within 30 days. Engineering change management process refreshed; sign-off by Site Reliability Manager. Cost: MYR 380k (parts + labour for full bearing replacement programme).')
    _docx_h2(doc, '2. Site Epsilon — Multi-Driver Underperformance')
    _docx_p(doc, 'Symptom: OEE 78% vs plan 72%-target-not-met; defect rate 15.8% vs benchmark 5%; LTIFR 2.4 vs target 1.5. Three-driver problem.')
    _docx_p(doc, 'Root causes: (1) raw-material supplier delivering off-spec material (28% of defects); (2) operator turnover 22% in last 12m has degraded skill base; (3) maintenance backlog 380 work orders (vs target <100).')
    _docx_p(doc, 'Corrective actions: (a) escalation to supplier with quality penalty clauses being negotiated; (b) accelerated training programme + retention package for shift leaders; (c) maintenance recovery sprint — 90-day plan to halve backlog. Total cost: MYR 2.2M.')
    _docx_h2(doc, '3. Site Gamma — Changeover Time')
    _docx_p(doc, 'Symptom: OEE 88% with most of the gap from changeover time between product runs (averaging 4.2 hours vs benchmark 2.5 hours).')
    _docx_p(doc, 'Root cause: SMED (Single Minute Exchange of Die) discipline lapsed; only 28% of changeover steps documented as external setup (vs 60% achievable). Operator confidence on rapid changeover low after recent fatigue-related quality holds.')
    _docx_p(doc, 'Corrective action: SMED workshop with operations engineering + operator training. Target 60% external setup within 90 days. Estimated impact: 1.5 hours saved per changeover × 14 changeovers/week = 21 hours/week recovered. Cost MYR 120k.')
    _docx_h2(doc, '4. Group-Wide — Material Quality')
    _docx_p(doc, 'Cross-site theme: 14% of defects across all sites trace back to raw-material variability from 2 suppliers. Procurement and Operations jointly running supplier consolidation programme. Expected to recover 0.8pp group-wide defect rate over 6 months.')
    _docx_h2(doc, '5. Recommendations Summary')
    _docx_table(doc,
        ['Site','Lever','Impact (12m)','Cost MYR','Owner'],
        [['Alpha','PM refresh + ECM','+3pp OEE','380k','Site Reliability Mgr'],
         ['Epsilon','Supplier escalation','-3pp defects','—','Procurement'],
         ['Epsilon','Workforce reset','-0.8 LTIFR','1,400k','Site GM'],
         ['Epsilon','Maintenance sprint','+5pp OEE','800k','Site Engineering'],
         ['Gamma','SMED programme','+1.5pp OEE','120k','Operations Eng'],
         ['Group','Supplier consolidation','-0.8pp defects','—','Procurement'],
         ['TOTAL','—','—','2,700k','—']])
    doc.save(str(d / 'Root_Cause_Investigations.docx'))

    # Capex_and_Maintenance_Plan.xlsx
    wb = _xl_new(d / 'Capex_and_Maintenance_Plan.xlsx')
    ws = wb.active; ws.title = 'Capex'
    cap = [
        ['CAP-01','Site Alpha — extruder upgrade','Reliability','Y1','480','Mandatory','Approved','Reliability/throughput'],
        ['CAP-02','Site Alpha — bearing PM tooling','Reliability','Y1','60','Mandatory','Approved','Cost avoidance'],
        ['CAP-03','Site Epsilon — quality lab','Quality','Y1-Y2','340','Strategic','Pending','Reduces defect rate 2pp'],
        ['CAP-04','Site Epsilon — operator training facility','HSE/Training','Y1','180','Strategic','Pending','LTIFR + skill retention'],
        ['CAP-05','Site Gamma — SMED tooling','Throughput','Y1','60','Strategic','Approved','Faster changeovers'],
        ['CAP-06','Group — predictive maintenance platform','Digital','Y1-Y2','420','Strategic','Approved','Catches 60% of failures earlier'],
        ['CAP-07','Site Delta — packaging line upgrade','Throughput','Y2','280','Strategic','Pending','Capacity +12%'],
        ['CAP-08','Group — quality data platform','Digital','Y1-Y2','220','Strategic','Approved','Defect traceability'],
        ['CAP-09','Site Beta — solar PV','ESG','Y2','340','Strategic','Approved','22% energy from renewable'],
        ['CAP-10','Group — safety automation','HSE','Y1','120','Mandatory','Approved','LTIFR reduction']
    ]
    _xl_write_sheet(ws,
        ['ID','Description','Category','Year','MYR k','Type','Status','Strategic Rationale'],
        cap, widths=[10,40,14,8,10,12,14,40])

    ws2 = wb.create_sheet('Maintenance Backlog')
    _xl_write_sheet(ws2,
        ['Site','Work Orders Open','Target','% Over Target','Avg Age (days)','Critical Items'],
        [['Alpha','142','100','42%','38','12'],
         ['Beta','86','100','-14%','22','3'],
         ['Gamma','118','100','18%','30','6'],
         ['Delta','94','100','-6%','24','4'],
         ['Epsilon','380','100','280%','62','38'],
         ['Zeta','108','100','8%','28','5'],
         ['Eta','82','100','-18%','20','2'],
         ['Group','1,010','700','44%','30','70']],
        widths=[12,18,12,16,16,14])

    ws3 = wb.create_sheet('Maintenance Plan')
    _xl_write_sheet(ws3,
        ['Site','Approach','Target Backlog','Investment MYR','Timing'],
        [['Alpha','Recovery sprint','100','120k','60 days'],
         ['Epsilon','Recovery sprint + outsource','100','680k','90 days'],
         ['Gamma','Steady reduction','100','40k','120 days'],
         ['Zeta','Steady reduction','100','40k','120 days'],
         ['Others','Maintain','—','—','Ongoing']],
        widths=[12,28,18,18,18])
    _xl_save(wb, d / 'Capex_and_Maintenance_Plan.xlsx')

    # Supplier_Performance_Scorecard.xlsx
    wb = _xl_new(d / 'Supplier_Performance_Scorecard.xlsx')
    ws = wb.active; ws.title = 'Suppliers'
    sup = [
        ['S-001','Raw Material A','Apex Materials Sdn Bhd','42,800','94%','3.2%','94%','A','Critical','Continue'],
        ['S-002','Raw Material B','Pacific Industrial Group','38,400','78%','11.4%','82%','C','Critical','At-risk — escalation'],
        ['S-003','Raw Material C','Asianova Chemicals','12,400','92%','4.4%','92%','A','Critical','Continue'],
        ['S-004','Packaging','Northwind Packaging','18,600','96%','1.2%','98%','A+','Important','Continue'],
        ['S-005','MRO Spares','Sundry Suppliers','8,400','88%','—','88%','B','Important','Consolidate'],
        ['S-006','Energy / Gas','Pacific Energy','22,400','99%','—','99%','A','Critical','Long-term contract'],
        ['S-007','Logistics','Apex Logistics','14,800','92%','1.6%','94%','A-','Critical','Continue + benchmark'],
        ['S-008','Raw Material D','Northbridge Industrial','24,200','82%','8.2%','84%','B-','Critical','Add second source'],
        ['S-009','Contract labour','Pacific Workforce','9,200','86%','—','86%','B','Important','—'],
        ['S-010','Engineering services','Specialist firms','12,400','94%','—','94%','A','Important','—']
    ]
    _xl_write_sheet(ws,
        ['Supplier ID','Category','Supplier Name','Spend MYR','On-Time %','Defect %','Quality Score','Rating','Criticality','Action'],
        sup, widths=[12,18,28,14,12,12,14,10,12,24])

    ws2 = wb.create_sheet('Supplier Trends')
    _xl_write_sheet(ws2,
        ['Supplier','Q1 FY25','Q2 FY25','Q3 FY25','Q4 FY25','Q1 FY26','Trend'],
        [['S-002 Pacific Industrial Group','94%','88%','84%','80%','78%','↓ Declining — driver of Site Epsilon defects'],
         ['S-008 Northbridge Industrial','92%','88%','86%','84%','82%','↓ Declining'],
         ['S-001 Apex Materials','92%','93%','94%','94%','94%','→ Stable'],
         ['S-003 Asianova Chemicals','90%','91%','92%','92%','92%','↑ Improving slightly']],
        widths=[30,10,10,10,10,10,38])
    _xl_save(wb, d / 'Supplier_Performance_Scorecard.xlsx')

    # Worker_Safety_Incidents.xlsx
    wb = _xl_new(d / 'Worker_Safety_Incidents.xlsx')
    ws = wb.active; ws.title = 'Incidents'
    si = [
        ['SAF-2026-001','2026-01-12','Epsilon','Lost Time','Slip and fall','3','Medium','Closed — RCA done','Floor coating + signage'],
        ['SAF-2026-002','2026-01-28','Alpha','First Aid','Cut hand on tooling','0','Low','Closed','PPE refresh; hand guard'],
        ['SAF-2026-003','2026-02-06','Epsilon','Lost Time','Manual handling injury','5','Medium','Closed — RCA done','Mechanical aid procured'],
        ['SAF-2026-004','2026-02-18','Gamma','Near Miss','Pedestrian / vehicle interaction','0','High','Closed','Pedestrian segregation lane'],
        ['SAF-2026-005','2026-03-04','Delta','First Aid','Welding spark','0','Low','Closed','Refresher training'],
        ['SAF-2026-006','2026-03-22','Epsilon','Lost Time','Strain — repetitive lifting','8','Medium','Closed — RCA done','Ergonomic assessment; rotation policy'],
        ['SAF-2026-007','2026-04-08','Alpha','Property Damage','Forklift impact','0','Medium','Closed','Operator retraining; floor markings'],
        ['SAF-2026-008','2026-04-19','Epsilon','Lost Time','Cut hand — guarding','2','High','Closed — RCA done','Machine guard refresh'],
        ['SAF-2026-009','2026-04-28','Beta','First Aid','Eye irritation — chemical splash','0','Low','Closed','PPE upgrade'],
        ['SAF-2026-010','2026-05-08','Epsilon','Near Miss','Falling object','0','High','Open — RCA in progress','Investigation under way']
    ]
    _xl_write_sheet(ws,
        ['ID','Date','Site','Type','Description','Days Lost','Severity','Status','Action'],
        si, widths=[14,12,10,14,30,12,12,22,32])

    ws2 = wb.create_sheet('Trends')
    _xl_write_sheet(ws2,
        ['Site','LTIFR YTD','LTIFR Target','LTIFR Q-1','Trend','Major Theme'],
        [['Alpha','1.8','1.5','1.6','↑ Slight worsening','Equipment-related'],
         ['Beta','0.9','1.5','0.8','→ Stable','Minor'],
         ['Gamma','1.2','1.5','1.4','↓ Improving','Pedestrian/vehicle'],
         ['Delta','1.1','1.5','1.0','→ Stable','Minor'],
         ['Epsilon','2.4','1.5','2.1','↑ Worsening','Manual handling, machine guarding'],
         ['Zeta','1.4','1.5','1.3','→ Stable','—'],
         ['Eta','1.0','1.5','1.0','→ Stable','—'],
         ['Group','1.4','1.5','1.3','↑ Stable but Epsilon dragging','—']],
        widths=[10,14,16,14,28,30])
    _xl_save(wb, d / 'Worker_Safety_Incidents.xlsx')


# ─── Archetype 12: ESG disclosure ───
def gen_esg_disclosure():
    d = OUT / 'nb-esg-disclosure'
    d.mkdir(parents=True, exist_ok=True)

    # Emissions_Inventory.xlsx
    wb = _xl_new(d / 'Emissions_Inventory.xlsx')
    ws = wb.active; ws.title = 'Scope 1+2+3'
    rows = [
        ['Scope 1 — Stationary combustion','tCO2e','148,400','124,200','-16.3%','Gas + diesel'],
        ['Scope 1 — Mobile combustion','tCO2e','62,800','58,400','-7.0%','Fleet'],
        ['Scope 1 — Fugitive (refrigerants)','tCO2e','8,400','7,200','-14.3%','Decommissioned units'],
        ['Scope 1 — Process emissions','tCO2e','42,600','42,000','-1.4%','Largely stable'],
        ['Scope 1 — Subtotal','tCO2e','262,200','231,800','-11.6%','—'],
        ['Scope 2 — Purchased electricity (market-based)','tCO2e','380,400','342,800','-9.9%','Renewable PPA Site Beta'],
        ['Scope 2 — Purchased heat / steam','tCO2e','18,200','17,400','-4.4%','Minor change'],
        ['Scope 2 — Subtotal','tCO2e','398,600','360,200','-9.6%','—'],
        ['Scope 3 — Cat 1 Purchased goods','tCO2e','1,820,000','1,742,000','-4.3%','Supplier engagement'],
        ['Scope 3 — Cat 3 Fuel & energy','tCO2e','94,000','88,000','-6.4%','—'],
        ['Scope 3 — Cat 4 Upstream transport','tCO2e','142,000','138,000','-2.8%','Modal shift'],
        ['Scope 3 — Cat 5 Waste','tCO2e','22,400','20,800','-7.1%','Recycling rate up'],
        ['Scope 3 — Cat 6 Business travel','tCO2e','18,400','21,200','+15.2%','Post-COVID rebound'],
        ['Scope 3 — Cat 7 Commuting','tCO2e','38,000','38,400','+1.1%','Flat'],
        ['Scope 3 — Cat 9 Downstream transport','tCO2e','62,000','58,400','-5.8%','—'],
        ['Scope 3 — Cat 11 Use of sold products','tCO2e','248,000','238,000','-4.0%','—'],
        ['Scope 3 — Cat 12 End of life','tCO2e','24,000','22,800','-5.0%','—'],
        ['Scope 3 — Cat 15 Investments','tCO2e','420,000','394,000','-6.2%','Portfolio carbon — banking, insurance'],
        ['Scope 3 — Subtotal','tCO2e','2,888,800','2,761,600','-4.4%','—'],
        ['TOTAL Group','tCO2e','3,549,600','3,353,600','-5.5%','—']
    ]
    _xl_write_sheet(ws,
        ['Category','Unit','FY2024','FY2025','% Change','Driver'],
        rows, widths=[40,8,14,14,12,40])

    ws2 = wb.create_sheet('By Division')
    _xl_write_sheet(ws2,
        ['Division','Scope 1 tCO2e','Scope 2 tCO2e','Scope 3 tCO2e','Total tCO2e','% of Group'],
        [['Plantation','82,000','38,400','622,000','742,400','22.1%'],
         ['O&G Downstream','62,000','42,400','318,000','422,400','12.6%'],
         ['Utilities','42,400','22,400','82,000','146,800','4.4%'],
         ['Manufacturing','22,400','98,400','382,000','502,800','15.0%'],
         ['Banking','—','22,400','620,000','642,400','19.2%'],
         ['Healthcare','12,800','58,400','142,000','213,200','6.4%'],
         ['Property','3,400','42,400','438,000','483,800','14.4%'],
         ['Telco','—','24,400','118,000','142,400','4.2%'],
         ['Retail','6,800','12,400','38,000','57,200','1.7%'],
         ['Other / Holding','—','—','1,600','1,600','0.0%'],
         ['Group Total','231,800','360,200','2,761,600','3,353,600','100%']],
        widths=[24,14,14,16,16,12])

    ws3 = wb.create_sheet('Targets vs Actuals')
    _xl_write_sheet(ws3,
        ['Year','Scope 1+2 Target tCO2e','Actual tCO2e','Variance','On Track?'],
        [['FY2024','660,800','660,800','0%','Baseline'],
         ['FY2025','594,720','592,000','-0.5%','Yes'],
         ['FY2026F','528,640','—','—','—'],
         ['FY2027F','462,560','—','—','—'],
         ['FY2030 SBTi','264,320','—','—','60% reduction by 2030 vs 2024 baseline']],
        widths=[10,22,18,12,40])
    _xl_save(wb, d / 'Emissions_Inventory.xlsx')

    # Sustainability_Strategy.docx
    doc = _docx_init('Group Sustainability Strategy — "Zava Tomorrow"', 'Zava Conglomerate · Sustainability Office · Board-approved 2025-09-30')
    _docx_h2(doc, '1. Net Zero Commitment')
    _docx_p(doc, 'Zava has committed to net-zero greenhouse gas emissions across Scope 1 + Scope 2 by 2040, and Scope 3 by 2050. Interim target: 60% reduction in Scope 1 + Scope 2 by 2030 versus 2024 baseline, aligned with Science Based Targets initiative (SBTi).')
    _docx_h2(doc, '2. Pillars')
    _docx_bullets(doc, [
        'Decarbonise operations — energy efficiency, electrification, renewable PPAs (target: 100% renewable electricity by 2030)',
        'Decarbonise value chain — supplier engagement programme covering 80% of Scope 3 Category 1 by 2027',
        'Sustainable finance — channel MYR 12bn into sustainability-linked transactions by 2030 (banking + insurance divisions)',
        'Climate resilience — physical-risk assessment + adaptation plans for all 11 divisions by 2026',
        'Nature & biodiversity — TNFD-aligned reporting from FY2026; deforestation-free supply chain (Plantation) by 2030',
        'Just transition — workforce upskilling fund MYR 80M over 5 years',
    ])
    _docx_h2(doc, '3. Governance')
    _docx_p(doc, 'Board Sustainability Committee oversees the strategy. Sustainability progress is a Board KPI with 15% weighting in annual executive compensation. Group Sustainability Officer reports to the CEO and Board Committee.')
    _docx_h2(doc, '4. Reporting Framework')
    _docx_bullets(doc, [
        'IFRS S1 + IFRS S2 — global baseline (from FY2026 reporting cycle)',
        'TCFD — climate-related financial disclosures (legacy framework, integrated into IFRS S2)',
        'GRI — broader sustainability disclosures',
        'TNFD — nature-related disclosures (from FY2026)',
        'Bursa Sustainability Reporting — Malaysian regulatory baseline',
        'OJK POJK 51 — Indonesian sustainability reporting (banking + insurance divisions)',
    ])
    _docx_h2(doc, '5. Roadmap Highlights (next 24 months)')
    _docx_table(doc,
        ['Initiative','Target Date','Owner','Status'],
        [['Site Beta solar PV','2026-Q3','Manufacturing','Construction'],
         ['Renewable PPA expansion','2026-Q4','Treasury + Sustainability','In negotiation'],
         ['Supplier engagement Scope 3 Cat 1','2026-Q4','Procurement','Pilot live'],
         ['Plantation deforestation policy','2026-Q3','Plantation','Drafted'],
         ['TNFD pilot','2026-Q4','Sustainability','In progress'],
         ['IFRS S2 first disclosure','2027-Q1','IR + Sustainability','Preparing'],
         ['SBTi target validation','2026-Q4','Sustainability','Submitted'],
         ['Just transition fund kick-off','2026-Q3','HR + Sustainability','Approved']])
    doc.save(str(d / 'Sustainability_Strategy.docx'))

    # Audit_Assurance_Findings.docx
    doc = _docx_init('Sustainability Assurance — Findings', 'Independent Assurance Provider · Limited Assurance Engagement · FY2025 · Final Report')
    _docx_h2(doc, '1. Scope of Assurance')
    _docx_p(doc, 'Limited assurance engagement covering Scope 1, Scope 2, and Scope 3 emissions inventory; energy consumption; LTIFR; voluntary attrition; gender representation; and the sustainability disclosure narrative.')
    _docx_h2(doc, '2. Conclusions')
    _docx_p(doc, '(a) Nothing has come to our attention to suggest that the Scope 1 + Scope 2 emissions data is materially misstated.')
    _docx_p(doc, '(b) For Scope 3 emissions, we note 3 categories rely on industry-average factors; the methodology is appropriate for current maturity but warrants enhancement (see findings below).')
    _docx_p(doc, '(c) Narrative disclosures fairly reflect the underlying data with limited exceptions noted.')
    _docx_h2(doc, '3. Findings')
    _docx_table(doc,
        ['ID','Category','Finding','Severity','Action Required'],
        [['AF-ESG-01','Scope 3 Cat 1 methodology','Spend-based proxy used for 38% of purchased-goods emissions; transition to supplier-specific data recommended','Medium','Pilot in FY26, scale to 80% by FY27'],
         ['AF-ESG-02','Scope 3 Cat 11 (sold products)','End-use assumption modelling needs refresh','Medium','Refresh by FY26 disclosure cycle'],
         ['AF-ESG-03','Internal controls','Emissions data flow lacks segregation of duties in 2 divisions','Medium','Strengthen access controls by FY26 close'],
         ['AF-ESG-04','Data lineage','Source systems for Scope 3 not centralised; manual aggregation creates risk','Medium','ETL pipeline build — FY26 project'],
         ['AF-ESG-05','Disclosure consistency','One narrative claim ("100% deforestation-free") goes beyond current evidence base — recommend qualifying language','High','Revise disclosure language'],
         ['AF-ESG-06','LTIFR — Epsilon site','Workings reconcile but methodology differs slightly from group standard','Low','Align methodology'],
         ['AF-ESG-07','Voluntary attrition','Sample testing — 3 of 40 records misclassified (involuntary vs voluntary)','Low','Process refresh + retraining']])
    _docx_h2(doc, '4. Recommendations Summary')
    _docx_p(doc, 'Maturity of GHG emissions reporting is strong for Scope 1 + Scope 2; Scope 3 maturity matches IFRS S2 Year-1 expectations but should advance through FY26-FY27 to enable reasonable assurance in FY27 or FY28.')
    _docx_p(doc, 'Disclosure language must be tightened to match the underlying evidence base (see AF-ESG-05 — HIGH priority before publication).')
    _docx_p(doc, 'Internal controls and data lineage need investment — propose enterprise sustainability data platform under FY26 capex.')
    doc.save(str(d / 'Audit_Assurance_Findings.docx'))

    # Disclosure_Framework_Mapping.xlsx
    wb = _xl_new(d / 'Disclosure_Framework_Mapping.xlsx')
    ws = wb.active; ws.title = 'Framework Mapping'
    fm = [
        ['Topic','IFRS S2','TCFD','GRI','Bursa SR','OJK POJK 51','Internal Tag','Status'],
        ['Governance','Strategy.1','Governance','GRI 2-12','Disclosure 2.1','Article 5','GOV','Disclosed FY25'],
        ['Strategy — climate scenarios','Strategy.4','Strategy.c','—','—','—','STRAT-SCEN','Pilot — FY26'],
        ['Climate risks (physical)','Risk.1','Risk.a','—','Disclosure 3.2','Article 6','CLIM-PHYS','Disclosed (partial)'],
        ['Climate risks (transition)','Risk.1','Risk.a','—','Disclosure 3.2','Article 6','CLIM-TRANS','Disclosed (partial)'],
        ['GHG emissions Scope 1','Metrics.3','Metrics.a','GRI 305-1','Disclosure 4.1','—','GHG-S1','Disclosed + assured'],
        ['GHG emissions Scope 2','Metrics.3','Metrics.a','GRI 305-2','Disclosure 4.1','—','GHG-S2','Disclosed + assured'],
        ['GHG emissions Scope 3','Metrics.3','Metrics.a','GRI 305-3','Disclosure 4.1','—','GHG-S3','Disclosed (limited assurance)'],
        ['Energy','Metrics.4','Metrics.b','GRI 302','Disclosure 4.2','—','ENERGY','Disclosed'],
        ['Water','—','—','GRI 303','Disclosure 4.3','—','WATER','Disclosed'],
        ['Biodiversity','—','—','GRI 304','—','—','BIO','TNFD pilot FY26'],
        ['Waste','—','—','GRI 306','Disclosure 4.4','—','WASTE','Disclosed'],
        ['Employment / DEI','—','—','GRI 401','Disclosure 5.1','—','DEI','Disclosed'],
        ['OH&S','—','—','GRI 403','Disclosure 5.2','—','OHS','Disclosed'],
        ['Training','—','—','GRI 404','Disclosure 5.3','—','TRAIN','Disclosed'],
        ['Human rights','—','—','GRI 412','—','—','HR','Disclosed (high-level)'],
        ['Community','—','—','GRI 413','Disclosure 6.1','—','COMM','Disclosed'],
        ['Customer privacy','—','—','GRI 418','—','—','CUST-PRIV','Disclosed'],
        ['Anti-corruption','—','—','GRI 205','—','Article 5.4','ANTI-CORR','Disclosed']
    ]
    _xl_write_sheet(ws,
        ['Topic','IFRS S2','TCFD','GRI','Bursa SR','OJK POJK 51','Internal Tag','Status'],
        fm[1:], widths=[28,12,10,10,16,18,14,30])
    _xl_save(wb, d / 'Disclosure_Framework_Mapping.xlsx')

    # Prior_Year_Disclosure.pdf
    _pdf(d / 'Prior_Year_Disclosure.pdf', 'Prior Year Disclosure', [
        ('h1','Zava Conglomerate — Sustainability Report FY2024 (Excerpt)'),
        ('small','For reference — current cycle disclosure must build on this baseline and address audit findings'),
        ('h2','Letter from the CEO'),
        ('p','"FY2024 has been a year of significant progress for Zava on our sustainability journey. We achieved our first material reduction in Scope 1 emissions through site-level efficiency programmes, and embarked on a strategic renewable-energy partnership covering 22% of our electricity demand. We acknowledge that our Scope 3 measurement is still maturing — and we are investing to address this. Looking ahead to FY2025 and beyond, our focus will be on translating our SBTi-aligned targets into operational milestones, accelerating supplier engagement, and embedding climate scenario analysis into capital allocation decisions."'),
        ('h2','Headline Figures FY2024'),
        ('p','• Scope 1 emissions: 262,200 tCO2e (baseline year)'),
        ('p','• Scope 2 emissions (market-based): 398,600 tCO2e'),
        ('p','• Scope 3 emissions: 2,888,800 tCO2e (limited assurance)'),
        ('p','• Renewable electricity share: 14% (target 100% by 2030)'),
        ('p','• Water withdrawal: 6.4M m³'),
        ('p','• Waste to landfill: 22,400 tonnes'),
        ('p','• Employees: 18,400 (44% female; 23% in management roles female)'),
        ('p','• LTIFR: 1.6'),
        ('p','• Voluntary attrition: 13.4%'),
        ('p','• Community investment: MYR 24M'),
        ('h2','Governance Highlights'),
        ('p','Board Sustainability Committee established 2024-Q3. Group Sustainability Officer appointed 2024-Q4. First sustainability-linked bond issuance MYR 800M in 2024-Q4 with 4 KPI-linked covenants.'),
        ('h2','Strategy & Targets'),
        ('p','Net-zero Scope 1+2 by 2040; net-zero Scope 3 by 2050; SBTi-aligned interim target of 60% reduction in Scope 1+2 by 2030 vs FY2024 baseline.'),
        ('h2','Climate Risk Highlights'),
        ('p','Physical-risk assessment completed for 4 of 11 divisions in FY2024 (Plantation, Healthcare, Manufacturing, Utilities). Transition risk assessment ongoing.'),
        ('h2','Acknowledged Limitations'),
        ('p','Scope 3 measurement is at early stage of maturity; in particular, Categories 1 (Purchased Goods) and 11 (Sold Products) rely on industry-average factors and require enhancement. Scenario analysis under TCFD is qualitative; quantitative analysis is planned for FY2026.'),
        ('h2','Looking Ahead to FY2025'),
        ('p','Priorities — accelerated supplier engagement programme; PPA expansion; Plantation deforestation policy operationalisation; first IFRS S2-aligned disclosure preparation; community investment portfolio refresh.'),
    ])


# ─── Archetype 13: HR onboarding pack ───
def gen_hr_onboarding_pack():
    d = OUT / 'nb-hr-onboarding-pack'
    d.mkdir(parents=True, exist_ok=True)

    # Employee_Handbook.pdf
    _pdf(d / 'Employee_Handbook.pdf', 'Employee Handbook', [
        ('h1','Zava Conglomerate — Employee Handbook'),
        ('small','Document ID: HR-HBK-2026 · Version 6.2 · Effective 2026-01-01 · Owner: Group HR · CONFIDENTIAL'),
        ('h2','1. Welcome'),
        ('p','Welcome to Zava Conglomerate. This Handbook describes our policies, expectations, and the support available to you. It applies to all permanent employees across our 11 divisions and 4 ASEAN markets. Where local statute or collective agreements provide greater benefits, those apply.'),
        ('h2','2. Code of Conduct'),
        ('p','2.1 Honesty & Integrity — act with integrity in all dealings. Do not misrepresent facts to colleagues, customers, regulators, or auditors.'),
        ('p','2.2 Confidentiality — protect Zava\'s confidential information and that entrusted to us by customers and partners. Use confidential information only for legitimate business purposes.'),
        ('p','2.3 Anti-bribery & corruption — Zava operates a zero-tolerance policy. Do not offer, give, request or accept any gift, hospitality, or other benefit intended to improperly influence a business decision.'),
        ('p','2.4 Conflicts of interest — declare any actual, perceived, or potential conflict to your line manager. Review the Group Conflicts Policy annually.'),
        ('p','2.5 Harassment-free workplace — zero tolerance for harassment, bullying, or discrimination on any protected ground.'),
        ('h2','3. Working Hours, Leave & Flexibility'),
        ('p','3.1 Standard working week — 40 hours / 5 days. Flexible working available subject to role and team needs.'),
        ('p','3.2 Annual leave — 18 days (years 1-3), 22 days (years 4-7), 24 days (year 8+).'),
        ('p','3.3 Sick leave — 14 days fully paid; medical certificate required if exceeding 2 consecutive days.'),
        ('p','3.4 Parental leave — Maternity: 98 days fully paid. Paternity: 14 days fully paid. Adoption: 60 days fully paid.'),
        ('p','3.5 Compassionate leave — up to 5 days paid for immediate family bereavement.'),
        ('h2','4. Compensation & Benefits'),
        ('p','4.1 Salary review — annual cycle in March. Performance-linked.'),
        ('p','4.2 Annual bonus — discretionary, tied to Group + Division + Individual performance.'),
        ('p','4.3 Long-term incentive — eligible employees from job grade 4 upwards. Vesting 3 years.'),
        ('p','4.4 Medical — comprehensive plan + dental + optical. Detail in the Benefits Guide.'),
        ('p','4.5 EPF / KWSP contributions — at statutory rates + 1% top-up from Zava (subject to maximum).'),
        ('h2','5. Performance & Development'),
        ('p','5.1 Performance reviews — twice yearly (mid-year + end-year). Continuous feedback encouraged.'),
        ('p','5.2 Development — each employee has an annual development plan including 40 hours minimum learning.'),
        ('p','5.3 Internal mobility — open job postings every quarter. Apply with line manager awareness.'),
        ('h2','6. Health, Safety & Well-being'),
        ('p','6.1 Workplace safety — adhere to site rules. Report incidents and near-misses immediately.'),
        ('p','6.2 Mental health — Employee Assistance Programme (24x7 confidential helpline + 8 free sessions per year).'),
        ('p','6.3 Work-from-home — eligible roles may work from home up to 2 days / week. Ergonomic setup support available.'),
        ('h2','7. Information Security & Data Protection'),
        ('p','7.1 Acceptable use — use Zava systems for legitimate business purposes. Personal use should be incidental.'),
        ('p','7.2 Information classification — handle data per Information Classification Standard (Public / Internal / Confidential / Restricted).'),
        ('p','7.3 PDPA — Malaysia (PDPA 2010) and equivalent — comply with data subject rights and breach reporting timelines.'),
        ('p','7.4 Generative AI tools — only Zava-sanctioned tools (e.g., Microsoft 365 Copilot) may be used on Restricted / Confidential data.'),
        ('h2','8. Speaking Up'),
        ('p','8.1 Whistleblowing — Zava operates a confidential whistleblowing channel (independent third-party hotline). No retaliation for good-faith reporting.'),
        ('p','8.2 Grievance procedure — informal resolution first, then formal grievance with HR if unresolved.'),
        ('h2','9. End of Employment'),
        ('p','9.1 Notice period — 1 month (job grades 1-3), 2 months (4-6), 3 months (7+).'),
        ('p','9.2 Exit interview — mandatory; feedback used for continuous improvement.'),
        ('p','9.3 Garden leave / restrictive covenants — applicable as per individual contract.'),
        ('h2','10. Annexes'),
        ('p','Annex A — Information Classification Standard'),
        ('p','Annex B — Travel & Expense Policy'),
        ('p','Annex C — Speaking-Up Policy + Whistleblowing Hotline contact details'),
        ('p','Annex D — Conflicts of Interest Declaration Form'),
        ('p','Annex E — Gifts & Hospitality Register Form'),
    ])

    # Benefits_Guide.pdf
    _pdf(d / 'Benefits_Guide.pdf', 'Benefits Guide', [
        ('h1','Zava Conglomerate — Benefits Guide FY2026'),
        ('small','Effective 2026-01-01 · Owner: Group HR · For permanent employees · Refer to your local HR partner for local-jurisdiction variations'),
        ('h2','1. Compensation Overview'),
        ('p','Your compensation comprises a fixed monthly salary, a discretionary annual bonus, and (where eligible) long-term incentives. The annual salary review takes effect each April with retrospective application to January.'),
        ('h2','2. Health & Medical'),
        ('p','2.1 Outpatient medical — fully covered at panel clinics. Choice of consultation at any registered medical practitioner with reimbursement.'),
        ('p','2.2 Hospitalisation — comprehensive medical insurance covering employee + spouse + up to 4 children. Room category: 2-bed for grades 1-3, 1-bed for grades 4-6, 1-bed deluxe for grade 7+. Annual limit RM 250,000 / per family.'),
        ('p','2.3 Dental — annual limit RM 1,500 / family. Includes preventive + restorative care.'),
        ('p','2.4 Optical — annual limit RM 1,000 / family. Includes lenses + frames + contact lenses.'),
        ('p','2.5 Critical illness — lump sum RM 100,000 (grade 1-3) up to RM 300,000 (grade 7+) on diagnosis.'),
        ('p','2.6 Maternity & paternity care — fully covered with no co-pay.'),
        ('h2','3. Life & Disability'),
        ('p','3.1 Group term life — sum assured 24x monthly salary.'),
        ('p','3.2 Personal accident — sum assured 24x monthly salary.'),
        ('p','3.3 Total permanent disability — 24x monthly salary.'),
        ('p','3.4 Income protection — 75% of basic salary for up to 24 months after illness/disability.'),
        ('h2','4. Retirement & Long-Term'),
        ('p','4.1 EPF / KWSP — at statutory rates with 1% Zava top-up (subject to maximum).'),
        ('p','4.2 Long-service award — tax-efficient cash + leave at 5/10/15/20/25/30 years of service.'),
        ('p','4.3 Retirement gratuity — for employees retiring after 10+ years of service.'),
        ('h2','5. Time Off'),
        ('p','5.1 Annual leave: 18 / 22 / 24 days depending on tenure.'),
        ('p','5.2 Public holidays as per jurisdiction. If a public holiday falls on a weekend, in-lieu day is given.'),
        ('p','5.3 Examination leave — up to 5 days / year for approved study.'),
        ('p','5.4 Volunteer leave — 2 paid days / year for approved community work.'),
        ('h2','6. Family Support'),
        ('p','6.1 Childcare allowance — up to RM 500 / month / child under 6 years.'),
        ('p','6.2 Education subsidy — 50% reimbursement of approved continuing education up to RM 8,000 / year.'),
        ('p','6.3 Adoption leave — 60 days fully paid.'),
        ('p','6.4 Family emergency leave — 3 days / year fully paid.'),
        ('h2','7. Lifestyle & Wellness'),
        ('p','7.1 Gym/wellness allowance — RM 1,200 / year reimbursable.'),
        ('p','7.2 Mental health support — Employee Assistance Programme; 8 confidential counselling sessions / year.'),
        ('p','7.3 Annual health screening — fully covered at panel hospitals.'),
        ('p','7.4 Flexible working — eligible roles up to 2 days / week WFH; flexible start/end times.'),
        ('h2','8. Mobility & Lifestyle'),
        ('p','8.1 Phone allowance — RM 200 / month for eligible roles.'),
        ('p','8.2 Parking — corporate parking discount at HQ.'),
        ('p','8.3 Transport — shuttle service between major LRT stations and HQ.'),
        ('p','8.4 Cafeteria — subsidised meals at HQ + main sites.'),
        ('h2','9. How to Claim'),
        ('p','All claims submitted via the HR system within 60 days of incurring the expense. Claims are reimbursed via payroll in the following month.'),
        ('h2','10. Contact'),
        ('p','For benefits-related questions: hr.benefits@zava.com.my · HR helpdesk 24/7 chatbot via Teams · or your local HR partner.'),
    ])

    # Role_Specific_Onboarding_Plan.docx
    doc = _docx_init('Role-Specific Onboarding Plan — Finance Analyst', 'Zava Group Finance · 30-60-90 Day Plan · Effective 2026-04 · Owner: Finance Director')
    _docx_h2(doc, 'Position Context')
    _docx_p(doc, 'Finance Analyst position reporting to the Group FP&A Manager. Focus area: Group consolidation, board reporting cycle, and divisional variance analysis. The role is part of the Group CFO\'s direct-reporting line.')
    _docx_h2(doc, 'First 30 Days — Orient & Connect')
    _docx_bullets(doc, [
        'Week 1: HR orientation (1 day), IT setup, mandatory training (Code of Conduct, Anti-Bribery, Data Protection, Information Security)',
        'Week 1: 1:1 with line manager — agree expectations, deliverables, and development plan',
        'Week 1-2: Meet 12 stakeholders (Group FP&A team x4, divisional CFOs x5, IR team x2, Treasury x1)',
        'Week 2: Read foundational materials — Board pack last 4 quarters, ZAVA Forward 2030 strategy, Group financial close calendar',
        'Week 3: Shadow the monthly close — observe end-to-end',
        'Week 3-4: Complete Group FP&A onboarding modules — consolidation tool, BI dashboards, financial KPI library',
        'Week 4: First written deliverable — divisional variance commentary for one division (under supervision)',
        'Week 4: 30-day check-in with line manager and HR business partner',
    ])
    _docx_h2(doc, 'Days 31-60 — Build Capability')
    _docx_bullets(doc, [
        'Take ownership of one division\'s monthly close + variance commentary',
        'Complete advanced consolidation tool certification',
        'Lead the monthly variance review meeting for assigned division',
        'Contribute to the quarterly Board pack — produce one division\'s pages',
        'Begin development plan: identify 2 stretch skills + agreed learning resources',
        'Build relationship with the divisional CFO (bi-weekly catch-up)',
        '60-day check-in with line manager (formal); HR business partner conversation',
    ])
    _docx_h2(doc, 'Days 61-90 — Operate Independently')
    _docx_bullets(doc, [
        'Take ownership of 2 divisions (consolidation + commentary)',
        'Deliver first standalone Board memo on assigned topic',
        'Identify 1 process-improvement opportunity and propose to line manager',
        'Complete first deep-dive analytical project (manager-assigned)',
        'Complete advanced Excel + Power BI + Copilot Studio training',
        '90-day formal probation review with manager + HR business partner',
        'Set goals for Year 1 in the formal performance system',
    ])
    _docx_h2(doc, 'Manager Checklist')
    _docx_bullets(doc, [
        'Day 1: Welcome + introductions + team lunch',
        'Day 7: Confirm IT setup + access to systems',
        'Day 14: Review onboarding plan with the hire',
        'Day 30: 30-day check-in',
        'Day 60: 60-day check-in',
        'Day 90: Probation review + Year 1 goal-setting',
    ])
    _docx_h2(doc, 'Resources')
    _docx_bullets(doc, [
        'Compliance Training Catalog (separate workbook)',
        'Buddy Programme Notes (separate document)',
        'Group FP&A Wiki on SharePoint',
        'Microsoft 365 Copilot — onboarding-licensed seat from Day 1',
    ])
    doc.save(str(d / 'Role_Specific_Onboarding_Plan.docx'))

    # Compliance_Training_Catalog.xlsx
    wb = _xl_new(d / 'Compliance_Training_Catalog.xlsx')
    ws = wb.active; ws.title = 'Training Catalog'
    tc = [
        ['CT-001','Code of Conduct','All employees','Mandatory','60 min','By Day 14','Annual refresh','In progress'],
        ['CT-002','Anti-Bribery & Corruption','All employees','Mandatory','45 min','By Day 14','Annual refresh','In progress'],
        ['CT-003','Data Protection (PDPA / GDPR)','All employees','Mandatory','30 min','By Day 14','Annual refresh','In progress'],
        ['CT-004','Information Security Awareness','All employees','Mandatory','45 min','By Day 14','Annual refresh','In progress'],
        ['CT-005','Workplace Conduct (Harassment, Discrimination)','All employees','Mandatory','30 min','By Day 30','Bi-annual refresh','Not started'],
        ['CT-006','Conflicts of Interest','All employees','Mandatory','20 min','By Day 14','Annual','Not started'],
        ['CT-007','Whistleblowing & Speak-Up','All employees','Mandatory','15 min','By Day 14','One-off','Not started'],
        ['CT-008','Generative AI — Responsible Use','All employees','Mandatory','30 min','By Day 14','One-off','Not started'],
        ['CT-009','Money Laundering Prevention','Banking + Insurance + Finance','Role-specific','60 min','By Day 30','Annual','Not started'],
        ['CT-010','Sanctions Compliance','Banking + Procurement + IR','Role-specific','45 min','By Day 30','Annual','Not started'],
        ['CT-011','Insider Trading & Market Abuse','Finance + IR + Strategy','Role-specific','30 min','By Day 30','Annual','Not started'],
        ['CT-012','Health & Safety Awareness','All employees','Mandatory','30 min','By Day 14','Annual','Not started'],
        ['CT-013','Site-Specific HSE','Operations + Plantation + O&G','Role-specific','120 min','By Day 7','Annual + before site entry','Not started'],
        ['CT-014','First Aid','Volunteer','Optional','120 min','—','Bi-annual','—'],
        ['CT-015','Mental Health First Aid','Volunteer','Optional','180 min','—','Bi-annual','—'],
        ['CT-016','Climate Risk Literacy','All employees','Optional','45 min','—','—','—'],
        ['CT-017','Microsoft 365 Copilot — Essentials','All employees','Recommended','60 min','By Day 30','—','Not started'],
        ['CT-018','Microsoft 365 Copilot — Advanced','Power-user roles','Optional','120 min','—','—','—']
    ]
    _xl_write_sheet(ws,
        ['Course ID','Title','Audience','Type','Duration','Deadline','Refresh','Status'],
        tc, widths=[10,42,28,16,12,16,22,14])

    ws2 = wb.create_sheet('Completion Tracking')
    _xl_write_sheet(ws2,
        ['Course','Group Target','Group Actual','Variance','Comment'],
        [['CT-001 Code of Conduct','100%','94%','-6%','Refresh due for 480 employees'],
         ['CT-002 ABC','100%','92%','-8%','—'],
         ['CT-003 PDPA','100%','88%','-12%','—'],
         ['CT-004 InfoSec','100%','86%','-14%','—'],
         ['CT-005 Workplace Conduct','100%','82%','-18%','Material gap — priority'],
         ['CT-008 GenAI','100%','46%','-54%','New course — rollout in progress'],
         ['CT-009 AML','100%','94%','-6%','—'],
         ['CT-010 Sanctions','100%','89%','-11%','—'],
         ['CT-013 Site HSE','100%','98%','-2%','—']],
        widths=[36,16,16,12,30])
    _xl_save(wb, d / 'Compliance_Training_Catalog.xlsx')

    # Buddy_Programme_Notes.docx
    doc = _docx_init('Buddy Programme — Onboarding Companion Notes', 'Zava Group HR · Buddy Guidance for New Hires + Their Buddies · Effective 2026')
    _docx_h2(doc, '1. Purpose')
    _docx_p(doc, 'The Buddy Programme pairs every new hire with a more experienced colleague to support orientation in the first 90 days. The buddy is a peer (not a line manager) who answers everyday questions, makes introductions, and helps the new hire find their feet.')
    _docx_h2(doc, '2. Time Commitment')
    _docx_p(doc, '~2 hours / week in weeks 1-4. ~1 hour / week in weeks 5-12. After 90 days, the formal programme ends but the friendship continues.')
    _docx_h2(doc, '3. Buddy Activities Checklist')
    _docx_bullets(doc, [
        'Week 1: Greet the new hire on Day 1; walk through the building / office tour',
        'Week 1: Coffee or lunch (twice in first week)',
        'Week 1-2: Introduce to 10-15 people across the team and adjacent teams',
        'Week 1-2: Share informal Zava cultural notes — how meetings work, how decisions get made, who to ask for what',
        'Week 2-3: Walk through the SharePoint sites the new hire will use most',
        'Week 3-4: Introduce them to social activities — Group sports clubs, ESG volunteering, lunch-and-learns',
        'Week 5-12: Weekly catch-up (15 min) — answer questions, identify any blockers',
        'Day 90: Buddy provides feedback to HR on programme effectiveness',
    ])
    _docx_h2(doc, '4. What\'s NOT in scope for the Buddy')
    _docx_bullets(doc, [
        'Performance feedback (line manager owns this)',
        'Compensation discussion (HR + line manager)',
        'Career planning (line manager + HR)',
        'Confidential business information not relevant to the new hire\'s role',
    ])
    _docx_h2(doc, '5. Tips for Buddies')
    _docx_bullets(doc, [
        'Be human, not corporate — share what you wish you\'d known on Day 1',
        'Make it easy to ask "dumb" questions',
        'Reassure them that the first 90 days are about learning, not delivering',
        'Help them prioritise — there\'s a lot of information thrown at new hires',
        'If they\'re struggling, raise it sensitively with the line manager and HR',
    ])
    _docx_h2(doc, '6. Tips for New Hires')
    _docx_bullets(doc, [
        'Use your buddy generously — they are there for you',
        'Keep a running list of questions, ask them in batches',
        'Note down jargon and acronyms — Zava has many',
        'Take notes during 1:1s, especially with your line manager',
        'In Week 4, do a self-reflection: what\'s working, what\'s not, what do you need?',
    ])
    _docx_h2(doc, '7. Buddy Programme Champions (FY2026)')
    _docx_p(doc, 'HR Programme owner: Sasha Ouellet (Group CoS — overall sponsorship) · Operational lead: HR business partners by division.')
    doc.save(str(d / 'Buddy_Programme_Notes.docx'))


# Run only the first archetype as a sanity check first
if __name__ == '__main__':
    GENERATORS = [
        gen_rfp_procurement,
        gen_clinical_case,
        gen_regulator_crossref,
        gen_ma_due_diligence,
        gen_board_prereread,
        gen_renewal_book,
        gen_incident_pmortem,
        gen_capex_business_case,
        gen_campaign_retro,
        gen_tax_treasury_close,
        gen_ops_turnaround,
        gen_esg_disclosure,
        gen_hr_onboarding_pack,
    ]
    for g in GENERATORS:
        g()
        print(f'✓ {g.__name__}')
    print(f'\nAll done. Files: ', sum(1 for _ in OUT.rglob("*") if _.is_file()))
