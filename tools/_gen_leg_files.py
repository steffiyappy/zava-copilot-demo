"""Generate LEG_07..LEG_13 supporting files for Site 5 Legal UCs.

Light-weight but realistic content. Files placed under zava-copilot-demo/files/
to match existing convention.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'files'
OUT.mkdir(exist_ok=True)


def H(d, text, level=1):
    h = d.add_heading(text, level=level)
    return h


def P(d, text, bold=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    if bold:
        r.bold = True
    return p


def make_leg_07():
    """Mining Regulation Brief — primer on Permen ESDM 7/2026."""
    d = Document()
    H(d, 'Mining Regulation Brief — Permen ESDM No. 7 Tahun 2026', level=0)
    P(d, 'Prepared by: Group Corporate Legal · Issue date: 15 May 2026 · Classification: Privileged & Confidential', bold=True)
    H(d, 'I. Executive Summary', 2)
    P(d, 'Permen ESDM No. 7 of 2026 (effective 1 July 2026) introduces three substantive changes affecting all PT Zava Mining (Persero) coal-mining subsidiaries: (1) accelerated coal-gasification facility deadlines under PP 96/2021 alignment; (2) new royalty (PNBP) escalator tied to HBA reference price brackets; (3) mandatory divestment-progress reporting on a quarterly basis. Non-compliance triggers IUPK suspension under UU 3/2020 (Minerba) Article 119A.')
    H(d, 'II. Key Provisions', 2)
    for k, v in [
        ('Article 4', 'Coal-gasification facility commissioning deadline brought forward by 18 months for IUPK holders > 10 Mt/year output.'),
        ('Article 9', 'PNBP escalator — 13.5% base + 1.5% per USD 10/t increment above HBA reference floor.'),
        ('Article 14', 'Quarterly divestment-progress reporting to ESDM with comparable disclosure to BEI.'),
        ('Article 21', 'Reclamation top-up bond requirement — 110% of disturbed-area liability (was 100%).'),
        ('Article 27', 'Material Adverse Change (MAC) reporting trigger lowered from 15% to 10% capacity reduction.'),
    ]:
        P(d, f'{k}: {v}')
    H(d, 'III. Subsidiary Exposure Map', 2)
    t = d.add_table(rows=7, cols=4)
    t.style = 'Light Grid Accent 1'
    hdr = ['Subsidiary', 'Article 4 Risk', 'Article 9 Risk', 'Article 21 Risk']
    for i, h in enumerate(hdr):
        t.cell(0, i).text = h
    for r, row in enumerate([
        ['PT Zava Coal Kalimantan', 'High', 'Medium', 'Low'],
        ['PT Zava Coal Sumatera', 'High', 'High', 'Medium'],
        ['PT Zava Mining Services', 'Low', 'Low', 'Low'],
        ['PT Zava Coal Trading', 'N/A', 'Medium', 'N/A'],
        ['PT Zava Tambang Tenggara', 'Medium', 'Medium', 'High'],
        ['PT Zava Energi Batubara', 'High', 'High', 'Medium'],
    ], start=1):
        for c, v in enumerate(row):
            t.cell(r, c).text = v
    H(d, 'IV. Source Authorities', 2)
    for s in [
        'Permen ESDM No. 7 Tahun 2026 (full text via JDIH-ESDM)',
        'UU No. 3 Tahun 2020 — Minerba (consolidated)',
        'PP No. 96 Tahun 2021 — IUP/IUPK Implementation',
        'PP No. 25 Tahun 2024 — Reclamation & Post-Mining',
        'POJK 17/2020 — Material Transactions',
        'Hukumonline — Pro analysis (April 2026)',
    ]:
        P(d, '• ' + s)
    P(d, '— End of brief —')
    d.save(OUT / 'LEG_07_Mining_Regulation_Brief.docx')


def make_leg_08():
    """Legal Knowledge Base — internal SOPs & policy index."""
    d = Document()
    H(d, 'Legal Knowledge Base — Group Corporate Legal', level=0)
    P(d, 'Master index of internal SOPs, policies, and standing legal authorities.', bold=True)
    sections = [
        ('1. Mining Regulation', [
            'SOP-LGL-001 — IUP/IUPK Compliance Calendar',
            'SOP-LGL-002 — Royalty (PNBP) Computation & Filing',
            'SOP-LGL-003 — Divestment Progress Reporting',
            'POL-LGL-101 — Reclamation & Post-Mining Liability',
        ]),
        ('2. Contracts & Commercial', [
            'SOP-LGL-010 — NDA Approval & Mark-up',
            'SOP-LGL-011 — Long-form Commercial Agreements',
            'POL-LGL-201 — Delegation of Authority Matrix',
        ]),
        ('3. M&A', [
            'M&A Playbook v4.0 — Deal Lifecycle (see LEG_11)',
            'SOP-LGL-020 — Due Diligence Standard',
            'SOP-LGL-021 — Material Transactions Disclosure (POJK 17/2020)',
        ]),
        ('4. Litigation', [
            'SOP-LGL-030 — Matter Intake & Triage',
            'SOP-LGL-031 — External Counsel Engagement',
            'POL-LGL-301 — Privileged Communications Protocol',
        ]),
        ('5. Regulatory & Disclosure', [
            'SOP-LGL-040 — BEI/Bursa Disclosure Drafting',
            'SOP-LGL-041 — Regulator Inquiry Response',
            'POL-LGL-401 — MAC Trigger & Notification',
        ]),
        ('6. Standing Legal Authorities (curated citations)', [
            'UU No. 3 Tahun 2020 (Minerba)',
            'UU No. 40 Tahun 2007 (Perseroan Terbatas)',
            'KUHPerdata — Pasal 1247–1248 (Ganti Rugi)',
            'POJK 17/2020 (Material Transactions)',
            'Putusan MA No. 2447 K/Pdt/2021',
            'Putusan MA No. 1812 K/Pdt.Sus-PHI/2023',
        ]),
    ]
    for title, items in sections:
        H(d, title, 2)
        for it in items:
            P(d, '• ' + it)
    d.save(OUT / 'LEG_08_Legal_Knowledge_Base.docx')


def make_leg_09():
    """Legal Memo Template — IRAC structure, 6 sections, ready to fill."""
    d = Document()
    H(d, 'Legal Memo Template — IRAC Structure', level=0)
    P(d, '[ON LETTERHEAD] · Privileged & Confidential — Attorney Work Product', bold=True)
    P(d, 'TO: [Recipient Title]')
    P(d, 'FROM: [Author], [Title], Group Corporate Legal')
    P(d, 'DATE: [DD MMM YYYY]')
    P(d, 'SUBJECT: [Concise legal question]')
    for h, body in [
        ('I. Latar Belakang', '[Faktual context — 2-4 paragraphs. Cite source documents.]'),
        ('II. Pertanyaan Hukum', '[Numbered list of 1-3 sharply scoped legal questions.]'),
        ('III. Analisis Hukum (IRAC)', 'Issue: [restate Q]\nRule: [cite UU/PP/Permen/POJK/Putusan MA]\nApplication: [apply rule to facts]\nConclusion: [direct answer]'),
        ('IV. Kesimpulan', '[Bullet conclusions, one per legal question.]'),
        ('V. Rekomendasi & Mitigasi Risiko', '[Action items + Risk Register table: Risk ID | Description | Severity | Owner | Mitigation | Deadline]'),
        ('VI. Limitasi & Disclaimer', 'Memo ini disusun berdasarkan dokumen yang dirujuk per tanggal di atas. Setiap perubahan fakta atau regulasi setelahnya dapat mengubah kesimpulan. Tidak boleh dikutip ke luar Grup tanpa izin tertulis dari General Counsel.'),
    ]:
        H(d, h, 2)
        P(d, body)
    d.save(OUT / 'LEG_09_Legal_Memo_Template.docx')


def make_leg_10():
    """NDA Draft — counterparty draft with red-flag clauses for review."""
    d = Document()
    H(d, 'MUTUAL NON-DISCLOSURE AGREEMENT', level=0)
    P(d, 'Between Pacific Resources Holdings Pte. Ltd. ("Party A") and PT Zava Mining (Persero) ("Party B")', bold=True)
    clauses = [
        ('1. Definitions', '"Confidential Information" means any non-public information disclosed by either party, in any form, marked or reasonably understood to be confidential.'),
        ('2. Scope', 'Each Receiving Party shall use Confidential Information solely for the purpose of evaluating a potential commercial transaction.'),
        ('3. Term', 'This Agreement shall remain in force for ten (10) years from the Effective Date and shall survive termination of any related transaction. [⚠ longer than Playbook standard of 3-5 years]'),
        ('4. Standard of Care', 'Each Receiving Party shall protect Confidential Information using the same degree of care it uses for its own confidential information, and not less than reasonable care.'),
        ('5. Indemnity', 'The Receiving Party shall indemnify, defend and hold harmless the Disclosing Party from any and all losses, including consequential, indirect and punitive damages, arising from any breach. [⚠ uncapped; Playbook caps at fees paid; excludes consequential damages]'),
        ('6. Governing Law & Disputes', 'This Agreement shall be governed by the laws of Singapore. Disputes shall be resolved by SIAC arbitration in Singapore. [⚠ Playbook requires Indonesian law + SIAC seat for Indonesian counterparty exposure]'),
        ('7. Exclusivity', 'For the duration of this Agreement, Party B shall not solicit, negotiate or transact with any party engaged in similar business in Indonesia. [⚠ overbroad — auto-reject under Playbook §4.2]'),
        ('8. Return / Destruction', 'Upon written request, the Receiving Party shall destroy or return all Confidential Information within thirty (30) days.'),
        ('9. Assignment', 'This Agreement may be assigned by either Party to any affiliate, successor or acquirer without consent. [⚠ Playbook requires written consent]'),
        ('10. Entire Agreement', 'This Agreement constitutes the entire agreement between the Parties on the subject matter hereof.'),
    ]
    for h, body in clauses:
        H(d, h, 2)
        P(d, body)
    P(d, '[Signature blocks omitted for draft.]')
    d.save(OUT / 'LEG_10_NDA_Draft.docx')


def make_leg_11():
    """M&A Playbook — internal counterparty red-flag standards."""
    d = Document()
    H(d, 'M&A Playbook v4.0 — Group Corporate Legal', level=0)
    P(d, 'Effective: 1 January 2026 · Owner: VP Corporate Legal · Classification: Internal Use Only', bold=True)
    sections = [
        ('§1. Purpose', 'Defines the standards and approval thresholds for all M&A activity at Zava Group, including NDAs, term sheets, MoUs, and definitive agreements.'),
        ('§2. NDA Standards', 'Standard term: 3-5 years (max 7 with deal-team approval). Indemnity: capped at fees paid; consequential damages excluded. Governing law: Indonesian law for Indonesian counterparty exposure; SIAC seat optional. Return/destruction: 30 days. No exclusivity unless explicitly negotiated and approved.'),
        ('§3. Material Transaction Threshold (POJK 17/2020)', 'Any transaction exceeding 20% of equity triggers Material Transaction disclosure. 50%+ triggers shareholder approval.'),
        ('§4. Auto-Reject Clauses', 'The following counterparty terms are automatic rejections requiring renegotiation: §4.1 uncapped indemnity; §4.2 overbroad exclusivity (any non-target market); §4.3 unilateral assignment without consent; §4.4 NDA term exceeding 7 years; §4.5 governing law of an offshore jurisdiction with no nexus to either party.'),
        ('§5. Risk Register Standards', 'Every deal must produce a Risk Register with at least: Operational, Commercial, Regulatory, HSE, Reputational categories. Severity scale: High / Medium / Low. Owner must be named individual, not function.'),
        ('§6. Citations', 'KUHPerdata Pasal 1247–1248 (Ganti Rugi); UU PT 40/2007; POJK 17/2020 (Material Transactions); UU 3/2020 (Minerba) for mining-asset M&A.'),
    ]
    for h, body in sections:
        H(d, h, 2)
        P(d, body)
    d.save(OUT / 'LEG_11_MA_Playbook.docx')


def make_leg_12():
    """Litigation Cases — 50 rows, 6 holding subsidiaries."""
    wb = Workbook()
    ws = wb.active
    ws.title = 'Litigation_Cases'
    headers = ['Case_ID', 'Anak_Holding', 'Case_Type', 'Forum_Level', 'Province',
               'Status', 'Outcome', 'Date_Opened', 'Date_Closed', 'Duration_Days',
               'Claim_Amount_IDR_Bn', 'Settlement_IDR_Bn']
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        ws.cell(row=1, column=c).font = Font(bold=True, color='FFFFFF')
        ws.cell(row=1, column=c).fill = PatternFill('solid', fgColor='1F2937')

    import random
    random.seed(7)
    subs = ['PT Zava Coal Kalimantan', 'PT Zava Coal Sumatera',
            'PT Zava Mining Services', 'PT Zava Coal Trading',
            'PT Zava Tambang Tenggara', 'PT Zava Energi Batubara']
    types = ['Sengketa Pertanahan', 'Sengketa Kontrak', 'Sengketa Ketenagakerjaan',
             'Sengketa Lingkungan', 'Sengketa Pajak/PNBP', 'Sengketa Korporasi']
    forums = ['PN Jakarta Pusat', 'PN Samarinda', 'PN Palembang', 'PN Banjarmasin',
              'PT DKI Jakarta', 'MA RI', 'BANI', 'PHI Samarinda', 'Pengadilan Pajak']
    provinces = ['DKI Jakarta', 'Kalimantan Timur', 'Sumatera Selatan',
                 'Kalimantan Selatan', 'Sulawesi Tenggara', 'Jambi']
    statuses = ['Ongoing', 'Closed']
    outcomes_closed = ['Menang', 'Kalah', 'Settle']
    for i in range(1, 51):
        sub = random.choice(subs)
        t = random.choice(types)
        forum = random.choice(forums)
        prov = random.choice(provinces)
        status = random.choices(statuses, weights=[55, 45])[0]
        opened = f'2023-{random.randint(1,12):02d}-{random.randint(1,28):02d}'
        if status == 'Closed':
            duration = random.randint(180, 1100)
            outcome = random.choice(outcomes_closed)
            closed = f'2025-{random.randint(1,12):02d}-{random.randint(1,28):02d}'
            settle = round(random.uniform(0, 80), 2) if outcome != 'Menang' else 0
        else:
            duration = random.randint(60, 900)
            outcome = ''
            closed = ''
            settle = 0
        claim = round(random.uniform(2.5, 250.0), 2)
        ws.append([f'LIT-2024-{i:03d}', sub, t, forum, prov, status, outcome,
                   opened, closed, duration, claim, settle])

    summary = wb.create_sheet('Summary_Hint')
    summary.append(['Hint', 'This is the raw dataset. Build Summary in Copilot per UC4 prompt.'])
    summary.append(['Total cases', '=COUNTA(Litigation_Cases!A2:A51)'])

    for col in 'ABCDEFGHIJKL':
        ws.column_dimensions[col].width = 18
    wb.save(OUT / 'LEG_12_Litigation_Cases.xlsx')


def make_leg_13():
    """Legal Sync Transcript — Teams meeting transcript with 7 action items."""
    d = Document()
    H(d, 'Legal Weekly Sync — Teams Meeting Transcript', level=0)
    P(d, 'Date: 12 May 2026 · 09:00–10:00 WIB · Channel: Group Corporate Legal · Transcript via Teams Recap', bold=True)
    P(d, 'Attendees: Ratna Sari (VP Corporate Legal, Chair), Andi Pratama (Head of Litigation), Dewi Kusuma (Head of Commercial), Budi Santoso (Senior Counsel — M&A), Linda Halim (Compliance Officer), Hadar Caspit (CFO — observer)', bold=True)

    H(d, 'Agenda', 2)
    for i, a in enumerate([
        'Permen ESDM 7/2026 — readiness review',
        'Litigation portfolio Q1 2026 update (50 cases across 6 subsidiaries)',
        'NDA backlog with Pacific Resources counterparty',
        'M&A pipeline — 3 active deals',
        'BEI inquiry response status',
        'Open action items review',
    ], start=1):
        P(d, f'{i}. {a}')

    H(d, 'Discussion Highlights', 2)
    for line in [
        'Ratna: "Permen ESDM 7 mulai berlaku 1 Juli. Kita perlu legal memo IRAC final ke Direktur Utama dalam 10 hari."',
        'Andi: "Litigation portfolio Q1 menunjukkan win rate 58% — turun dari 62% di Q4 2025. Driver utamanya adalah kasus sengketa pertanahan di Kaltim."',
        'Budi: "NDA Pacific Resources punya 5 red flag — saya menargetkan mark-up + email pengantar bilingual besok."',
        'Dewi: "M&A pipeline punya 3 deal aktif. Material Transactions Threshold di POJK 17/2020 tersentuh oleh deal #2."',
        'Linda: "BEI request klarifikasi sudah dijawab. Tinggal monitoring."',
        'Hadar: "Saya perlu Executive Legal Dashboard tiap kuartal — bukan tahunan. Format HTML self-contained untuk Direktur Utama."',
    ]:
        P(d, '• ' + line)

    H(d, 'Action Items', 2)
    t = d.add_table(rows=8, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['#', 'Action', 'Owner', 'Deadline']):
        t.cell(0, i).text = h
    actions = [
        ('A1', 'Finalise Legal Memo on Permen ESDM 7/2026 (IRAC, 1.5 pages)', 'Ratna Sari', '22 May 2026'),
        ('A2', 'Submit NDA mark-up + bilingual cover email to Pacific Resources', 'Budi Santoso', '15 May 2026'),
        ('A3', 'Reach out to ESDM for clarification on Article 21 reclamation top-up bond', 'Linda Halim', '20 May 2026'),
        ('A4', 'Litigation Q1 root-cause analysis — top 3 drivers + 6 preventive actions', 'Andi Pratama', '25 May 2026'),
        ('A5', 'Build Executive Legal Dashboard (HTML, self-contained) for Group MD', 'Ratna Sari', '28 May 2026'),
        ('A6', 'M&A Playbook v4.0 — refresh §4 Auto-Reject list with new counterparty patterns', 'Budi Santoso', '5 June 2026'),
        ('A7', 'Onboard "Group Legal Counsel" agent (Agent Builder, no-code)', 'Ratna Sari', '10 June 2026'),
    ]
    for r, (idn, act, own, dl) in enumerate(actions, start=1):
        for c, v in enumerate([idn, act, own, dl]):
            t.cell(r, c).text = v
    P(d, '— End of transcript —')
    d.save(OUT / 'LEG_13_Legal_Sync_Transcript.docx')


if __name__ == '__main__':
    print('Generating LEG_07..LEG_13 supporting files...')
    make_leg_07(); print('  ✓ LEG_07_Mining_Regulation_Brief.docx')
    make_leg_08(); print('  ✓ LEG_08_Legal_Knowledge_Base.docx')
    make_leg_09(); print('  ✓ LEG_09_Legal_Memo_Template.docx')
    make_leg_10(); print('  ✓ LEG_10_NDA_Draft.docx')
    make_leg_11(); print('  ✓ LEG_11_MA_Playbook.docx')
    make_leg_12(); print('  ✓ LEG_12_Litigation_Cases.xlsx')
    make_leg_13(); print('  ✓ LEG_13_Legal_Sync_Transcript.docx')
    print('\nDone. Files in:', OUT)
