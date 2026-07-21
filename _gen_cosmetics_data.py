# -*- coding: utf-8 -*-
"""Generate the 5 COSMETICS_0N sample data files for the Contoso Beauty Innovation demo."""
import os, random, datetime as dt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(42)

OUT = r"C:\Users\peiyiyap\OneDrive - Microsoft\Documents\Customers\Wave 3 Demo\Data\Cosmetics"
os.makedirs(OUT, exist_ok=True)

BRAND_HOUSES = ["Wardah", "Make Over", "Emina", "Kahf", "Instaperfect", "OMG!"]
CATEGORIES   = ["Colour Cosmetics", "Skincare", "Haircare", "Personal Wash", "Fragrance"]
GMs = {"Wardah":"Salma Prameswari", "Make Over":"Rifqi Ardiansyah",
       "Emina":"Cindy Halim", "Kahf":"Ustaz Faizal Rahim",
       "Instaperfect":"Nadya Kusuma", "OMG!":"Bryan Tanujaya"}
RETAILERS = ["Sociolla ID","Watsons ID","Guardian ID","Alfamart","Indomaret",
             "TikTok Shop ID","Shopee Live ID","Tokopedia Beauty",
             "Watsons MY","Guardian MY","Watsons SG","Sephora SEA"]

# ---------- Header helpers ----------
HDR = Font(bold=True, color="FFFFFF", size=11)
HDR_FILL = PatternFill("solid", fgColor="EC4899")
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)
BORDER = Border(*[Side(style="thin", color="D0D0D0")]*4)

def style_header(ws, ncols):
    for c in range(1, ncols+1):
        cell = ws.cell(row=1, column=c)
        cell.font = HDR; cell.fill = HDR_FILL; cell.alignment = CENTER; cell.border = BORDER
    ws.row_dimensions[1].height = 32
    ws.freeze_panes = "A2"

def autosize(ws):
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        maxlen = max((len(str(c.value)) for c in col if c.value is not None), default=10)
        ws.column_dimensions[letter].width = min(max(maxlen+2, 12), 42)

# ============================================================
# FILE 1 — BPOM Notifikasi Kosmetik Tracker
# ============================================================
def gen_bpom_tracker():
    wb = Workbook()
    ws = wb.active; ws.title = "SKU Master"
    cols = ["SKU Code","Product Name","Brand House","Category","Pack Size",
            "BPOM Notifikasi No.","Notif. Issue Date","Notif. Expiry Date","Days To Expiry",
            "12M Sales Vol (units)","12M Revenue (Rp M)","Gross Margin %",
            "LPPOM MUI Halal Cert No.","LPPOM Expiry","JAKIM Cert (MY)","JAKIM Expiry",
            "Manufacturing Site","INCI On File","Safety Assessment","Stability Study",
            "Claim Substantiation","Dossier Readiness","Renewal Priority","Status Flag","Owner (RA)"]
    ws.append(cols); style_header(ws, len(cols))

    today = dt.date(2026, 7, 16)
    rows = []
    sku_id = 1000
    for house in BRAND_HOUSES:
        for i in range(45):  # 270 SKUs total
            sku_id += 1
            cat = random.choice(CATEGORIES)
            issue = today - dt.timedelta(days=random.randint(200, 3*365 - 30))
            expiry = issue + dt.timedelta(days=3*365)
            dte = (expiry - today).days
            vol = random.randint(2_000, 480_000)
            rev = round(vol * random.uniform(28_000, 260_000) / 1_000_000, 1)
            gm = round(random.uniform(-4, 62), 1)
            dossier = random.choices(["Ready","Gap","Missing"], weights=[65,25,10])[0]
            halal_no = f"LPPOM-00{random.randint(100000,999999)}-{random.randint(20,25)}"
            halal_exp = today + dt.timedelta(days=random.randint(-90, 700))
            jakim_no = "" if random.random() < 0.6 else f"JAKIM/MY/{random.randint(1000,9999)}/{random.randint(23,25)}"
            jakim_exp = "" if not jakim_no else today + dt.timedelta(days=random.randint(-30, 900))
            if dte < 0:
                flag = "EXPIRED — DO NOT SELL"
            elif dte < 60 and dossier != "Ready":
                flag = "STOCK-OUT RISK"
            elif dte < 90:
                flag = "URGENT RENEWAL"
            elif dte < 180:
                flag = "PLAN RENEWAL"
            else:
                flag = "OK"
            priority = round((vol * (gm if gm > 0 else 1)) / 1000, 0)
            rows.append([
                f"CBI-{sku_id}",
                f"{house} {cat.split()[0]} {random.choice(['Perfect','Bright','Matte','Glow','Hydra','Ultra','Pro'])} {random.choice(['01','02','03','04','05'])}",
                house, cat,
                random.choice(["15 ml","30 ml","50 ml","75 ml","100 ml","150 ml","200 ml"]),
                f"NA{random.randint(18,45)}{random.randint(100000000,999999999)}",
                issue, expiry, dte, vol, rev, gm,
                halal_no, halal_exp, jakim_no, jakim_exp,
                random.choice(["Jababeka 1","Jababeka 2","Cikande","Bekasi Plant 4"]),
                "Yes" if dossier != "Missing" else "No",
                "On File" if dossier == "Ready" else ("Draft" if dossier == "Gap" else "Missing"),
                "On File" if dossier == "Ready" else random.choice(["Draft","Pending","Missing"]),
                "On File" if dossier == "Ready" else random.choice(["Draft","Pending"]),
                dossier, priority, flag,
                random.choice(["Ani Ratnasari","Dinda Prawira","Rezky Hakim","Meilani Utari","Fajar Wibowo"])
            ])

    # Two forced hydroquinone recall SKUs on Wardah Perfect Bright (matches the scenario)
    rows[0] = ["CBI-WPB-101","Wardah Perfect Bright Serum 30ml","Wardah","Skincare","30 ml",
               "NA18240611234","2023-08-01","2026-08-01",16,412_000,68.5,52.3,
               "LPPOM-00456789-24", today + dt.timedelta(days=180), "JAKIM/MY/4521/24", today + dt.timedelta(days=240),
               "Jababeka 1","Yes","Under Review","On File","Draft","Gap",21540,
               "STOCK-OUT RISK — HYDROQUINONE RECALL", "Ani Ratnasari"]
    rows[1] = ["CBI-WPB-102","Wardah Perfect Bright Night Cream 50ml","Wardah","Skincare","50 ml",
               "NA18240611235","2023-08-01","2026-08-01",16,368_000,71.2,54.8,
               "LPPOM-00456790-24", today + dt.timedelta(days=180), "JAKIM/MY/4522/24", today + dt.timedelta(days=240),
               "Jababeka 1","Yes","Under Review","On File","Draft","Gap",20180,
               "STOCK-OUT RISK — HYDROQUINONE RECALL", "Ani Ratnasari"]

    for r in rows:
        ws.append(r)
    autosize(ws)

    # Summary sheet
    ws2 = wb.create_sheet("Summary by Brand House")
    ws2.append(["Brand House","SKUs","SKUs OK","Urgent (<90d)","Stock-Out Risk","Expired",
                "12M Revenue (Rp M)","Avg Gross Margin %"])
    style_header(ws2, 8)
    for house in BRAND_HOUSES:
        h = [r for r in rows if r[2] == house]
        ok = sum(1 for r in h if r[23] == "OK")
        urg = sum(1 for r in h if "URGENT" in str(r[23]) or "PLAN" in str(r[23]))
        stk = sum(1 for r in h if "STOCK-OUT" in str(r[23]))
        exp = sum(1 for r in h if "EXPIRED" in str(r[23]))
        rev = round(sum(r[10] for r in h),1)
        gm  = round(sum(r[11] for r in h)/len(h),1) if h else 0
        ws2.append([house, len(h), ok, urg, stk, exp, rev, gm])
    autosize(ws2)

    ws3 = wb.create_sheet("Notes")
    ws3.append(["Contoso Beauty Innovation — BPOM Notifikasi Kosmetik Tracker"])
    ws3.append(["Refresh cadence: monthly (1st of each month)"])
    ws3.append(["Regulator: Badan Pengawas Obat dan Makanan (BPOM), notifkos.pom.go.id"])
    ws3.append(["Halal: LPPOM MUI (ID) + JAKIM (MY) tracked separately from Notifikasi"])
    ws3.append(["Owner: Group Head of Regulatory & Halal Compliance"])
    ws3.append(["Note: hydroquinone-recall SKUs (CBI-WPB-101 / 102) require PJT sign-off before any release action"])
    ws3.column_dimensions['A'].width = 90

    path = os.path.join(OUT, "COSMETICS_01_BPOM_Notifikasi_Tracker.xlsx")
    wb.save(path); return path

# ============================================================
# FILE 2 — SKU Margin + Live-Shopping Tracker
# ============================================================
def gen_sku_margin():
    wb = Workbook()
    ws = wb.active; ws.title = "SKU x Creator GMV"
    cols = ["Creator","Creator Tier","Platform","Brand House","SKU","Category",
            "Live Sessions (Q4)","GMV (Rp M)","Commission %","Commission (Rp M)",
            "Ad Spend (Rp M)","Return Rate %","Refunds (Rp M)","Net Margin %",
            "Fulfilment Cost (Rp M)","Sell-Through %","Repeat Purchase %",
            "Score","Recommendation","Target Commission %","Notes"]
    ws.append(cols); style_header(ws, len(cols))

    creators = []
    for i in range(140):
        tier = random.choices(["Nano","Micro","Mid","Macro"], weights=[35,35,20,10])[0]
        creators.append((f"@{random.choice(['bella','riri','nadia','tasya','icha','yuki','kevin','dita','shaza','mira','tio','ayla','farah','lulu','indra'])}_{random.randint(10,999)}", tier))

    def commission_range(tier):
        return {"Nano":(5,10),"Micro":(8,15),"Mid":(12,22),"Macro":(18,32)}[tier]

    rows = []
    for c, tier in creators:
        n_skus = random.randint(1, 4)
        for _ in range(n_skus):
            house = random.choice(BRAND_HOUSES)
            cat = random.choice(CATEGORIES)
            platform = random.choices(["TikTok Shop","Shopee Live"], weights=[65,35])[0]
            gmv = round({"Nano":random.uniform(8,80),"Micro":random.uniform(40,320),
                         "Mid":random.uniform(180,1200),"Macro":random.uniform(600,4200)}[tier], 1)
            comm_pct = round(random.uniform(*commission_range(tier)), 1)
            comm = round(gmv * comm_pct/100, 1)
            ad = round(gmv * random.uniform(0.04, 0.18), 1)
            rr  = round(random.uniform(4, 38), 1)
            ref = round(gmv * rr/100, 1)
            fulf = round(gmv * random.uniform(0.06, 0.14), 1)
            revenue_kept = gmv - comm - ad - ref - fulf
            net_pct = round(revenue_kept / gmv * 100 if gmv else 0, 1)
            sellt = round(random.uniform(55, 98), 1)
            rep = round(random.uniform(6, 42), 1)
            if net_pct >= 12:
                rec = "KEEP"; target = comm_pct
            elif net_pct >= 4:
                rec = "RATIONALISE"
                target = max(commission_range(tier)[0], round(comm_pct - random.uniform(3,7),1))
            else:
                rec = "CUT"; target = 0
            if rr > 25:
                rec = "CUT"; target = 0
            note = ""
            if "CUT" in rec: note = "Return-rate or margin below threshold"
            elif "RATIONALISE" in rec: note = f"Renegotiate to {target}% commission band"
            else: note = "Retain in FY2026 roster"
            score = round(net_pct * 0.6 + sellt * 0.2 + rep * 0.2, 1)
            rows.append([c, tier, platform, house,
                         f"CBI-{random.randint(1001,1270)}",
                         cat, random.randint(2, 42), gmv, comm_pct, comm, ad,
                         rr, ref, net_pct, fulf, sellt, rep, score, rec, target, note])

    for r in rows: ws.append(r)
    autosize(ws)

    # Summary
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric","Value"]); style_header(ws2, 2)
    total_gmv = round(sum(r[7] for r in rows),1)
    total_comm = round(sum(r[9] for r in rows),1)
    keep = sum(1 for r in rows if r[18]=="KEEP")
    rat  = sum(1 for r in rows if r[18]=="RATIONALISE")
    cut  = sum(1 for r in rows if r[18]=="CUT")
    ws2.append(["Total Q4 live-shopping GMV (Rp M)", total_gmv])
    ws2.append(["Total creator commission paid (Rp M)", total_comm])
    ws2.append(["Trade-spend overrun vs budget (Rp M)", 60_000])
    ws2.append(["Creators — KEEP", keep])
    ws2.append(["Creators — RATIONALISE", rat])
    ws2.append(["Creators — CUT", cut])
    ws2.append(["Rationalise share of roster", f"{round(rat/(keep+rat+cut)*100,1)}%"])
    ws2.append(["Cut share of roster", f"{round(cut/(keep+rat+cut)*100,1)}%"])
    autosize(ws2)

    # Brand-house cut
    ws3 = wb.create_sheet("Brand House Cut")
    ws3.append(["Brand House","Rows","GMV (Rp M)","Commission (Rp M)","Avg Net Margin %","Avg Return Rate %"])
    style_header(ws3, 6)
    for h in BRAND_HOUSES:
        subset = [r for r in rows if r[3] == h]
        ws3.append([h, len(subset),
                    round(sum(r[7] for r in subset),1),
                    round(sum(r[9] for r in subset),1),
                    round(sum(r[13] for r in subset)/len(subset),1) if subset else 0,
                    round(sum(r[11] for r in subset)/len(subset),1) if subset else 0])
    autosize(ws3)

    path = os.path.join(OUT, "COSMETICS_02_SKU_Margin_Tracker.xlsx")
    wb.save(path); return path

# ============================================================
# FILE 3 — Ingredient Recall Programme (Word SOP)
# ============================================================
def gen_recall_sop():
    doc = Document()
    # margins
    for section in doc.sections:
        section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)

    t = doc.add_heading("Contoso Beauty Innovation", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ingredient Contamination — Recall Response Programme (SOP RA-CR-014)"); r.bold = True; r.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Version 4.2 · Effective 1 April 2026 · Owner: Group Head of Regulatory & Halal Compliance"); r2.italic = True

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This SOP governs Contoso Beauty Innovation's response to any confirmed or suspected ingredient "
        "contamination event affecting a marketed cosmetics SKU across the six brand houses (Wardah, Make Over, "
        "Emina, Kahf, Instaperfect, OMG!). It is the master document invoked whenever BPOM (Indonesia), NPRA/KKM "
        "(Malaysia), HSA (Singapore) or US MoCRA authorities issue a notice, or when internal QC identifies an "
        "out-of-specification (OOS) result on a released batch. The programme aligns to BPOM Peraturan Nomor 12 "
        "Tahun 2020 on cosmetics recall and to LPPOM MUI's HAS 23000 halal-integrity obligations."
    )

    doc.add_heading("2. Scope & Trigger Events", level=1)
    doc.add_paragraph(
        "The SOP is triggered by any one of: (a) a BPOM / NPRA / HSA / US FDA notice referencing a Contoso Beauty "
        "Innovation SKU; (b) an internal QC finding of a Schedule-2 or Schedule-3 prohibited ingredient (including "
        "hydroquinone > 0.02%, mercury, lead > 20 ppm, dexamethasone, corticosteroids); (c) a validated consumer "
        "complaint alleging adverse reaction with medical evidence; (d) a third-party lab report circulated by "
        "media or influencers naming a Contoso Beauty SKU."
    )

    doc.add_heading("3. Crisis War-Room Composition", level=1)
    table = doc.add_table(rows=1, cols=2); table.style = "Light Grid Accent 5"
    hdr = table.rows[0].cells; hdr[0].text = "Role"; hdr[1].text = "Named Person"
    war_room = [
        ("Group CEO (chair)", "Vania Ardhani"),
        ("Group Head of Regulatory & Halal Compliance", "Dr. Retno Kusumawati"),
        ("Penanggung Jawab Teknis (PJT)", "Apt. Widiastuti Prasetyo, S.Farm"),
        ("Group Head of Quality", "Bramantyo Nugraha"),
        ("Group Head of Corporate Affairs", "Sanya Halim"),
        ("Group General Counsel", "Rangga Sanjaya, S.H."),
        ("Head of Retail Partnerships", "Ferdinand Wibisono"),
        ("Head of Digital & Live Commerce", "Amanda Sudirjo"),
        ("Wardah GM", GMs["Wardah"]),
        ("Make Over GM", GMs["Make Over"]),
        ("Emina GM", GMs["Emina"]),
        ("Kahf GM", GMs["Kahf"]),
        ("Instaperfect GM", GMs["Instaperfect"]),
        ("OMG! GM", GMs["OMG!"]),
    ]
    for role, person in war_room:
        row = table.add_row().cells; row[0].text = role; row[1].text = person

    doc.add_heading("4. 72-Hour Response Timeline", level=1)
    doc.add_paragraph("The 72-hour clock starts at T0 = time of BPOM notice receipt or internal OOS confirmation, whichever is earlier.")
    tl = [
        ("T + 0h to T + 2h", "PJT + Head of Quality confirm trigger. Group Head of R&HC convenes the war-room. Batch codes are frozen in SAP; further release blocked."),
        ("T + 2h to T + 8h", "Retail partnerships issues the withdrawal notice to Sociolla, Watsons, Guardian, Alfamart, Indomaret, Watsons MY, Guardian MY, Watsons SG, Sephora SEA. Live-commerce lead pauses the two affected SKUs on TikTok Shop and Shopee Live and notifies all fronting creators."),
        ("T + 8h to T + 24h", "PJT drafts and sends the formal reply to BPOM (and NPRA / HSA / US FDA as applicable). Root-cause investigation initiated at the manufacturing site. Consumer FAQ published on brand-house websites."),
        ("T + 24h to T + 48h", "Board update circulated. First reverse-logistics wave in market. Internal Teams post to all six brand houses. Creator brief re-issued with updated Q&A."),
        ("T + 48h to T + 72h", "Interim root-cause report ready for BPOM. Media statement pre-cleared. Halal-integrity assessment (LPPOM MUI) if the trigger touches a halal-certified SKU."),
    ]
    tl_tbl = doc.add_table(rows=1, cols=2); tl_tbl.style = "Light List Accent 5"
    hdr = tl_tbl.rows[0].cells; hdr[0].text = "Window"; hdr[1].text = "Action"
    for w, a in tl:
        row = tl_tbl.add_row().cells; row[0].text = w; row[1].text = a

    doc.add_heading("5. Communication Cascade — Named Recipients", level=1)
    doc.add_paragraph(
        "All external communication is signed by the PJT (regulator-facing) or Group Head of Corporate Affairs "
        "(media-facing). Internal communication is signed by the Group Head of Regulatory & Halal Compliance. "
        "No brand-house GM may issue an independent statement without war-room clearance."
    )
    for line in [
        "Regulator: BPOM (Deputi 2 — Bidang Pengawasan Obat Tradisional, Suplemen Kesehatan dan Kosmetik), NPRA / KKM, HSA Singapore, US FDA (MoCRA-registered contact).",
        "Retail: Sociolla (Beauty Category Head), Watsons Indonesia (Buying), Guardian Indonesia, Alfamart Commercial, Indomaret Commercial, Watsons MY, Guardian MY, Watsons SG, Sephora SEA.",
        "Digital: TikTok Shop ID / MY Category Manager, Shopee Live Beauty Vertical Lead, Tokopedia Beauty.",
        "Creators: the ~40 macro and mid creators currently fronting the affected SKUs (list maintained in COSMETICS_02).",
        "Consumers: brand-house customer-care team + the 24/7 hotline (0800-1-CONTOSO-BI)."
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("6. Reverse Logistics & Financial Reserve", level=1)
    doc.add_paragraph(
        "The reverse-logistics playbook designates PT Sinar Logistik Utama and PT Jaya Prima Express as the two "
        "pre-authorised 3PL partners for recall pickup. Retailers issue return-goods notes referencing this SOP "
        "number. Financial reserve for a Class-1 recall is Rp 45 billion, ring-fenced in the Group balance sheet "
        "under contingency provisions (SAP GL 2451-000). Any recall event over Rp 60 billion requires Group CFO "
        "and Board Audit Committee escalation within 24 hours."
    )

    doc.add_heading("7. Halal Integrity Assessment", level=1)
    doc.add_paragraph(
        "For any recall touching a halal-certified SKU (LPPOM MUI or JAKIM), the SJH (Sistem Jaminan Halal) "
        "internal auditor must confirm within 48 hours that (a) the contamination pathway does not compromise "
        "halal integrity for other SKUs sharing the same production line and (b) no cross-contamination has "
        "occurred. If either fails, LPPOM MUI must be notified in parallel with BPOM."
    )

    doc.add_heading("8. Language, Register & Off-Topic List", level=1)
    doc.add_paragraph(
        "All BPOM-facing correspondence is in Bahasa Indonesia in formal regulatory register (Menimbang / "
        "Mengingat / Menetapkan structure where applicable). Consumer-facing communication is plain Bahasa "
        "Indonesia and English (for Singapore and export markets), no legal jargon. Creators and staff are "
        "explicitly instructed not to (i) speculate on root cause; (ii) name suppliers; (iii) discuss internal "
        "batch-release procedures; (iv) offer opinions on regulator behaviour; (v) engage with media without "
        "clearance from Group Corporate Affairs."
    )

    doc.add_heading("9. Post-Recall Review", level=1)
    doc.add_paragraph(
        "Within 30 days of recall closure, the Group Head of R&HC chairs a lessons-learned session with the "
        "war-room and issues an updated CAPA (corrective and preventive action) plan to the Board Audit "
        "Committee. Metrics tracked: time to first BPOM reply, retailer response coverage %, creator response "
        "coverage %, cumulative financial impact, consumer sentiment recovery arc (12-month), and repeat-batch "
        "OOS incidence."
    )

    doc.add_heading("10. Annex A — Batch-Freeze Runbook (SAP)", level=1)
    doc.add_paragraph(
        "Transaction MSC2N — set deletion flag on affected batch numbers. Transaction VL06O — block outbound "
        "deliveries. Transaction MB1B — move on-hand stock from Unrestricted (101) to Blocked (150). PJT "
        "signs off. Backup: Head of Quality has parallel authorisation."
    )

    path = os.path.join(OUT, "COSMETICS_03_Ingredient_Recall_Programme.docx")
    doc.save(path); return path

# ============================================================
# FILE 4 — TiO2 + Surfactant Hedge Book
# ============================================================
def gen_hedge_book():
    wb = Workbook()
    ws = wb.active; ws.title = "Hedge Positions"
    cols = ["Position ID","Ingredient","Supplier","Contract Type","Notional (kg)",
            "Strike Price (USD/kg)","Current Spot (USD/kg)","Trade Date","Expiry",
            "MTM (USD)","Coverage % of FY2026 Demand","Delta vs FY2025 Cost %",
            "Counterparty Rating","Counterparty Country","Status","Owner"]
    ws.append(cols); style_header(ws, len(cols))
    today = dt.date(2026, 7, 16)
    ingredients = [
        ("Titanium Dioxide (TiO2, cosmetic grade)", 6.20, 7.32),
        ("Palm-derived Sodium Lauryl Sulphate (SLS)", 1.85, 2.18),
        ("Palm-derived Sodium Laureth Sulphate (SLES)", 2.10, 2.48),
        ("Cocamidopropyl Betaine (CAPB)", 2.65, 3.05),
        ("Cetearyl Alcohol (palm-derived)", 3.10, 3.62),
        ("Glycerin (palm-derived, USP)", 1.45, 1.72),
    ]
    suppliers = [("Kronos Worldwide","BBB+","US"),
                 ("Venator Materials","BBB","UK"),
                 ("PT Sinar Mas Chemical","BBB-","ID"),
                 ("KLK Oleo","A-","MY"),
                 ("Wilmar Oleochemicals","A","SG"),
                 ("Musim Mas","BBB+","SG"),
                 ("BASF SE","A+","DE"),
                 ("Croda International","A-","UK")]
    rows = []
    pid = 500
    for ing, y2025, spot in ingredients:
        for _ in range(random.randint(6, 10)):
            pid += 1
            sup, rating, country = random.choice(suppliers)
            ct  = random.choices(["Forward","Cap","Collar","Physical Fwd"], weights=[45,20,15,20])[0]
            notional = random.choice([25_000, 50_000, 100_000, 200_000, 300_000])
            strike = round(y2025 * random.uniform(0.98, 1.12), 3)
            spot_v = round(spot * random.uniform(0.94, 1.06), 3)
            days_out = random.randint(30, 330)
            expiry = today + dt.timedelta(days=days_out)
            trade = today - dt.timedelta(days=random.randint(15, 200))
            mtm = round((spot_v - strike) * notional, 0)
            cov = round(random.uniform(6, 22), 1)
            delta = round((spot_v / y2025 - 1)*100, 1)
            status = "Open" if expiry >= today else "Closed"
            rows.append([f"HP-{pid}", ing, sup, ct, notional, strike, spot_v,
                         trade, expiry, mtm, cov, delta, rating, country, status,
                         random.choice(["Treasury Desk 1","Treasury Desk 2","Group Procurement"])])
    for r in rows: ws.append(r)
    autosize(ws)

    # Coverage summary
    ws2 = wb.create_sheet("FY2026 Coverage")
    ws2.append(["Ingredient","FY2026 Demand (kg)","Hedged Coverage %","Open MTM (USD)","Cost Volatility (12M range %)","Board Guardrail"])
    style_header(ws2, 6)
    for ing, y2025, spot in ingredients:
        subset = [r for r in rows if r[1] == ing and r[14]=="Open"]
        cov = round(sum(r[10] for r in subset), 1)
        mtm = round(sum(r[9] for r in subset), 0)
        ws2.append([ing,
                    random.choice([2_400_000, 1_800_000, 3_200_000, 640_000, 850_000, 1_100_000]),
                    cov, mtm,
                    f"+{round(random.uniform(12,28),1)}%",
                    ">= 55% hedged by Q2, >= 70% by Q4"])
    autosize(ws2)

    # Guardrails sheet
    ws3 = wb.create_sheet("Hedging Guardrails")
    ws3.append(["Guardrail","Value"]); style_header(ws3, 2)
    gr = [
        ("Max single-supplier concentration", "≤ 35% of ingredient hedged volume"),
        ("Max unhedged spot exposure", "≤ 25% of FY2026 forecast demand"),
        ("Minimum counterparty rating", "BBB- (S&P) or equivalent"),
        ("Escalation threshold (MTM loss)", "USD 2.5m per ingredient — CFO sign-off"),
        ("Board reporting cadence", "Quarterly, in the FY board pack"),
        ("Palm-derived sustainability", "100% RSPO Mass Balance minimum; Segregated preferred"),
        ("TiO2 EU labelling","Compliant with EU 2022/63 (removed as food additive; cosmetics permitted with restrictions)"),
    ]
    for a,b in gr: ws3.append([a,b])
    autosize(ws3)

    path = os.path.join(OUT, "COSMETICS_04_TiO2_Surfactant_Hedge_Book.xlsx")
    wb.save(path); return path

# ============================================================
# FILE 5 — FY2026 Live-Shopping Guardrails (Word policy)
# ============================================================
def gen_live_guardrails():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
    doc.add_heading("Contoso Beauty Innovation", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FY2026 Live-Shopping & Creator-Affiliate Guardrails"); r.bold = True; r.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Approved by Group Brand Council · Effective 1 August 2026 · Owner: Head of Digital & Live Commerce (Amanda Sudirjo)"); r2.italic = True

    doc.add_heading("1. Context", level=1)
    doc.add_paragraph(
        "Q4 FY2025 saw live-shopping trade-spend on TikTok Shop and Shopee Live overrun by Rp 60 billion "
        "across the six brand houses, driven by aggressive creator-affiliate commissions, unbudgeted platform "
        "co-invest, and unmanaged return rates on colour-cosmetics SKUs. This document sets the FY2026 "
        "operating envelope. It applies to all live-shopping activity across Wardah, Make Over, Emina, Kahf, "
        "Instaperfect, OMG! on TikTok Shop ID/MY, Shopee Live ID/MY, and Tokopedia Beauty Live."
    )

    doc.add_heading("2. Commission Bands (Non-Negotiable Ceilings)", level=1)
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 5"
    for i, h in enumerate(["Creator Tier","Followers","Max Commission %","Additional Ad-Boost %"]):
        t.rows[0].cells[i].text = h
    for row in [("Nano","1k–10k","8%","≤ 2%"),
                ("Micro","10k–100k","12%","≤ 3%"),
                ("Mid","100k–500k","18%","≤ 5%"),
                ("Macro","500k–5m","25%","≤ 7%"),
                ("Mega / Celebrity","5m+","Bespoke — GM + CFO sign-off","≤ 10%")]:
        r_ = t.add_row().cells
        for i, v in enumerate(row): r_[i].text = v

    doc.add_heading("3. ROI Gates per Session", level=1)
    doc.add_paragraph(
        "Every scheduled live-shopping session must clear these gates before booking: (a) expected GMV ≥ "
        "Rp 40m for TikTok Shop and Rp 25m for Shopee Live; (b) expected net margin after commission, ad, "
        "return, fulfilment ≥ 12% (macro/celebrity ≥ 8% permitted for launch windows only); (c) return-rate "
        "assumption ≤ 20% modelled off the same SKU's trailing 90-day baseline; (d) BPOM Notifikasi and "
        "LPPOM MUI cert both valid for the entire 30-day window post-live. Sessions that miss any gate "
        "require Head of Digital & Live Commerce sign-off in writing."
    )

    doc.add_heading("4. SKU Eligibility", level=1)
    for line in [
        "SKUs on BPOM Notifikasi within 90 days of expiry: EXCLUDED — no live-shopping activity.",
        "SKUs with dossier readiness 'Missing' or 'Gap' on COSMETICS_01: EXCLUDED.",
        "SKUs in active recall investigation: EXCLUDED — hard block in SAP + platform SKU catalogue.",
        "Halal-branded SKUs (Wardah, Emina, Kahf halal lines): must display valid LPPOM MUI cert on the live overlay.",
        "New-launch SKUs: eligible only in the first 60 days post-launch with Brand-House GM co-sign.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("5. Creator Vetting", level=1)
    doc.add_paragraph(
        "Every macro and mid creator must complete the annual Contoso Beauty Innovation creator induction "
        "before their first FY2026 session. Induction covers: BPOM cosmetics claim rules (no medicinal / "
        "curative claims), halal-integrity language, off-topic list during recall events, and the platform's "
        "own live-selling policies. Creators who breach twice in a rolling 12 months are removed from the "
        "roster and reported to the platform trust & safety team."
    )

    doc.add_heading("6. Off-Limits Claim Language", level=1)
    for line in [
        "No 'medical', 'cure', 'treat', 'obat' or equivalent language.",
        "No comparative claims naming competitor brands (Mustika Ratu, L\u2019Oréal Indonesia, Mandom, Unilever, etc.).",
        "No unsubstantiated 'BPOM-approved' language — BPOM notifies, does not approve.",
        "No halal claim without a valid LPPOM MUI or JAKIM certificate number displayed.",
        "No skin-lightening claim on any SKU that has not passed the enhanced safety-assessment protocol post the FY2025 hydroquinone review.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("7. Budget Envelope & Escalation", level=1)
    doc.add_paragraph(
        "Total FY2026 live-shopping trade spend is capped at Rp 320 billion (down from Rp 380 billion FY2025 "
        "run-rate). Brand-house sub-caps: Wardah Rp 110 bn, Make Over Rp 60 bn, Emina Rp 45 bn, Kahf Rp 40 bn, "
        "Instaperfect Rp 35 bn, OMG! Rp 30 bn. Any overrun ≥ 5% requires Group CFO written approval; ≥ 10% "
        "requires Group Brand Council escalation."
    )

    doc.add_heading("8. Reporting Cadence", level=1)
    for line in [
        "Weekly: platform-level GMV, commission, return-rate dashboard to Brand Council.",
        "Monthly: creator-level ROI cut using COSMETICS_02 methodology.",
        "Quarterly: Board-level summary in the FY board pack.",
        "Ad-hoc: any single-session net-margin miss ≥ Rp 500m reported to Head of Digital & Live Commerce within 24 hours.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("9. Named Approvers", level=1)
    t2 = doc.add_table(rows=1, cols=2); t2.style = "Light List Accent 5"
    t2.rows[0].cells[0].text = "Decision"; t2.rows[0].cells[1].text = "Approver"
    for a, b in [
        ("Roster additions (macro creator)", "Brand-House GM + Head of Digital & Live Commerce"),
        ("Commission-band exceptions", "Head of Digital & Live Commerce + Group CFO"),
        ("Recall-window pause on any SKU", "Group Head of Regulatory & Halal Compliance"),
        ("Total budget overrun ≥ 5%", "Group CFO"),
        ("Total budget overrun ≥ 10%", "Group Brand Council (chaired by Group CEO Vania Ardhani)"),
    ]:
        row = t2.add_row().cells; row[0].text = a; row[1].text = b

    doc.add_heading("10. Review Cycle", level=1)
    doc.add_paragraph(
        "This document is reviewed quarterly by the Head of Digital & Live Commerce and re-approved annually "
        "by the Group Brand Council. Interim amendments require Group Head of Regulatory & Halal Compliance "
        "concurrence for any change touching claim language, SKU eligibility, or halal-integrity provisions."
    )

    path = os.path.join(OUT, "COSMETICS_05_FY2026_Live_Shopping_Guardrails.docx")
    doc.save(path); return path

# ============================================================
# RUN
# ============================================================
if __name__ == "__main__":
    for fn in [gen_bpom_tracker, gen_sku_margin, gen_recall_sop, gen_hedge_book, gen_live_guardrails]:
        p = fn(); print("wrote", p)
